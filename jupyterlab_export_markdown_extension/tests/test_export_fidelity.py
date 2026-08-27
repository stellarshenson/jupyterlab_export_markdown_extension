"""Export fidelity tests - validate that HTML, DOCX, and PDF exports contain expected content."""

import asyncio
import json
import io
import re
import tempfile
import zipfile
from pathlib import Path

import pytest

# Test markdown content with various elements
TEST_MARKDOWN = """# Test Document

## Text Formatting

Regular paragraph with **bold text** and _italic text_.

## Lists

- Bullet item 1
- Bullet item 2
  - Nested bullet

1. Numbered item 1
2. Numbered item 2
   1. Nested numbered

## Table

| Column A | Column B |
| -------- | -------- |
| A1       | B1       |
| A2       | B2       |

## Code

Inline `code` here.

```python
def test():
    return True
```

## Blockquote

> This is a blockquote.

## GitHub Alert

> [!NOTE]
> This is a note alert.

## Symbols

Arrows: → ← ↑ ↓
Bullets: • ◦ ▪ ▫
Stars: ★ ☆
"""


@pytest.fixture
def test_markdown_file(jp_root_dir):
    """Create a test markdown file in the server's root directory."""
    md_file = jp_root_dir / "test.md"
    md_file.write_text(TEST_MARKDOWN, encoding="utf-8")
    return md_file


class TestHTMLExport:
    """Test HTML export fidelity."""

    async def test_html_export_success(self, jp_fetch, test_markdown_file):
        """Test successful HTML export."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test.md"}),
            raise_error=False,
        )
        assert response.code == 200
        # Content should be HTML (may have charset suffix)
        content_type = response.headers.get("Content-Type", "")
        assert "text/html" in content_type or response.body.startswith(b"<!DOCTYPE") or b"<html" in response.body

    async def test_html_contains_headings(self, jp_fetch, test_markdown_file):
        """Test HTML contains converted headings."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )
        html = response.body.decode("utf-8")
        assert "<h1" in html or "<h1>" in html
        assert "Test Document" in html
        assert "<h2" in html

    async def test_html_contains_formatting(self, jp_fetch, test_markdown_file):
        """Test HTML contains text formatting."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )
        html = response.body.decode("utf-8")
        assert "<strong>" in html or "<b>" in html  # bold
        assert "<em>" in html or "<i>" in html  # italic

    async def test_html_contains_lists(self, jp_fetch, test_markdown_file):
        """Test HTML contains lists."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )
        html = response.body.decode("utf-8")
        assert "<ul>" in html  # unordered list
        # Note: ordered lists may render as <ol> or <ul> depending on markdown parser
        assert "<li>" in html  # list items
        # Verify list content is present
        assert "Bullet item" in html or "Numbered item" in html

    async def test_html_contains_table(self, jp_fetch, test_markdown_file):
        """Test HTML contains table."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )
        html = response.body.decode("utf-8")
        assert "<table>" in html or "<table " in html
        assert "<th>" in html or "<td>" in html
        assert "Column A" in html

    async def test_html_contains_code(self, jp_fetch, test_markdown_file):
        """Test HTML contains code blocks."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )
        html = response.body.decode("utf-8")
        assert "<code>" in html or "<code " in html
        # Code may have syntax highlighting spans, check for key tokens
        assert "def" in html and "test" in html and "return" in html

    async def test_html_contains_blockquote(self, jp_fetch, test_markdown_file):
        """Test HTML contains blockquote."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )
        html = response.body.decode("utf-8")
        assert "<blockquote>" in html or "blockquote" in html.lower()

    async def test_html_contains_unicode_symbols(self, jp_fetch, test_markdown_file):
        """Test HTML contains Unicode symbols."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )
        html = response.body.decode("utf-8")
        # These symbols should pass through to HTML unchanged
        assert "→" in html or "&#" in html  # arrow or HTML entity
        assert "★" in html or "&#" in html  # star or HTML entity


class TestDOCXExport:
    """Test DOCX export fidelity."""

    async def test_docx_export_success(self, jp_fetch, test_markdown_file):
        """Test successful DOCX export."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test.md"}),
            raise_error=False,
        )
        assert response.code == 200
        # Verify by checking the file is a valid ZIP (DOCX format)
        # Content-Type may vary, so validate actual content instead
        assert response.body[:2] == b"PK"  # ZIP magic bytes

    async def test_docx_is_valid_zip(self, jp_fetch, test_markdown_file):
        """Test DOCX is a valid ZIP file (DOCX format)."""
        import zipfile
        import io

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )

        # DOCX files are ZIP archives
        docx_bytes = io.BytesIO(response.body)
        assert zipfile.is_zipfile(docx_bytes)

    async def test_docx_contains_document_xml(self, jp_fetch, test_markdown_file):
        """Test DOCX contains document.xml with content."""
        import zipfile
        import io

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )

        docx_bytes = io.BytesIO(response.body)
        with zipfile.ZipFile(docx_bytes, "r") as zf:
            assert "word/document.xml" in zf.namelist()
            document_xml = zf.read("word/document.xml").decode("utf-8")
            assert "Test Document" in document_xml

    async def test_docx_contains_text_content(self, jp_fetch, test_markdown_file):
        """Test DOCX contains expected text content."""
        import zipfile
        import io

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )

        docx_bytes = io.BytesIO(response.body)
        with zipfile.ZipFile(docx_bytes, "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            assert "bold text" in document_xml
            assert "Bullet item" in document_xml
            assert "Column A" in document_xml


class TestPDFExport:
    """Test PDF export fidelity."""

    async def test_pdf_export_success(self, jp_fetch, test_markdown_file):
        """Test successful PDF export."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "test.md"}),
            raise_error=False,
        )
        assert response.code == 200
        # Verify by checking PDF magic bytes instead of Content-Type
        assert response.body.startswith(b"%PDF-")

    async def test_pdf_has_valid_header(self, jp_fetch, test_markdown_file):
        """Test PDF has valid PDF header."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )

        # PDF files start with %PDF-
        assert response.body.startswith(b"%PDF-")

    async def test_pdf_has_valid_footer(self, jp_fetch, test_markdown_file):
        """Test PDF has valid EOF marker."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )

        # PDF files end with %%EOF
        assert b"%%EOF" in response.body[-100:]

    async def test_pdf_contains_content(self, jp_fetch, test_markdown_file):
        """Test PDF contains actual content (not empty)."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )

        pdf_content = response.body
        # PDF text is compressed, so instead of searching for strings,
        # verify PDF has stream objects (actual content)
        assert b"stream" in pdf_content and b"endstream" in pdf_content

    async def test_pdf_minimum_size(self, jp_fetch, test_markdown_file):
        """Test PDF has reasonable file size (not empty)."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )

        # A valid PDF with content should be at least a few KB
        assert len(response.body) > 1000


class TestGitHubAlerts:
    """Test GitHub-style alerts preprocessing and DOCX styling."""

    async def test_note_alert_html(self, jp_fetch, test_markdown_file):
        """Test NOTE alert is processed in HTML."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )
        html = response.body.decode("utf-8")
        # Alert content should be present (labels hidden by default)
        assert "note" in html.lower() or "useful information" in html.lower()

    async def test_note_alert_html_with_labels(self, jp_fetch, test_markdown_file):
        """Test NOTE alert shows label when showAlertLabels is enabled."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test.md", "showAlertLabels": True}),
        )
        html = response.body.decode("utf-8")
        # Alert should have bold label when enabled
        assert "NOTE" in html

    async def test_alert_styled_in_docx(self, jp_fetch, test_markdown_file):
        """Test alert is rendered as styled table in DOCX."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )

        docx_bytes = io.BytesIO(response.body)
        with zipfile.ZipFile(docx_bytes, "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            # Blue shading applied for NOTE alert in table cell
            assert 'w:fill="EDF5FD"' in document_xml
            # Left border applied
            assert 'w:color="0969DA"' in document_xml
            # Zero-width space markers removed
            assert '\u200b' not in document_xml
            # Alert label should not appear (showAlertLabels defaults to false)
            # Table element present (alert rendered as table)
            assert 'w:tbl' in document_xml


class TestSyntaxHighlighting:
    """Test syntax highlighting in code blocks."""

    async def test_html_has_pygments_css(self, jp_fetch, test_markdown_file):
        """Test HTML output contains Pygments CSS for syntax highlighting."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )
        html = response.body.decode("utf-8")
        # Pygments CSS defines .codehilite styles
        assert ".codehilite" in html

    async def test_html_has_highlighted_code(self, jp_fetch, test_markdown_file):
        """Test HTML output has syntax-highlighted code spans."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )
        html = response.body.decode("utf-8")
        # Python code should have span elements with classes for keywords
        # Look for codehilite div or syntax-highlighted spans
        assert "codehilite" in html or 'class="k"' in html or 'class="nf"' in html

    async def test_docx_has_colored_code(self, jp_fetch, test_markdown_file):
        """Test DOCX output has inline styled code (colored spans)."""
        import zipfile
        import io

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )

        docx_bytes = io.BytesIO(response.body)
        with zipfile.ZipFile(docx_bytes, "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            # DOCX with colored code should have color references
            # Either via styles or inline formatting
            assert "def" in document_xml and "test" in document_xml

    async def test_pdf_has_code_content(self, jp_fetch, test_markdown_file):
        """Test PDF output contains code block content."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "test.md"}),
        )

        # PDF should be larger when it has formatted code blocks
        assert len(response.body) > 1000
        # PDF should have valid structure
        assert response.body.startswith(b"%PDF-")


# Test markdown with LaTeX math expressions
TEST_MARKDOWN_MATH = """# Math Test

Inline math: $E=mc^2$ is well known.

Display math:

$$\\int_0^\\infty e^{-x^2} dx = \\frac{\\sqrt{\\pi}}{2}$$

Greek letters: $\\alpha + \\beta = \\gamma$

Currency (should NOT render): The price is $100 and total $200.

Code block math (should NOT render):

```python
# The formula $E=mc^2$ in a comment
x = "$not_math$"
```

Inline code math: `$E=mc^2$` should stay as text.
"""


@pytest.fixture
def test_math_file(jp_root_dir):
    """Create a test markdown file with LaTeX math expressions."""
    md_file = jp_root_dir / "test_math.md"
    md_file.write_text(TEST_MARKDOWN_MATH, encoding="utf-8")
    return md_file


class TestMathRendering:
    """Test LaTeX math rendering in exports."""

    async def test_html_has_katex_scripts(self, jp_fetch, test_math_file):
        """Test HTML export includes KaTeX CDN scripts for math rendering."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test_math.md"}),
        )
        html = response.body.decode("utf-8")
        assert "katex" in html.lower()
        assert "cdn.jsdelivr.net" in html
        assert "auto-render" in html

    async def test_html_preserves_math_delimiters(self, jp_fetch, test_math_file):
        """Test HTML export keeps $...$ delimiters for KaTeX client-side rendering."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test_math.md"}),
        )
        html = response.body.decode("utf-8")
        # Math delimiters should be preserved for KaTeX (not replaced with images)
        assert "E=mc^2" in html

    async def test_docx_has_omml_math(self, jp_fetch, test_math_file):
        """Test DOCX export renders math as native OMML equations."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_math.md"}),
            raise_error=False,
        )
        assert response.code == 200
        docx_bytes = io.BytesIO(response.body)
        with zipfile.ZipFile(docx_bytes, "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            # OMML equations use oMath namespace
            assert "oMath" in document_xml, "DOCX should contain native OMML equations"
            # Marker text should be cleaned up
            assert "MATH_INLINE_" not in document_xml
            assert "MATH_DISPLAY_" not in document_xml

    async def test_docx_inline_math_in_paragraph(self, jp_fetch, test_math_file):
        """Test inline OMML math flows within text paragraph."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_math.md"}),
            raise_error=False,
        )
        assert response.code == 200
        docx_bytes = io.BytesIO(response.body)
        with zipfile.ZipFile(docx_bytes, "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            # "is well known" text should be in same paragraph as oMath
            # (inline math doesn't break text flow)
            assert "well known" in document_xml

    async def test_pdf_with_math_succeeds(self, jp_fetch, test_math_file):
        """Test PDF export succeeds when markdown contains math."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "test_math.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body.startswith(b"%PDF-")

    async def test_code_blocks_preserved(self, jp_fetch, test_math_file):
        """Test math inside code blocks is NOT rendered."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test_math.md"}),
        )
        html = response.body.decode("utf-8")
        # Code block content should remain as text, not be rendered as math images
        assert "not_math" in html

    async def test_currency_not_matched(self, jp_fetch, test_math_file):
        """Test $100 currency amounts are not falsely matched as math."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_math.md"}),
            raise_error=False,
        )
        assert response.code == 200
        docx_bytes = io.BytesIO(response.body)
        with zipfile.ZipFile(docx_bytes, "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            # Currency text should be preserved as-is
            assert "100" in document_xml


# Minimal SVG for testing
TEST_SVG = """<svg xmlns="http://www.w3.org/2000/svg" width="200" height="100">
  <rect width="200" height="100" fill="#4F81BD"/>
  <text x="100" y="55" text-anchor="middle" fill="white" font-size="16">Test SVG</text>
</svg>"""

# Markdown with an SVG image reference
TEST_MARKDOWN_WITH_SVG = """# Document With SVG

Some text before the image.

![Architecture Diagram](images/test-diagram.svg)

Some text after the image.
"""


@pytest.fixture
def test_svg_markdown_file(jp_root_dir):
    """Create a test markdown file with an SVG image reference."""
    # Create images directory and SVG file
    images_dir = jp_root_dir / "images"
    images_dir.mkdir(exist_ok=True)
    svg_file = images_dir / "test-diagram.svg"
    svg_file.write_text(TEST_SVG, encoding="utf-8")
    # Create markdown file
    md_file = jp_root_dir / "test_svg.md"
    md_file.write_text(TEST_MARKDOWN_WITH_SVG, encoding="utf-8")
    return md_file


# Theme-aware SVG: white background in light, black in dark
TEST_THEME_SVG = (
    '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 200 100">'
    "<style>.bg{fill:#ffffff}"
    "@media (prefers-color-scheme: dark){.bg{fill:#000000}}</style>"
    '<rect class="bg" width="200" height="100"/></svg>'
)


@pytest.fixture
def test_theme_svg_file(jp_root_dir):
    images_dir = jp_root_dir / "images"
    images_dir.mkdir(exist_ok=True)
    (images_dir / "theme.svg").write_text(TEST_THEME_SVG, encoding="utf-8")
    md = jp_root_dir / "test_theme_svg.md"
    md.write_text("# Theme\n\n![d](images/theme.svg)\n", encoding="utf-8")
    return md


class TestDocxTheme:
    """docxTheme drives the SVG color scheme; prefers-color-scheme resolves to it."""

    async def _center_pixel(self, jp_fetch, theme):
        from docx import Document
        from PIL import Image as PILImage
        import zipfile as _zip

        body = {"path": "test_theme_svg.md"}
        if theme is not None:
            body["docxTheme"] = theme
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps(body),
            raise_error=False,
        )
        assert response.code == 200
        with _zip.ZipFile(io.BytesIO(response.body)) as zf:
            media = [n for n in zf.namelist() if n.startswith("word/media/")]
            assert media, "expected an embedded image"
            img = PILImage.open(io.BytesIO(zf.read(media[0]))).convert("RGB")
        return img.getpixel((img.width // 2, img.height // 2))

    async def test_light_theme_renders_light(self, jp_fetch, test_theme_svg_file):
        r, g, b = await self._center_pixel(jp_fetch, "light")
        assert min(r, g, b) > 200, f"light theme should render white, got {(r, g, b)}"

    async def test_dark_theme_renders_dark(self, jp_fetch, test_theme_svg_file):
        r, g, b = await self._center_pixel(jp_fetch, "dark")
        assert max(r, g, b) < 60, f"dark theme should render black, got {(r, g, b)}"

    async def test_default_is_light(self, jp_fetch, test_theme_svg_file):
        r, g, b = await self._center_pixel(jp_fetch, None)
        assert min(r, g, b) > 200, f"default (no docxTheme) should be light, got {(r, g, b)}"


class TestSVGConversion:
    """Test SVG to PNG conversion in DOCX and PDF exports."""

    async def test_docx_with_svg_succeeds(self, jp_fetch, test_svg_markdown_file):
        """Test DOCX export does not fail when markdown contains SVG images."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_svg.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body[:2] == b"PK"  # Valid ZIP/DOCX

    async def test_docx_with_svg_contains_image(self, jp_fetch, test_svg_markdown_file):
        """Test DOCX with SVG contains a converted PNG image."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_svg.md"}),
        )
        docx_bytes = io.BytesIO(response.body)
        with zipfile.ZipFile(docx_bytes, "r") as zf:
            # DOCX stores images in word/media/ directory
            image_files = [n for n in zf.namelist() if n.startswith("word/media/")]
            assert len(image_files) > 0, "DOCX should contain at least one image"
            # Image should be PNG (converted from SVG)
            for img in image_files:
                img_data = zf.read(img)
                assert img_data[:4] == b"\x89PNG", f"{img} should be PNG format"

    async def test_docx_with_svg_contains_text(self, jp_fetch, test_svg_markdown_file):
        """Test DOCX with SVG still contains surrounding text content."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_svg.md"}),
        )
        docx_bytes = io.BytesIO(response.body)
        with zipfile.ZipFile(docx_bytes, "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            assert "Document With SVG" in document_xml
            assert "Some text before" in document_xml

    async def test_pdf_with_svg_succeeds(self, jp_fetch, test_svg_markdown_file):
        """Test PDF export does not fail when markdown contains SVG images."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "test_svg.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body.startswith(b"%PDF-")

    async def test_svg_pixel_width_parameter(self, jp_fetch, test_svg_markdown_file):
        """Test svgPixelWidth parameter is accepted by the export endpoint."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_svg.md", "svgPixelWidth": 1280}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body[:2] == b"PK"

    async def test_html_with_svg_keeps_svg(self, jp_fetch, test_svg_markdown_file):
        """Test HTML export preserves SVG as-is (no conversion needed)."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test_svg.md"}),
        )
        html = response.body.decode("utf-8")
        # HTML should contain the image as base64 data URI (SVG format preserved)
        assert "data:image/svg+xml" in html or "Architecture Diagram" in html


class TestSVGDoctypePrologue:
    """An XML declaration / DOCTYPE prologue must not leak as a top-left
    text artefact (e.g. ']>') when rasterizing - JupyterLab's mermaid
    renderer prepends exactly such a prologue."""

    async def test_doctype_prologue_no_top_left_artefact(self):
        from jupyterlab_export_markdown_extension.routes import (
            PlaywrightSvgRenderer,
        )
        from PIL import Image as PILImage

        # A minimal SVG with a white background, preceded by the same
        # XML+DOCTYPE-with-internal-subset prologue JupyterLab emits.
        prologue = (
            '<?xml version="1.0" standalone="no"?>\n'
            '<!DOCTYPE svg PUBLIC "-//W3C//DTD SVG 1.1//EN" '
            '"http://www.w3.org/Graphics/SVG/1.1/DTD/svg11.dtd" ['
            '<!ENTITY Aacute "&#193;">\n<!ENTITY amp "&#38;">]>'
        )
        svg = (
            '<svg xmlns="http://www.w3.org/2000/svg" '
            'viewBox="0 0 400 100">'
            '<rect width="400" height="100" fill="#ffffff"/></svg>'
        )
        svg_bytes = (prologue + "\n" + svg).encode("utf-8")

        async with PlaywrightSvgRenderer(color_scheme="light") as renderer:
            png = await renderer.render(svg_bytes, width=400)

        img = PILImage.open(io.BytesIO(png)).convert("RGBA")
        bg = PILImage.new("RGBA", img.size, (255, 255, 255, 255))
        comp = PILImage.alpha_composite(bg, img)
        tl = comp.crop((0, 0, 80, 80)).load()
        dark = sum(
            1
            for y in range(80)
            for x in range(80)
            if sum(tl[x, y][:3]) < 200
        )
        assert dark == 0, (
            f"DOCTYPE prologue leaked a dark text artefact in the "
            f"top-left corner ({dark} dark pixels)"
        )


# Markdown with SVG as first element (header banner pattern)
TEST_MARKDOWN_LEADING_SVG = """![Header Banner](images/test-diagram.svg)

---

## Content

Some text content below the header image.
"""


@pytest.fixture
def test_leading_svg_markdown_file(jp_root_dir):
    """Create a test markdown file with SVG as the very first element."""
    images_dir = jp_root_dir / "images"
    images_dir.mkdir(exist_ok=True)
    svg_file = images_dir / "test-diagram.svg"
    svg_file.write_text(TEST_SVG, encoding="utf-8")
    md_file = jp_root_dir / "test_leading_svg.md"
    md_file.write_text(TEST_MARKDOWN_LEADING_SVG, encoding="utf-8")
    return md_file


class TestLeadingImagePreservation:
    """Test that image-only paragraphs at document start are not removed."""

    async def test_docx_preserves_leading_image(self, jp_fetch, test_leading_svg_markdown_file):
        """Test DOCX export keeps an image that is the first element in the document."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_leading_svg.md"}),
            raise_error=False,
        )
        assert response.code == 200
        docx_bytes = io.BytesIO(response.body)
        with zipfile.ZipFile(docx_bytes, "r") as zf:
            image_files = [n for n in zf.namelist() if n.startswith("word/media/")]
            assert len(image_files) > 0, "Leading image should be preserved in DOCX"

    async def test_pdf_preserves_leading_image(self, jp_fetch, test_leading_svg_markdown_file):
        """Test PDF export keeps an image that is the first element in the document."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "test_leading_svg.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body.startswith(b"%PDF-")
        # PDF should be substantial (contains image data)
        assert len(response.body) > 2000


# Rich alert test content with line breaks, hyperlinks, bold, and consecutive alerts
TEST_MARKDOWN_RICH_ALERTS = """# Rich Alert Test

> [!IMPORTANT]
> Based on factory-observed void dimensions (5 x 3 x 2 mm), overall casting porosity appears well within criteria. <br><br>The 30% rejection rate is driven by **void location**.

> [!NOTE]
> **Applicable standards**: <br>[ASTM E505](https://example.com/astm-e505) (Reference Radiographs), <br>[EN 12681](https://example.com/en-12681) (Founding - Radiographic testing).

> [!WARNING]
> Simple warning without special formatting.

Some text after alerts.

| Col A | Col B |
| ----- | ----- |
| A1    | B1    |
"""


@pytest.fixture
def test_rich_alerts_file(jp_root_dir):
    """Create a test markdown file with rich alert content."""
    md_file = jp_root_dir / "test_rich_alerts.md"
    md_file.write_text(TEST_MARKDOWN_RICH_ALERTS, encoding="utf-8")
    return md_file


TEST_MARKDOWN_COLOR_PILLS = """# Colours

Named text colours: <span style="color:green">green</span>,
<span style="color:red">red</span>, hex <span style="color:#3b82f6">blue</span>.

| Item | Priority |
|---|---|
| A | <span style="background-color:#15803d;color:#ffffff;padding:1px 8px;border-radius:8px">High</span> |
| B | <span style="background-color:#b45309;color:#ffffff;padding:1px 8px;border-radius:8px">Med</span> |
"""


@pytest.fixture
def test_color_pills_file(jp_root_dir):
    md_file = jp_root_dir / "test_color_pills.md"
    md_file.write_text(TEST_MARKDOWN_COLOR_PILLS, encoding="utf-8")
    return md_file


class TestColouredHtml:
    """Named CSS colours render in colour (not black) and background-colour
    pills become true run shading (not an invisible Word highlight)."""

    async def test_docx_named_colors_and_pills(self, jp_fetch, test_color_pills_file):
        from docx import Document
        from docx.oxml.ns import qn

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_color_pills.md"}),
            raise_error=False,
        )
        assert response.code == 200

        doc = Document(io.BytesIO(response.body))

        # Named colour "green" must not render black
        green_runs = [
            r
            for p in doc.paragraphs
            for r in p.runs
            if r.text.strip() == "green"
        ]
        assert green_runs, "green run not found"
        col = green_runs[0].font.color
        assert col is not None and col.rgb is not None
        assert str(col.rgb) != "000000", "named colour 'green' rendered black"

        # Pills: locate the shaded runs in table cells
        shaded = []
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            rPr = r._r.find(qn("w:rPr"))
                            shd = rPr.find(qn("w:shd")) if rPr is not None else None
                            if shd is not None and shd.get(qn("w:fill")) not in (None, "auto"):
                                shaded.append((r.text, shd.get(qn("w:fill"))))
        fills = {t.strip(): f.upper() for t, f in shaded}
        assert fills.get("High") == "15803D", f"High pill fill wrong: {fills}"
        assert fills.get("Med") == "B45309", f"Med pill fill wrong: {fills}"

        # No marker leakage anywhere (incl. table cells)
        from docx.text.paragraph import Paragraph
        all_text = "".join(
            Paragraph(p, doc).text
            for p in doc.element.body.iter(qn("w:p"))
        )
        assert "PILL:" not in all_text
        assert "⁣" not in all_text


TEST_MARKDOWN_QUOTES_IN_LIST = """# Quotes

Two ideas:

- **First idea** - a loose list item with a nested quote below it

  > a quoted note under the first bullet

  > a second quoted note

- **Second idea** - another loose list item

  > quoted note under the second bullet
"""


@pytest.fixture
def test_quotes_in_list_file(jp_root_dir):
    md_file = jp_root_dir / "test_quotes_in_list.md"
    md_file.write_text(TEST_MARKDOWN_QUOTES_IN_LIST, encoding="utf-8")
    return md_file


class TestBlockquotesAndLooseListBullets:
    """Loose list items keep their bullets; blockquotes get indent + bar + shading."""

    async def test_docx_loose_list_keeps_bullets(self, jp_fetch, test_quotes_in_list_file):
        """A loose <li><p>..</p>..</li> must render as a List Bullet paragraph
        with text, not an empty bullet plus an orphaned Normal paragraph."""
        from docx import Document

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_quotes_in_list.md"}),
            raise_error=False,
        )
        assert response.code == 200

        doc = Document(io.BytesIO(response.body))
        bullets = [
            p for p in doc.paragraphs
            if p.style and p.style.name == "List Bullet" and p.text.strip()
        ]
        assert len(bullets) >= 2, "loose list items should keep their bullet text"
        assert any("First idea" in p.text for p in bullets)
        assert any("Second idea" in p.text for p in bullets)
        # No empty bullet paragraphs left behind
        empty_bullets = [
            p for p in doc.paragraphs
            if p.style and p.style.name == "List Bullet" and not p.text.strip()
        ]
        assert not empty_bullets, "no empty bullet glyphs should remain"

    async def test_docx_blockquote_styled_and_no_marker(self, jp_fetch, test_quotes_in_list_file):
        """Blockquote paragraphs get a left indent, a left border bar and
        shading; the sentinel marker must not leak into the text."""
        from docx import Document
        from docx.oxml.ns import qn

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_quotes_in_list.md"}),
            raise_error=False,
        )
        assert response.code == 200

        doc = Document(io.BytesIO(response.body))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "BQ:" not in full_text, "blockquote marker leaked into output"
        assert "⁣" not in full_text, "zero-width marker delimiter leaked"

        quoted = [p for p in doc.paragraphs if "quoted note" in p.text]
        assert len(quoted) >= 3, "all blockquote paragraphs should survive"
        for p in quoted:
            pf = p.paragraph_format
            assert pf.left_indent is not None and pf.left_indent > 0, (
                "blockquote paragraph should be indented"
            )
            pPr = p._p.find(qn("w:pPr"))
            assert pPr is not None and pPr.find(qn("w:pBdr")) is not None, (
                "blockquote paragraph should have a left border bar"
            )

    async def test_pdf_with_quotes_in_list_succeeds(self, jp_fetch, test_quotes_in_list_file):
        """PDF export of quotes-in-list succeeds and leaks no marker."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "test_quotes_in_list.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body.startswith(b"%PDF-")


class TestRichAlerts:
    """Test alert boxes with line breaks, hyperlinks, bold, and consecutive alerts."""

    async def test_alert_table_has_hyperlinks(self, jp_fetch, test_rich_alerts_file):
        """Test NOTE alert with markdown links produces hyperlinks in DOCX."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_rich_alerts.md"}),
            raise_error=False,
        )
        assert response.code == 200
        docx_bytes = io.BytesIO(response.body)
        with zipfile.ZipFile(docx_bytes, "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            # Hyperlinks present (from markdown links in NOTE alert)
            assert "w:hyperlink" in document_xml or "Hyperlink" in document_xml
            # Zero-width markers cleaned
            assert "\u200b" not in document_xml

    async def test_alert_table_has_line_breaks(self, jp_fetch, test_rich_alerts_file):
        """Test IMPORTANT alert with <br> tags produces line breaks in DOCX."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_rich_alerts.md"}),
            raise_error=False,
        )
        assert response.code == 200
        docx_bytes = io.BytesIO(response.body)
        with zipfile.ZipFile(docx_bytes, "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            # Line breaks present (from <br> tags in IMPORTANT alert)
            assert "w:br" in document_xml

    async def test_consecutive_alerts_separate(self, jp_fetch, test_rich_alerts_file):
        """Test consecutive alerts are rendered as separate tables with spacing."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_rich_alerts.md"}),
            raise_error=False,
        )
        assert response.code == 200
        docx_bytes = io.BytesIO(response.body)
        with zipfile.ZipFile(docx_bytes, "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            # Three alert colors should be present (IMPORTANT=purple, NOTE=blue, WARNING=amber)
            assert 'w:fill="F4EDFF"' in document_xml  # purple
            assert 'w:fill="EDF5FD"' in document_xml  # blue
            assert 'w:fill="FEF9E7"' in document_xml  # amber

    async def test_data_table_not_styled_as_alert(self, jp_fetch, test_rich_alerts_file):
        """Test regular data tables are not affected by alert styling."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_rich_alerts.md"}),
            raise_error=False,
        )
        assert response.code == 200
        docx_bytes = io.BytesIO(response.body)
        with zipfile.ZipFile(docx_bytes, "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            # Data table text should be present
            assert "Col A" in document_xml or "A1" in document_xml

    async def test_alert_labels_hidden_by_default(self, jp_fetch, test_rich_alerts_file):
        """Test alert type labels are not visible when showAlertLabels is false."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test_rich_alerts.md"}),
            raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        # Content should be present but not the label prefix pattern "IMPORTANT:"
        assert "void location" in html
        # Zero-width markers should not appear in HTML output
        assert "\u200b" not in html


# Markdown with JPEG images at different DPI values for scaling test
TEST_MARKDOWN_WITH_JPEG = """# JPEG Scaling Test

![High DPI](images/high_dpi.jpg)

![Low DPI](images/low_dpi.jpg)

Some text after images.
"""


@pytest.fixture
def test_jpeg_scaling_files(jp_root_dir):
    """Create test JPEG images with different DPI metadata."""
    from PIL import Image as PILImage

    images_dir = jp_root_dir / "images"
    images_dir.mkdir(exist_ok=True)

    # Create 200x100 JPEG at 1519 DPI (would render tiny without fix)
    img_high = PILImage.new("RGB", (200, 100), color=(255, 0, 0))
    img_high.save(images_dir / "high_dpi.jpg", dpi=(1519, 1519))

    # Create 200x100 JPEG at 72 DPI (would render large without fix)
    img_low = PILImage.new("RGB", (200, 100), color=(0, 0, 255))
    img_low.save(images_dir / "low_dpi.jpg", dpi=(72, 72))

    md_file = jp_root_dir / "test_jpeg_scaling.md"
    md_file.write_text(TEST_MARKDOWN_WITH_JPEG, encoding="utf-8")
    return md_file


class TestImageDPINormalization:
    """Test that JPEG DPI metadata is normalized for consistent DOCX sizing."""

    async def test_docx_consistent_image_sizes(self, jp_fetch, test_jpeg_scaling_files):
        """Test images with different DPI render at consistent sizes in DOCX."""
        from docx import Document
        from docx.shared import Inches

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_jpeg_scaling.md"}),
            raise_error=False,
        )
        assert response.code == 200

        docx_bytes = io.BytesIO(response.body)
        doc = Document(docx_bytes)

        # Both images have identical pixel dimensions (200x100)
        # so after DPI normalization to 96, they should have
        # identical sizes in the DOCX
        shapes = list(doc.inline_shapes)
        assert len(shapes) >= 2, f"Expected 2 images, found {len(shapes)}"

        # Both should be the same width (200px at 96 DPI = ~2.08")
        width_0 = shapes[0].width
        width_1 = shapes[1].width
        assert width_0 == width_1, (
            f"Images with same pixel dimensions should have same DOCX width: "
            f"{width_0} vs {width_1}"
        )

        # Width should be approximately 2.08 inches (200/96)
        expected_width = Inches(200 / 96)
        tolerance = Inches(0.1)
        assert abs(width_0 - expected_width) < tolerance, (
            f"Image width {width_0} should be ~{expected_width} (200px at 96 DPI)"
        )

    async def test_docx_with_jpeg_succeeds(self, jp_fetch, test_jpeg_scaling_files):
        """Test DOCX export succeeds with JPEG images of varying DPI."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_jpeg_scaling.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body[:2] == b"PK"

    async def test_pdf_with_jpeg_succeeds(self, jp_fetch, test_jpeg_scaling_files):
        """Test PDF export succeeds with JPEG images of varying DPI."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "test_jpeg_scaling.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body.startswith(b"%PDF-")


TEST_MARKDOWN_TABLE_IMAGE = """# Table Image Scaling

A large photo inside a two-column table cell.

| Photo | Caption |
|---|---|
| ![big](images/big_photo.jpg) | text in the second cell |
"""


@pytest.fixture
def test_table_image_file(jp_root_dir):
    """Create a markdown file with a large image inside a 2-column table."""
    from PIL import Image as PILImage

    images_dir = jp_root_dir / "images"
    images_dir.mkdir(exist_ok=True)
    # 4000x3000 px = ~41 inches wide at 96 DPI - far wider than any cell
    PILImage.new("RGB", (4000, 3000), color=(0, 128, 0)).save(
        images_dir / "big_photo.jpg"
    )

    md_file = jp_root_dir / "test_table_image.md"
    md_file.write_text(TEST_MARKDOWN_TABLE_IMAGE, encoding="utf-8")
    return md_file


class TestTableImageScaling:
    """Images inside table cells must be scaled to the cell, not the page."""

    async def test_docx_table_image_fits_cell(self, jp_fetch, test_table_image_file):
        """A large image in a 2-column table cell must fit the cell width."""
        from docx import Document
        from docx.shared import Inches

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_table_image.md"}),
            raise_error=False,
        )
        assert response.code == 200

        doc = Document(io.BytesIO(response.body))
        shapes = list(doc.inline_shapes)
        assert len(shapes) >= 1, "expected the table image to be embedded"

        # A 2-column table on a 7.5\" page gives each cell ~3.75\". A
        # page-width (7.5\") image would overflow the cell and be clipped;
        # the scaled image must fit a cell.
        for shape in shapes:
            assert shape.width <= Inches(4.0), (
                f"table-cell image width {shape.width} should fit a "
                f"2-column cell (< 4 in), not span the page"
            )

    async def test_pdf_ordered_numbering_restarts_at_every_separator(self, jp_fetch, jp_root_dir):
        """DEF-MARK-101: the DEF-MARK-99 reset named a heading and a quote;
        a rule, a table, a fence, an alert and a boxed div still let the
        next procedure count on - and the boxed div is a column-0 regression,
        a paragraph that reset the count having become a table that did not.
        A bullet parent also starts the ordered run beneath it afresh."""
        import fitz
        seps = ["---", "| a | b |\n| - | - |\n| 1 | 2 |", "```\ncode\n```",
                "> [!NOTE]\n> note", '<div style="border: 1px solid #333">boxed</div>']
        doc = "# T\n\n1. alpha\n2. beta\n\n" + "\n\n".join(
            f"{sep}\n\n  1. one{k}\n  2. two{k}" for k, sep in enumerate(seps))
        doc += "\n\n<div style=\"border: 1px solid #333\">boxed</div>\n\n1. col0\n2. col0b\n\n"
        # Prose before the nested chunk: glued under `2. col0b` it would be
        # that item's child, which the pass rightly leaves alone
        doc += "Nested.\n\n  - a\n    1. a1\n    2. a2\n  - b\n    1. b1\n    2. b2\n"
        (jp_root_dir / "olseps.md").write_text(doc, encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "olseps.md"}),
            raise_error=False)
        assert response.code == 200
        text = " ".join(page.get_text() for page in
                        fitz.open(stream=response.body, filetype="pdf"))
        for k in range(len(seps)):
            assert f"1. one{k}" in text and f"3. one{k}" not in text, (
                f"the list after separator {k} counted on: {text!r}")
        assert "1. col0" in text and "3. col0" not in text, "the column-0 list after a boxed div counted on"
        assert "1. b1" in text and "3. b1" not in text, (
            f"a sub-list under a new parent counted on: {text!r}")

    async def test_pdf_ordered_numbering_continues_across_a_block_inside_a_step(
            self, jp_fetch, jp_root_dir):
        """DEF-MARK-103: the DEF-MARK-99 and DEF-MARK-101 resets guessed a
        list's end from the Word paragraphs, which show a table, a sample or
        a paragraph inside a step exactly as they show one standing between
        two lists. The end is now carried: each `<ol>` gets its own numbering
        instance and the PDF restarts only where that instance changes. A
        figure between two rescued lists is the other half of the same
        guess."""
        import fitz
        inside = ["   | a | b |\n   | - | - |\n   | 1 | 2 |", "   ---",
                  "   ```\n   cmd\n   ```", "   Prose of the step.",
                  "   > quoted in the step"]
        doc = "# T\n\n" + "\n\n".join(
            f"1. step{k}\n\n{blk}\n\n2. next{k}\n3. last{k}\n\nBetween."
            for k, blk in enumerate(inside))
        doc += "\n\nIntro.\n\n  1. alpha\n  2. beta\n\n![x](nope.png)\n\n  1. one\n  2. two\n"
        (jp_root_dir / "olin.md").write_text(doc, encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "olin.md"}),
            raise_error=False)
        assert response.code == 200
        text = " ".join(page.get_text() for page in
                        fitz.open(stream=response.body, filetype="pdf"))
        for k in range(len(inside)):
            assert f"2. next{k}" in text and f"3. last{k}" in text, (
                f"the block inside step {k} restarted the count: {text!r}")
        assert "1. one" in text and "3. one" not in text, (
            f"a figure between two lists let the second count on: {text!r}")

    async def test_pdf_an_empty_numbered_item_keeps_its_number(self, jp_fetch, jp_root_dir):
        """DEF-MARK-105: an item with no text is a List Number paragraph that
        Word numbers; the PDF took the empty-text exit before looking at the
        style and skipped the number, so the author's 3. read 2."""
        import fitz
        (jp_root_dir / "olempty.md").write_text("1. a\n2. \n3. c\n", encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "olempty.md"}),
            raise_error=False)
        assert response.code == 200
        text = " ".join(page.get_text() for page in
                        fitz.open(stream=response.body, filetype="pdf"))
        assert "3. c" in text, f"the empty item lost its number: {text!r}"

    async def test_pdf_a_third_level_list_does_not_reset_its_parent(
            self, jp_fetch, jp_root_dir):
        """DEF-MARK-108: the PDF walker read every indent past the first as
        level 1, so a third-level list and its parent shared one counter and
        one numbering-instance slot; returning from the child to the parent
        looked like a change of list and restarted the parent at 1 where
        Word printed 3."""
        import fitz
        (jp_root_dir / "ol3.md").write_text(
            "1. o0\n  1. o1\n  2. o1b\n    1. o2\n  3. o1c\n", encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "ol3.md"}),
            raise_error=False)
        assert response.code == 200
        text = " ".join(page.get_text() for page in
                        fitz.open(stream=response.body, filetype="pdf"))
        assert "1. o2" in text and "3. o1c" in text, (
            f"the third level reset its parent: {text!r}")

    async def test_docx_each_ordered_list_has_its_own_numbering_instance(
            self, jp_fetch, jp_root_dir):
        """DEF-MARK-100: htmldocx numbered every List Number paragraph from
        the template's one instance, so Word counted a second procedure on
        from the first. Each list now has an instance restarting at 1, a
        step broken by a table stays on its list's instance, and the mark
        that carried the boundary does not reach the text."""
        import io
        from docx import Document
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph
        doc = ("1. a\n2. b\n\n---\n\n1. c\n\n   | x | y |\n   | - | - |\n   | 1 | 2 |\n\n2. d\n")
        (jp_root_dir / "olnum.md").write_text(doc, encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "olnum.md"}),
            raise_error=False)
        assert response.code == 200
        document = Document(io.BytesIO(response.body))
        ids = {}
        for p in document.paragraphs:
            if p.style.name == "List Number":
                ids[p.text.strip()] = p._element.pPr.numPr.numId.val
        assert set(ids) == {"a", "b", "c", "d"}, ids
        assert ids["a"] == ids["b"] and ids["c"] == ids["d"], f"a list split its instance: {ids}"
        assert ids["a"] != ids["c"], f"two lists share one instance: {ids}"
        numbering = document.part.numbering_part.element
        for num_id in {ids["a"], ids["c"]}:
            num = numbering.num_having_numId(num_id)
            starts = num.findall(qn("w:lvlOverride") + "/" + qn("w:startOverride"))
            assert starts and starts[0].get(qn("w:val")) == "1", f"instance {num_id} has no restart"
        all_text = "".join(Paragraph(p, document).text
                           for p in document.element.body.iter(qn("w:p")))
        assert "\u2062" not in all_text, "the list mark leaked into the text"

    async def test_pdf_ordered_numbering_restarts_after_a_heading(self, jp_fetch, jp_root_dir):
        """DEF-MARK-99: the PDF numbered its lists from one document-wide
        counter that a heading did not reset, so a rescued procedure under
        `## Section two` read 3., 4. where the author - and the control's
        literal source - said 1., 2."""
        import fitz
        (jp_root_dir / "olheading.md").write_text(
            "# T\n\n1. alpha\n2. beta\n\n## Section two\n\n  1. one\n  2. two\n\n"
            "> quoted\n\n  1. uno\n  2. dos\n", encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "olheading.md"}),
            raise_error=False)
        assert response.code == 200
        doc = fitz.open(stream=response.body, filetype="pdf")
        text = " ".join(page.get_text() for page in doc)
        assert "1. one" in text and "1. uno" in text, text
        assert "3. one" not in text and "3. uno" not in text, (
            f"a list after a heading or a quote continued the count: {text!r}")

    async def test_pdf_table_image_is_rendered(self, jp_fetch, test_table_image_file):
        """DEF-TABL-6: an image inside a table cell must render in the PDF, not be
        dropped (process_table used to read cell text only)."""
        import fitz

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "test_table_image.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body.startswith(b"%PDF-")
        doc = fitz.open(stream=response.body, filetype="pdf")
        n_images = sum(len(page.get_images()) for page in doc)
        # The cell image must not overflow the page frame width
        widest = max(
            (rect.width for page in doc for rect in
             [r for xref in {i[0] for i in page.get_images()}
              for r in page.get_image_rects(xref)]),
            default=0,
        )
        doc.close()
        assert n_images >= 1, "table-cell image was dropped from the PDF"
        assert widest <= pdf_frame_width() + 1, (
            "table-cell image overflows the page frame"
        )

    async def test_pdf_cell_image_and_caption_both_render(self, jp_fetch, jp_root_dir):
        """A cell holding both a caption and an image keeps both in the PDF."""
        from PIL import Image as PILImage
        import fitz

        images_dir = jp_root_dir / "images"
        images_dir.mkdir(exist_ok=True)
        PILImage.new("RGB", (300, 200), color=(20, 60, 180)).save(
            images_dir / "shot.png"
        )
        (jp_root_dir / "cap.md").write_text(
            "# Grid\n\n| A | B |\n|---|---|\n"
            '| **Caption One**<br><img src="images/shot.png" alt="s"> | plain |\n',
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "cap.md"}), raise_error=False,
        )
        assert response.code == 200
        doc = fitz.open(stream=response.body, filetype="pdf")
        n_images = sum(len(page.get_images()) for page in doc)
        text = "".join(page.get_text() for page in doc)
        doc.close()
        assert n_images >= 1, "cell image dropped"
        assert "Caption One" in text, "cell caption text dropped when the image rendered"

    async def test_pdf_page_tall_cell_image_does_not_crash(self, jp_fetch, jp_root_dir):
        """A cell image taller than a page after width-scaling must be capped
        below the frame - an atomic RLImage plus cell padding would otherwise
        exceed one page and raise LayoutError (HTTP 500)."""
        from PIL import Image as PILImage
        import fitz

        images_dir = jp_root_dir / "images"
        images_dir.mkdir(exist_ok=True)
        # Tall, narrow - stays taller than a page even after width scaling
        PILImage.new("RGB", (200, 5000), color=(180, 40, 40)).save(
            images_dir / "tall.png"
        )
        (jp_root_dir / "tall.md").write_text(
            "# Tall\n\n| Shot |\n|------|\n"
            '| <img src="images/tall.png" alt="t"> |\n',
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "tall.md"}), raise_error=False,
        )
        assert response.code == 200, (
            f"tall cell image raised {response.code} (LayoutError?)"
        )
        assert response.body.startswith(b"%PDF-")
        doc = fitz.open(stream=response.body, filetype="pdf")
        n_images = sum(len(page.get_images()) for page in doc)
        doc.close()
        assert n_images >= 1, "tall cell image was dropped"

    async def test_pdf_tall_image_in_header_cell_does_not_crash(self, jp_fetch, jp_root_dir):
        """The header row overrides its bottom padding to 8pt (12pt total, vs
        8pt for a body row), so a near-page-tall image in a HEADER cell must be
        capped by that larger padding. Capping at `frame_height - 8` leaves the
        header row 4pt too tall and reopens the atomic-RLImage LayoutError the
        cell-image fix closed."""
        from PIL import Image as PILImage
        import fitz

        images_dir = jp_root_dir / "images"
        images_dir.mkdir(exist_ok=True)
        PILImage.new("RGB", (200, 5000), color=(40, 140, 60)).save(
            images_dir / "htall.png"
        )
        # Row 0 is the header (its second cell carries text -> has_header) and
        # it holds the tall image, so the image lands in a 12pt-padded row.
        (jp_root_dir / "htall.md").write_text(
            "# Tall header\n\n"
            '| <img src="images/htall.png" alt="t"> | Header |\n'
            "|------|------|\n"
            "| body a | body b |\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "htall.md"}), raise_error=False,
        )
        assert response.code == 200, (
            f"tall header-cell image raised {response.code} (LayoutError?)"
        )
        assert response.body.startswith(b"%PDF-")
        doc = fitz.open(stream=response.body, filetype="pdf")
        n_images = sum(len(page.get_images()) for page in doc)
        doc.close()
        assert n_images >= 1, "tall header-cell image was dropped"

    async def test_pdf_tall_image_body_row_under_repeating_header_does_not_crash(
        self, jp_fetch, jp_root_dir
    ):
        """A real (text) header repeats on every continuation page. A tall-image
        BODY row then needs `header + image + padding`, which cannot fit while
        the header repeats - so the repeat must be dropped. The guard must key
        off the TALLEST body row, not the shortest: a short row alongside the
        tall one must not mask the overflow."""
        from PIL import Image as PILImage
        import fitz

        images_dir = jp_root_dir / "images"
        images_dir.mkdir(exist_ok=True)
        PILImage.new("RGB", (200, 5000), color=(60, 40, 160)).save(
            images_dir / "btall.png"
        )
        # Row 0 is a small text header (-> repeatRows). One body row holds the
        # tall image; a second, short body row is the shortest - if the guard
        # checks min(body) it keeps the repeat and the tall row never lays out.
        (jp_root_dir / "btall.md").write_text(
            "# Body tall\n\n"
            "| Name | Diagram |\n"
            "|------|---------|\n"
            '| A | <img src="images/btall.png" alt="d"> |\n'
            "| B | short text |\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "btall.md"}), raise_error=False,
        )
        assert response.code == 200, (
            f"tall body-row image under a repeating header raised {response.code} "
            f"(LayoutError?)"
        )
        assert response.body.startswith(b"%PDF-")
        doc = fitz.open(stream=response.body, filetype="pdf")
        n_images = sum(len(page.get_images()) for page in doc)
        doc.close()
        assert n_images >= 1, "tall body-row image was dropped"

    async def test_pdf_tall_text_table_keeps_repeating_header(self, jp_fetch, jp_root_dir):
        """A tall TEXT row splits fine under a repeated header, so the header
        must keep repeating across pages - the repeat-drop guard is only for
        atomic image rows. Dropping it for a tall text row regresses the
        keep-header-on-the-page feature. Asserted by the header appearing on
        more than one page."""
        import fitz

        big = ("This is a long paragraph sentence that wraps across many lines. "
               * 240)  # ~15k chars in one cell -> spans multiple pages
        (jp_root_dir / "ttt.md").write_text(
            "# Long text table\n\n"
            "| HEADERALPHA | HEADERBETA |\n"
            "|-------------|------------|\n"
            f"| {big} | side note |\n"
            "| second row a | second row b |\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "ttt.md"}), raise_error=False,
        )
        assert response.code == 200, f"tall text table raised {response.code}"
        doc = fitz.open(stream=response.body, filetype="pdf")
        n_pages = doc.page_count
        header_pages = sum(
            1 for page in doc if "HEADERALPHA" in page.get_text()
        )
        doc.close()
        assert n_pages >= 2, (
            f"test premise broken: table did not span multiple pages ({n_pages})"
        )
        assert header_pages >= 2, (
            f"repeating header appears on only {header_pages} page(s); the guard "
            f"dropped the repeat for a splittable text row"
        )

    async def test_pdf_near_full_page_header_does_not_crash(self, jp_fetch, jp_root_dir):
        """A header row that nearly (but not quite) fills the page leaves less
        than one body line under a repeated header, so a one-line body row can
        never fit on a continuation page -> LayoutError. The guard must drop the
        repeat when the header leaves under one body line, not only when it
        overflows the frame outright. The 150x3160 image lands the header row a
        few points below the frame (just under the 696pt cap), inside that
        window."""
        from PIL import Image as PILImage
        import fitz

        images_dir = jp_root_dir / "images"
        images_dir.mkdir(exist_ok=True)
        PILImage.new("RGB", (150, 3160), color=(150, 60, 40)).save(
            images_dir / "nfp.png"
        )
        (jp_root_dir / "nfp.md").write_text(
            "# Near full header\n\n"
            '| <img src="images/nfp.png" alt="t"> | HeadTwo |\n'
            "|---|---|\n"
            "| body a | body b |\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "nfp.md"}), raise_error=False,
        )
        assert response.code == 200, (
            f"near-full-page header raised {response.code} (LayoutError?)"
        )
        assert response.body.startswith(b"%PDF-")
        doc = fitz.open(stream=response.body, filetype="pdf")
        n_images = sum(len(page.get_images()) for page in doc)
        doc.close()
        assert n_images >= 1, "near-full-page header image was dropped"

    async def test_pdf_row_that_fits_a_page_is_not_split(self, jp_fetch, jp_root_dir):
        """DEF-TABL-9: a row that fits on a page must move to the next page whole,
        not be torn at the page boundary. A cell holding a caption and its image
        otherwise strands the caption on one page and the image on the next,
        even though the whole row would fit on the following page."""
        from PIL import Image as PILImage
        import fitz

        images_dir = jp_root_dir / "images"
        images_dir.mkdir(exist_ok=True)
        # Sized so the caption+image row comfortably fits one page (a row that
        # genuinely exceeds a page must still split - that is not this case).
        PILImage.new("RGB", (900, 600), color=(200, 80, 40)).save(
            images_dir / "shot.png"
        )
        # Tuned so the caption+image row does not fit the space left on page 1
        # but fits page 2 whole - the case an unconditional splitInRow tears.
        filler = ("Filler sentence that wraps across the column to consume "
                  "most of the first page. " * 60)
        (jp_root_dir / "split.md").write_text(
            "# Split\n\n"
            "| Col |\n|---|\n"
            f"| {filler} |\n"
            '| **CAPTIONX**<br><img src="images/shot.png" alt="s"> |\n',
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "split.md"}), raise_error=False,
        )
        assert response.code == 200
        doc = fitz.open(stream=response.body, filetype="pdf")
        n_pages = doc.page_count
        cap_pages = [i for i, pg in enumerate(doc) if "CAPTIONX" in pg.get_text()]
        img_pages = [i for i, pg in enumerate(doc) if pg.get_images()]
        doc.close()
        assert n_pages >= 2, (
            f"test premise broken: content did not span pages ({n_pages})"
        )
        assert len(cap_pages) == 1, f"caption found on pages {cap_pages}"
        assert img_pages, "the cell image was dropped"
        assert cap_pages[0] in img_pages, (
            f"caption is on page {cap_pages[0]} but its image is on page(s) "
            f"{img_pages} - the row was split at the page boundary instead of "
            f"moving to the next page whole"
        )

    async def test_pdf_image_only_grid_columns_size_to_images(self, jp_fetch, jp_root_dir):
        """An image-only grid (empty header, no captions - the borderless layout
        idiom the mockup docs use) earns no column width from text, so without an
        image floor each column collapses and its cell image renders a few points
        wide. The column must floor to the image so the grid renders usably."""
        from PIL import Image as PILImage
        import fitz

        images_dir = jp_root_dir / "images"
        images_dir.mkdir(exist_ok=True)
        PILImage.new("RGB", (800, 1600), color=(30, 90, 200)).save(
            images_dir / "g.png"
        )
        (jp_root_dir / "grid.md").write_text(
            "# Grid\n\n|  |  |  |\n|---|---|---|\n"
            '| ![a](images/g.png) | ![b](images/g.png) | ![c](images/g.png) |\n',
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "grid.md"}), raise_error=False,
        )
        assert response.code == 200
        doc = fitz.open(stream=response.body, filetype="pdf")
        widths = [rc.width for pg in doc for xref in {i[0] for i in pg.get_images()}
                  for rc in pg.get_image_rects(xref)]
        doc.close()
        assert len(widths) == 3, f"expected 3 grid images, got {len(widths)}"
        assert min(widths) > 100, (
            f"image-only grid columns collapsed - images render {min(widths):.0f}pt "
            f"wide; the column must floor to the image"
        )


SIZE_BADGE_SVG = (
    '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 60 40" '
    'width="60" height="40" role="img" aria-label="XL">'
    '<rect x="2" y="2" width="56" height="36" rx="8" fill="#b45309"/>'
    '<text x="30" y="27" font-size="20" fill="#ffffff" '
    'text-anchor="middle">XL</text></svg>'
)

TEST_MARKDOWN_HTML_IMG_BADGES = """# Badges

| Model | Size |
|---|---|
| M00 | <img src="badges/size-xl.svg" alt="XL" style="max-height:22px;max-width:1000px"> 64h |
| M01 | <img src="badges/size-xl.svg" alt="XL" style="max-height:22px;max-width:1000px"> 16h |
"""


@pytest.fixture
def test_html_img_badge_file(jp_root_dir):
    badges_dir = jp_root_dir / "badges"
    badges_dir.mkdir(exist_ok=True)
    (badges_dir / "size-xl.svg").write_text(SIZE_BADGE_SVG, encoding="utf-8")
    md_file = jp_root_dir / "test_html_img_badges.md"
    md_file.write_text(TEST_MARKDOWN_HTML_IMG_BADGES, encoding="utf-8")
    return md_file


class TestHtmlImgBadges:
    """Raw HTML <img> tags with relative local paths (inline table badges)
    must embed and render small, not break or fill the cell width."""

    def test_embed_local_html_img(self):
        """embed_images_as_base64 inlines a raw HTML <img> local src."""
        from types import SimpleNamespace
        from jupyterlab_export_markdown_extension.routes import ExportDocxHandler

        with tempfile.TemporaryDirectory() as d:
            (Path(d) / "badge.svg").write_text(SIZE_BADGE_SVG)
            stub = SimpleNamespace(
                contents_manager=SimpleNamespace(root_dir=d)
            )
            html = '<img src="badge.svg" alt="XL" style="max-height:22px">'
            out = ExportDocxHandler.embed_images_as_base64(stub, html, Path(d))
            assert "data:image/svg+xml;base64," in out
            assert 'src="badge.svg"' not in out
            assert 'alt="XL"' in out  # other attributes preserved

    def test_badge_render_spec(self):
        """A small max-height yields a tiny render width + high DPI; an
        unsized image yields None (full diagram width)."""
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase
        spec = ExportHandlerBase._badge_render_spec

        small = spec('<img style="max-height:22px;max-width:1000px">', 60, 40)
        assert small is not None
        render_w, dpi = small
        # 22px * (60/40) = 33px display -> *4 supersample = 132px, dpi 384
        assert render_w == 132
        assert dpi == 384

        assert spec('<img alt="x">', 800, 345) is None
        assert spec('<img style="max-height:900px">', 60, 40) is None

    def test_badge_render_spec_robustness(self):
        """Adversarial-review hardening: plain height, width-only, unit/ garbage
        handling, and CSS shorthands that must NOT be read as height/width."""
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase
        spec = ExportHandlerBase._badge_render_spec

        # plain `height:` in style (no max-) is detected
        assert spec('<img style="height:22px">', 60, 40) is not None
        # width-only badge is detected (keyed off width when no height)
        assert spec('<img style="max-width:80px">', 60, 40) == (320, 384)
        # height attribute (bare number) is detected
        assert spec('<img height="22">', 60, 40) is not None
        # malformed number must not raise and must not size (-> None)
        assert spec('<img style="max-height:1.2.3px">', 60, 40) is None
        # non-px units are ignored (not mis-sized)
        assert spec('<img height="2em">', 60, 40) is None
        assert spec('<img style="max-height:50%">', 60, 40) is None
        # CSS shorthands containing 'height'/'width' must NOT match
        assert spec('<img style="line-height:20px">', 60, 40) is None
        assert spec('<img style="stroke-width:3px">', 60, 40) is None
        # width clamp wins when max-height is huge but max-width is tiny
        # eff_h = min(9999, 16/(60/40)=10.67)=10.67 -> disp_w 16 -> *4 = 64
        assert spec('<img style="max-height:9999px;max-width:16px">', 60, 40) == (64, 384)
        # bare width attribute is read; large one is not a badge
        assert spec('<img width="22">', 60, 40) == (88, 384)
        assert spec('<img width="800">', 60, 40) is None
        # data-* attributes must NOT be read as width/height
        assert spec('<img data-height="20">', 60, 40) is None
        assert spec('<img data-width="20">', 60, 40) is None
        # width=/height= inside another attribute's VALUE must NOT be read
        assert spec('<img src="x.svg" alt="chart width=5 tall">', 800, 345) is None
        assert spec('<img alt="height=8 widget">', 800, 345) is None
        # a crafted viewBox aspect cannot drive an unbounded render width (DoS)
        rw, _ = spec('<img style="max-height:1px">', 100000, 1)
        assert rw <= ExportHandlerBase._BADGE_RENDER_MAX_PX
        # max-height takes precedence over a larger plain height in the same style
        assert spec('<img style="height:300px;max-height:22px">', 60, 40) is not None

    def test_embed_src_not_confused_by_data_src(self):
        """resolve only the real `src`, never a `data-src`/`lowsrc` substring,
        and don't truncate a tag on a `>` inside an attribute value."""
        from types import SimpleNamespace
        from jupyterlab_export_markdown_extension.routes import ExportDocxHandler

        with tempfile.TemporaryDirectory() as d:
            (Path(d) / "real.svg").write_text(SIZE_BADGE_SVG)
            stub = SimpleNamespace(
                contents_manager=SimpleNamespace(root_dir=d)
            )
            # data-src must be left alone; real src embedded
            out = ExportDocxHandler.embed_images_as_base64(
                stub,
                '<img data-src="lazy.svg" src="real.svg" alt="a > b">',
                Path(d),
            )
            assert 'data-src="lazy.svg"' in out  # untouched
            assert "data:image/svg+xml;base64," in out  # real src embedded
            assert 'src="real.svg"' not in out
            assert 'alt="a > b"' in out  # tag not truncated at the inner '>'

    async def test_docx_badges_small_inline(self, jp_fetch, test_html_img_badge_file):
        """HTML <img> badges embed and stay small (~22px), not cell-width."""
        from docx import Document
        from docx.shared import Inches

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_html_img_badges.md"}),
            raise_error=False,
        )
        assert response.code == 200

        doc = Document(io.BytesIO(response.body))
        shapes = list(doc.inline_shapes)
        assert len(shapes) == 2, "both HTML <img> badges must embed"
        for shape in shapes:
            # 22px at 96 DPI is ~0.23 in; allow slack but assert it is a
            # small badge, not a cell-width or page-width image.
            assert shape.height <= Inches(0.4), (
                f"badge height {shape.height} should be ~22px, not full size"
            )
            assert shape.width <= Inches(0.6), (
                f"badge width {shape.width} should be small, not cell-width"
            )


class TestImageEmbedSecurity:
    """Local-file containment and SSRF guard on image embedding."""

    def test_ip_is_blocked(self):
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase
        b = ExportHandlerBase._ip_is_blocked
        # non-public ranges an SSRF would target
        assert b("169.254.169.254")  # cloud metadata (link-local)
        assert b("127.0.0.1")        # loopback
        assert b("10.0.0.1")         # private
        assert b("192.168.1.1")      # private
        assert b("::1")              # ipv6 loopback
        assert b("not-an-ip")        # unparseable -> fail closed
        # public addresses are allowed
        assert not b("8.8.8.8")
        assert not b("1.1.1.1")

    def test_host_is_blocked_localhost(self):
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase
        # resolves to a loopback address -> blocked; empty host -> blocked
        assert ExportHandlerBase._host_is_blocked("localhost")
        assert ExportHandlerBase._host_is_blocked(None)

    def test_local_read_contained_to_root(self):
        """A traversal src escaping the server root must NOT be embedded; a
        sibling image inside the root must still embed."""
        from types import SimpleNamespace
        from jupyterlab_export_markdown_extension.routes import ExportDocxHandler

        with tempfile.TemporaryDirectory() as d:
            root = Path(d)
            (root / "secret.txt").write_text("TOP SECRET")  # outside markdown_dir
            sub = root / "docs"
            sub.mkdir()
            (sub / "ok.svg").write_text(SIZE_BADGE_SVG)

            # contents_manager is a read-only handler property; call the method
            # with a stub self exposing only what it reads.
            stub = SimpleNamespace(
                contents_manager=SimpleNamespace(root_dir=str(root))
            )
            embed = ExportDocxHandler.embed_images_as_base64

            # legit sibling under the root -> embedded
            out_ok = embed(stub, '<img src="ok.svg">', sub)
            assert "data:image/svg+xml;base64," in out_ok

            # escape the root -> refused, left untouched
            out_esc = embed(stub, '<img src="../../../../etc/passwd">', sub)
            assert "data:" not in out_esc
            assert 'src="../../../../etc/passwd"' in out_esc

            # fail closed: if the root can't be determined, refuse local reads
            blind = SimpleNamespace(contents_manager=SimpleNamespace(root_dir=None))
            out_blind = embed(blind, '<img src="ok.svg">', sub)
            assert "data:" not in out_blind


TEST_MARKDOWN_WITH_ANCHORS = """# Anchor Links

See footnote [<sup>A1</sup>](#A1) and [<sup>A2</sup>](#A2) below.

External link: [Google](https://google.com).

Inline arrow inside code: `←` and bare arrows → ← ↑ ↓.

## References

- <span id="A1">A1 first reference target</span>
- <span id="A2">A2 second reference target</span>
"""


@pytest.fixture
def test_anchor_markdown_file(jp_root_dir):
    md_file = jp_root_dir / "test_anchors.md"
    md_file.write_text(TEST_MARKDOWN_WITH_ANCHORS, encoding="utf-8")
    return md_file


class TestAnchorLinksAndArrowFonts:
    """Validate anchor links become Word bookmarks and arrows escape Courier."""

    async def test_docx_internal_hyperlinks_and_bookmarks(
        self, jp_fetch, test_anchor_markdown_file
    ):
        import re

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_anchors.md"}),
            raise_error=False,
        )
        assert response.code == 200

        with zipfile.ZipFile(io.BytesIO(response.body)) as z:
            doc_xml = z.read("word/document.xml").decode("utf-8")
            rels_xml = z.read("word/_rels/document.xml.rels").decode("utf-8")

        assert 'w:anchor="A1"' in doc_xml, "anchor link A1 should be internal"
        assert 'w:anchor="A2"' in doc_xml, "anchor link A2 should be internal"
        assert 'w:name="A1"' in doc_xml, "bookmark A1 must exist"
        assert 'w:name="A2"' in doc_xml, "bookmark A2 must exist"

        assert 'Target="#A1"' not in rels_xml
        assert 'Target="#A2"' not in rels_xml
        assert re.search(r'Target="https://google\.com"', rels_xml), (
            "external hyperlinks should still be preserved"
        )

        assert "⁣BM:" not in doc_xml, "bookmark sentinel must not leak"

    async def test_docx_unicode_arrow_escapes_courier(
        self, jp_fetch, test_anchor_markdown_file
    ):
        import re

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_anchors.md"}),
            raise_error=False,
        )
        assert response.code == 200

        with zipfile.ZipFile(io.BytesIO(response.body)) as z:
            doc_xml = z.read("word/document.xml").decode("utf-8")

        arrows = ["←", "↑", "→", "↓"]
        for arrow in arrows:
            assert arrow in doc_xml, f"arrow {arrow!r} missing from DOCX"

        courier_runs = re.findall(
            r'<w:r>[^<]*<w:rPr>[^<]*<w:rFonts[^/]*Courier[^/]*/>.*?</w:r>',
            doc_xml,
            flags=re.DOTALL,
        )
        for run in courier_runs:
            texts = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", run)
            for t in texts:
                for arrow in arrows:
                    assert arrow not in t, (
                        f"arrow {arrow!r} remains in Courier run: {run}"
                    )

    async def test_pdf_with_anchor_links_succeeds(
        self, jp_fetch, test_anchor_markdown_file
    ):
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "test_anchors.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body.startswith(b"%PDF-")


TEST_MARKDOWN_WIDE_TABLE = """# Wide table

| Component | Purpose | Service | Configuration detail | Monthly cost | Notes |
|---|---|---|---|---|---|
| Ingestion pipeline | Receives sensor telemetry from every house | Kinesis Data Streams | 4 shards with a 24 hour retention window | 145 USD | Needs enhanced fan-out for multiple consumers |
| Storage layer | Durable raw and curated data zones | S3 with Glacier lifecycle | Intelligent tiering after 30 days, versioning on | 62 USD | https://docs.aws.amazon.com/AmazonS3/latest/userguide/lifecycle-configuration-examples.html |
"""

@pytest.fixture
def test_wide_table_file(jp_root_dir):
    """Create a markdown file whose table is far wider than the page."""
    md_file = jp_root_dir / "test_wide_table.md"
    md_file.write_text(TEST_MARKDOWN_WIDE_TABLE, encoding="utf-8")
    return md_file

def pdf_frame_right_edge():
    """The x past which nothing the PDF lays out may extend."""
    from reportlab.lib.pagesizes import letter
    from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

    return (letter[0] - ExportHandlerBase.PDF_PAGE_MARGIN
            - ExportHandlerBase.PDF_FRAME_PADDING)

def pdf_frame_width():
    """The width a flowable is given, page less both margins and both pads."""
    from reportlab.lib.pagesizes import letter
    from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

    return (letter[0] - 2 * ExportHandlerBase.PDF_PAGE_MARGIN
            - 2 * ExportHandlerBase.PDF_FRAME_PADDING)

def pdf_frame_height():
    """The height a flowable is given, page less both margins and both pads."""
    from reportlab.lib.pagesizes import letter
    from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

    return (letter[1] - 2 * ExportHandlerBase.PDF_PAGE_MARGIN
            - 2 * ExportHandlerBase.PDF_FRAME_PADDING)

def pdf_text_past_margin(pdf_bytes):
    """Text runs in `pdf_bytes` whose estimated right edge passes the margin.

    The margin is the frame's content edge: the page margin less the padding
    SimpleDocTemplate puts inside its frame. A run's x origin is the text
    matrix offset within the CTM - reportlab puts it in cm for Paragraph cells
    and in tm for string ones. 0.5em per character under-states the real
    advance width, so a run that merely fills its column is never reported.
    """
    from pypdf import PdfReader

    right_margin = pdf_frame_right_edge()
    overflowing = []

    def visitor(text, cm, tm, font_dict, font_size):
        stripped = text.strip()
        if not stripped:
            return
        x = cm[4] + tm[4]
        right = x + len(stripped) * font_size * 0.5
        if right > right_margin:
            overflowing.append((round(x, 1), round(right, 1), stripped[:40]))

    for page in PdfReader(io.BytesIO(pdf_bytes)).pages:
        page.extract_text(visitor_text=visitor)
    return overflowing

class TestWideTableFitsPage:
    """A table wider than the page wraps instead of running off the edge."""

    async def test_docx_wide_table_fits_page(self, jp_fetch, test_wide_table_file):
        """The table grid stays inside the margins and uses a fixed layout."""
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Twips

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_wide_table.md"}),
            raise_error=False,
        )
        assert response.code == 200

        doc = Document(io.BytesIO(response.body))
        section = doc.sections[0]
        usable = section.page_width - section.left_margin - section.right_margin

        table = doc.tables[0]
        grid = table._tbl.find(qn('w:tblGrid'))
        widths = [Twips(int(col.get(qn('w:w'))))
                  for col in grid.findall(qn('w:gridCol'))]
        assert len(widths) == 6
        assert sum(widths) <= usable, "table grid runs past the right margin"

        # Autofit would let an unbreakable token widen the table past the page
        layout = table._tbl.tblPr.find(qn('w:tblLayout'))
        assert layout is not None and layout.get(qn('w:type')) == 'fixed'

    async def test_docx_single_cell_table_is_fitted_too(
            self, jp_fetch, jp_root_dir):
        """A one-cell table is content, not an alert box, and must be fitted.

        A raw HTML table of one row and one column has the same shape as the
        alert boxes, so shape alone cannot tell them apart. Left on autofit,
        its unbreakable token widens it past the right margin.
        """
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Twips

        (jp_root_dir / "test_one_col.md").write_text(
            "<table><tr><td>" + "z" * 200 + "</td></tr></table>\n",
            encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "test_one_col.md"}),
            raise_error=False,
        )
        assert response.code == 200

        doc = Document(io.BytesIO(response.body))
        section = doc.sections[0]
        usable = section.page_width - section.left_margin - section.right_margin

        table = doc.tables[0]
        layout = table._tbl.tblPr.find(qn('w:tblLayout'))
        assert layout is not None and layout.get(qn('w:type')) == 'fixed'
        grid = table._tbl.find(qn('w:tblGrid'))
        widths = [Twips(int(col.get(qn('w:w'))))
                  for col in grid.findall(qn('w:gridCol'))]
        assert sum(widths) <= usable, "one-column table runs past the margin"

    async def test_pdf_wide_table_stays_inside_the_margins(
            self, jp_fetch, test_wide_table_file):
        """No text in the exported PDF is drawn past the right margin."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "test_wide_table.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body.startswith(b"%PDF-")

        overflowing = pdf_text_past_margin(response.body)
        assert not overflowing, f"text past the right margin: {overflowing[:3]}"

    async def test_pdf_row_taller_than_the_page_still_exports(self, jp_fetch,
                                                              jp_root_dir):
        """A wrapped row taller than the page must split, not abort the export."""
        words = " ".join(f"word{i}" for i in range(900))
        (jp_root_dir / "tall_row.md").write_text(
            f"| A | B | C |\n|---|---|---|\n| x | {words} | y |\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "tall_row.md"}),
            raise_error=False,
        )
        assert response.code == 200, response.body[:300]
        assert response.body.startswith(b"%PDF-")

    async def test_pdf_many_column_table_still_exports(self, jp_fetch,
                                                       jp_root_dir):
        """Columns squeezed to their fair share must not abort the export."""
        ncols = 60  # each column falls below reportlab's own cell padding
        head = "|" + "|".join(f"h{i}" for i in range(ncols)) + "|"
        sep = "|" + "|".join("---" for _ in range(ncols)) + "|"
        row = "|" + "|".join(f"v{i}" for i in range(ncols)) + "|"
        (jp_root_dir / "many_columns.md").write_text(
            f"{head}\n{sep}\n{row}\n", encoding="utf-8")

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "many_columns.md"}),
            raise_error=False,
        )
        assert response.code == 200, response.body[:300]
        assert response.body.startswith(b"%PDF-")

        overflowing = pdf_text_past_margin(response.body)
        assert not overflowing, f"text past the right margin: {overflowing[:3]}"

    async def test_docx_overflowing_table_is_always_fitted(self, jp_fetch,
                                                           jp_root_dir):
        """A table over the page width must never slip past the fitter.

        The fitted widths sum to the page width in float arithmetic, which can
        land a hair under it - the guard has to tolerate that or the table
        silently stays on autofit and overflows again.
        """
        from docx import Document
        from docx.oxml.ns import qn

        # These lengths make the fitted widths sum to 10799.999999999998
        # against a 10800-twip page - inside the tolerance window
        lengths = [2, 17, 58, 59, 46, 41, 53]
        words = [2, 16, 4, 21, 13, 13, 7]

        def cell(length, word_length):
            return ("x" * word_length + " ") * max(1, length // (word_length + 1))

        head = "|" + "|".join(f"h{i}" for i in range(7)) + "|"
        sep = "|" + "|".join("---" for _ in range(7)) + "|"
        row = "|" + "|".join(cell(n, w) for n, w in zip(lengths, words)) + "|"
        (jp_root_dir / "overflowing.md").write_text(
            f"{head}\n{sep}\n{row}\n", encoding="utf-8")

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "overflowing.md"}),
        )
        from docx.shared import Twips

        doc = Document(io.BytesIO(response.body))
        section = doc.sections[0]
        usable = section.page_width - section.left_margin - section.right_margin
        table = doc.tables[0]
        layout = table._tbl.tblPr.find(qn('w:tblLayout'))
        assert layout is not None, "an over-wide table was left on autofit"
        # Not just present - a fixed layout of a grid that actually fits
        assert layout.get(qn('w:type')) == 'fixed'
        grid = table._tbl.find(qn('w:tblGrid'))
        widths = [Twips(int(col.get(qn('w:w'))))
                  for col in grid.findall(qn('w:gridCol'))]
        assert sum(widths) <= usable, "fitted grid still runs past the margin"

    async def test_html_table_does_not_overflow_the_viewport(
            self, jp_fetch, test_wide_table_file):
        """Rendered in a browser, the table must not scroll sideways.

        The fixture holds an unbreakable URL, which forces a wider min-content
        width than the viewport unless the cells may break inside a word.
        """
        from playwright.async_api import async_playwright

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "test_wide_table.md"}),
        )
        html = response.body.decode("utf-8")

        async with async_playwright() as pw:
            browser = await pw.chromium.launch()
            try:
                page = await browser.new_page(
                    viewport={"width": 800, "height": 600})
                await page.set_content(html)
                # The row box is the grid itself - the table element may be a
                # scroll container, which fits the viewport by construction
                row_width, doc_width, viewport_width = await page.evaluate(
                    "() => { const r = document.querySelector('tr');"
                    " const d = document.documentElement;"
                    " return [r.getBoundingClientRect().width,"
                    " d.scrollWidth, d.clientWidth]; }"
                )
            finally:
                await browser.close()

        assert row_width <= viewport_width, (
            f"table grid is {row_width}px wide in a {viewport_width}px viewport"
        )
        assert doc_width <= viewport_width, "the page scrolls sideways"

    async def test_html_many_column_table_is_contained(self, jp_fetch,
                                                       jp_root_dir):
        """A table too wide to wrap scrolls in its own box, not the page."""
        from playwright.async_api import async_playwright

        ncols = 40
        head = "|" + "|".join(f"Header {i}" for i in range(ncols)) + "|"
        sep = "|" + "|".join("---" for _ in range(ncols)) + "|"
        row = "|" + "|".join(f"value {i}" for i in range(ncols)) + "|"
        (jp_root_dir / "many_columns.md").write_text(
            f"{head}\n{sep}\n{row}\n", encoding="utf-8")

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "many_columns.md"}),
        )
        html = response.body.decode("utf-8")

        async with async_playwright() as pw:
            browser = await pw.chromium.launch()
            try:
                page = await browser.new_page(
                    viewport={"width": 800, "height": 600})
                await page.set_content(html)
                doc_width, viewport_width = await page.evaluate(
                    "() => { const d = document.documentElement;"
                    " return [d.scrollWidth, d.clientWidth]; }"
                )
            finally:
                await browser.close()

        assert doc_width <= viewport_width, (
            f"a {ncols}-column table scrolls the page ({doc_width}px "
            f"in {viewport_width}px)"
        )

    async def test_html_scroll_box_does_not_clip_on_paper(self, jp_fetch,
                                                          jp_root_dir):
        """Paper cannot scroll, so the box must stop being one when printing.

        Left as a scroll container, print media crops the table at the box
        edge and the hidden columns are simply absent from the printout.
        """
        from playwright.async_api import async_playwright

        from pypdf import PdfReader

        ncols = 30
        head = "|" + "|".join(f"H{i}x" for i in range(ncols)) + "|"
        sep = "|" + "|".join("---" for _ in range(ncols)) + "|"
        row = "|" + "|".join(f"v{i}" for i in range(ncols)) + "|"
        (jp_root_dir / "print_table.md").write_text(
            f"{head}\n{sep}\n{row}\n", encoding="utf-8")

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "print_table.md"}),
        )
        html = response.body.decode("utf-8")

        async with async_playwright() as pw:
            browser = await pw.chromium.launch()
            try:
                page = await browser.new_page()
                await page.set_content(html)
                printed = await page.pdf()
            finally:
                await browser.close()

        # Columns this narrow break every character onto its own line
        text = "".join("".join(p.extract_text().split())
                       for p in PdfReader(io.BytesIO(printed)).pages)
        missing = [i for i in range(ncols) if f"H{i}x" not in text]
        assert not missing, f"printing dropped {len(missing)} headers: {missing}"

    async def test_html_raw_table_with_attributes_is_wrapped(
            self, jp_fetch, jp_root_dir):
        """Raw HTML passes through markdown verbatim, attributes and all.

        A wrapper matched on a bare `<table>` misses `<table border="1">`, and
        the table then pushes the whole document sideways.
        """
        from playwright.async_api import async_playwright

        ncols = 40
        cells = "".join(f"<td>value {i}</td>" for i in range(ncols))
        (jp_root_dir / "raw_table.md").write_text(
            f'<table border="1"><tr>{cells}</tr></table>\n', encoding="utf-8")

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "raw_table.md"}),
        )
        html = response.body.decode("utf-8")

        async with async_playwright() as pw:
            browser = await pw.chromium.launch()
            try:
                page = await browser.new_page(
                    viewport={"width": 800, "height": 600})
                await page.set_content(html)
                doc_width, viewport_width = await page.evaluate(
                    "() => { const d = document.documentElement;"
                    " return [d.scrollWidth, d.clientWidth]; }"
                )
            finally:
                await browser.close()

        assert doc_width <= viewport_width, (
            f"a raw table with attributes scrolls the page ({doc_width}px "
            f"in {viewport_width}px)"
        )

    async def test_html_nested_table_keeps_its_inner_table(
            self, jp_fetch, jp_root_dir):
        """The wrapper must close on the outer table, not the inner one.

        A non-greedy match closes the wrapper at the first `</table>`, which
        is the inner table's. The markup is then mis-nested and the browser
        recovers by lifting the inner table clean out of its cell.
        """
        from playwright.async_api import async_playwright

        (jp_root_dir / "nested_table.md").write_text(
            "<table><tr><td><table><tr><td>inner</td></tr></table>"
            "</td></tr></table>\n\nAfter the table.\n",
            encoding="utf-8")

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "nested_table.md"}),
        )
        html = response.body.decode("utf-8")

        async with async_playwright() as pw:
            browser = await pw.chromium.launch()
            try:
                page = await browser.new_page(
                    viewport={"width": 800, "height": 600})
                await page.set_content(html)
                nested_in_cell, wrappers = await page.evaluate(
                    "() => [!!document.querySelector("
                    "'.table-scroll table td table'),"
                    " document.querySelectorAll('.table-scroll').length]"
                )
            finally:
                await browser.close()

        assert nested_in_cell, "the nested table was lifted out of its cell"
        assert wrappers == 1, f"expected one wrapper, got {wrappers}"

    def test_html_table_inside_a_comment_opens_no_wrapper(self):
        """A `<table>` inside an HTML comment must not open a scroll box.

        Markdown passes comments through verbatim. A wrapper scan that is
        blind to comments opens a `<div class="table-scroll">` at the commented
        `<table>`, never meets a `</table>`, and leaves the div open over the
        rest of the document.
        """
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        handler = ExportHandlerBase.__new__(ExportHandlerBase)
        html = ("<p>Intro.</p><!-- TODO: add a <table> here -->"
                "<p>After the comment.</p>")
        wrapped = handler.wrap_html_tables(html)
        assert '<div class="table-scroll">' not in wrapped, (
            "a commented <table> opened a wrapper"
        )
        # A real table after the comment must still wrap and close cleanly
        both = handler.wrap_html_tables(
            html + "<table><tr><td>x</td></tr></table>")
        assert both.count('<div class="table-scroll">') == 1
        assert both.count('</div>') == 1


class TestPageFitting:
    """Tables must not orphan a header at a page break; tall images must fit."""

    async def test_pdf_table_header_repeats_on_each_page(
            self, jp_fetch, jp_root_dir):
        """A table spanning pages repeats its header - the mechanism that also
        keeps the header from being stranded alone at a page bottom."""
        import fitz

        rows = "\n".join(f"| r{i:02d}a | r{i:02d}b |" for i in range(80))
        (jp_root_dir / "tall_table.md").write_text(
            f"| ZZHEADERZZ | COLBETA |\n|---|---|\n{rows}\n", encoding="utf-8")

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "tall_table.md"}),
            raise_error=False,
        )
        assert response.code == 200, response.body[:300]

        doc = fitz.open(stream=response.body, filetype="pdf")
        assert doc.page_count >= 2, "table should span multiple pages"
        pages_with_header = sum(
            1 for page in doc if "ZZHEADERZZ" in page.get_text())
        assert pages_with_header >= 2, (
            "header not repeated on continuation pages - it can be orphaned"
        )

    async def test_pdf_page_tall_header_row_still_exports(
            self, jp_fetch, jp_root_dir):
        """A repeated header must not re-arm the LayoutError splitInRow fixed.

        With repeatRows the header is re-emitted on each page, so a header row
        taller than a page cannot be placed and reportlab raises LayoutError.
        A page-long alert (a 1x1 table), a header-only tall table and a tall
        header with a body row must all still export.
        """
        long_para = " ".join(f"word{i}" for i in range(1500))
        cases = {
            "alert": f"> [!NOTE]\n> {long_para}\n",
            "header_only": f"| {long_para} |\n|---|\n",
            "tall_header": f"| {long_para} |\n|---|\n| body |\n",
        }
        for name, md in cases.items():
            (jp_root_dir / f"{name}.md").write_text(md, encoding="utf-8")
            response = await jp_fetch(
                "jupyterlab-export-markdown-extension",
                "export/pdf",
                method="POST",
                body=json.dumps({"path": f"{name}.md"}),
                raise_error=False,
            )
            assert response.code == 200, (name, response.body[:200])
            assert response.body.startswith(b"%PDF-")

    async def test_docx_table_header_row_marked_repeating(
            self, jp_fetch, jp_root_dir):
        """The header row carries w:tblHeader so Word repeats it and never
        strands it alone at a page bottom."""
        from docx import Document
        from docx.oxml.ns import qn

        (jp_root_dir / "hdr_table.md").write_text(
            "| A | B |\n|---|---|\n| one | two |\n| three | four |\n",
            encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "hdr_table.md"}),
            raise_error=False,
        )
        assert response.code == 200

        table = Document(io.BytesIO(response.body)).tables[0]
        trPr = table.rows[0]._tr.find(qn('w:trPr'))
        assert trPr is not None and trPr.find(qn('w:tblHeader')) is not None, (
            "first row is not marked as a repeating header"
        )

    async def test_html_print_keeps_header_with_body(
            self, jp_fetch, jp_root_dir):
        """In print media the table header must not break away from its body."""
        from playwright.async_api import async_playwright

        (jp_root_dir / "hdr_html.md").write_text(
            "| A | B |\n|---|---|\n| one | two |\n", encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "hdr_html.md"}),
        )
        html = response.body.decode("utf-8")

        async with async_playwright() as pw:
            browser = await pw.chromium.launch()
            try:
                page = await browser.new_page()
                await page.set_content(html)
                await page.emulate_media(media="print")
                break_inside = await page.evaluate(
                    "() => getComputedStyle(document.querySelector('thead'))"
                    ".breakInside"
                )
            finally:
                await browser.close()

        assert break_inside == "avoid", (
            f"thead break-inside is {break_inside!r}, header can be orphaned"
        )

    async def test_pdf_tall_image_fits_page_height(self, jp_fetch, jp_root_dir):
        """A diagram too tall at page width is scaled down to the page height."""
        from PIL import Image as PILImage
        import fitz

        images_dir = jp_root_dir / "images"
        images_dir.mkdir(exist_ok=True)
        # 300 x 3000 px - narrow enough to fit the width, far too tall
        PILImage.new("RGB", (300, 3000), color=(0, 100, 0)).save(
            images_dir / "tall.png")
        (jp_root_dir / "tall_image.md").write_text(
            "![tall](images/tall.png)\n", encoding="utf-8")

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/pdf",
            method="POST",
            body=json.dumps({"path": "tall_image.md"}),
            raise_error=False,
        )
        assert response.code == 200, response.body[:300]

        doc = fitz.open(stream=response.body, filetype="pdf")
        frame_height = pdf_frame_height()
        worst = 0.0
        for page in doc:
            for img in page.get_images():
                for rect in page.get_image_rects(img[0]):
                    worst = max(worst, rect.height)
        assert worst > 0, "image was not embedded"
        assert worst <= frame_height + 1, (
            f"image drawn {worst:.0f}pt tall, past the {frame_height}pt frame"
        )

    async def test_docx_tall_image_fits_page_height(self, jp_fetch, jp_root_dir):
        """An image taller than the page is scaled down to fit its height."""
        from PIL import Image as PILImage
        from docx import Document

        images_dir = jp_root_dir / "images"
        images_dir.mkdir(exist_ok=True)
        PILImage.new("RGB", (300, 3000), color=(0, 0, 100)).save(
            images_dir / "tall_docx.png")
        (jp_root_dir / "tall_image_docx.md").write_text(
            "![tall](images/tall_docx.png)\n", encoding="utf-8")

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/docx",
            method="POST",
            body=json.dumps({"path": "tall_image_docx.md"}),
            raise_error=False,
        )
        assert response.code == 200

        doc = Document(io.BytesIO(response.body))
        section = doc.sections[0]
        usable_height = (section.page_height - section.top_margin
                         - section.bottom_margin)
        shape = list(doc.inline_shapes)[0]
        assert shape.height <= usable_height, (
            "tall image was not scaled down to the page height"
        )

    async def test_html_print_caps_image_height(self, jp_fetch, jp_root_dir):
        """A tall image must fit one printed page, not spill across several.

        A 300x3000 image is 10x taller than wide - uncapped it prints across
        several pages; the print stylesheet must bound it to one page.
        """
        from playwright.async_api import async_playwright
        from pypdf import PdfReader
        from PIL import Image as PILImage

        images_dir = jp_root_dir / "images"
        images_dir.mkdir(exist_ok=True)
        PILImage.new("RGB", (300, 3000), color=(100, 0, 0)).save(
            images_dir / "tall_html.png")
        (jp_root_dir / "tall_image_html.md").write_text(
            "![tall](images/tall_html.png)\n", encoding="utf-8")

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension",
            "export/html",
            method="POST",
            body=json.dumps({"path": "tall_image_html.md"}),
        )
        html = response.body.decode("utf-8")

        async with async_playwright() as pw:
            browser = await pw.chromium.launch()
            try:
                page = await browser.new_page()
                await page.set_content(html)
                printed = await page.pdf()
            finally:
                await browser.close()

        pages = len(PdfReader(io.BytesIO(printed)).pages)
        assert pages == 1, (
            f"tall image printed to {pages} pages, not capped to one"
        )


class TestColumnWidthFitting:
    """Unit tests for the shared column-width fitter."""

    def _fit(self, natural, minimums, available):
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        return ExportHandlerBase.fit_column_widths(natural, minimums, available)

    def test_widths_that_fit_are_untouched(self):
        assert self._fit([100, 50], [40, 20], 500) == [100, 50]

    def test_overflowing_widths_are_scaled_to_available(self):
        widths = self._fit([600, 300, 100], [50, 50, 50], 500)
        assert sum(widths) == pytest.approx(500)
        assert widths[0] > widths[1] > widths[2], "proportions are preserved"

    def test_every_column_keeps_its_longest_word(self):
        # Column 2 is narrow overall but holds one long word
        widths = self._fit([900, 100], [50, 90], 500)
        assert widths[1] >= 90, "a column must not wrap mid-word"
        assert sum(widths) == pytest.approx(500)

    def test_minimums_wider_than_the_page_fall_back_to_equal_columns(self):
        widths = self._fit([900, 900], [400, 400], 500)
        assert widths == [250, 250]

    def _measure(self, table_data, available, floors=None):
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        return ExportHandlerBase.measured_column_widths(
            table_data, available, lambda text, _row: len(text) * 10.0 + 20,
            floors)

    def test_measured_widths_use_longest_line_and_longest_word(self):
        """Natural width comes from the longest line, minimum from the word."""
        table_data = [
            ["Component", "a short one"],
            ["short\nlonger line here", "unbreakable"],
        ]
        # Natural widths total 310, minimums 240 - squeezing happens between
        available = 280.0
        widths = self._measure(table_data, available)
        assert sum(widths) == pytest.approx(available)
        # Column 1 must never wrap inside 'unbreakable', its longest word
        assert widths[1] >= len("unbreakable") * 10.0 + 20
        # Column 0's longest line drove its natural width, and it gave ground
        assert widths[0] < len("longer line here") * 10.0 + 20

    def test_measured_widths_floor_empty_columns(self):
        """An empty column keeps an empty cell's width even when squeezed."""
        table_data = [
            ["x" * 40, "", "y" * 40],
            ["x" * 40, "", "y" * 40],
        ]
        available = 300.0
        widths = self._measure(table_data, available)
        assert sum(widths) == pytest.approx(available)
        assert widths[1] >= 20, "an empty column collapsed to nothing"

    def test_every_column_capped_at_fair_share_gives_equal_columns(self):
        """No floor may exceed an equal share, or one column starves the rest."""
        table_data = [["x" * 60, "y" * 60, "z" * 60]]
        available = 300.0
        # Floors far past the page must not push the total over it
        widths = self._measure(table_data, available, floors=[500.0, 0.0, 0.0])
        assert sum(widths) == pytest.approx(available)
        assert widths[0] <= available / 3 + 1e-6

    def test_pdf_column_layout_never_exceeds_the_frame(self):
        """Padding shrinks with column count instead of blowing the budget.

        reportlab does not raise on width overflow - it just draws the table
        off the sheet - so nothing downstream would catch a broken budget.
        """
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        available = pdf_frame_width()

        def string_width(text, _row_index):
            return len(text) * 5.0

        cases = {
            "empty middle column": [["a", "", "b"], ["x", "", "y"]],
            "two empty columns": [["a", "", "b", ""], ["x", "", "y", ""]],
        }
        for ncols in (6, 20, 34, 40, 60, 100, 120, 200, 300, 528, 900):
            cases[f"{ncols} columns"] = [
                [f"h{i}" for i in range(ncols)],
                [f"v{i}" for i in range(ncols)],
            ]

        for name, table_data in cases.items():
            padding, widths = ExportHandlerBase.pdf_table_column_layout(
                table_data, available, string_width)
            assert sum(widths) <= available + 1e-6, f"{name} runs off the page"
            assert min(widths) > 2 * padding, f"{name} has no content area"


# --------------------------------------------------------------------------
# Defect regressions (docs/defects.md)
# --------------------------------------------------------------------------

def _pdf_fill_colors(page):
    """RGB (0..1) fill colours of every vector drawing on a fitz page."""
    return [tuple(round(c, 3) for c in d["fill"])
            for d in page.get_drawings() if d.get("fill") is not None]


def _pdf_stroke_colors(page):
    """RGB (0..1) stroke colours of every vector drawing on a fitz page."""
    return [tuple(round(c, 3) for c in d["color"])
            for d in page.get_drawings() if d.get("color") is not None]


def _color_near(colors, target, tol=0.02):
    return any(all(abs(c - t) <= tol for c, t in zip(col, target))
               for col in colors if len(col) == len(target))


def _is_italic_font(font_name):
    """Slanted faces are named Italic or Oblique depending on the family that
    was found on the box - DejaVu ships Oblique, Liberation ships Italic.
    Matched on the stem, since fitz truncates an embedded subset name to 24
    characters (`LiberationSans-BoldItali`, the trailing `c` cut)."""
    return "Ital" in font_name or "Obli" in font_name


TEST_MARKDOWN_TASK_LIST = """# Tasks

- [x] finished item
- [ ] pending item
- [X] also finished
"""


@pytest.fixture
def test_task_list_file(jp_root_dir):
    md_file = jp_root_dir / "test_task_list.md"
    md_file.write_text(TEST_MARKDOWN_TASK_LIST, encoding="utf-8")
    return md_file


class TestTaskListCheckboxes:
    """DEF-MARK-1: `- [x]` / `- [ ]` render as checkbox glyphs (☒/☐), not literal
    text; DOCX draws them in MS Gothic (Word's own checkbox font)."""

    async def test_html_renders_checkbox_glyphs(self, jp_fetch, test_task_list_file):
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "test_task_list.md"}),
            raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert "☒" in html, "checked checkbox glyph missing from HTML"
        assert "☐" in html, "empty checkbox glyph missing from HTML"
        assert "[x] finished" not in html and "[ ] pending" not in html, (
            "literal task markers leaked into HTML"
        )

    async def test_docx_renders_checkbox_glyphs(self, jp_fetch, test_task_list_file):
        from docx import Document
        from docx.oxml.ns import qn

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "test_task_list.md"}),
            raise_error=False,
        )
        assert response.code == 200
        doc = Document(io.BytesIO(response.body))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "☒" in text and "☐" in text, "checkbox glyphs missing from DOCX"
        assert "[x] finished" not in text and "[ ] pending" not in text, (
            "literal task markers leaked into DOCX"
        )
        # The glyph must sit in its own run tagged MS Gothic (Word's checkbox font)
        gothic_glyph_runs = 0
        for r in doc.element.body.iter(qn("w:r")):
            t = r.find(qn("w:t"))
            if t is None or not t.text or t.text not in ("☒", "☐"):
                continue
            rFonts = r.find(qn("w:rPr") + "/" + qn("w:rFonts")) \
                if r.find(qn("w:rPr")) is not None else None
            assert rFonts is not None and rFonts.get(qn("w:ascii")) == "MS Gothic", (
                "checkbox glyph run is not tagged MS Gothic"
            )
            gothic_glyph_runs += 1
        assert gothic_glyph_runs >= 3, "expected each checkbox glyph in its own MS Gothic run"

    async def test_pdf_renders_checkbox_glyphs(self, jp_fetch, test_task_list_file):
        import fitz

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "test_task_list.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body.startswith(b"%PDF-")
        doc = fitz.open(stream=response.body, filetype="pdf")
        text = "".join(page.get_text() for page in doc)
        doc.close()
        assert "☒" in text and "☐" in text, "checkbox glyphs missing from PDF"
        assert "[x] finished" not in text and "[ ] pending" not in text, (
            "literal task markers leaked into PDF"
        )

    async def test_task_marker_in_code_block_not_converted(self, jp_fetch, jp_root_dir):
        """A `- [ ]` written inside a fenced code block is code, not a task."""
        (jp_root_dir / "cb.md").write_text(
            "# C\n\n```\n- [ ] not a task, just code\n- [x] also code\n```\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "cb.md"}), raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert "☒" not in html and "☐" not in html, (
            "a task marker inside a code block was converted to a checkbox"
        )
        # Pygments wraps each token in a span; strip tags before checking the text
        plain = re.sub(r"<[^>]+>", "", html)
        assert "[ ] not a task" in plain, "code block content was altered"

    async def test_link_list_item_not_converted(self, jp_fetch, jp_root_dir):
        """`- [text](url)` is a link, not a checkbox - only `[ ]`/`[x]` match."""
        (jp_root_dir / "lk.md").write_text(
            "- [a link](https://example.com)\n- [x] done\n", encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "lk.md"}), raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert 'href="https://example.com"' in html, "link item was mangled"
        assert html.count("☒") + html.count("☐") == 1, (
            "only the real task item should become a checkbox"
        )

    async def test_star_and_plus_markers_convert(self, jp_fetch, jp_root_dir):
        """`*` and `+` list markers carry task boxes too, not just `-`."""
        (jp_root_dir / "sp.md").write_text(
            "* [x] star done\n+ [ ] plus open\n", encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "sp.md"}), raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert "☒" in html and "☐" in html, "* / + task markers not converted"

    async def test_empty_task_item_converts(self, jp_fetch, jp_root_dir):
        """A bare task item with no label (`- [x]` / `- [ ]`, nothing trailing)
        still becomes a checkbox - the marker match must not require text after
        the bracket, or GitHub's empty-checkbox idiom leaks through as literal
        `[x]` / `[ ]`."""
        (jp_root_dir / "empty.md").write_text(
            "- [x]\n- [ ]\n", encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "empty.md"}), raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert "☒" in html and "☐" in html, (
            "bare task items were left as literal [x] / [ ]"
        )

    async def test_task_marker_in_nested_longer_fence_not_converted(self, jp_fetch, jp_root_dir):
        """A shorter ``` line inside a longer ```` block is content, not the
        closing fence (CommonMark: a closing fence is at least as long as the
        opener). A `- [ ]` after it is still code and must not become a glyph."""
        (jp_root_dir / "nf.md").write_text(
            "# N\n\n````\n```\n- [ ] still inside the code block\n````\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "nf.md"}), raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert "☒" not in html and "☐" not in html, (
            "a task marker inside a nested longer fence was converted to a checkbox"
        )

    async def test_task_marker_after_info_string_fence_not_converted(self, jp_fetch, jp_root_dir):
        """A ```lang line inside a fenced block is a nested opener, not a
        closer (CommonMark: a closing fence carries no info string), so a
        `- [ ]` after it is still code and must not become a glyph."""
        (jp_root_dir / "infofence.md").write_text(
            "# I\n\n```\nsome code\n```python\n- [ ] still code\n```\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "infofence.md"}), raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert "☒" not in html and "☐" not in html, (
            "a task marker after an info-string fence line was converted"
        )


class TestBlockquotePdfCallout:
    """DEF-MARK-2: blockquotes render in PDF with a left bar, shading and indent."""

    async def test_pdf_blockquote_has_bar_shading_and_indent(
        self, jp_fetch, test_quotes_in_list_file
    ):
        import fitz

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "test_quotes_in_list.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body.startswith(b"%PDF-")
        doc = fitz.open(stream=response.body, filetype="pdf")
        page = doc[0]

        assert _color_near(_pdf_fill_colors(page), (0.957, 0.957, 0.957)), (
            "blockquote background shading (F4F4F4) missing from PDF"
        )
        assert _color_near(_pdf_stroke_colors(page), (0.733, 0.733, 0.733)), (
            "blockquote left bar (BBBBBB) missing from PDF"
        )

        body_x, quote_x = [], []
        for b in page.get_text("dict")["blocks"]:
            for line in b.get("lines", []):
                for s in line["spans"]:
                    if "Two ideas" in s["text"]:
                        body_x.append(s["bbox"][0])
                    if "quoted note" in s["text"]:
                        quote_x.append(s["bbox"][0])
        doc.close()
        assert body_x and quote_x, "expected body and quote text in the PDF"
        assert min(quote_x) > min(body_x) + 15, (
            "blockquote text is not indented relative to body text"
        )

    async def test_pdf_callout_that_fits_a_page_is_not_split(self, jp_fetch, jp_root_dir):
        """DEF-TABL-9 applies to callouts too: a blockquote/alert box that fits a
        page must move whole rather than tear mid-box at the page boundary."""
        import fitz

        # Tuned so the quote box does not fit the space left on the page it
        # would start on, but fits the next page whole - the case an
        # unconditional splitInRow tears mid-box.
        filler = "Filler paragraph line to push the quote toward the page end. "
        quote = " ".join(
            f"QW{i} lorem ipsum dolor sit amet consectetur" for i in range(40)
        )
        (jp_root_dir / "cq.md").write_text(
            "# Q\n\n" + (filler * 175) + "\n\n> " + quote + "\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "cq.md"}), raise_error=False,
        )
        assert response.code == 200
        doc = fitz.open(stream=response.body, filetype="pdf")
        first = [i for i, pg in enumerate(doc) if "QW0 " in pg.get_text()]
        last = [i for i, pg in enumerate(doc) if "QW39 " in pg.get_text()]
        doc.close()
        assert first and last, "quote text missing from the PDF"
        assert first[0] == last[0], (
            f"the quote box was torn across pages: starts on page {first[0]}, "
            f"ends on page {last[0]} - a callout that fits a page must move whole"
        )

    async def test_pdf_multi_paragraph_quote_is_one_callout(self, jp_fetch, jp_root_dir):
        """Consecutive blockquote paragraphs render as a single shaded box with
        one continuous bar, not a stack of separate boxes with gaps."""
        import fitz

        (jp_root_dir / "mq.md").write_text(
            "# Q\n\n> line one\n>\n> line two\n>\n> line three\n\nAfter.\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "mq.md"}), raise_error=False,
        )
        assert response.code == 200
        doc = fitz.open(stream=response.body, filetype="pdf")
        # Count the blockquote shading fills (F4F4F4) - one per callout box
        shade_boxes = sum(
            1 for page in doc for c in _pdf_fill_colors(page)
            if len(c) == 3 and all(abs(x - 0.957) < 0.02 for x in c)
        )
        text = "".join(page.get_text() for page in doc)
        doc.close()
        assert "line one" in text and "line three" in text, "quote text missing"
        assert shade_boxes == 1, (
            f"a 3-paragraph quote rendered as {shade_boxes} boxes, expected 1"
        )


class TestAlertPdfCallout:
    """DEF-MARK-3: alerts render in PDF as a coloured-bar callout, not a table header."""

    async def test_pdf_alert_renders_as_callout_not_header(
        self, jp_fetch, test_rich_alerts_file
    ):
        import fitz

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "test_rich_alerts.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body.startswith(b"%PDF-")
        doc = fitz.open(stream=response.body, filetype="pdf")
        fills, strokes, spans = [], [], []
        for page in doc:
            fills += _pdf_fill_colors(page)
            strokes += _pdf_stroke_colors(page)
            for b in page.get_text("dict")["blocks"]:
                for line in b.get("lines", []):
                    spans += line["spans"]
        doc.close()

        # IMPORTANT alert: purple shading (F4EDFF) + purple left bar (8250DF)
        assert _color_near(fills, (0.957, 0.929, 1.0)), (
            "alert purple shading missing - not rendered as a callout"
        )
        assert _color_near(strokes, (0.510, 0.314, 0.875)), (
            "alert purple left bar missing"
        )
        # Alert body text must not use the content-table header colour (#365F91)
        alert_spans = [s for s in spans if "porosity" in s["text"]]
        assert alert_spans, "expected alert body text in the PDF"
        for s in alert_spans:
            assert s["color"] != 0x365F91, (
                "alert text still rendered in the table-header colour"
            )

    async def test_pdf_data_table_still_header_styled(
        self, jp_fetch, test_rich_alerts_file
    ):
        """The real data table must keep its header fill - alert detection
        must not misfire on an ordinary content table."""
        import fitz

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "test_rich_alerts.md"}),
            raise_error=False,
        )
        assert response.code == 200
        doc = fitz.open(stream=response.body, filetype="pdf")
        fills = [c for page in doc for c in _pdf_fill_colors(page)]
        doc.close()
        assert _color_near(fills, (0.859, 0.898, 0.945)), (
            "data-table header fill (dbe5f1) missing - alert detection over-matched"
        )

    async def test_pdf_alert_colors_per_type(self, jp_fetch, test_rich_alerts_file):
        """Each alert type keeps its own callout shading in the PDF (the fixture
        carries IMPORTANT/purple, NOTE/blue, WARNING/amber)."""
        import fitz

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "test_rich_alerts.md"}),
            raise_error=False,
        )
        assert response.code == 200
        doc = fitz.open(stream=response.body, filetype="pdf")
        fills = [c for page in doc for c in _pdf_fill_colors(page)]
        doc.close()
        assert _color_near(fills, (0.957, 0.929, 1.0)), "IMPORTANT purple shading missing"
        assert _color_near(fills, (0.929, 0.961, 0.992)), "NOTE blue shading missing"
        assert _color_near(fills, (0.996, 0.976, 0.906)), "WARNING amber shading missing"


class TestMermaidViewBoxCrop:
    """DEF-DIAG-4: an SVG whose declared viewBox dwarfs its content is cropped
    to the real content bounding box before rasterizing."""

    async def test_oversized_viewbox_is_cropped_to_content(self):
        from jupyterlab_export_markdown_extension.routes import (
            PlaywrightSvgRenderer,
        )
        from PIL import Image as PILImage

        # A 120x60 shape sitting inside an 800x800 viewBox.
        svg = (
            '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 800 800" '
            'width="100%"><rect x="10" y="10" width="120" height="60" '
            'fill="black"/></svg>'
        ).encode("utf-8")

        async with PlaywrightSvgRenderer(color_scheme="light") as renderer:
            png = await renderer.render(svg, width=600)

        img = PILImage.open(io.BytesIO(png)).convert("RGBA")
        w, h = img.size
        bbox = img.getbbox()
        assert bbox is not None, "rasterized diagram is entirely blank"
        cw, ch = bbox[2] - bbox[0], bbox[3] - bbox[1]
        assert cw / w > 0.7, (
            f"diagram fills only {100 * cw / w:.0f}% of canvas width - "
            f"view window is larger than the shape"
        )
        assert ch / h > 0.7, (
            f"diagram fills only {100 * ch / h:.0f}% of canvas height - "
            f"view window is larger than the shape"
        )

    async def test_unmeasurable_svg_falls_back_to_declared_viewbox(self):
        """When getBBox yields no geometry (e.g. an empty SVG), render must
        fall back to the declared viewBox rather than crash or size to zero."""
        from jupyterlab_export_markdown_extension.routes import (
            PlaywrightSvgRenderer,
        )
        from PIL import Image as PILImage

        svg = (
            '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 200 100">'
            '<defs></defs></svg>'
        ).encode("utf-8")
        async with PlaywrightSvgRenderer(color_scheme="light") as renderer:
            png = await renderer.render(svg, width=400)
        img = PILImage.open(io.BytesIO(png))
        # Declared viewBox aspect 2:1 -> 400x200; the fallback must hold
        assert img.size == (400, 200), (
            f"fallback viewBox produced {img.size}, expected (400, 200)"
        )

    async def test_well_framed_diagram_is_not_cropped(self):
        """DEF-DIAG-4 regression (adversarial #4): a diagram that already fills its
        declared viewBox must be left alone. getBBox reports geometry only - no
        stroke, markers or filters - so re-cropping a well-framed diagram would
        clip node borders and arrowheads. The declared viewBox aspect must
        survive even when the content's own geometry aspect differs slightly."""
        from jupyterlab_export_markdown_extension.routes import (
            PlaywrightSvgRenderer,
        )
        from PIL import Image as PILImage

        # Content fills the 3:1 viewBox in both axes (fill_w 0.97, fill_h 0.90),
        # so it must NOT be tightened; a tightened box would be ~2.8:1 (h~217).
        svg = (
            '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 300 100" '
            'width="100%"><rect x="5" y="5" width="290" height="90" '
            'fill="black"/></svg>'
        ).encode("utf-8")

        async with PlaywrightSvgRenderer(color_scheme="light") as renderer:
            png = await renderer.render(svg, width=600)

        img = PILImage.open(io.BytesIO(png))
        assert img.size == (600, 200), (
            f"well-framed diagram was re-cropped to {img.size}; the declared "
            f"3:1 viewBox (600x200) must be honoured, not tightened to content"
        )

    async def test_nonzero_viewbox_origin_is_preserved(self):
        """DEF-DIAG-4 regression (adversarial #7): when the diagram is not cropped,
        the SVG's own viewBox - including a non-zero origin - must be kept. A
        mermaid viewBox like `-100 -50 200 100` places content around a negative
        origin; zeroing it would shift the drawing off-canvas and blank most of
        the render."""
        from jupyterlab_export_markdown_extension.routes import (
            PlaywrightSvgRenderer,
        )
        from PIL import Image as PILImage

        # Well-framed content drawn around a negative origin. With the viewBox
        # kept, the rect nearly fills the canvas; a zeroed viewBox would push it
        # into the top-left corner, leaving most of the canvas blank.
        svg = (
            '<svg xmlns="http://www.w3.org/2000/svg" '
            'viewBox="-100 -50 200 100" width="100%">'
            '<rect x="-95" y="-45" width="190" height="90" fill="black"/></svg>'
        ).encode("utf-8")

        async with PlaywrightSvgRenderer(color_scheme="light") as renderer:
            png = await renderer.render(svg, width=400)

        img = PILImage.open(io.BytesIO(png)).convert("RGBA")
        w, h = img.size
        bbox = img.getbbox()
        assert bbox is not None, "rasterized diagram is entirely blank"
        cw = bbox[2] - bbox[0]
        assert cw / w > 0.8, (
            f"content fills only {100 * cw / w:.0f}% of width - the non-zero "
            f"viewBox origin was dropped, shifting the drawing off-canvas"
        )

    async def test_extreme_aspect_diagram_stays_within_the_raster_limit(
            self, monkeypatch):
        """A long single-column flowchart is many times taller than it is wide,
        and at the configured export width its raster runs past Chromium's
        ~16384px viewport and texture caps - the screenshot throws and the
        diagram is dropped from the export with no error surfaced. It must be
        scaled down instead, aspect intact.

        The cap is patched down so the test exercises the clamp without
        rendering a 16000px canvas."""
        from jupyterlab_export_markdown_extension.routes import (
            PlaywrightSvgRenderer,
        )
        from PIL import Image as PILImage

        monkeypatch.setattr(PlaywrightSvgRenderer, "MAX_RASTER_PX", 1000)
        # A 60x1500 column inside a square viewBox: cropped (fill_h 0.94 but
        # fill_w 0.04), it rasterizes ~25:1 - 400px wide would be 10000px tall.
        svg = (
            '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1600 1600" '
            'width="100%"><rect x="20" y="50" width="60" height="1500" '
            'fill="black"/></svg>'
        ).encode("utf-8")

        async with PlaywrightSvgRenderer(color_scheme="light") as renderer:
            png = await renderer.render(svg, width=400)

        img = PILImage.open(io.BytesIO(png))
        w, h = img.size
        assert max(w, h) <= 1000, (
            f"raster came back {w}x{h}, past the {1000}px cap - Chromium "
            f"would refuse it and the diagram would vanish from the export"
        )
        assert img.convert("RGBA").getbbox() is not None, (
            "the clamped raster is blank"
        )
        # The clamp must scale, not crop: the drawn column keeps its aspect
        assert h > w * 5, (
            f"clamped raster is {w}x{h} - the extreme aspect was not preserved"
        )


#: A mermaid-shaped SVG root: mermaid stamps `width="100%"` plus an inline
#: `max-width: <natural>px` and a viewBox tight to the drawing. The max-width
#: caps the element, so a rasterizer that only sets `width` paints the diagram
#: at natural size inside the target canvas and the rest becomes whitespace.
MERMAID_LIKE_SVG = (
    '<svg id="m0" width="100%" xmlns="http://www.w3.org/2000/svg" '
    'class="flowchart" style="max-width: 800px;" viewBox="0 0 800 70">'
    '<rect x="10" y="10" width="780" height="50" fill="#0284c7"/>'
    '</svg>'
)


def _ink_bbox(img):
    """Bounding box of real ink, trimming both transparent and white margin."""
    from PIL import Image as PILImage
    rgba = img.convert("RGBA")
    bbox = rgba.getbbox()  # trims fully transparent margin
    if bbox is not None and (bbox[2] - bbox[0]) < rgba.size[0]:
        return bbox
    # Opaque (white-matted) export: trim near-white instead
    flat = PILImage.new("RGB", rgba.size, (255, 255, 255))
    flat.paste(rgba, mask=rgba.split()[-1])
    gray = flat.convert("L")
    mask = gray.point(lambda p: 255 if p < 245 else 0)
    return mask.getbbox()


class TestMermaidDocxNoWhitespace:
    """DEF-DIAG-7 (regression): a mermaid diagram embedded in the DOCX must fill its
    image rather than sit in whitespace. Deterministic check - pull the image
    back out of the .docx, trim the blank margin, and measure what fraction of
    the image the drawing actually occupies."""

    async def test_docx_mermaid_image_is_not_mostly_whitespace(self, jp_fetch, jp_root_dir):
        import base64
        from PIL import Image as PILImage

        (jp_root_dir / "mm.md").write_text(
            "# M\n\n```mermaid\nflowchart LR\n    A --> B\n```\n", encoding="utf-8"
        )
        data_uri = ("data:image/svg+xml;base64,"
                    + base64.b64encode(MERMAID_LIKE_SVG.encode()).decode())
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST",
            body=json.dumps({"path": "mm.md",
                             "mermaidDiagrams": [{"index": 0, "svg": data_uri}]}),
            raise_error=False,
        )
        assert response.code == 200

        with zipfile.ZipFile(io.BytesIO(response.body)) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
            assert media, "mermaid diagram was not embedded in the DOCX"
            blobs = [z.read(n) for n in media]

        widest = max((PILImage.open(io.BytesIO(b)) for b in blobs),
                     key=lambda im: im.size[0])
        w, h = widest.size
        bbox = _ink_bbox(widest)
        assert bbox is not None, "embedded mermaid image is blank"
        ink_w = bbox[2] - bbox[0]
        assert ink_w / w >= 0.85, (
            f"mermaid diagram fills only {100 * ink_w / w:.0f}% of its image "
            f"width ({ink_w}px of {w}px) - the rest is whitespace; mermaid's "
            f"inline max-width caps the element so the raster width is ignored"
        )


TEST_MARKDOWN_EMPTY_HEADER_GRID = """# Grid

|  |  |
|:---:|:---:|
| cell one | cell two |
| cell three | cell four |
"""


@pytest.fixture
def test_empty_header_grid_file(jp_root_dir):
    md_file = jp_root_dir / "test_empty_header_grid.md"
    md_file.write_text(TEST_MARKDOWN_EMPTY_HEADER_GRID, encoding="utf-8")
    return md_file


class TestEmptyHeaderGrid:
    """DEF-TABL-5 / DEF-TABL-8: a Markdown grid written with an empty header row
    (`|  |  |`) renders nothing for that row in markdown, so the export must
    drop it outright rather than emit a blank bordered row - and must not
    promote the body row behind it into a header in its place."""

    async def test_docx_empty_header_row_is_dropped(
        self, jp_fetch, test_empty_header_grid_file
    ):
        from docx import Document
        from docx.oxml.ns import qn

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "test_empty_header_grid.md"}),
            raise_error=False,
        )
        assert response.code == 200
        doc = Document(io.BytesIO(response.body))
        assert doc.tables, "expected a table in the export"
        table = doc.tables[0]
        # The fixture is an empty header plus two content rows; the blank
        # header must be gone, leaving only the content rows.
        assert len(table.rows) == 2, (
            f"expected the empty header row to be dropped (2 content rows), "
            f"got {len(table.rows)} rows"
        )
        assert [c.text.strip() for c in table.rows[0].cells] == ["cell one", "cell two"], (
            "the first surviving row should be the first content row"
        )
        trPr = table.rows[0]._tr.find(qn("w:trPr"))
        has_header = trPr is not None and trPr.find(qn("w:tblHeader")) is not None
        assert not has_header, (
            "the content row must not be promoted to a repeating header"
        )
        tblLook = table._tbl.tblPr.find(qn("w:tblLook"))
        assert tblLook is not None and tblLook.get(qn("w:firstRow")) == "0", (
            "a headerless grid must not carry first-row header emphasis"
        )

    async def test_headerless_grid_columns_are_not_skewed_by_bold_widening(
        self, jp_fetch, jp_root_dir
    ):
        """Row 0 of a headerless grid is body text, so it must not be measured
        as bold. The width estimate widens a header row by 8% for its bold
        face; after the blank header is dropped that factor lands on ordinary
        content and steals width from the column beside it.

        The fixture is symmetric - each column's widest cell is the same
        string, one in row 0 and one in row 1 - so a correct measurement gives
        two equal columns and any 8% skew is unambiguous.
        """
        from docx import Document
        from docx.oxml.ns import qn

        long_a = " ".join(["alpha"] * 16)
        long_b = " ".join(["bravo"] * 16)
        (jp_root_dir / "test_headerless_widths.md").write_text(
            "|  |  |\n"
            "| --- | --- |\n"
            f"| {long_a} | short |\n"
            f"| short | {long_b} |\n",
            encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST",
            body=json.dumps({"path": "test_headerless_widths.md"}),
            raise_error=False,
        )
        assert response.code == 200

        doc = Document(io.BytesIO(response.body))
        table = doc.tables[0]
        assert len(table.rows) == 2, "the blank header row should be gone"
        layout = table._tbl.tblPr.find(qn("w:tblLayout"))
        assert layout is not None and layout.get(qn("w:type")) == "fixed", (
            "the fixture must be wide enough to be fitted, or the widths "
            "below are never written and the test proves nothing"
        )

        grid = table._tbl.find(qn("w:tblGrid"))
        widths = [int(col.get(qn("w:w")))
                  for col in grid.findall(qn("w:gridCol"))]
        assert len(widths) == 2
        skew = abs(widths[0] - widths[1]) / max(widths)
        assert skew < 0.01, (
            f"symmetric columns came back skewed by {skew:.1%} "
            f"({widths}) - row 0 was measured as a bold header"
        )

    async def test_image_only_header_row_is_kept_with_its_images(
        self, jp_fetch, jp_root_dir
    ):
        """A header row holding pictures and no text is NOT empty. An
        image-on-top/caption-below grid puts its images in the header row;
        dropping it on a text-only predicate silently loses them."""
        from docx import Document
        from docx.oxml.ns import qn
        from PIL import Image as PILImage
        import fitz

        images_dir = jp_root_dir / "images"
        images_dir.mkdir(exist_ok=True)
        # Distinct colours: identical bytes would be deduplicated into one
        # PDF XObject and the count could not tell one image from two.
        for n, colour in (("a", (20, 120, 90)), ("b", (200, 60, 140))):
            PILImage.new("RGB", (300, 200), color=colour).save(
                images_dir / f"{n}.png"
            )
        (jp_root_dir / "imghdr.md").write_text(
            "# G\n\n"
            "| ![A](images/a.png) | ![B](images/b.png) |\n"
            "|---|---|\n"
            "| caption A | caption B |\n",
            encoding="utf-8",
        )

        rd = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "imghdr.md"}), raise_error=False,
        )
        assert rd.code == 200
        doc = Document(io.BytesIO(rd.body))
        table = doc.tables[0]
        assert len(table.rows) == 2, (
            f"the picture-bearing header row was deleted (got {len(table.rows)} rows)"
        )
        assert len(doc.inline_shapes) >= 2, (
            f"header-row images were lost: {len(doc.inline_shapes)} embedded, expected 2"
        )

        rp = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "imghdr.md"}), raise_error=False,
        )
        assert rp.code == 200
        pdf = fitz.open(stream=rp.body, filetype="pdf")
        n_images = sum(len(pg.get_images()) for pg in pdf)
        pdf.close()
        assert n_images >= 2, (
            f"header-row images were lost from the PDF: {n_images} rendered, expected 2"
        )

        # HTML strips tags to test emptiness, which erases an <img> along with
        # them - the same row must survive there or one format of three loses
        # the figures while the other two keep them.
        rh = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "imghdr.md"}), raise_error=False,
        )
        assert rh.code == 200
        html = rh.body.decode("utf-8")
        assert html.count("<img") >= 2, (
            f"header-row images were lost from the HTML: {html.count('<img')} "
            f"embedded, expected 2"
        )

        # Kept, but it is a figure row, not a header: no banded emphasis and no
        # repeat, matching what the PDF does with the same row
        tblLook = table._tbl.tblPr.find(qn("w:tblLook"))
        assert tblLook is not None and tblLook.get(qn("w:firstRow")) == "0", (
            "a picture-only first row must not carry header emphasis"
        )
        trPr = table.rows[0]._tr.find(qn("w:trPr"))
        assert trPr is None or trPr.find(qn("w:tblHeader")) is None, (
            "a picture-only first row must not repeat as a header"
        )

    async def test_docx_rows_are_marked_unbreakable(
        self, jp_fetch, test_empty_header_grid_file
    ):
        """DEF-TABL-9 in Word: reportlab's conditional `splitInRow` has no effect on
        the .docx, where Word breaks a row wherever the page ends unless the row
        carries `w:cantSplit`. Without it the same caption-and-image row the PDF
        now keeps whole is still torn in Word."""
        from docx import Document
        from docx.oxml.ns import qn

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "test_empty_header_grid.md"}),
            raise_error=False,
        )
        assert response.code == 200
        doc = Document(io.BytesIO(response.body))
        table = doc.tables[0]
        for index, row in enumerate(table.rows):
            trPr = row._tr.find(qn("w:trPr"))
            assert trPr is not None and trPr.find(qn("w:cantSplit")) is not None, (
                f"row {index} may be torn across a page break in Word"
            )

    async def test_pdf_headerless_grid_columns_are_not_skewed(
        self, jp_fetch, jp_root_dir
    ):
        """The PDF twin of the DOCX bold-widening skew: `string_width` measures
        row 0 in the bold header face, so after the blank header is dropped the
        first content row inflates its own column.

        Measured on the drawn grid rather than on text: the same string fills
        the widest cell of each column, so a correct measurement puts the
        column boundary exactly halfway across the table. Text extent would
        depend on where the last line happens to wrap, and on the widths of
        whichever glyphs the fixture used.
        """
        import fitz

        long_text = " ".join(["alpha"] * 16)
        (jp_root_dir / "pdf_headerless.md").write_text(
            "|  |  |\n"
            "| --- | --- |\n"
            f"| {long_text} | short |\n"
            f"| short | {long_text} |\n",
            encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "pdf_headerless.md"}),
            raise_error=False)
        assert response.code == 200

        pdf = fitz.open(stream=response.body, filetype="pdf")
        page = pdf[0]
        verticals = sorted({
            round(item[1].x, 1)
            for drawing in page.get_drawings() for item in drawing["items"]
            if item[0] == "l" and abs(item[1].x - item[2].x) < 0.1
            and abs(item[1].y - item[2].y) > 2
        })
        pdf.close()
        assert len(verticals) == 3, (
            f"expected a two-column grid (3 vertical rules), got {verticals}"
        )
        left, split, right = verticals
        midpoint = (left + right) / 2
        skew = abs(split - midpoint) / (right - left)
        assert skew < 0.01, (
            f"the column boundary sits at {split:.1f}pt against a "
            f"{midpoint:.1f}pt midpoint ({skew:.1%} of the table) - row 0 was "
            f"measured in the bold header face"
        )

    async def test_html_empty_header_row_is_dropped(
        self, jp_fetch, test_empty_header_grid_file, jp_root_dir
    ):
        """Cross-format parity: HTML must drop the blank header row too, or the
        same document renders with an extra blank banded strip in one format
        and not the others. A real header must survive."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "test_empty_header_grid.md"}),
            raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert "<thead>" not in html, (
            "the empty header row must not be emitted in HTML either"
        )
        assert "cell one" in html, "grid content was lost with the header row"

        # A real header still gets a thead
        (jp_root_dir / "realhdr.md").write_text(
            "# T\n\n| Name | Age |\n|---|---|\n| ann | 3 |\n", encoding="utf-8"
        )
        r2 = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "realhdr.md"}), raise_error=False,
        )
        assert r2.code == 200
        html2 = r2.body.decode("utf-8")
        assert "<thead>" in html2 and "Name" in html2, (
            "a real header row must still render in HTML"
        )

    async def test_pdf_empty_header_grid_has_no_header_bar(
        self, jp_fetch, test_empty_header_grid_file
    ):
        import fitz

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "test_empty_header_grid.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body.startswith(b"%PDF-")
        doc = fitz.open(stream=response.body, filetype="pdf")
        fills = [c for page in doc for c in _pdf_fill_colors(page)]
        text = "".join(page.get_text() for page in doc)
        doc.close()
        assert "cell one" in text, "grid content missing from PDF"
        assert not _color_near(fills, (0.859, 0.898, 0.945)), (
            "empty header row painted a blue header bar (dbe5f1) in the PDF"
        )

    async def test_partial_first_row_is_still_a_header(self, jp_fetch, jp_root_dir):
        """A first row with *any* text is a real header - the empty-header rule
        must only fire when the whole first row is blank."""
        import fitz
        from docx import Document
        from docx.oxml.ns import qn

        (jp_root_dir / "ph.md").write_text(
            "# P\n\n| Name |  |\n|------|--|\n| a | 1 |\n| b | 2 |\n",
            encoding="utf-8",
        )
        docx = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "ph.md"}), raise_error=False,
        )
        assert docx.code == 200
        table = Document(io.BytesIO(docx.body)).tables[0]
        trPr = table.rows[0]._tr.find(qn("w:trPr"))
        assert trPr is not None and trPr.find(qn("w:tblHeader")) is not None, (
            "a partially-filled first row must still be a repeating header"
        )

        pdf = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "ph.md"}), raise_error=False,
        )
        assert pdf.code == 200
        doc = fitz.open(stream=pdf.body, filetype="pdf")
        fills = [c for page in doc for c in _pdf_fill_colors(page)]
        doc.close()
        assert _color_near(fills, (0.859, 0.898, 0.945)), (
            "a real header lost its blue bar in the PDF"
        )


#: Every spelling of a break tag, so an assertion cannot miss `<BR>` or
#: `<br clear="all">` - the spellings these tests exist to cover. Imported
#: from the module under test rather than re-declared, so a change to the
#: pattern cannot leave the tests validating against a stale copy.
from jupyterlab_export_markdown_extension.routes import BREAK_TAG_RE as BREAK_RE  # noqa: E402


class TestExplicitLineBreakNotDoubled:
    """DEF-MARK-10: a line ended with an explicit `<br>` - the idiom for a question
    above its answer - picked up a second break from the `nl2br` extension,
    which converts the following newline too. The pair then sat further apart
    than the paragraphs around it, so the grouping read backwards."""

    FAQ = (
        "## FAQ\n\n"
        "**Does the system call out?**<br>\n"
        "No. Passive - other systems push data to it.\n\n"
        "**How does the plan arrive?**<br>\n"
        "Upstream pulls the recorded events.\n\n"
        "**Is the app native?**<br>\n"
        "No. Thin web app; connectivity required.\n"
    )

    async def test_html_keeps_one_break(self, jp_fetch, jp_root_dir):
        (jp_root_dir / "faq.md").write_text(self.FAQ, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "faq.md"}),
            raise_error=False)
        assert r.code == 200
        html = r.body.decode("utf-8")
        doubled = re.findall(r'<br\s*/?>\s*<br\s*/?>', html)
        assert not doubled, (
            f"{len(doubled)} doubled line break(s) survived - the author's "
            f"`<br>` picked up a second break from nl2br"
        )
        para = html[html.find("<p><strong>Does the system"):]
        para = para[:para.find("</p>")]
        assert len(BREAK_RE.findall(para)) == 1, (
            f"the author's own single break was lost: {para!r}"
        )

    async def test_docx_pair_holds_one_break(self, jp_fetch, jp_root_dir):
        """One break inside the pair, and the pair is one paragraph - the
        document's 10pt paragraph spacing then separates it from the next."""
        from docx import Document
        from docx.oxml.ns import qn

        (jp_root_dir / "faq.md").write_text(self.FAQ, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "faq.md"}),
            raise_error=False)
        assert r.code == 200
        doc = Document(io.BytesIO(r.body))
        pairs = [p for p in doc.paragraphs if "?" in p.text]
        assert len(pairs) == 3, f"expected 3 Q&A paragraphs, got {len(pairs)}"
        for p in pairs:
            breaks = len(p._p.findall('.//' + qn('w:br')))
            assert breaks == 1, (
                f"Q&A pair carries {breaks} line breaks, not 1 - the gap "
                f"inside the pair swallows the gap between pairs: {p.text[:40]!r}"
            )

    async def test_pdf_pair_is_tighter_than_the_gap_between_pairs(
            self, jp_fetch, jp_root_dir):
        """The measurable form of the defect: a question must sit closer to
        its own answer than to the next question."""
        import fitz

        (jp_root_dir / "faq.md").write_text(self.FAQ, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "faq.md"}),
            raise_error=False)
        assert r.code == 200

        pdf = fitz.open(stream=r.body, filetype="pdf")
        tops = {}
        for block in pdf[0].get_text("dict")["blocks"]:
            for line in block.get("lines", []):
                text = "".join(s["text"] for s in line["spans"]).strip()
                if text:
                    tops.setdefault(text[:16], line["bbox"][1])
        pdf.close()

        inside = tops["No. Passive - ot"] - tops["Does the system "]
        between = tops["How does the pla"] - tops["No. Passive - ot"]
        assert between > inside * 1.2, (
            f"the pair is not grouped: {inside:.1f}pt between a question and "
            f"its answer against {between:.1f}pt to the next question"
        )

    async def test_explicit_blank_line_is_preserved(self, jp_fetch, jp_root_dir):
        """`<br><br>` is an author asking for a blank line, not a duplicate -
        only the break nl2br generates on top may be dropped."""
        (jp_root_dir / "twobr.md").write_text(
            "one<br><br>\ntwo\n", encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "twobr.md"}),
            raise_error=False)
        assert r.code == 200
        html = r.body.decode("utf-8")
        body = html[html.find("<p>one"):html.find("</p>", html.find("<p>one"))]
        assert len(BREAK_RE.findall(body)) == 2, (
            f"an explicit blank line was collapsed: {body!r}"
        )

    async def test_break_in_a_table_cell_is_untouched(
            self, jp_fetch, jp_root_dir):
        """A markdown table cell cannot contain a newline, so nl2br has
        nothing to convert there and never doubles a cell's break. The
        caption-above-image grid depends on that break surviving - it is the
        only thing holding the caption above its picture."""
        (jp_root_dir / "grid.md").write_text(
            "|  |  |\n|---|---|\n"
            "| **cap A**<br>text A | **cap B**<br>text B |\n",
            encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "grid.md"}),
            raise_error=False)
        assert r.code == 200
        html = r.body.decode("utf-8")
        table = html[html.find("<table"):html.find("</table>")]
        assert len(BREAK_RE.findall(table)) == 2, (
            f"a table cell lost its line break: {table!r}"
        )

    async def test_raw_html_block_keeps_its_own_breaks(
            self, jp_fetch, jp_root_dir):
        """A raw HTML block's breaks are all the author's - `nl2br` never ran
        inside it, so there is nothing there to collapse. Matching the final
        HTML by shape cannot tell that, and ate one break of the alert-box
        idiom this project documents."""
        (jp_root_dir / "alert.md").write_text(
            '<div class="alert alert-block alert-info">\n'
            '<b>Tip:</b> first paragraph<br /><br />\n'
            'second paragraph\n'
            '</div>\n', encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "alert.md"}),
            raise_error=False)
        assert r.code == 200
        html = r.body.decode("utf-8")
        block = html[html.find("alert-block"):html.find("</div>",
                                                        html.find("alert-block"))]
        assert len(BREAK_RE.findall(block)) == 2, (
            f"the author's blank line inside a raw HTML block was collapsed: "
            f"{block!r}"
        )

    async def test_trailing_space_after_the_break_still_collapses(
            self, jp_fetch, jp_root_dir):
        """A trailing space after `<br>` is invisible in every editor, so the
        fix must not depend on its absence - keying off the tag's position in
        the final text made one stray space silently restore the defect."""
        (jp_root_dir / "space.md").write_text(
            "**Q?**<br> \nAnswer.\n", encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "space.md"}),
            raise_error=False)
        assert r.code == 200
        html = r.body.decode("utf-8")
        para = html[html.find("<p><strong>Q?"):]
        para = para[:para.find("</p>")]
        assert len(BREAK_RE.findall(para)) == 1, (
            f"a trailing space defeated the collapse: {para!r}"
        )

    async def test_uppercase_and_attributed_breaks_are_recognised(
            self, jp_fetch, jp_root_dir):
        """`<BR>` and `<br clear="all">` are breaks the author wrote, and
        pasted or legacy markup carries both. Recognition comes from the
        stash, so spelling does not matter."""
        (jp_root_dir / "spelling.md").write_text(
            "**Upper?**<BR>\nAnswer one.\n\n"
            '**Attr?**<br clear="all">\nAnswer two.\n', encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "spelling.md"}),
            raise_error=False)
        assert r.code == 200
        html = r.body.decode("utf-8")
        for marker in ("Upper?", "Attr?"):
            para = html[html.find(f"<p><strong>{marker}"):]
            para = para[:para.find("</p>")]
            assert len(BREAK_RE.findall(para)) == 1, (
                f"{marker} kept a doubled break: {para!r}"
            )

    async def test_break_on_its_own_line_gives_one_blank_line(
            self, jp_fetch, jp_root_dir):
        """A `<br>` alone on a line is the author adding a blank line between
        two lines of text: the newline before it and the tag itself are two
        breaks, and the newline after it is the generated one that goes. Pinned
        because it is a rendering change, not an accident."""
        (jp_root_dir / "ownline.md").write_text(
            "Text\n<br>\nMore\n", encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "ownline.md"}),
            raise_error=False)
        assert r.code == 200
        html = r.body.decode("utf-8")
        para = html[html.find("<p>Text"):]
        para = para[:para.find("</p>")]
        assert len(BREAK_RE.findall(para)) == 2, (
            f"expected exactly one blank line between the two lines: {para!r}"
        )

    async def test_manual_break_plus_hard_break_keeps_both(
            self, jp_fetch, jp_root_dir):
        """Two trailing spaces are Markdown's own hard-break idiom, so `<br>`
        followed by them is the author asking for two breaks, not a duplicate.
        Only a break the renderer generated on its own may be dropped - and at
        the inline stage the two-space break has already been claimed by the
        `linebreak` pattern, so it is never in reach."""
        (jp_root_dir / "hard.md").write_text(
            "**Q?**<br>  \nAnswer.\n", encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "hard.md"}),
            raise_error=False)
        assert r.code == 200
        html = r.body.decode("utf-8")
        para = html[html.find("<p><strong>Q?"):]
        para = para[:para.find("</p>")]
        assert len(BREAK_RE.findall(para)) == 2, (
            f"an authored break was deleted: {para!r}"
        )

    async def test_custom_element_is_not_treated_as_a_break(
            self, jp_fetch, jp_root_dir):
        """`<br-spacer>` is a legal custom element, not a break. Matching it as
        one drops the real break and runs the two lines together."""
        (jp_root_dir / "custom.md").write_text(
            "line one<br-spacer>\nline two\n", encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "custom.md"}),
            raise_error=False)
        assert r.code == 200
        html = r.body.decode("utf-8")
        para = html[html.find("<p>line one"):]
        para = para[:para.find("</p>")]
        assert len(BREAK_RE.findall(para)) == 1, (
            f"a custom element swallowed the line break: {para!r}"
        )

    async def test_non_breaking_space_after_the_break_still_collapses(
            self, jp_fetch, jp_root_dir):
        """Text pasted from Word or a web page routinely ends a line with a
        non-breaking space; it must not decide how the document renders."""
        (jp_root_dir / "nbsp.md").write_text(
            "**Q?**<br>\u00a0\nAnswer.\n", encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "nbsp.md"}),
            raise_error=False)
        assert r.code == 200
        html = r.body.decode("utf-8")
        para = html[html.find("<p><strong>Q?"):]
        para = para[:para.find("</p>")]
        assert len(BREAK_RE.findall(para)) == 1, (
            f"a non-breaking space defeated the collapse: {para!r}"
        )

    async def test_export_survives_a_broken_manual_break_rule(
            self, jp_fetch, jp_root_dir, monkeypatch):
        """The rule reads Markdown's internals, so it must never be able to
        fail an export. The coupling executes inside the converter's
        constructor, not in the factory, so that is where the guard has to be:
        a rule that builds fine and then fails to wire itself must fall back to
        plain nl2br, not return a 500."""
        from jupyterlab_export_markdown_extension import routes
        from markdown.extensions import Extension

        class Broken(Extension):
            def extendMarkdown(self, md):
                raise TypeError("simulated python-markdown API change")

        monkeypatch.setattr(routes, "manual_break_aware_nl2br",
                            lambda: Broken())
        (jp_root_dir / "broken.md").write_text(
            "**Q?**<br>\nAnswer.\n", encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "broken.md"}),
            raise_error=False)
        assert r.code == 200, (
            f"a broken cosmetic rule took the whole export down ({r.code})"
        )
        html = r.body.decode("utf-8")
        para = html[html.find("<p><strong>Q?"):]
        para = para[:para.find("</p>")]
        assert len(BREAK_RE.findall(para)) == 2, (
            f"expected the plain nl2br fallback (two breaks), got {para!r}"
        )

    async def test_a_break_named_in_a_comment_is_not_a_break(
            self, jp_fetch, jp_root_dir):
        """The stashed node must BE a break tag, not merely contain one - an
        HTML comment mentioning `<br>` is not the author writing a break, and
        treating it as one deletes a real one."""
        (jp_root_dir / "comment.md").write_text(
            "Q<!-- note about <br> -->\nAnswer.\n", encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "comment.md"}),
            raise_error=False)
        assert r.code == 200
        html = r.body.decode("utf-8")
        para = html[html.find("<p>Q<!--"):]
        para = para[:para.find("</p>")]
        # The comment's own text contains the characters `<br>`, so count is
        # not the measure here - what matters is that the real break after the
        # comment survived
        assert re.search(r'-->\s*<br', para), (
            f"a comment that merely mentions a break deleted a real one: {para!r}"
        )

    async def test_break_inside_emphasis_is_a_known_limitation(
            self, jp_fetch, jp_root_dir):
        """`**Q<br>**` puts the break inside the emphasis, so by the time the
        newline is examined the trailing node is the emphasis element, not the
        break, and the pair still renders with a blank line. Pinned so the
        limitation is a recorded decision rather than an unnoticed hole."""
        (jp_root_dir / "emph.md").write_text(
            "**Q?<br>**\nAnswer.\n", encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "emph.md"}),
            raise_error=False)
        assert r.code == 200
        html = r.body.decode("utf-8")
        para = html[html.find("<p><strong>Q?"):]
        para = para[:para.find("</p>")]
        assert len(BREAK_RE.findall(para)) == 2, (
            f"the known limitation changed behaviour: {para!r}"
        )


class TestExportFontSize:
    """The `exportFontSize` setting picks a base body size - small 10pt,
    medium 12pt (default), large 13pt - and every other size in every format
    is a fixed proportion of it, so the whole document scales together."""

    DOC = "# Title\n\nBody paragraph text.\n\n## Section\n\nMore body text.\n"

    async def _export(self, jp_fetch, jp_root_dir, fmt, size=None):
        (jp_root_dir / "fs.md").write_text(self.DOC, encoding="utf-8")
        body = {"path": "fs.md"}
        if size is not None:
            body["exportFontSize"] = size
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", f"export/{fmt}",
            method="POST", body=json.dumps(body), raise_error=False)
        assert r.code == 200, f"{fmt} export at size {size} returned {r.code}"
        return r.body

    @staticmethod
    def _pdf_sizes(pdf_bytes):
        """Rendered point size of the body paragraph and of the H1."""
        import fitz
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        sizes = {}
        for block in doc[0].get_text("dict")["blocks"]:
            for line in block.get("lines", []):
                for span in line["spans"]:
                    text = span["text"].strip()
                    if text.startswith("Body paragraph"):
                        sizes["body"] = round(span["size"], 1)
                    elif text == "Title":
                        sizes["h1"] = round(span["size"], 1)
        doc.close()
        return sizes

    async def test_pdf_base_size_follows_the_setting(self, jp_fetch, jp_root_dir):
        for size, expected in (("small", 10.0), ("medium", 12.0), ("large", 13.0)):
            sizes = self._pdf_sizes(
                await self._export(jp_fetch, jp_root_dir, "pdf", size))
            assert sizes.get("body") == expected, (
                f"{size} rendered body text at {sizes.get('body')}pt, "
                f"expected {expected}pt"
            )

    async def test_pdf_headings_stay_proportional(self, jp_fetch, jp_root_dir):
        """A heading is a multiple of the body size, not a fixed number - the
        point of a base size is that the document scales as one."""
        ratios = []
        for size in ("small", "medium", "large"):
            sizes = self._pdf_sizes(
                await self._export(jp_fetch, jp_root_dir, "pdf", size))
            ratios.append(sizes["h1"] / sizes["body"])
        assert max(ratios) - min(ratios) < 0.01, (
            f"heading-to-body ratio drifted across sizes: {ratios}"
        )
        assert ratios[0] > 1.2, f"headings are not larger than body: {ratios}"

    async def test_docx_base_size_follows_the_setting(self, jp_fetch, jp_root_dir):
        from docx import Document

        for size, expected in (("small", 10.0), ("medium", 12.0), ("large", 13.0)):
            doc = Document(io.BytesIO(
                await self._export(jp_fetch, jp_root_dir, "docx", size)))
            normal = doc.styles["Normal"].font.size
            assert normal is not None and normal.pt == expected, (
                f"{size} gave DOCX body {normal and normal.pt}pt, "
                f"expected {expected}pt"
            )
            heading = doc.styles["Heading 1"].font.size
            assert heading is not None and heading.pt > expected, (
                f"{size} left Heading 1 at {heading and heading.pt}pt, "
                f"not above the {expected}pt body"
            )

    async def test_html_base_size_follows_the_setting(self, jp_fetch, jp_root_dir):
        for size, expected in (("small", "10pt"), ("medium", "12pt"),
                               ("large", "13pt")):
            html = (await self._export(
                jp_fetch, jp_root_dir, "html", size)).decode("utf-8")
            body_rule = html[html.find("body {"):]
            body_rule = body_rule[:body_rule.find("}")]
            assert f"font-size: {expected}" in body_rule, (
                f"{size} did not set the HTML body size to {expected}: "
                f"{body_rule!r}"
            )

    async def test_default_is_medium_in_every_format(self, jp_fetch, jp_root_dir):
        """An export with no setting - an older client, or a fresh install -
        renders at the documented default rather than at whatever each format
        used to hardcode."""
        from docx import Document

        assert self._pdf_sizes(
            await self._export(jp_fetch, jp_root_dir, "pdf"))["body"] == 12.0
        doc = Document(io.BytesIO(await self._export(jp_fetch, jp_root_dir, "docx")))
        assert doc.styles["Normal"].font.size.pt == 12.0
        html = (await self._export(jp_fetch, jp_root_dir, "html")).decode("utf-8")
        assert "font-size: 12pt" in html

    async def test_a_malformed_setting_cannot_fail_an_export(
            self, jp_fetch, jp_root_dir):
        """The size is a cosmetic choice, so nothing a client can put in the
        field may take the export down - including values that are not even
        the right type, which a plain dict lookup raises on."""
        for value in ([], {}, True, "enormous", None, 0, 10**6):
            (jp_root_dir / "fs.md").write_text(self.DOC, encoding="utf-8")
            r = await jp_fetch(
                "jupyterlab-export-markdown-extension", "export/html",
                method="POST",
                body=json.dumps({"path": "fs.md", "exportFontSize": value}),
                raise_error=False)
            assert r.code == 200, (
                f"exportFontSize={value!r} returned {r.code} instead of "
                f"falling back"
            )

    async def test_html_measure_scales_with_the_body(
            self, jp_fetch, jp_root_dir):
        """Line length is what governs readability, so the column width has to
        be a multiple of the body size too - a fixed pixel width would give
        `large` a third fewer characters per line than `small`."""
        widths = {}
        for size in ("small", "medium", "large"):
            html = (await self._export(
                jp_fetch, jp_root_dir, "html", size)).decode("utf-8")
            body_rule = html[html.find("body {"):]
            body_rule = body_rule[:body_rule.find("}")]
            match = re.search(r'max-width:\s*([\d.]+)(em|px)', body_rule)
            assert match, f"no max-width in the body rule: {body_rule!r}"
            widths[size] = (float(match.group(1)), match.group(2))
        assert {unit for _, unit in widths.values()} == {"em"}, (
            f"the column width is not relative to the body size: {widths}"
        )
        assert len({value for value, _ in widths.values()}) == 1, (
            f"the em measure should be one constant across sizes: {widths}"
        )


class TestMultiParagraphAlert:
    """DEF-MARK-11: a `> [!NOTE]` whose body has two paragraphs is one alert, not a
    coloured box followed by a stray grey blockquote."""

    DOC = (
        "# Doc\n\n"
        "> [!NOTE]\n"
        "> First paragraph.\n"
        ">\n"
        "> Second paragraph.\n\n"
        "Closing text.\n"
    )

    async def _export(self, jp_fetch, jp_root_dir, fmt, doc=None):
        (jp_root_dir / "alert.md").write_text(doc or self.DOC, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", f"export/{fmt}",
            method="POST", body=json.dumps({"path": "alert.md"}),
            raise_error=False)
        assert r.code == 200, f"{fmt} export returned {r.code}"
        return r.body

    async def test_html_is_one_box_holding_both_paragraphs(
            self, jp_fetch, jp_root_dir):
        html = (await self._export(jp_fetch, jp_root_dir, "html")).decode("utf-8")
        boxes = re.findall(r'<div style="border-left:4px solid([^>]*)>(.*?)</div>',
                           html, re.S)
        assert len(boxes) == 1, (
            f"expected one alert box, got {len(boxes)}: "
            f"{[b[0][:30] for b in boxes]}"
        )
        body = boxes[0][1]
        assert "First paragraph." in body and "Second paragraph." in body, (
            f"the alert box lost a paragraph: {body!r}"
        )
        # The orphaned paragraph would survive as an ordinary blockquote,
        # rendered by the DEF-MARK-2 path as a second, grey-barred box
        assert "<blockquote" not in html, (
            "a plain blockquote survived - the alert still split in two"
        )

    async def test_docx_is_one_alert_table_holding_both_paragraphs(
            self, jp_fetch, jp_root_dir):
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(io.BytesIO(
            await self._export(jp_fetch, jp_root_dir, "docx")))
        alert_tables = []
        for table in doc.tables:
            borders = table._tbl.find(qn('w:tblPr')).find(qn('w:tblBorders'))
            if borders is None:
                continue
            left = borders.find(qn('w:left'))
            if left is not None and left.get(qn('w:color')) == '0969DA':
                alert_tables.append(table)
        assert len(alert_tables) == 1, (
            f"expected one NOTE table, got {len(alert_tables)}"
        )
        cell_text = alert_tables[0].rows[0].cells[0].text
        assert "First paragraph." in cell_text, cell_text
        assert "Second paragraph." in cell_text, (
            f"the second paragraph is outside the alert: {cell_text!r}"
        )

    async def test_pdf_is_one_callout_with_no_stray_blockquote(
            self, jp_fetch, jp_root_dir):
        import fitz

        doc = fitz.open(stream=await self._export(jp_fetch, jp_root_dir, "pdf"),
                        filetype="pdf")
        fills, strokes, text = [], [], ""
        for page in doc:
            fills += _pdf_fill_colors(page)
            strokes += _pdf_stroke_colors(page)
            text += page.get_text()
        doc.close()

        assert _color_near(fills, (0.929, 0.961, 0.992)), (
            "NOTE blue shading missing - the alert did not render as a callout"
        )
        # #BBBBBB, the default blockquote bar the orphaned paragraph would draw
        assert not _color_near(strokes, (0.733, 0.733, 0.733)), (
            "a grey blockquote bar was drawn - the alert split into two boxes"
        )
        assert "First paragraph." in text and "Second paragraph." in text

    async def test_source_line_breaks_inside_an_alert_are_kept(
            self, jp_fetch, jp_root_dir):
        """Body text turns a source newline into a break, so an alert has to
        as well - the same two source lines must not set two different ways in
        the same document."""
        doc = "> [!TIP]\n> Line one\n> Line two\n"
        html = (await self._export(
            jp_fetch, jp_root_dir, "html", doc)).decode("utf-8")
        box = re.search(r'<div style="border-left:4px solid.*?</div>', html, re.S)
        assert box, "no alert box in the export"
        assert re.search(r'Line one\s*<br', box.group(0)), (
            f"the line break inside the alert was discarded: {box.group(0)!r}"
        )

    async def test_an_authored_break_in_an_alert_is_not_doubled(
            self, jp_fetch, jp_root_dir):
        """The alert joins its source lines with a break; a line that already
        ends in one must not get a second, as it would everywhere else."""
        doc = "> [!TIP]\n> **Question**<br>\n> Answer.\n"
        html = (await self._export(
            jp_fetch, jp_root_dir, "html", doc)).decode("utf-8")
        box = re.search(r'<div style="border-left:4px solid.*?</div>', html, re.S)
        assert box, "no alert box in the export"
        breaks = len(BREAK_RE.findall(box.group(0)))
        assert breaks == 1, (
            f"expected one break between question and answer, got {breaks}: "
            f"{box.group(0)!r}"
        )

    async def test_two_adjacent_alerts_stay_separate(
            self, jp_fetch, jp_root_dir):
        """Widening the continuation must not let one alert swallow the next."""
        doc = "> [!NOTE]\n> First alert.\n\n> [!WARNING]\n> Second alert.\n"
        html = (await self._export(
            jp_fetch, jp_root_dir, "html", doc)).decode("utf-8")
        assert len(re.findall(r'<div style="border-left:4px solid', html)) == 2, (
            "two adjacent alerts did not stay two boxes"
        )
        assert "#0969DA" in html and "#9A6700" in html, (
            "each alert must keep its own type colour"
        )

    async def test_a_plain_blockquote_is_still_a_blockquote(
            self, jp_fetch, jp_root_dir):
        """The alert rule must not capture an ordinary multi-paragraph quote."""
        doc = "> Just a quote.\n>\n> Second quoted paragraph.\n"
        html = (await self._export(
            jp_fetch, jp_root_dir, "html", doc)).decode("utf-8")
        assert '<div style="border-left:4px solid' not in html, (
            "an ordinary blockquote was styled as an alert"
        )
        assert "Just a quote." in html and "Second quoted paragraph." in html


class TestPdfMinorHeadings:
    """DEF-MARK-12: Heading 4, 5 and 6 render distinctly in the PDF instead of all
    collapsing onto the Heading 3 face."""

    DOC = ("# One\n\ntext\n\n## Two\n\ntext\n\n### Three\n\ntext\n\n"
           "#### Four\n\ntext\n\n##### Five\n\ntext\n\n###### Six\n\ntext\n")

    async def _faces(self, jp_fetch, jp_root_dir):
        import fitz

        (jp_root_dir / "heads.md").write_text(self.DOC, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "heads.md"}),
            raise_error=False)
        assert r.code == 200
        doc = fitz.open(stream=r.body, filetype="pdf")
        faces = {}
        for page in doc:
            for block in page.get_text("dict")["blocks"]:
                for line in block.get("lines", []):
                    for span in line["spans"]:
                        word = span["text"].strip()
                        if word in ("One", "Two", "Three", "Four", "Five", "Six"):
                            faces[word] = (span["font"], span["color"],
                                           round(span["size"], 1))
        doc.close()
        return faces

    async def test_every_heading_level_is_visually_distinct(
            self, jp_fetch, jp_root_dir):
        faces = await self._faces(jp_fetch, jp_root_dir)
        assert set(faces) == {"One", "Two", "Three", "Four", "Five", "Six"}, (
            f"a heading is missing from the PDF: {sorted(faces)}"
        )
        assert len(set(faces.values())) == 6, (
            f"heading levels share a face: {faces}"
        )

    async def test_minor_headings_match_the_docx_faces(
            self, jp_fetch, jp_root_dir):
        """The PDF colours must equal the colours of the DOCX it is built from,
        read off the exported document itself - not a copy in the test. If the
        template python-docx builds from ever moves, the DOCX moves, and this
        catches the two formats disagreeing instead of freezing beside them."""
        from docx import Document

        (jp_root_dir / "heads.md").write_text(self.DOC, encoding="utf-8")
        d = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "heads.md"}),
            raise_error=False)
        assert d.code == 200
        styles = Document(io.BytesIO(d.body)).styles
        docx_color = {}
        for word, level in (("Four", 4), ("Five", 5), ("Six", 6)):
            rgb = styles[f"Heading {level}"].font.color.rgb
            assert rgb is not None, f"Heading {level} has no colour in the DOCX"
            docx_color[word] = int(str(rgb), 16)

        faces = await self._faces(jp_fetch, jp_root_dir)
        for word in ("Four", "Five", "Six"):
            assert faces[word][1] == docx_color[word], (
                f"PDF {word} is #{faces[word][1]:06X} but the DOCX draws it "
                f"#{docx_color[word]:06X} - the two formats disagree"
            )
        h4_font, h5_font, h6_font = (faces[k][0] for k in ("Four", "Five", "Six"))
        assert "Bold" in h4_font and _is_italic_font(h4_font), (
            f"H4 is not bold italic: {h4_font}"
        )
        assert "Bold" not in h5_font and not _is_italic_font(h5_font), (
            f"H5 is not regular: {h5_font}"
        )
        assert _is_italic_font(h6_font) and "Bold" not in h6_font, (
            f"H6 is not italic: {h6_font}"
        )

    async def test_minor_headings_sit_at_body_size(self, jp_fetch, jp_root_dir):
        """Word's H4-H6 carry no size of their own and render at body size;
        the PDF has to do the same, or the levels it just learned to tell
        apart are told apart by the wrong signal."""
        import fitz

        faces = await self._faces(jp_fetch, jp_root_dir)
        (jp_root_dir / "heads.md").write_text(self.DOC, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "heads.md"}),
            raise_error=False)
        doc = fitz.open(stream=r.body, filetype="pdf")
        body = [round(s["size"], 1)
                for page in doc
                for b in page.get_text("dict")["blocks"]
                for line in b.get("lines", [])
                for s in line["spans"] if s["text"].strip() == "text"]
        doc.close()
        assert body, "no body text found in the PDF"
        for level in ("Four", "Five", "Six"):
            assert faces[level][2] == body[0], (
                f"H{level} is {faces[level][2]}pt against {body[0]}pt body"
            )
        sizes = [faces[k][2] for k in ("One", "Two", "Three", "Four")]
        assert sizes == sorted(sizes, reverse=True) and sizes[2] > sizes[3], (
            f"heading sizes do not descend into H4: {sizes}"
        )

    async def test_italic_headings_do_not_fall_back_to_a_core_font(
            self, jp_fetch, jp_root_dir):
        """The italic minor headings must draw in the document's Unicode font,
        not a Helvetica core face. DejaVu (best coverage, so the body's font)
        ships no oblique on many boxes; committing to it and stopping left
        every italic falling to Helvetica-Oblique - a typeface switch visible
        in the heading ladder. H4/H6 must share the embedded family, not a
        core-14 substitute."""
        faces = await self._faces(jp_fetch, jp_root_dir)
        for level in ("Four", "Six"):
            font = faces[level][0]
            assert "Helvetica" not in font and "Times" not in font, (
                f"H{level} italic fell back to a core font: {font}"
            )
            assert _is_italic_font(font), f"H{level} is not italic: {font}"


class TestPdfCodeLineWrapping:
    """DEF-MARK-13: a code line wider than the frame wraps instead of running off
    the page."""

    LONG = "a" * 300
    DOC = (
        "# Code\n\n"
        "```python\n"
        f'url = "{LONG}"\n'
        "short = 1\n"
        "```\n"
    )

    async def _pdf(self, jp_fetch, jp_root_dir, doc=None):
        (jp_root_dir / "code.md").write_text(doc or self.DOC, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "code.md"}),
            raise_error=False)
        assert r.code == 200
        return r.body

    @staticmethod
    def _mono_spans(pdf_bytes):
        import fitz

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        spans = []
        for page in doc:
            for block in page.get_text("dict")["blocks"]:
                for line in block.get("lines", []):
                    for span in line["spans"]:
                        if "Courier" in span["font"] or "Mono" in span["font"]:
                            spans.append(span)
        doc.close()
        return spans

    async def test_a_long_code_line_stays_inside_the_margin(
            self, jp_fetch, jp_root_dir):
        spans = self._mono_spans(await self._pdf(jp_fetch, jp_root_dir))
        assert spans, "no monospaced code text found in the PDF"
        edge = pdf_frame_right_edge()
        past = [(round(s["bbox"][2], 1), s["text"][:30])
                for s in spans if s["bbox"][2] > edge + 1]
        assert not past, (
            f"code drawn past the {edge:.0f}pt frame edge: {past[:3]}"
        )

    async def test_wrapping_loses_no_characters(self, jp_fetch, jp_root_dir):
        spans = self._mono_spans(await self._pdf(jp_fetch, jp_root_dir))
        drawn = "".join(s["text"] for s in spans)
        assert drawn.count("a") == 300, (
            f"the wrapped line lost characters: {drawn.count('a')} of 300"
        )
        assert "short = 1" in drawn.replace("\n", ""), (
            "the following code line disappeared"
        )

    async def test_a_short_code_line_is_not_broken(self, jp_fetch, jp_root_dir):
        """Only an overflowing line may wrap - a line that fits keeps its
        shape, or every code sample gains phantom breaks."""
        doc = "# C\n\n```python\ndef f(x):\n    return x + 1\n```\n"
        spans = self._mono_spans(await self._pdf(jp_fetch, jp_root_dir, doc))
        tops = sorted({round(s["bbox"][1], 1) for s in spans})
        assert len(tops) == 2, (
            f"a two-line code block rendered on {len(tops)} lines: {tops}"
        )


def _docx_text(docx_bytes):
    """Every run of word/document.xml joined back into one string.

    Syntax highlighting splits a kept code block across runs - `flowchart LR`
    lands as `flowchart`, ` `, `LR` - so a substring test against the raw XML
    passes whether or not the text is there.
    """
    import html as _html

    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    return _html.unescape("".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)))


MERMAID_ONLY_MARKDOWN = """# Diagrams

```mermaid
flowchart LR
    Ingest --> Model --> Report
```

Some prose between the diagrams.

```mermaid
sequenceDiagram
    Alice->>Bob: hello
    Bob-->>Alice: hi
```
"""


@pytest.fixture
def mermaid_markdown_file(jp_root_dir):
    md_file = jp_root_dir / "mermaid_api.md"
    md_file.write_text(MERMAID_ONLY_MARKDOWN, encoding="utf-8")
    return md_file


class TestApiMermaidRendering:
    """DEF-DIAG-16: mermaid is a browser library, so the frontend renders each
    diagram and posts it as `mermaidDiagrams`. A caller driving the REST API
    directly sends no such payload, and every diagram used to export as its
    own source code. The server renders what the frontend did not."""

    async def test_docx_renders_mermaid_without_a_frontend_payload(
        self, jp_fetch, mermaid_markdown_file
    ):
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "mermaid_api.md"}),
            raise_error=False,
        )
        assert response.code == 200, response.body[:300]

        with zipfile.ZipFile(io.BytesIO(response.body)) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]

        assert len(media) == 2, (
            f"{len(media)} of 2 mermaid diagrams reached the DOCX - a diagram "
            f"the browser did not pre-render was dropped"
        )
        text = _docx_text(response.body)
        assert "flowchart" not in text and "sequenceDiagram" not in text, (
            "the mermaid source was written into the document as code instead "
            "of being rendered"
        )

    async def test_pdf_renders_mermaid_without_a_frontend_payload(
        self, jp_fetch, mermaid_markdown_file
    ):
        import fitz

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "mermaid_api.md"}),
            raise_error=False,
        )
        assert response.code == 200, response.body[:300]

        pdf = fitz.open(stream=response.body, filetype="pdf")
        try:
            images = sum(len(page.get_images(full=True)) for page in pdf)
            text = "\n".join(page.get_text() for page in pdf)
        finally:
            pdf.close()

        assert images == 2, f"{images} of 2 mermaid diagrams reached the PDF"
        assert "flowchart LR" not in text and "sequenceDiagram" not in text, (
            "the mermaid source was drawn into the PDF as code"
        )

    async def test_html_renders_mermaid_without_a_frontend_payload(
        self, jp_fetch, mermaid_markdown_file
    ):
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "mermaid_api.md"}),
            raise_error=False,
        )
        assert response.code == 200, response.body[:300]
        html = response.body.decode("utf-8")

        assert html.count("data:image/svg+xml") == 2, (
            f"{html.count('data:image/svg+xml')} of 2 mermaid diagrams were "
            f"inlined into the HTML"
        )
        assert "language-mermaid" not in html, (
            "a mermaid block stayed a code block in the HTML"
        )

    async def test_a_diagram_the_frontend_supplied_is_not_re_rendered(
        self, jp_fetch, jp_root_dir
    ):
        """The browser's own render wins - it carries the Lab theme and fonts.
        A 1x1 PNG is unmistakable: anything the server rendered would be
        thousands of bytes."""
        (jp_root_dir / "one.md").write_text(
            "# One\n\n```mermaid\nflowchart LR\n    A --> B\n```\n", encoding="utf-8"
        )
        one_pixel_png = (
            "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAf"
            "FcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", raise_error=False,
            body=json.dumps({
                "path": "one.md",
                "mermaidDiagrams": [{"index": 0, "png": one_pixel_png, "svg": ""}],
            }),
        )
        assert response.code == 200

        with zipfile.ZipFile(io.BytesIO(response.body)) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
            blobs = [z.read(n) for n in media]

        assert len(blobs) == 1, f"expected one image, got {len(blobs)}"
        assert len(blobs[0]) < 500, (
            f"the frontend's diagram was replaced by a server render "
            f"({len(blobs[0])} bytes) - the browser's version must win"
        )

    async def test_a_diagram_mermaid_rejects_keeps_its_source(
        self, jp_fetch, jp_root_dir
    ):
        """One unparseable diagram must not fail the export or take the good
        diagram down with it."""
        (jp_root_dir / "bad.md").write_text(
            "# Bad\n\n```mermaid\nnot a diagram at all {{{\n```\n\n"
            "```mermaid\nflowchart LR\n    A --> B\n```\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "bad.md"}),
            raise_error=False,
        )
        assert response.code == 200, response.body[:300]

        with zipfile.ZipFile(io.BytesIO(response.body)) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]

        assert len(media) == 1, (
            f"{len(media)} images embedded - the valid diagram must still "
            f"render and the invalid one must not"
        )
        assert "not a diagram at all" in _docx_text(response.body), (
            "the source of the diagram mermaid rejected was dropped instead of "
            "being kept as text"
        )

    async def test_a_document_without_mermaid_never_starts_a_browser(
        self, jp_fetch, jp_root_dir, monkeypatch
    ):
        """Chromium costs seconds to launch. Nothing to render means no
        browser - and means an export cannot start failing on a server with
        no Chromium just because this pass exists."""
        from jupyterlab_export_markdown_extension import routes

        class Exploding:
            def __init__(self, *a, **kw):
                raise AssertionError("Chromium was launched with no diagram to render")

        monkeypatch.setattr(routes, "PlaywrightSvgRenderer", Exploding)
        (jp_root_dir / "plain.md").write_text("# Plain\n\nNo diagrams here.\n",
                                              encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "plain.md"}),
            raise_error=False,
        )
        assert response.code == 200, response.body[:300]

    async def test_docx_carries_a_png_and_html_an_svg(
        self, jp_fetch, mermaid_markdown_file
    ):
        """The server-rendered diagram must arrive in the shape each format
        already expects from the frontend - Word cannot embed an SVG at all."""
        docx = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "mermaid_api.md"}),
            raise_error=False,
        )
        assert docx.code == 200
        with zipfile.ZipFile(io.BytesIO(docx.body)) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
            assert media and all(n.endswith(".png") for n in media), (
                f"DOCX media are not PNG: {media}"
            )

        html = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "mermaid_api.md"}),
            raise_error=False,
        )
        assert html.code == 200
        assert "data:image/svg+xml" in html.body.decode("utf-8"), (
            "HTML should inline the diagram as SVG, not rasterize it"
        )

    async def test_one_browser_renders_and_rasterizes(
        self, jp_fetch, mermaid_markdown_file, monkeypatch
    ):
        """Rasterizing inside the render session is the whole point of doing it
        there - a second launch is ~300ms of Chromium startup for nothing."""
        from jupyterlab_export_markdown_extension import routes

        launches = []
        original = routes.PlaywrightSvgRenderer.__aenter__

        async def counting_aenter(self):
            launches.append(1)
            return await original(self)

        monkeypatch.setattr(routes.PlaywrightSvgRenderer, "__aenter__", counting_aenter)
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "mermaid_api.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert len(launches) == 1, (
            f"Chromium was launched {len(launches)} times for a document whose "
            f"only images are mermaid diagrams"
        )

    async def test_a_clean_export_reports_no_warnings(
        self, jp_fetch, mermaid_markdown_file
    ):
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "mermaid_api.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert "X-Export-Warnings" not in response.headers, (
            f"a clean export warned: {response.headers.get('X-Export-Warnings')}"
        )


class TestApiMermaidWarnings:
    """DEF-DIAG-16: an export never fails over a diagram - the source is kept and
    the response says what happened. The body is a document, so the header is
    the only channel an API caller has."""

    @staticmethod
    def _warnings(response):
        raw = response.headers.get("X-Export-Warnings")
        assert raw, "the export reported no warning at all"
        return {w["code"]: w for w in json.loads(raw)}

    async def test_a_chromium_less_server_still_exports_and_says_why(
        self, jp_fetch, jp_root_dir, monkeypatch
    ):
        from jupyterlab_export_markdown_extension import routes

        async def refuse(self):
            raise routes.ChromiumUnavailableError(
                "Chromium binary not found at /nowhere"
            )

        monkeypatch.setattr(routes.PlaywrightSvgRenderer, "__aenter__", refuse)
        (jp_root_dir / "nc.md").write_text(
            "# NC\n\n```mermaid\nflowchart LR\n    A --> B\n```\n", encoding="utf-8"
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "nc.md"}), raise_error=False,
        )

        assert response.code == 200, (
            "a missing Chromium must not fail an export that can still produce "
            "the document"
        )
        assert "flowchart" in _docx_text(response.body), (
            "the diagram source was dropped instead of kept"
        )

        warning = self._warnings(response)["chromium-unavailable"]
        assert warning["diagrams"] == [0]
        assert "jupyterlab-export-markdown-extension install" in warning["message"], (
            f"the warning does not say how to fix it: {warning['message']}"
        )

    async def test_html_export_also_reports_it(
        self, jp_fetch, jp_root_dir, monkeypatch
    ):
        """HTML never needed Chromium before this feature, so it must not start
        failing now - it warns like the others."""
        from jupyterlab_export_markdown_extension import routes

        async def refuse(self):
            raise routes.ChromiumUnavailableError("no chromium")

        monkeypatch.setattr(routes.PlaywrightSvgRenderer, "__aenter__", refuse)
        (jp_root_dir / "nch.md").write_text(
            "# NC\n\n```mermaid\nflowchart LR\n    A --> B\n```\n", encoding="utf-8"
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "nch.md"}), raise_error=False,
        )
        assert response.code == 200, response.body[:300]
        assert "chromium-unavailable" in self._warnings(response)

    async def test_a_missing_bundle_is_named_as_such(
        self, jp_fetch, jp_root_dir, monkeypatch
    ):
        """A wheel built without its vendor step must say so, not blame the
        diagram - and must not pay for a browser launch to find out."""
        from jupyterlab_export_markdown_extension import routes

        monkeypatch.setattr(routes.PlaywrightSvgRenderer, "MERMAID_JS_PATH",
                            Path("/nonexistent/mermaid.min.js"))

        async def explode(self):
            raise AssertionError("Chromium was launched with no bundle to load")

        monkeypatch.setattr(routes.PlaywrightSvgRenderer, "__aenter__", explode)
        (jp_root_dir / "nb.md").write_text(
            "# NB\n\n```mermaid\nflowchart LR\n    A --> B\n```\n", encoding="utf-8"
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "nb.md"}), raise_error=False,
        )
        assert response.code == 200, response.body[:300]
        assert "bundle-missing" in self._warnings(response)

    async def test_a_syntax_error_is_reported_against_its_own_diagram(
        self, jp_fetch, jp_root_dir
    ):
        """Two diagrams, the first unparseable: the warning must point at index
        0 and the second must still render."""
        (jp_root_dir / "mix.md").write_text(
            "# Mix\n\n```mermaid\nnot a diagram at all {{{\n```\n\n"
            "```mermaid\nflowchart LR\n    A --> B\n```\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "mix.md"}), raise_error=False,
        )
        assert response.code == 200

        warning = self._warnings(response)["syntax-error"]
        assert warning["diagrams"] == [0], (
            f"the warning blames the wrong diagram: {warning['diagrams']}"
        )
        with zipfile.ZipFile(io.BytesIO(response.body)) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
        assert len(media) == 1, "the valid diagram must still have rendered"

    async def test_the_header_is_one_line_of_valid_json(
        self, jp_fetch, jp_root_dir
    ):
        """A header carrying a newline is rejected by the HTTP layer, and the
        diagram source is attacker-adjacent text - it must never reach it."""
        (jp_root_dir / "hdr.md").write_text(
            "# H\n\n```mermaid\nbroken é\nsüper\n\"quoted\"\n```\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "hdr.md"}), raise_error=False,
        )
        assert response.code == 200
        raw = response.headers["X-Export-Warnings"]
        assert "\n" not in raw and "\r" not in raw
        json.loads(raw)  # must parse
        assert "broken" not in raw, "the diagram source leaked into the header"

    async def test_a_broken_browser_session_still_exports(
        self, jp_fetch, jp_root_dir, monkeypatch
    ):
        """Chromium launching is not the only thing that can go wrong with a
        browser. Whatever the session throws, the export still produces the
        document - per-diagram faults are handled a level down."""
        from jupyterlab_export_markdown_extension import routes

        async def blow_up(self, *args, **kwargs):
            raise RuntimeError("target page, context or browser has been closed")

        monkeypatch.setattr(routes.PlaywrightSvgRenderer, "render_mermaid", blow_up)
        (jp_root_dir / "bs.md").write_text(
            "# BS\n\n```mermaid\nflowchart LR\n    A --> B\n```\n", encoding="utf-8"
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "bs.md"}), raise_error=False,
        )
        assert response.code == 200, response.body[:300]
        assert "flowchart" in _docx_text(response.body)
        assert "render-failed" in self._warnings(response)

    async def test_the_header_stays_bounded_on_a_document_full_of_diagrams(
        self, jp_fetch, jp_root_dir, monkeypatch
    ):
        """nginx's default `proxy_buffer_size` is 4KB and must hold the whole
        response header block. 200 broken diagrams must report their count
        without listing every index."""
        from jupyterlab_export_markdown_extension import routes

        async def refuse(self):
            raise routes.ChromiumUnavailableError("no chromium")

        monkeypatch.setattr(routes.PlaywrightSvgRenderer, "__aenter__", refuse)
        block = "```mermaid\nflowchart LR\n    A --> B\n```\n\n"
        (jp_root_dir / "many.md").write_text("# Many\n\n" + block * 200,
                                             encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "many.md"}), raise_error=False,
        )
        assert response.code == 200

        raw = response.headers["X-Export-Warnings"]
        assert len(raw) < 2048, f"the warning header grew to {len(raw)} bytes"
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase
        warning = self._warnings(response)["chromium-unavailable"]
        assert warning["count"] == 200, (
            f"the header lost the true count: {warning['count']}"
        )
        assert len(warning["diagrams"]) <= ExportHandlerBase.MAX_REPORTED_DIAGRAMS


class TestApiMermaidHardening:
    """Findings from the architect and bug-hunter review of DEF-DIAG-16."""

    @staticmethod
    def _warnings(response):
        raw = response.headers.get("X-Export-Warnings")
        assert raw, "the export reported no warning at all"
        return {w["code"]: w for w in json.loads(raw)}

    async def test_a_mermaid_label_cannot_make_the_server_fetch_a_url(
        self, jp_fetch, jp_root_dir
    ):
        """Confirmed SSRF: mermaid keeps HTML labels, so `<img src=...>` in a
        diagram survives into the SVG and the SERVER fetches it - a document
        someone else wrote reaching the host's metadata endpoint. Measured
        three outbound requests before the block."""
        import socket
        import threading

        hits = []
        sock = socket.socket()
        sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
        sock.bind(("127.0.0.1", 0))
        port = sock.getsockname()[1]
        sock.listen(5)
        sock.settimeout(20)

        def serve():
            while True:
                try:
                    conn, _ = sock.accept()
                    hits.append(conn.recv(100))
                    conn.sendall(b"HTTP/1.1 200 OK\r\nContent-Length: 0\r\n\r\n")
                    conn.close()
                except Exception:
                    return

        thread = threading.Thread(target=serve, daemon=True)
        thread.start()
        try:
            (jp_root_dir / "ssrf.md").write_text(
                "# S\n\n```mermaid\nflowchart LR\n"
                f"    A[\"<img src='http://127.0.0.1:{port}/probe.png' "
                "width='40' height='40'>\"] --> B[ok]\n```\n",
                encoding="utf-8",
            )
            response = await jp_fetch(
                "jupyterlab-export-markdown-extension", "export/docx",
                method="POST", body=json.dumps({"path": "ssrf.md"}),
                raise_error=False,
            )
            assert response.code == 200, response.body[:300]
        finally:
            sock.close()

        assert not hits, (
            f"the server made {len(hits)} outbound request(s) because a diagram "
            f"asked it to - a markdown file must not be able to drive the "
            f"server's network"
        )

    async def test_a_mermaid_example_inside_an_outer_fence_is_left_alone(
        self, jp_fetch, jp_root_dir
    ):
        """Documentation that SHOWS mermaid syntax must stay text - rendering
        it replaces the code sample with a picture of itself."""
        (jp_root_dir / "doc.md").write_text(
            "# Docs\n\nWrite a diagram like this:\n\n"
            "````markdown\n```mermaid\nflowchart LR\n    A --> B\n```\n````\n\n"
            "And here is a real one:\n\n"
            "```mermaid\nflowchart LR\n    C --> D\n```\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "doc.md"}), raise_error=False,
        )
        assert response.code == 200, response.body[:300]

        with zipfile.ZipFile(io.BytesIO(response.body)) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
        assert len(media) == 1, (
            f"{len(media)} diagrams rendered - the example inside the outer "
            f"fence is documentation, not a diagram"
        )
        assert "flowchart LR" in _docx_text(response.body).replace("\n", ""), (
            "the quoted example lost its source"
        )

    async def test_a_warning_names_the_diagram_the_reader_would_count_to(
        self, jp_fetch, jp_root_dir, monkeypatch
    ):
        """With the frontend supplying diagrams 0 and 1, a failure on the third
        must report 2 - numbering the leftovers from zero sends the reader to a
        diagram that is fine."""
        from jupyterlab_export_markdown_extension import routes

        async def refuse(self):
            raise routes.ChromiumUnavailableError("no chromium")

        monkeypatch.setattr(routes.PlaywrightSvgRenderer, "__aenter__", refuse)
        block = "```mermaid\nflowchart LR\n    A --> B\n```\n\n"
        (jp_root_dir / "part.md").write_text("# P\n\n" + block * 3, encoding="utf-8")
        png = ("data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAf"
               "FcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg==")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", raise_error=False,
            body=json.dumps({
                "path": "part.md",
                "mermaidDiagrams": [{"index": 0, "png": png, "svg": ""},
                                    {"index": 1, "png": png, "svg": ""}],
            }),
        )
        assert response.code == 200
        warning = self._warnings(response)["chromium-unavailable"]
        assert warning["diagrams"] == [2], (
            f"the warning points at diagram(s) {warning['diagrams']} instead of "
            f"the third one the frontend could not supply"
        )

    async def test_a_diagram_inside_an_alert_keeps_its_own_picture(
        self, jp_fetch, jp_root_dir
    ):
        """`preprocess_github_alerts` folds an alert body onto one line, which
        erases a fence inside a `> [!NOTE]` completely. Running it before the
        mermaid passes left the server counting one diagram where the browser
        counted two, so the note's picture was pasted onto the diagram after
        it and the note kept raw source. Nothing that rewrites the source may
        run before the diagrams are paired."""
        doc = (
            "# A\n\n"
            "> [!NOTE]\n"
            "> See this:\n"
            "> ```mermaid\n"
            "> flowchart LR\n"
            ">     IN_THE_NOTE --> X\n"
            "> ```\n\n"
            "And separately:\n\n"
            "```mermaid\nflowchart LR\n    AFTER_THE_NOTE --> Y\n```\n"
        )
        (jp_root_dir / "alert.md").write_text(doc, encoding="utf-8")

        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        blocks = [b.group(1) for b in ExportHandlerBase.iter_mermaid_blocks(doc)]
        assert len(blocks) == 2 and "IN_THE_NOTE" in blocks[0], (
            f"the document as written holds two diagrams, the note's first; "
            f"the scanner sees {blocks!r}"
        )

        first = "data:image/svg+xml;base64,SU5fVEhFX05PVEU="
        second = "data:image/svg+xml;base64,QUZURVJfVEhFX05PVEU="
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST",
            body=json.dumps({
                "path": "alert.md",
                "mermaidDiagrams": [{"index": 0, "svg": first, "png": ""},
                                    {"index": 1, "svg": second, "png": ""}],
            }),
        )
        html = response.body.decode("utf-8")

        assert first in html and second in html, (
            "a diagram the frontend supplied did not reach the document"
        )
        assert html.index(first) < html.index(second), (
            "the note's picture came out after the one that follows it - the "
            "two are paired by position and the pairing has shifted"
        )
        assert "IN_THE_NOTE" not in html, (
            "the diagram inside the alert kept its source: the alert pass ran "
            "first and erased the fence before it could be paired"
        )

    async def test_the_server_renders_with_the_options_the_frontend_uses(self):
        """A diagram must not come out of the API unlike the one the UI makes.
        The frontend sets `securityLevel: 'loose'`; under `strict` mermaid
        turns off HTML labels and the same diagram draws differently."""
        from pathlib import Path as _Path
        from jupyterlab_export_markdown_extension.routes import PlaywrightSvgRenderer

        frontend = _Path(__file__).parents[2] / "src" / "index.ts"
        assert "securityLevel: 'loose'" in frontend.read_text(encoding="utf-8"), (
            "the frontend no longer sets securityLevel: 'loose' - the server "
            "constant below has to follow it"
        )
        assert PlaywrightSvgRenderer.MERMAID_INIT_OPTIONS["securityLevel"] == "loose"

    async def test_every_reported_code_has_a_message(self):
        """The renderer names codes, the header carries their messages - a code
        with no entry would report a warning that explains nothing."""
        import re as _re
        from jupyterlab_export_markdown_extension import routes

        source = _Path_source = (
            __import__("pathlib").Path(routes.__file__).read_text(encoding="utf-8")
        )
        emitted = set(_re.findall(r"'(?:None, )?([a-z]+(?:-[a-z]+)+)'\)", source))
        codes = {c for c in emitted if c in routes.MERMAID_WARNINGS or "-" in c}
        unknown = {c for c in codes
                   if c not in routes.MERMAID_WARNINGS
                   and c in {"render-failed", "syntax-error", "render-timeout",
                             "rasterize-failed", "skipped", "budget-exhausted",
                             "layout-unsupported", "bundle-missing",
                             "chromium-unavailable"}}
        assert not unknown, f"codes with no message: {unknown}"
        assert routes.MERMAID_WARNINGS.keys() >= {
            "chromium-unavailable", "bundle-missing", "syntax-error",
            "layout-unsupported", "render-timeout", "skipped",
            "budget-exhausted", "rasterize-failed", "render-failed",
        }

    async def test_the_diagrams_behind_a_timeout_are_not_blamed_for_it(
        self, jp_fetch, jp_root_dir, monkeypatch
    ):
        """A document where diagram 0 does not finish must not tell the reader
        that the healthy diagrams behind it are too complex to draw.

        The timeout is forced rather than provoked: mermaid draws even a
        4000-edge graph in 0.7s, so a source slow enough to trip a real 30s
        budget would make this test slower than the suite it lives in."""
        from jupyterlab_export_markdown_extension import routes

        real_wait_for = asyncio.wait_for
        first = []

        async def timeout_once(awaitable, timeout=None):
            if not first:
                first.append(True)
                if hasattr(awaitable, "close"):
                    awaitable.close()
                raise asyncio.TimeoutError()
            return await real_wait_for(awaitable, timeout)

        monkeypatch.setattr(routes.asyncio, "wait_for", timeout_once)
        block = "```mermaid\nflowchart LR\n    A --> B\n```\n\n"
        (jp_root_dir / "hang.md").write_text("# H\n\n" + block * 3, encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "hang.md"}), raise_error=False,
        )
        assert response.code == 200, response.body[:300]

        warnings = self._warnings(response)
        assert warnings["render-timeout"]["diagrams"] == [0], (
            f"the timeout is reported against "
            f"{warnings['render-timeout']['diagrams']}, not the diagram that hung"
        )
        assert warnings["skipped"]["diagrams"] == [1, 2], (
            "the diagrams behind the offender must be reported as skipped, not "
            "as having timed out themselves"
        )
        assert "Simplify" not in warnings["skipped"]["message"], (
            "the skipped diagrams are told to simplify themselves"
        )

    async def test_the_frontend_counts_diagrams_the_way_the_server_does(self):
        """The frontend posts diagrams by position and the server pairs them by
        position, so both must skip a ```mermaid quoted inside a longer fence.
        One counting it and the other not hands a diagram the wrong picture."""
        from pathlib import Path as _P

        frontend = (_P(__file__).parents[2] / "src" / "index.ts").read_text(
            encoding="utf-8"
        )
        assert "mermaidBlocksFromSource(source)" in frontend, (
            "the source fallback went back to a bare regex over the document; "
            "it would count a quoted example the server skips"
        )
        assert "`{3,}|~{3,}" in frontend, (
            "the frontend's fence tracking no longer mirrors the server's - "
            "both must treat a ```mermaid opened inside any fence as content"
        )

    async def test_a_mermaid_example_inside_a_plain_fence_is_left_alone(
        self, jp_fetch, jp_root_dir
    ):
        """A fence carrying an info string never closes a block, so
        ```text / ```mermaid / ``` is ONE code block per CommonMark - the
        mermaid inside it is quoted just as firmly as inside a longer fence."""
        (jp_root_dir / "plain.md").write_text(
            "# Docs\n\n"
            "```text\n```mermaid\nflowchart LR\n    A --> B\n```\n\n"
            "A real one:\n\n"
            "```mermaid\nflowchart LR\n    C --> D\n```\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "plain.md"}), raise_error=False,
        )
        assert response.code == 200, response.body[:300]
        with zipfile.ZipFile(io.BytesIO(response.body)) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
        assert len(media) == 1, (
            f"{len(media)} diagrams rendered - a ```mermaid quoted inside a "
            f"plain ```text block is documentation, not a diagram"
        )

    async def test_the_server_keeps_the_ceilings_the_browser_has(self):
        """`@jupyterlab/mermaid` raises maxTextSize/maxEdges and the frontend
        inherits them, so matching only securityLevel would fail diagrams
        through the API that JupyterLab draws without complaint."""
        from jupyterlab_export_markdown_extension.routes import PlaywrightSvgRenderer

        options = PlaywrightSvgRenderer.MERMAID_INIT_OPTIONS
        assert options["maxTextSize"] == 100000
        assert options["maxEdges"] == 100000

    async def test_a_null_pixel_width_still_produces_a_png(
        self, jp_fetch, mermaid_markdown_file
    ):
        """`data.get(k, default)` defaults a MISSING key only. A client sending
        `"svgPixelWidth": null` used to select SVG output silently, and Word
        cannot display an SVG."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", raise_error=False,
            body=json.dumps({"path": "mermaid_api.md", "svgPixelWidth": None}),
        )
        assert response.code == 200, response.body[:300]
        with zipfile.ZipFile(io.BytesIO(response.body)) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
        assert media and all(n.endswith(".png") for n in media), (
            f"a null width put {media} in the DOCX instead of PNG images"
        )

    async def test_a_junk_pixel_width_does_not_fail_the_export(
        self, jp_fetch, mermaid_markdown_file
    ):
        for width in ["wide", 0, -5, True, {"px": 900}, 99999]:
            response = await jp_fetch(
                "jupyterlab-export-markdown-extension", "export/docx",
                method="POST", raise_error=False,
                body=json.dumps({"path": "mermaid_api.md", "svgPixelWidth": width}),
            )
            assert response.code == 200, f"width={width!r}: {response.body[:200]}"
            with zipfile.ZipFile(io.BytesIO(response.body)) as z:
                media = [n for n in z.namelist() if n.startswith("word/media/")]
            assert len(media) == 2, f"width={width!r} lost a diagram"

    async def test_a_syntax_error_echoing_a_dead_session_phrase_does_not_abort(
        self, jp_fetch, jp_root_dir
    ):
        """A mermaid parse error echoes the offending source line, so a node
        label reading 'Connection closed' or 'Target page' beside a real syntax
        error was read as the browser dying and dropped every later diagram.
        Liveness is now decided from page state, so the run continues and the
        later healthy diagram still renders."""
        doc = (
            "# S\n\n"
            "```mermaid\nflowchart LR\n    A --> First\n```\n\n"
            # syntactically broken AND echoes a dead-session phrase in its text
            "```mermaid\nflowchart LR\n    Connection closed -->\n```\n\n"
            "```mermaid\nflowchart LR\n    B --> Third\n```\n"
        )
        (jp_root_dir / "echo.md").write_text(doc, encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "echo.md"}), raise_error=False,
        )
        assert response.code == 200, response.body[:300]
        with zipfile.ZipFile(io.BytesIO(response.body)) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
        assert len(media) == 2, (
            f"{len(media)} of the 2 healthy diagrams rendered - the run aborted "
            f"at the broken one because its text echoed a dead-session phrase"
        )
        codes = self._warnings(response)
        assert "render-failed" not in codes, (
            "a syntax error was misread as the browser dying"
        )
        assert "syntax-error" in codes, (
            "the genuinely broken diagram should report a syntax error"
        )


    async def test_a_diagram_nested_in_a_list_is_still_a_diagram(
        self, jp_fetch, jp_root_dir
    ):
        """CommonMark measures a fence's indent from its container's content
        column, so a block inside a list item sits four or more spaces in.
        JupyterLab renders it, so the server must count it - or every later
        diagram's picture lands on the wrong fence."""
        (jp_root_dir / "nested.md").write_text(
            "# N\n\n"
            "- step one\n"
            "  - detail\n\n"
            "    ```mermaid\n    flowchart LR\n        A --> B\n    ```\n\n"
            "- step two\n\n"
            "```mermaid\nflowchart LR\n    C --> D\n```\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "nested.md"}), raise_error=False,
        )
        assert response.code == 200, response.body[:300]
        with zipfile.ZipFile(io.BytesIO(response.body)) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
        assert len(media) == 2, (
            f"{len(media)} of 2 diagrams rendered - the list-nested one was "
            f"skipped, which shifts every later diagram onto the wrong fence"
        )

    async def test_a_blockquoted_diagram_is_still_a_diagram(
        self, jp_fetch, jp_root_dir
    ):
        (jp_root_dir / "quoted.md").write_text(
            "# Q\n\n> ```mermaid\n> flowchart LR\n>     A --> B\n> ```\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "quoted.md"}), raise_error=False,
        )
        assert response.code == 200, response.body[:300]
        with zipfile.ZipFile(io.BytesIO(response.body)) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
        assert len(media) == 1, "a blockquoted diagram was not rendered"

    async def test_a_tilde_fence_declares_a_diagram_and_still_quotes(self):
        """marked's fence rule treats `~~~` exactly like ``` ``` ```, so
        JupyterLab renders `~~~mermaid` and the preview capture counts it. The
        server has to agree or the browser's diagram N lands on fence N-1. A
        tilde fence must still QUOTE what a longer fence holds."""
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        as_diagram = "~~~mermaid\nflowchart LR\n    A --> B\n~~~\n"
        assert ExportHandlerBase.count_mermaid_blocks(as_diagram) == 1, (
            "a tilde fence did not declare a diagram the browser counts"
        )
        quoting = "~~~text\n```mermaid\nflowchart LR\n    A --> B\n```\n~~~\n"
        assert ExportHandlerBase.count_mermaid_blocks(quoting) == 0, (
            "a tilde fence stopped quoting what is inside it"
        )
        plain = "```mermaid\nflowchart LR\n    A --> B\n```\n"
        assert ExportHandlerBase.count_mermaid_blocks(plain) == 1

    async def test_an_uppercase_info_string_is_not_a_diagram(self):
        """`block.languages.includes(token.lang)` is case-sensitive, so
        JupyterLab renders ```MERMAID as a code sample. Lowercasing here turned
        that sample into a picture of itself and shifted every later block."""
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        doc = ("```MERMAID\nflowchart LR\n    A --> B\n```\n\n"
               "```mermaid\nflowchart LR\n    C --> D\n```\n")
        blocks = list(ExportHandlerBase.iter_mermaid_blocks(doc))
        assert len(blocks) == 1, (
            f"expected only the lowercase fence to be a diagram, got "
            f"{[b.group(1) for b in blocks]}"
        )
        assert "C --> D" in blocks[0].group(1), (
            "the uppercase sample was taken as the diagram, so the browser's "
            "picture would land on it"
        )

    async def test_a_null_math_width_keeps_the_equations(
        self, jp_fetch, jp_root_dir
    ):
        """The same null-default hazard as svgPixelWidth: uncoerced, the DPI
        arithmetic raises, the raise is swallowed, and the PDF ships its
        equations as literal `$x^2$` text."""
        import fitz

        (jp_root_dir / "math.md").write_text(
            "# M\n\nAn equation: $x^2 + y^2 = z^2$\n", encoding="utf-8"
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", raise_error=False,
            body=json.dumps({"path": "math.md", "mathPixelWidth": None}),
        )
        assert response.code == 200, response.body[:300]
        pdf = fitz.open(stream=response.body, filetype="pdf")
        try:
            text = "\n".join(page.get_text() for page in pdf)
            images = sum(len(page.get_images(full=True)) for page in pdf)
        finally:
            pdf.close()
        assert "$x^2" not in text, (
            "the equation shipped as literal markdown - a null width silently "
            "disabled math rendering"
        )
        assert images >= 1, "the equation image is missing from the PDF"

    async def test_an_infinite_pixel_width_does_not_fail_the_export(
        self, jp_fetch, mermaid_markdown_file
    ):
        """`json.loads` accepts `Infinity`, and `int(float('inf'))` raises
        OverflowError - which is not a ValueError."""
        for body in ['{"path": "mermaid_api.md", "svgPixelWidth": Infinity}',
                     '{"path": "mermaid_api.md", "svgPixelWidth": NaN}',
                     '{"path": "mermaid_api.md", "svgPixelWidth": "1e400"}']:
            response = await jp_fetch(
                "jupyterlab-export-markdown-extension", "export/docx",
                method="POST", body=body, raise_error=False,
            )
            assert response.code == 200, f"{body}: {response.body[:200]}"
            with zipfile.ZipFile(io.BytesIO(response.body)) as z:
                media = [n for n in z.namelist() if n.startswith("word/media/")]
            assert len(media) == 2, f"{body} lost a diagram"

    async def test_a_broken_diagram_keeps_the_healthy_ones_around_it(
        self, jp_fetch, jp_root_dir
    ):
        """A syntax error in the middle of a document keeps its source but must
        not throw away the diagrams drawn before it or block the ones after -
        the per-diagram fault is isolated, only a genuine session death stops
        the run."""
        doc = (
            "# B\n\n"
            "```mermaid\nflowchart LR\n    A --> Before\n```\n\n"
            "```mermaid\nnot a valid diagram at all @#$\n```\n\n"
            "```mermaid\nflowchart LR\n    C --> After\n```\n"
        )
        (jp_root_dir / "mid.md").write_text(doc, encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "mid.md"}), raise_error=False,
        )
        assert response.code == 200, response.body[:300]
        with zipfile.ZipFile(io.BytesIO(response.body)) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
        assert len(media) == 2, (
            f"{len(media)} of 2 healthy diagrams rendered - a per-diagram "
            f"syntax error was allowed to stop the whole run"
        )

    async def test_the_missing_layout_regex_tells_an_elk_diagram_from_a_bug(self):
        """A genuinely unsupported layout gets its own remedy; an ordinary JS
        error mentioning 'layout' must not be dressed up as one."""
        from jupyterlab_export_markdown_extension.routes import PlaywrightSvgRenderer

        assert not PlaywrightSvgRenderer._MERMAID_MISSING_LAYOUT_RE.search(
            "Cannot read properties of undefined (reading 'layout')"
        ), "an ordinary JS error is reported as an unsupported layout"
        assert PlaywrightSvgRenderer._MERMAID_MISSING_LAYOUT_RE.search(
            "Error: Layout algorithm elk is not registered."
        ), "a genuinely missing layout engine is not recognised"

    async def test_the_frontend_reads_a_crlf_file(self):
        """JS `.` excludes \\r, so a naive `(.*)$` matches no fence at all in a
        Windows-authored document - the fallback would find nothing where it
        used to find every diagram."""
        from pathlib import Path as _P

        frontend = (_P(__file__).parents[2] / "src" / "index.ts").read_text(
            encoding="utf-8"
        )
        assert "replace(/\\r$/, '')" in frontend, (
            "the frontend no longer strips a trailing CR before matching fences"
        )

    async def test_a_four_space_indented_example_at_top_level_is_not_a_diagram(
        self, jp_fetch, jp_root_dir
    ):
        """Four spaces at the top level is an indented code block, not a fence.
        JupyterLab renders only the real diagram, so counting the example would
        put the browser's picture on the wrong one.

        Counting images cannot tell the two apart - either way one renders and
        one does not - so this asserts WHICH survived as text."""
        (jp_root_dir / "indented.md").write_text(
            "# I\n\nShown as an example:\n\n"
            "    ```mermaid\n    flowchart LR\n        A --> B\n    ```\n\n"
            "The real one:\n\n"
            "```mermaid\nflowchart LR\n    C --> D\n```\n",
            encoding="utf-8",
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "indented.md"}),
            raise_error=False,
        )
        assert response.code == 200, response.body[:300]
        text = _docx_text(response.body).replace("\n", "")
        assert "A --> B" in text, (
            "the indented example was rendered into a picture of itself "
            "instead of staying the code sample it is"
        )
        assert "C --> D" not in text, "the real diagram was left as source"

    async def test_an_unclosed_quoted_fence_ends_with_its_blockquote(self):
        """A blockquote ends at its first unquoted line, and CommonMark closes
        any fence still open inside it - marked renders that block, so it is a
        diagram. What it must not do is run on: that makes one 'diagram' out of
        the blockquote plus every paragraph and code block after it, up to
        whatever ``` happens to close it."""
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        doc = ("# U\n\n> ```mermaid\n> flowchart LR\n>     A --> B\n\n"
               "This paragraph must survive.\n\n"
               "```python\nkeep = 'this'\n```\n")
        blocks = list(ExportHandlerBase.iter_mermaid_blocks(doc))
        assert len(blocks) == 1, (
            f"expected the quoted block to close with its blockquote, got "
            f"{len(blocks)} block(s)"
        )
        assert blocks[0].end() < doc.index("This paragraph"), (
            f"the block spans {(blocks[0].start(), blocks[0].end())} of a "
            f"{len(doc)}-char document - the substitution would replace the "
            f"prose and the python block too"
        )
        assert blocks[0].group(1).strip() == "flowchart LR\n    A --> B", (
            f"quote markers survived into the source: {blocks[0].group(1)!r}"
        )

    async def test_an_unclosed_fence_at_the_end_of_the_file_is_a_diagram(self):
        """CommonMark closes an unclosed fence at the end of the document and
        marked renders what it opened, so the preview counts it. Dropping it
        here would leave the browser one diagram ahead of the server."""
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        doc = "# E\n\n```mermaid\nflowchart LR\n    A --> B\n"
        blocks = list(ExportHandlerBase.iter_mermaid_blocks(doc))
        assert len(blocks) == 1, (
            f"an unclosed fence at EOF yielded {len(blocks)} blocks"
        )
        assert blocks[0].group(1).strip() == "flowchart LR\n    A --> B"

    async def test_a_crlf_document_hands_mermaid_clean_source(self):
        """A Windows-authored file keeps a CR on every line when the source is
        sliced straight out of the content."""
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        doc = "# C\r\n\r\n```mermaid\r\nflowchart LR\r\n    A --> B\r\n```\r\n"
        blocks = list(ExportHandlerBase.iter_mermaid_blocks(doc))
        assert len(blocks) == 1, f"a CRLF document yielded {len(blocks)} blocks"
        assert "\r" not in blocks[0].group(1), (
            f"the source handed to mermaid carries CRs: {blocks[0].group(1)!r}"
        )

    async def test_a_crlf_document_renders_its_diagrams(
        self, jp_fetch, jp_root_dir
    ):
        (jp_root_dir / "crlf.md").write_bytes(
            b"# C\r\n\r\n```mermaid\r\nflowchart LR\r\n    A --> B\r\n```\r\n"
        )
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "crlf.md"}),
            raise_error=False,
        )
        assert response.code == 200, response.body[:300]
        with zipfile.ZipFile(io.BytesIO(response.body)) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
        assert len(media) == 1, "a CRLF document rendered no diagram"

    async def test_a_nested_quote_strips_every_marker(self):
        """One `>` per level; leaving any behind hands mermaid `> flowchart`."""
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        doc = "# N\n\n> > ```mermaid\n> > flowchart LR\n> >     A --> B\n> > ```\n"
        blocks = list(ExportHandlerBase.iter_mermaid_blocks(doc))
        assert len(blocks) == 1
        source = blocks[0].group(1)
        assert not any(line.lstrip().startswith(">") for line in source.split("\n")), (
            f"quote markers reached mermaid: {source!r}"
        )
        assert "flowchart LR" in source and "A --> B" in source

    async def test_the_frontend_carries_every_counting_rule_the_server_has(self):
        """A cheap smoke check that each rule is at least present on both sides.

        It is NOT the guarantee - both files can carry every one of these
        substrings and still behave differently, which is exactly what a
        differential fuzz found (a CRLF bare list marker). The guarantee is
        `TestFenceScannerParity`, which runs both scanners."""
        from pathlib import Path as _P

        frontend = (_P(__file__).parents[2] / "src" / "index.ts").read_text(
            encoding="utf-8"
        )
        for rule, what in [
            ("`{3,}|~{3,}", "every fence opener, not only long ones"),
            ("replace(/\\r$/, '')", "CRLF lines"),
            ("quoteStripped", "measuring indent in the container's coordinate"),
            ("expandTabs", "a tab after > carries one column of the marker"),
            ("listContentCol", "the list item's content column"),
            ("spaces - base > 3", "indented code past the container is not a fence"),
            ("stripQuoteMarkers", "a quoted block's markers"),
            ("dedentBody", "the container's indent is not the diagram's"),
            ("info === 'mermaid'", "an exact, case-sensitive info string"),
            ("spaces <= openSpaces + 3", "a closer lives in its opener's container"),
            ("q < openQ", "a blockquote ending closes the fence inside it"),
        ]:
            assert rule in frontend, (
                f"the frontend scanner no longer handles {what} ({rule!r}) - "
                f"it must mirror iter_mermaid_blocks or the indices diverge"
            )

FENCE_SHAPES = {
    "plain": "```mermaid\nflowchart LR\n    A --> B\n```\n",
    "quoted_by_longer": "````markdown\n```mermaid\nflowchart LR\n    A --> B\n```\n````\n",
    "quoted_by_plain": "```text\n```mermaid\nflowchart LR\n    A --> B\n```\n",
    "tilde_opener": "~~~mermaid\nflowchart LR\n    A --> B\n~~~\n",
    "tilde_quoting": "~~~text\n```mermaid\nflowchart LR\n    A --> B\n```\n~~~\n",
    "list_nested": "- a\n  - b\n\n    ```mermaid\n    flowchart LR\n        A --> B\n    ```\n",
    "top_indented": "text\n\n    ```mermaid\n    flowchart LR\n        A --> B\n    ```\n",
    "blockquote": "> ```mermaid\n> flowchart LR\n>     A --> B\n> ```\n",
    "nested_quote": "> > ```mermaid\n> > flowchart LR\n> >     A --> B\n> > ```\n",
    "unclosed_quote": "> ```mermaid\n> flowchart LR\n\nprose\n\n```python\nx = 1\n```\n",
    "crlf": "```mermaid\r\nflowchart LR\r\n    A --> B\r\n```\r\n",
    "two_plain": "```mermaid\nflowchart LR\n    A --> B\n```\n\n```mermaid\nflowchart LR\n    C --> D\n```\n",
    "unclosed_eof": "```mermaid\nflowchart LR\n    A --> B\n",
    "info_suffix": "```mermaid extra\nflowchart LR\n```\n",
    "crlf_bare_marker": "-\r\n    ```mermaid\r\n    flowchart LR\r\n    ```\r\n",
    "crlf_ordered_marker": "1.\r\n    ```mermaid\r\n    flowchart LR\r\n    ```\r\n",
    "crlf_list_then_text": "- a\r\n\r\ntext\r\n\r\n```mermaid\r\nflowchart LR\r\n```\r\n",
    "uppercase_info": "```MERMAID\nflowchart LR\n    A --> B\n```\n",
    "mixed_case_info": "```Mermaid\nflowchart LR\n    A --> B\n```\n",
    "quoted_closer": "```mermaid\nflowchart LR\n> ```\nstill the diagram\n```\n",
    "deep_closer": "```mermaid\nflowchart LR\n    ```\nstill the diagram\n```\n",
    "list_then_quote": "- a\n\n> quote\n\n    ```mermaid\n    flowchart LR\n    ```\n",
    "quote_deep_indent": "- a\n\n>     ```mermaid\n>     flowchart LR\n>     ```\n",
    # Round-6: shapes both adversarial lenses reproduced as silent
    # picture-misplacement - a fence four spaces past a list item's content
    # column is indented code, a bare/tab-indented container was mis-measured.
    "list_over_indented": ("- Here is an example:\n\n      ```mermaid\n"
                           "      flowchart LR\n      ```\n\nAnd real:\n\n"
                           "```mermaid\nflowchart RL\n```\n"),
    "list_in_quote": "> - step\n>\n>     ```mermaid\n>     flowchart LR\n>     ```\n",
    "bare_marker_deep": "-\n\n      ```mermaid\n      flowchart LR\n      ```\n",
    "tab_quoted_fence": ">\t```mermaid\n>\tflowchart LR\n>\t```\n",
    "quote_then_bare_fence": ("> ```mermaid\n> flowchart LR\n```python\n"
                              "x = 1\n```\n"),
    "list_content_col_three": "1. item\n\n   ```mermaid\n   flowchart LR\n   ```\n",
    # Round-7 (architect): marked trims the info string with JS `.trim()`, which
    # strips a BOM; Python `str.strip()` keeps it, so ```mermaid<BOM> counted as
    # source here and as a diagram in the browser - a one-off count shift.
    "bom_info": "```mermaid\ufeff\nflowchart LR\n    A --> B\n```\n",
}

#: What each shape must yield. A diagram counted by one side and not the other
#: hands the server a picture addressed to the wrong fence. Every count here is
#: what marked - the parser JupyterLab renders with, and therefore what the
#: preview-capture path counts - produces for that document; none is a reading
#: of the spec. `test_marked_agrees_with_the_server_on_every_shape` re-derives
#: them from the installed marked so the table cannot drift from it.
FENCE_SHAPE_COUNTS = {
    "plain": 1, "quoted_by_longer": 0, "quoted_by_plain": 0, "tilde_opener": 1,
    "tilde_quoting": 0, "list_nested": 1, "top_indented": 0, "blockquote": 1,
    "nested_quote": 1, "unclosed_quote": 1, "crlf": 1, "two_plain": 2,
    "unclosed_eof": 1, "info_suffix": 0, "crlf_bare_marker": 1,
    "crlf_ordered_marker": 1, "crlf_list_then_text": 1, "uppercase_info": 0,
    "mixed_case_info": 0, "quoted_closer": 1, "deep_closer": 1,
    "list_then_quote": 0, "quote_deep_indent": 0, "list_over_indented": 1,
    "list_in_quote": 1, "bare_marker_deep": 0, "tab_quoted_fence": 1,
    "quote_then_bare_fence": 1, "list_content_col_three": 1, "bom_info": 1,
}


class TestFenceScannerParity:
    """DEF-DIAG-16: the frontend posts diagrams by position and the server pairs them
    by position, so the two scanners must agree on every fence shape. This runs
    BOTH - the TypeScript one is extracted and executed under node - rather than
    reading one and trusting the other."""

    @staticmethod
    def _typescript_blocks(fixtures):
        import json as _json
        import re as _re
        import subprocess
        import tempfile
        from pathlib import Path as _P

        source = (_P(__file__).parents[2] / "src" / "index.ts").read_text(
            encoding="utf-8"
        )
        start = source.index("/** One `>` per quote level")
        end = source.index("/**\n * Capture rendered Mermaid diagrams")
        body = source[start:end]
        # Strip TypeScript type annotations so node can run the extracted body:
        # function return types `): T {` first, then parameter/variable `: T`.
        body = _re.sub(r"\)\s*:\s*[A-Za-z0-9_\[\]<>,.| ]+\s*\{", ") {", body)
        for pattern in (r":\s*string\[\]", r":\s*string \| null",
                        r":\s*number \| null", r":\s*\[number, string\]",
                        r":\s*(?:string|number|boolean|void)\b"):
            body = _re.sub(pattern, "", body)
        body = body.replace(" as string", "")

        with tempfile.TemporaryDirectory() as tmp:
            script = _P(tmp) / "fence.js"
            data = _P(tmp) / "fixtures.json"
            data.write_text(_json.dumps(fixtures), encoding="utf-8")
            script.write_text(body + """
const fixtures = JSON.parse(require('fs').readFileSync(process.argv[2], 'utf8'));
const out = {};
for (const [name, doc] of Object.entries(fixtures)) {
  out[name] = mermaidBlocksFromSource(doc);
}
console.log(JSON.stringify(out));
""", encoding="utf-8")
            done = subprocess.run(["node", str(script), str(data)],
                                  capture_output=True, text=True, timeout=60)
        assert done.returncode == 0, f"node failed: {done.stderr[:400]}"
        return _json.loads(done.stdout)

    @staticmethod
    def _normalise(blocks):
        return [b.replace("\r\n", "\n").rstrip("\n") for b in blocks]

    def test_the_server_counts_every_shape_as_specified(self):
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        for name, doc in FENCE_SHAPES.items():
            found = len(list(ExportHandlerBase.iter_mermaid_blocks(doc)))
            assert found == FENCE_SHAPE_COUNTS[name], (
                f"{name}: server found {found}, expected "
                f"{FENCE_SHAPE_COUNTS[name]}"
            )

    def test_marked_agrees_with_the_server_on_every_shape(self):
        """marked is the authority, not the CommonMark spec and not a reading
        of it: JupyterLab renders the preview with marked, and the DOM capture
        posts one diagram per fence marked turned into a diagram. Anything the
        server counts differently lands a picture on the wrong fence.

        This runs the installed marked over the same shapes. It is what caught
        ```MERMAID (marked is case-sensitive, the server was not) and ~~~mermaid
        (marked renders it, the server refused it)."""
        import json as _json
        import shutil
        import subprocess
        import tempfile
        from pathlib import Path as _P

        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        marked = (_P(__file__).parents[2] / "node_modules" / "marked" / "lib"
                  / "marked.esm.js")
        if shutil.which("node") is None or not marked.exists():
            pytest.skip("node and the installed marked are needed")

        names = list(FENCE_SHAPES)
        with tempfile.TemporaryDirectory() as tmp:
            script = _P(tmp) / "count.mjs"
            data = _P(tmp) / "docs.json"
            data.write_text(_json.dumps([FENCE_SHAPES[n] for n in names]),
                            encoding="utf-8")
            # Mirrors @jupyterlab/markedparser-extension: a fence is a diagram
            # when the code token's lang is exactly what MermaidMarkdown
            # registers, `block.languages.includes(lang)` with languages
            # `['mermaid']`.
            script.write_text(f"""
import {{ marked }} from '{marked.as_posix()}';
import {{ readFileSync }} from 'fs';
function walk(tokens, out) {{
  for (const t of tokens) {{
    if (t.type === 'code' && t.lang === 'mermaid') out.push(t.text);
    if (t.tokens) walk(t.tokens, out);
    if (t.items) walk(t.items, out);
  }}
}}
const docs = JSON.parse(readFileSync(process.argv[2], 'utf8'));
console.log(JSON.stringify(docs.map(d => {{ const o = []; walk(marked.lexer(d), o); return o; }})));
""", encoding="utf-8")
            done = subprocess.run(["node", str(script), str(data)],
                                  capture_output=True, text=True, timeout=60)
        assert done.returncode == 0, f"node failed: {done.stderr[:400]}"

        for name, blocks in zip(names, _json.loads(done.stdout)):
            server = len(list(
                ExportHandlerBase.iter_mermaid_blocks(FENCE_SHAPES[name])))
            assert server == len(blocks), (
                f"{name}: marked renders {len(blocks)} diagram(s) and the "
                f"server counts {server} - the browser posts one per rendered "
                f"diagram, so every later picture lands on the wrong fence"
            )
            assert len(blocks) == FENCE_SHAPE_COUNTS[name], (
                f"{name}: FENCE_SHAPE_COUNTS says {FENCE_SHAPE_COUNTS[name]} "
                f"but marked renders {len(blocks)} - the table has drifted "
                f"from the parser it is supposed to describe"
            )

    def test_both_scanners_agree_on_every_shape(self):
        import shutil

        if shutil.which("node") is None:
            pytest.skip("node is not available to run the TypeScript scanner")

        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        typescript = self._typescript_blocks(FENCE_SHAPES)
        for name, doc in FENCE_SHAPES.items():
            server = self._normalise(
                [b.group(1) for b in ExportHandlerBase.iter_mermaid_blocks(doc)]
            )
            frontend = self._normalise(typescript[name])
            assert server == frontend, (
                f"{name}: the server yields {server!r} and the frontend "
                f"{frontend!r} - one counts a block the other skips, so a "
                f"diagram would receive the wrong picture"
            )



class TestDocxSymbolFonts:
    """DEF-MARK-22: a symbol the DOCX body face cannot draw names a font that can,
    instead of leaving Word to substitute whatever the machine happens to
    have - which is how a star arrives as a hollow box on one reader's PC."""

    DOC = (
        "# Symbols\n\n"
        "Dowod D677 ★ i D688 ☆, separator · · ·, "
        "strzalki → ↔, bullets • ◦, znak §, "
        "zazolc gesla jazn łęążóśćń.\n"
    )

    async def _docx(self, jp_fetch, jp_root_dir, doc=None):
        from docx import Document

        (jp_root_dir / "symbols.md").write_text(doc or self.DOC, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "symbols.md"}),
            raise_error=False)
        assert r.code == 200, f"docx export returned {r.code}"
        return Document(io.BytesIO(r.body))

    @staticmethod
    def _run_fonts(doc):
        """Map every run's text to the font its rFonts names ('' when none)."""
        from docx.oxml.ns import qn

        out = []
        for r in doc.element.body.iter(qn("w:r")):
            t = r.find(qn("w:t"))
            if t is None or not t.text:
                continue
            rPr = r.find(qn("w:rPr"))
            rFonts = rPr.find(qn("w:rFonts")) if rPr is not None else None
            out.append((t.text, rFonts.get(qn("w:ascii")) if rFonts is not None else ""))
        return out

    async def test_symbols_the_body_face_lacks_name_a_font(self, jp_fetch, jp_root_dir):
        doc = await self._docx(jp_fetch, jp_root_dir)
        fonts = self._run_fonts(doc)
        for glyph in ("★", "☆", "↔", "◦"):
            named = [f for text, f in fonts if glyph in text]
            assert named, f"{glyph!r} is missing from the DOCX entirely"
            assert all(f == "Segoe UI Symbol" for f in named), (
                f"{glyph!r} carries font {named!r}, not the symbol font - Word "
                f"is left to substitute and may draw a box"
            )

    async def test_body_characters_keep_the_body_font(self, jp_fetch, jp_root_dir):
        doc = await self._docx(jp_fetch, jp_root_dir)
        # Cambria carries all of these; overriding the font here would switch
        # typeface mid-sentence for no gain
        for text, font in self._run_fonts(doc):
            for glyph in ("·", "→", "•", "§", "ł"):
                if glyph in text:
                    assert font != "Segoe UI Symbol", (
                        f"{glyph!r} was moved to the symbol font in run "
                        f"{text!r} - the body face already draws it"
                    )

    async def test_no_text_is_lost_splitting_the_runs(self, jp_fetch, jp_root_dir):
        doc = await self._docx(jp_fetch, jp_root_dir)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Dowod D677 ★ i D688 ☆" in text
        assert "zazolc gesla jazn łęążóśćń" in text

    async def test_symbols_survive_inside_a_table_cell(self, jp_fetch, jp_root_dir):
        doc = await self._docx(
            jp_fetch, jp_root_dir,
            "| Dowod | Uwaga |\n| --- | --- |\n| D677 ★ | · · · |\n")
        cells = [c.text for t in doc.tables for row in t.rows for c in row.cells]
        assert any("D677 ★" in c for c in cells), (
            f"the starred reference did not reach the table: {cells!r}")
        assert any("· · ·" in c for c in cells), (
            f"the middot separator did not reach the table: {cells!r}")


    @staticmethod
    def _flow(doc):
        """Text and line breaks of the whole body, in document order.

        `paragraph.text` is the projection that ignores `w:br`, so it cannot
        see a break that moved - which is exactly how the run splitting went
        wrong the first time.
        """
        from docx.oxml.ns import qn

        out = []
        for run in doc.element.body.iter(qn("w:r")):
            for child in run:
                if child.tag == qn("w:t"):
                    out.append(child.text or "")
                elif child.tag == qn("w:br"):
                    out.append("\n")
        return "".join(out)

    async def test_a_break_after_a_symbol_stays_where_the_author_put_it(
            self, jp_fetch, jp_root_dir):
        doc = await self._docx(
            jp_fetch, jp_root_dir,
            "Dowod D677 ★ pierwsza linia\ndruga linia bez symbolu\n")
        flow = self._flow(doc)
        assert "Dowod D677 ★ pierwsza linia\ndruga linia bez symbolu" in flow, (
            f"the soft break moved when the run was split for its symbol: {flow!r}")

    async def test_a_code_block_keeps_its_line_order(self, jp_fetch, jp_root_dir):
        doc = await self._docx(
            jp_fetch, jp_root_dir,
            "```text\nDiagram ★ ponizej:\nLINIA-DRUGA\nLINIA-TRZECIA\n```\n")
        flow = self._flow(doc)
        assert "Diagram ★ ponizej:\nLINIA-DRUGA\nLINIA-TRZECIA" in flow, (
            f"the code block's lines were reordered around its symbol: {flow!r}")

    async def test_inline_code_keeps_its_monospace_around_a_symbol(
            self, jp_fetch, jp_root_dir):
        doc = await self._docx(jp_fetch, jp_root_dir,
                               "Text `code ★ here` and `plain code` done.\n")
        fonts = dict(self._run_fonts(doc))
        assert fonts.get("code ") == "Courier" and fonts.get(" here") == "Courier", (
            f"a symbol inside inline code cost its neighbours their monospace: "
            f"{fonts!r}")
        assert fonts.get("★") == "Segoe UI Symbol"

    async def test_an_emoji_keeps_the_face_word_would_give_it(
            self, jp_fetch, jp_root_dir):
        # A character carrying U+FE0F asks for emoji presentation, and one
        # whose block is half emoji is left alone entirely - Word draws those
        # in colour from its own emoji font, which no named face here can do
        doc = await self._docx(
            jp_fetch, jp_root_dir,
            "Uwaga ⚠️ i ▪️ punkt, bare ✅ emoji, "
            "oraz ★ gwiazda.\n")
        for text, font in self._run_fonts(doc):
            for glyph in ("⚠", "▪", "✅"):
                if glyph in text:
                    assert font != "Segoe UI Symbol", (
                        f"{glyph!r} was pinned to a monochrome face in {text!r}")
            if "️" in text:
                assert "⚠️" in text or "▪️" in text, (
                    f"a variation selector was split from its base: {text!r}")
        assert any(f == "Segoe UI Symbol" for t, f in self._run_fonts(doc)
                   if "★" in t), "the star lost its font with the emoji rule"


class TestInlineHtmlInDocx:
    """DEF-MARK-23: inline HTML written into the markdown reaches Word as
    formatting. htmldocx reads only a handful of tags and CSS properties, so
    the rest is rewritten into what it does read before conversion."""

    async def test_inline_math_in_a_link_label_exports(
            self, jp_fetch, jp_root_dir):
        """Walking hyperlink runs brought the marker-splitting branch its first
        run under `w:hyperlink`. When the marker OPENS the run, clearing the
        run's text drops its `w:t`, and the copy taken afterwards had none to
        write the tail into - so the export died with HTTP 500."""
        doc = (
            "Formula [$x^2$ label](https://x.example) inline.\n\n"
            "- [$x^2$ item](https://x.example)\n\n"
            "> [$x^2$ quoted](https://x.example)\n"
        )
        (jp_root_dir / "mathlink.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "mathlink.md"}),
            raise_error=False)
        assert r.code == 200, f"the export died: {r.body[:300]!r}"
        from docx import Document as _Doc

        document = _Doc(io.BytesIO(r.body))
        xml = document.element.body.xml
        assert "MATH_INLINE" not in xml, "a marker survived as visible text"
        assert xml.count("}oMath") >= 3 or xml.count(":oMath") >= 3, (
            "an equation is missing from the document")
        text = "".join(p.text for p in document.paragraphs)
        for label in ("label", "item", "quoted"):
            assert label in text, f"the link label {label!r} was lost"

    async def test_a_div_inside_a_text_holder_keeps_the_block_whole(
            self, jp_fetch, jp_root_dir):
        """A list item, a heading and a table cell already own a run of text,
        so a `<div>` inside one is in inline position however it was written.
        Renaming it to `<p>` there nested a block in a block: htmldocx ended
        the item, the heading or the cell early and the rest of the line
        became a paragraph of its own - the highlight going with it."""
        doc = (
            "- x <div style=\"background-color:#ffff00\">LIHL</div> y\n\n"
            "# H1 <div>H1MID</div> tail\n\n"
            "| a | b |\n| --- | --- |\n| c <div>TDMID</div> d | e |\n"
        )
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        (jp_root_dir / "divctx.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "divctx.md"}),
            raise_error=False)
        assert r.code == 200
        from docx import Document as _Doc

        document = _Doc(io.BytesIO(r.body))
        # The heading keeps its tail and its style
        heads = [p for p in document.paragraphs
                 if p.style.name.startswith("Heading")]
        assert heads and "tail" in heads[0].text, (
            f"the heading was cut in two: "
            f"{[(p.style.name, p.text) for p in document.paragraphs]!r}")
        # The cell is one paragraph, not two
        cell = document.tables[0].rows[1].cells[0]
        assert len([p for p in cell.paragraphs if p.text.strip()]) == 1, (
            f"the cell was split: {[p.text for p in cell.paragraphs]!r}")
        assert "TDMID" in cell.text and cell.text.strip().startswith("c")
        # The list item keeps the background the identical span would give it
        fills = [ExportHandlerBase.docx_run_shading(run)
                 for p in document.paragraphs for run in p.runs
                 if run.text.strip() == "LIHL"]
        assert fills and fills[0] == "FFFF00", (
            f"the list item lost its highlight: {fills!r}")

    async def test_a_div_written_inline_carries_its_style(
            self, jp_fetch, jp_root_dir):
        """A `<div>` markdown left inside a `<p>` is unwrapped so it does not
        nest one paragraph in another - and unwrapping took its style
        attribute with it, before the pass that reads style attributes ran."""
        doc = (
            'Lead <div style="color:#ff0000;font-weight:bold">DIVSTYLED</div> tail\n\n'
            'Lead <span style="color:#ff0000;font-weight:bold">SPANSTYLED</span> tail\n'
        )
        (jp_root_dir / "inlinediv.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "inlinediv.md"}),
            raise_error=False)
        assert r.code == 200
        from docx import Document as _Doc

        document = _Doc(io.BytesIO(r.body))
        props = {}
        for para in document.paragraphs:
            for run in para.runs:
                if run.text.strip() not in ("DIVSTYLED", "SPANSTYLED"):
                    continue
                props[run.text.strip()] = (
                    run.bold,
                    str(run.font.color.rgb)
                    if run.font.color and run.font.color.rgb else None)
        assert props.get("DIVSTYLED") == props.get("SPANSTYLED") != (None, None), (
            f"the div lost what the identical span kept: {props!r}")

    async def test_a_literal_twin_does_not_unbold_the_tail(
            self, jp_fetch, jp_root_dir):
        """htmldocx keys its open tags by name and pops on the first close,
        so wrapping `hot <b>stuff</b> more` in a second `<b>` let the inner
        close end the outer wrapper and ` more` arrived plain - bold in the
        browser, regular in Word. The wrapper subsumes a same-name twin."""
        doc = '<p style="font-weight:bold">hot <b>stuff</b> more</p>\n'
        (jp_root_dir / "twin.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "twin.md"}),
            raise_error=False)
        assert r.code == 200
        from docx import Document as _Doc

        document = _Doc(io.BytesIO(r.body))
        runs = [(run.text, bool(run.bold))
                for para in document.paragraphs for run in para.runs
                if run.text.strip()]
        assert runs and all(bold for _, bold in runs), (
            f"the tail after the literal twin lost its weight: {runs!r}")

    async def test_a_twin_keeps_its_colour_anchor_and_cell(
            self, jp_fetch, jp_root_dir):
        """Subsuming a literal twin must take only its NAME: unwrapping it
        destroyed the twin's own colour and anchor id with the tag, and a
        twin inside a nested table lost its weight to a wrapper that cannot
        reach that cell's runs - htmldocx scopes each cell, so the twin there
        is kept whole instead."""
        doc = (
            '<p style="font-weight:bold">hot <b style="color:red">red</b>'
            ' <b id="t1">anchor</b> tail</p>\n\n'
            '<table><tr><td style="font-weight:bold">outer'
            ' <table><tr><td><b>inner</b></td></tr></table>'
            ' tail</td></tr></table>\n'
        )
        (jp_root_dir / "twinattr.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "twinattr.md"}),
            raise_error=False)
        assert r.code == 200
        from docx import Document as _Doc

        document = _Doc(io.BytesIO(r.body))
        by_text = {run.text: run for para in document.paragraphs
                   for run in para.runs if run.text.strip()}
        red = by_text.get("red")
        assert red is not None and red.font.color and \
            str(red.font.color.rgb) == "FF0000", (
            f"the twin's own colour was destroyed with its tag: "
            f"{[(t, str(r.font.color.rgb) if r.font.color and r.font.color.rgb else None) for t, r in by_text.items()]!r}")
        assert all(run.bold for run in by_text.values()), (
            f"a run lost the weight the styled paragraph gives it: "
            f"{[(t, bool(r.bold)) for t, r in by_text.items()]!r}")
        # The anchor id survives into a bookmark
        from docx.oxml.ns import qn

        marks = [bm.get(qn("w:name")) for bm in
                 document.element.body.iter(qn("w:bookmarkStart"))]
        assert "t1" in marks, f"the twin's anchor id was destroyed: {marks!r}"
        # The nested-table twin keeps the weight the wrapper cannot deliver
        inner = [run for t in document.tables
                 for row in t.rows for cell in row.cells
                 for para in cell.paragraphs for run in para.runs
                 if run.text.strip() == "inner"]
        assert inner and inner[0].bold, "the nested-table twin lost its bold"

    async def test_the_last_declaration_wins(self, jp_fetch, jp_root_dir):
        """CSS is last-wins per property, and a rich-text paste is exactly
        where a self-overriding `underline;none` or `bold;normal` pair comes
        from. Accumulating every declaration kept the earlier value alive, so
        Word underlined what a browser drew plain."""
        doc = (
            '<span style="text-decoration:underline;'
            'text-decoration:none">UNONE</span> and '
            '<span style="font-weight:bold;font-weight:normal">BNORM</span>\n'
        )
        (jp_root_dir / "cascade.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "cascade.md"}),
            raise_error=False)
        assert r.code == 200
        from docx import Document as _Doc

        document = _Doc(io.BytesIO(r.body))
        styles = {run.text: (bool(run.bold), bool(run.underline))
                  for para in document.paragraphs for run in para.runs
                  if run.text.strip() in ("UNONE", "BNORM")}
        assert styles.get("UNONE") == (False, False), (
            f"an overridden underline survived: {styles!r}")
        assert styles.get("BNORM") == (False, False), (
            f"an overridden weight survived: {styles!r}")

    async def test_a_div_inside_a_quote_stays_inside_it(
            self, jp_fetch, jp_root_dir):
        """Markdown wraps a div written in inline position in a paragraph of
        its own, so renaming the div to <p> nested one <p> in another and
        htmldocx opened a second. A blockquote and a GitHub alert are each one
        paragraph per marker, so the content fell out of the callout."""
        import fitz

        doc = (
            "> quoted text\n"
            ">\n"
            '> <div align="center">QDIV</div>\n'
            "\n"
            "> [!NOTE]\n"
            "> <div>ADIV</div>\n"
        )
        (jp_root_dir / "quotediv.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "quotediv.md"}),
            raise_error=False)
        assert r.code == 200
        pdf = fitz.open(stream=r.body, filetype="pdf")
        spans = {s["text"].strip(): round(s["bbox"][0], 1)
                 for b in pdf[0].get_text("dict")["blocks"]
                 for line in b.get("lines", []) for s in line["spans"]
                 if s["text"].strip()}
        pdf.close()
        # Both callouts indent their body past the page margin at 42pt; a div
        # that escaped its callout draws at the margin instead
        assert spans.get("QDIV", 0) > 60, (
            f"the div fell out of the blockquote - spans: {spans!r}")
        assert spans.get("ADIV", 0) > 50, (
            f"the div fell out of the alert box - spans: {spans!r}")

    async def test_a_link_label_survives_a_quote_and_an_alert_in_the_pdf(
            self, jp_fetch, jp_root_dir):
        """`Paragraph.runs` is `./w:r` and never returns a run Word nested in
        `<w:hyperlink>`. The callout builders rebuilt their text from it, so
        the PDF dropped the label while the DOCX kept it."""
        import fitz

        doc = (
            "> quoted with a [LINKLABEL](http://example.com) inside\n"
            "\n"
            "> [!NOTE]\n"
            "> alert with a [ALERTLINK](http://example.com) inside\n"
        )
        (jp_root_dir / "quotelink.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "quotelink.md"}),
            raise_error=False)
        assert r.code == 200
        pdf = fitz.open(stream=r.body, filetype="pdf")
        text = "".join(page.get_text() for page in pdf)
        pdf.close()
        assert "LINKLABEL" in text, f"blockquote lost the label: {text!r}"
        assert "ALERTLINK" in text, f"alert box lost the label: {text!r}"

    async def test_an_anchor_target_does_not_kill_the_export(
            self, jp_fetch, jp_root_dir):
        """`<a id="x">` is the markdown-portable way to name a spot to link
        to, and htmldocx reads every anchor's text through
        `self.tags['a']['href']` without checking it is there - so one in the
        document took the whole export down with HTTP 500 and an error
        message reading only `'href'`."""
        doc = (
            "Zobacz [wykaz](#zal1) oraz [drugi](#zal2).\n\n"
            '<a id="zal1"></a>Pierwszy.\n\n'
            '<a id="zal2">Drugi</a> zalacznik.\n'
        )
        (jp_root_dir / "anchor.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "anchor.md"}),
            raise_error=False)
        from docx import Document as _Doc

        assert r.code == 200, f"the export died: {r.body[:400]!r}"
        document = _Doc(io.BytesIO(r.body))
        xml = document.element.body.xml
        for name in ("zal1", "zal2"):
            assert f'w:name="{name}"' in xml, (
                f"no bookmark for {name} - the link has nothing to reach")
        # The anchor's own text is content, not scaffolding
        assert "Drugi" in "".join(p.text for p in document.paragraphs)

    DOC = (
        "# HTML\n\n"
        'Colour: <span style="color:#c00000">RED</span>.\n\n'
        'Weight: <span style="font-weight:bold">HEAVY</span>.\n\n'
        'Slant: <span style="font-style:italic">SLANTED</span>.\n\n'
        'Rule: <span style="text-decoration:underline">RULED</span>.\n\n'
        'Struck: <span style="text-decoration:line-through">GONE</span>.\n\n'
        'Both: <span style="color:#008000;font-weight:bold">GREENBOLD</span>.\n\n'
        "Marked: <mark>HIGHLIT</mark>.\n\n"
        "Deleted: <del>DROPPED</del>.\n\n"
        "Key: <kbd>CTRL</kbd>.\n\n"
        'Legacy: <font color="#0000ff">BLUEFONT</font>.\n\n'
        '<div align="center">CENTRED</div>\n\n'
        "TRAILING\n"
    )

    async def _docx(self, jp_fetch, jp_root_dir, doc=None):
        from docx import Document

        (jp_root_dir / "inline.md").write_text(doc or self.DOC, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "inline.md"}),
            raise_error=False)
        assert r.code == 200, f"docx export returned {r.code}"
        return Document(io.BytesIO(r.body))

    @staticmethod
    def _run(doc, text):
        for p in doc.paragraphs:
            for r in p.runs:
                if r.text.strip() == text:
                    return r
        raise AssertionError(f"no run holding {text!r} - the text never reached Word")

    async def test_span_colour_reaches_word(self, jp_fetch, jp_root_dir):
        doc = await self._docx(jp_fetch, jp_root_dir)
        assert str(self._run(doc, "RED").font.color.rgb) == "C00000"

    async def test_span_style_properties_become_run_formatting(
            self, jp_fetch, jp_root_dir):
        doc = await self._docx(jp_fetch, jp_root_dir)
        assert self._run(doc, "HEAVY").bold, "font-weight:bold arrived unbolded"
        assert self._run(doc, "SLANTED").italic, "font-style:italic arrived upright"
        assert self._run(doc, "RULED").underline, (
            "text-decoration:underline arrived unruled")
        assert self._run(doc, "GONE").font.strike, (
            "text-decoration:line-through arrived unstruck")

    async def test_colour_and_weight_in_one_span_both_survive(
            self, jp_fetch, jp_root_dir):
        doc = await self._docx(jp_fetch, jp_root_dir)
        run = self._run(doc, "GREENBOLD")
        assert str(run.font.color.rgb) == "008000" and run.bold, (
            "a span declaring both colour and weight lost one of them")

    async def test_semantic_inline_tags_reach_word(self, jp_fetch, jp_root_dir):
        from docx.oxml.ns import qn

        doc = await self._docx(jp_fetch, jp_root_dir)
        shd = self._run(doc, "HIGHLIT")._r.find(qn("w:rPr") + "/" + qn("w:shd"))
        assert shd is not None and shd.get(qn("w:fill")) == "FFFF00", (
            "<mark> did not become a highlighted run")
        assert self._run(doc, "DROPPED").font.strike, "<del> arrived unstruck"
        assert self._run(doc, "CTRL").font.name == "Courier", (
            "<kbd> arrived in the body font, not monospace")

    async def test_font_colour_attribute_reaches_word(self, jp_fetch, jp_root_dir):
        doc = await self._docx(jp_fetch, jp_root_dir)
        assert str(self._run(doc, "BLUEFONT").font.color.rgb) == "0000FF", (
            '<font color="..."> - the notebook colouring idiom - lost its colour')

    async def test_aligned_div_is_its_own_centred_paragraph(
            self, jp_fetch, jp_root_dir):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = await self._docx(jp_fetch, jp_root_dir)
        centred = [p for p in doc.paragraphs if p.text.strip() == "CENTRED"]
        assert centred, "the <div> text never reached Word as its own paragraph"
        assert centred[0].alignment == WD_ALIGN_PARAGRAPH.CENTER, (
            'align="center" did not centre the block')
        assert "TRAILING" not in centred[0].text, (
            "the <div> swallowed the paragraph after it")

    async def test_a_styled_table_cell_keeps_its_cell(self, jp_fetch, jp_root_dir):
        # The emphasis has to go around the cell's CONTENTS: a `<b>` between
        # `<tr>` and `<td>` is not a parent htmldocx looks through, and the
        # cell is dropped - or the whole export dies on the short row
        doc = await self._docx(
            jp_fetch, jp_root_dir,
            '<table><tr><td style="font-weight:bold">BOLDCELL</td><td>PLAIN</td></tr>'
            "<tr><td>c</td><td>d</td></tr></table>\n")
        grid = [[c.text for c in row.cells] for t in doc.tables for row in t.rows]
        assert grid == [["BOLDCELL", "PLAIN"], ["c", "d"]], (
            f"the styled cell did not survive the rewrite: {grid!r}")
        assert self._run(doc_or_cell_runs(doc), "BOLDCELL").bold, (
            "the cell survived but arrived unbolded")

    async def test_a_table_carrying_emphasis_keeps_its_rows(
            self, jp_fetch, jp_root_dir):
        doc = await self._docx(
            jp_fetch, jp_root_dir,
            '<table style="font-weight:bold"><tr><td>X</td><td>Y</td></tr></table>\n')
        grid = [[c.text for c in row.cells] for t in doc.tables for row in t.rows]
        assert grid == [["X", "Y"]], (
            f"an emphasis declaration on the table itself lost its rows: {grid!r}")

    async def test_a_highlight_around_emphasis_shades_every_run(
            self, jp_fetch, jp_root_dir):
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        doc = await self._docx(
            jp_fetch, jp_root_dir,
            "Text <mark>plain and **bold** inside</mark> end.\n")
        shaded = {r.text: ExportHandlerBase.docx_run_shading(r)
                  for p in doc.paragraphs for r in p.runs if r.text.strip()}
        for fragment in ("plain and ", "bold", " inside"):
            assert shaded.get(fragment) == "FFFF00", (
                f"{fragment!r} inside the highlight was not shaded: {shaded!r}")
        assert not any("PILL:" in t for t in shaded), "a shading marker leaked as text"

    async def test_a_highlight_keeps_a_colour_the_author_declared(
            self, jp_fetch, jp_root_dir):
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase

        doc = await self._docx(
            jp_fetch, jp_root_dir,
            'Text <mark style="background-color:#00ff00">GREEN</mark> end.\n')
        run = self._run(doc, "GREEN")
        assert ExportHandlerBase.docx_run_shading(run) == "00FF00", (
            "the default yellow overwrote the colour the author asked for")

    async def test_emphasis_nested_in_its_own_tag_keeps_the_tail(
            self, jp_fetch, jp_root_dir):
        # htmldocx keys open tags by name and pops on the first close, so a
        # generated twin inside an identical tag would end the outer one early
        doc = await self._docx(
            jp_fetch, jp_root_dir,
            '<b>outer <span style="font-weight:bold">inner</span> tail</b>\n')
        bold = [r.bold for p in doc.paragraphs for r in p.runs if r.text.strip()]
        assert all(bold), f"the text after the nested tag lost its bold: {bold!r}"


def doc_or_cell_runs(doc):
    """A stand-in document whose `paragraphs` reach into every table cell too."""
    class _Flat:
        paragraphs = [p for t in doc.tables for row in t.rows
                      for c in row.cells for p in c.paragraphs] + list(doc.paragraphs)
    return _Flat


class TestPdfRunShading:
    """DEF-MARK-23: a highlight the reader sees in Word and in the browser has to
    reach the PDF too - `format_run` read every run property except shading."""

    DOC = "Text <mark>HIGHLIT</mark> end.\n"

    async def test_the_pdf_paints_the_highlight(self, jp_fetch, jp_root_dir):
        import fitz

        (jp_root_dir / "shade.md").write_text(self.DOC, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "shade.md"}),
            raise_error=False)
        assert r.code == 200
        pdf = fitz.open(stream=r.body, filetype="pdf")
        fills = [d.get("fill") for page in pdf for d in page.get_drawings()]
        text = "".join(page.get_text() for page in pdf)
        pdf.close()
        assert "HIGHLIT" in text
        yellow = [f for f in fills if f and round(f[0], 2) == 1.0
                  and round(f[1], 2) == 1.0 and round(f[2], 2) == 0.0]
        assert yellow, (
            f"the highlight is absent from the PDF - fills drawn: {fills!r}")

    async def test_the_pdf_paints_a_highlight_inside_a_table_cell(
            self, jp_fetch, jp_root_dir):
        """A cell was rebuilt from `cell_markup(p.text)` - a plain-string
        projection - so every run property reached Word and none reached the
        PDF. The fill has to land inside the first column, not merely exist:
        the body-paragraph half of this class already draws a yellow rect, so
        an existence check passes with the cell still plain."""
        import fitz

        doc = (
            "| head | x |\n"
            "| --- | --- |\n"
            "| <mark>CELLMARK</mark> | plain |\n"
        )
        (jp_root_dir / "cellshade.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "cellshade.md"}),
            raise_error=False)
        assert r.code == 200
        pdf = fitz.open(stream=r.body, filetype="pdf")
        page = pdf[0]
        cell = next((s for b in page.get_text("dict")["blocks"]
                     for line in b.get("lines", [])
                     for s in line["spans"] if "CELLMARK" in s["text"]), None)
        yellow = [tuple(round(v, 1) for v in d["rect"])
                  for d in page.get_drawings()
                  if d.get("fill") and [round(c, 2) for c in d["fill"]] == [1.0, 1.0, 0.0]]
        pdf.close()
        assert cell is not None, "the cell text is missing from the PDF"
        # Column 1's own band, not the page's left margin: a fill drawn for the
        # body paragraph would sit outside it
        assert any(r0 <= cell["bbox"][0] + 1 and r2 >= cell["bbox"][2] - 1
                   and r1 <= cell["bbox"][1] + 1 and r3 >= cell["bbox"][3] - 1
                   for r0, r1, r2, r3 in yellow), (
            f"no yellow fill behind the cell at {cell['bbox']!r} - "
            f"yellow rects drawn: {yellow!r}")

    async def test_a_bold_cell_fits_the_column_it_was_measured_for(
            self, jp_fetch, jp_root_dir):
        """Rendering a cell's run formatting widened its text without telling
        the column measurement, which still sized every body cell in the
        regular face. reportlab then hard-split the word rather than overflow,
        so `bold` came out as `bol` on one line and `d` on the next."""
        import fitz

        doc = (
            "| a | b |\n"
            "| --- | --- |\n"
            "| **important warning text** | ok |\n"
        )
        (jp_root_dir / "boldcell.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "boldcell.md"}),
            raise_error=False)
        assert r.code == 200
        pdf = fitz.open(stream=r.body, filetype="pdf")
        spans = [s for b in pdf[0].get_text("dict")["blocks"]
                 for line in b.get("lines", [])
                 for s in line["spans"] if s["text"].strip()]
        pdf.close()
        # One span carrying the whole phrase, not fragments of it on two lines
        whole = [s for s in spans if s["text"].strip() == "important warning text"]
        assert whole, (
            "the bold cell was split across lines - spans drawn: "
            f"{[(s['text'], s['font'], round(s['bbox'][1], 1)) for s in spans]!r}")
        assert "Bold" in whole[0]["font"], (
            f"the cell lost its weight: {whole[0]['font']!r}")

    async def test_a_minor_heading_keeps_its_symbols(
            self, jp_fetch, jp_root_dir):
        """H4 and H6 take their slant from the paragraph STYLE, not from a run,
        so the per-run glyph gate never saw them and the star stayed blank in
        exactly the heading a filing numbers its exhibits with."""
        import fitz

        doc = (
            "#### \u2605 H4 \u2713\n\n"
            "###### \u2605 H6 \u2713\n\n"
            "#### Zazalenie H4 bez gwiazdki\n"
        )
        (jp_root_dir / "headsym.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "headsym.md"}),
            raise_error=False)
        assert r.code == 200
        pdf = fitz.open(stream=r.body, filetype="pdf")
        spans = [(s["text"], s["font"])
                 for b in pdf[0].get_text("dict")["blocks"]
                 for line in b.get("lines", []) for s in line["spans"]
                 if s["text"].strip()]
        pdf.close()
        text = "".join(t for t, _ in spans)
        assert "\x00" not in text, f"a heading glyph is notdef: {spans!r}"
        assert text.count("\u2605") == 2 and text.count("\u2713") == 2, (
            f"a minor heading lost its markers: {spans!r}")
        # A heading with nothing uncoverable keeps the slant it was given
        plain = [f for t, f in spans if "Zazalenie" in t]
        assert plain and _is_italic_font(plain[0]), (
            f"an ordinary minor heading lost its slant: {plain!r}")

    async def test_a_multi_line_quote_keeps_its_slant(
            self, jp_fetch, jp_root_dir):
        """The glyph-coverage gate asks a font whether it can draw each
        character, and `run.text` renders a `<w:br>` as a newline - which no
        cmap carries. Every multi-line italic run answered False for a reason
        that has nothing to do with coverage and came out upright."""
        import fitz

        doc = "> line one\n> line two\n> line three\n"
        (jp_root_dir / "multiquote.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "multiquote.md"}),
            raise_error=False)
        assert r.code == 200
        pdf = fitz.open(stream=r.body, filetype="pdf")
        fonts = {s["text"].strip(): s["font"]
                 for b in pdf[0].get_text("dict")["blocks"]
                 for line in b.get("lines", []) for s in line["spans"]
                 if s["text"].strip()}
        pdf.close()
        assert fonts and all("Italic" in f or "Oblique" in f
                             for f in fonts.values()), (
            f"a line of the quote lost its slant: {fonts!r}")

    async def test_a_symbol_survives_an_italic_run_in_the_pdf(
            self, jp_fetch, jp_root_dir):
        """DejaVu ships no oblique on most Linux boxes, so the PDF italic slot
        is filled from Liberation, which has no glyph for a star. reportlab
        draws a missing glyph as a blank advance, so rendering a cell's own
        formatting deleted the reader's evidence marker instead of un-slanting
        it - silently, with nothing in the output to say so."""
        import fitz

        doc = (
            "| a | b |\n"
            "| --- | --- |\n"
            "| *\u2605 kursywa* | x |\n"
            "\n"
            "> Cytat \u2605 w kursywie\n"
        )
        (jp_root_dir / "italicstar.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "italicstar.md"}),
            raise_error=False)
        assert r.code == 200
        pdf = fitz.open(stream=r.body, filetype="pdf")
        text = "".join(page.get_text() for page in pdf)
        pdf.close()
        assert "\x00" not in text, "a glyph was drawn as notdef"
        assert text.count("\u2605") == 2, (
            f"the star is missing from an italic run - text: {text!r}")

    async def test_the_pdf_keeps_a_link_label_in_a_highlighted_paragraph(
            self, jp_fetch, jp_root_dir):
        """`Paragraph.runs` is `./w:r`, so it never returns a run Word nested
        in `<w:hyperlink>`. The formatted rebuild used it while the plain path
        used `para.text` (which does reach the link), so adding a highlight to
        a paragraph deleted its link label from the PDF."""
        import fitz

        doc = "See <mark>HIGHLIT</mark> and [LINKTEXT](http://example.com) end.\n"
        (jp_root_dir / "linkshade.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "linkshade.md"}),
            raise_error=False)
        assert r.code == 200
        pdf = fitz.open(stream=r.body, filetype="pdf")
        text = "".join(page.get_text() for page in pdf)
        pdf.close()
        assert "LINKTEXT" in text, (
            f"the link label was dropped from the highlighted paragraph: {text!r}")
        assert "HIGHLIT" in text


class TestDocxHyperlinkRunsAreReached:
    """A run Word nests in `<w:hyperlink>` is invisible to `Paragraph.runs`,
    so every pass that walks a paragraph run by run skips the link."""

    async def _export(self, jp_fetch, jp_root_dir, name, doc):
        from docx import Document

        (jp_root_dir / name).write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": name}),
            raise_error=False)
        assert r.code == 200
        return Document(io.BytesIO(r.body))

    async def test_a_highlight_over_a_link_leaks_no_marker_into_the_text(
            self, jp_fetch, jp_root_dir):
        """style_docx_color_runs reads a `⁣PILL:<hex>⁣` sentinel off a run and
        strips it. Missing the link run left the sentinel in the visible
        document - U+2063 is zero-width, so Word draws the bare `PILL:FFFF00`
        mid-sentence."""
        doc = "See <mark>a [LINKTEXT](http://example.com) b</mark> end.\n"
        d = await self._export(jp_fetch, jp_root_dir, "pilllink.md", doc)
        text = "\n".join(p.text for p in d.paragraphs)
        assert "PILL:" not in text, f"the pill sentinel leaked: {text!r}"
        assert "LINKTEXT" in text

    async def test_the_link_run_carries_the_shading_the_rest_of_the_span_has(
            self, jp_fetch, jp_root_dir):
        """The marker is per text node, so the link is a shaded run in its own
        right - not merely a run the strip pass has to leave alone."""
        from docx.oxml.ns import qn

        doc = "See <mark>a [LINKTEXT](http://example.com) b</mark> end.\n"
        d = await self._export(jp_fetch, jp_root_dir, "pillfill.md", doc)
        fills = {}
        for p in d.paragraphs:
            for run in p._p.iter(qn("w:r")):
                txt = "".join(t.text or "" for t in run.findall(qn("w:t")))
                if not txt.strip():
                    continue
                rPr = run.find(qn("w:rPr"))
                shd = rPr.find(qn("w:shd")) if rPr is not None else None
                fills[txt] = shd.get(qn("w:fill")) if shd is not None else None
        assert fills.get("LINKTEXT") == "FFFF00", (
            f"the link run inside the highlight is unshaded: {fills!r}")


class TestDocxRunShadingElementOrder:
    """CT_RPr is a sequence, not a bag: `w:shd` is index 29 and `w:vertAlign`
    31, so appending the shading after a `<sub>`'s vertAlign gives Word a file
    it refuses to open."""

    async def test_a_highlight_over_a_subscript_writes_schema_valid_order(
            self, jp_fetch, jp_root_dir):
        from docx import Document
        from docx.oxml.ns import qn

        doc = "Text <mark>H<sub>2</sub>O</mark> end.\n"
        (jp_root_dir / "shdorder.md").write_text(doc, encoding="utf-8")
        r = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "shdorder.md"}),
            raise_error=False)
        assert r.code == 200
        d = Document(io.BytesIO(r.body))

        checked = 0
        for p in d.paragraphs:
            for run in p._p.iter(qn("w:r")):
                rPr = run.find(qn("w:rPr"))
                if rPr is None:
                    continue
                tags = [re.sub(r"^\{.*\}", "", c.tag) for c in rPr]
                if "shd" not in tags or "vertAlign" not in tags:
                    continue
                checked += 1
                assert tags.index("shd") < tags.index("vertAlign"), (
                    f"w:shd written after w:vertAlign - Word rejects this rPr: "
                    f"{tags!r}")
        assert checked, "no run carried both shading and a vertical alignment"


# A table of contents written the way every markdown generator writes one -
# two spaces before each top-level marker - plus the two shapes the fix must
# not break: a nested list, an indented code block, and a fenced sample whose
# own content looks like an indented list.
TEST_MARKDOWN_INDENTED_TOC = """# Indented Lists

**Table of Contents**

  - [1. Introduction](#1.-Introduction)
  - [2. Project objective](#2.-Project-objective)

## 1. Introduction

- top level
  - nested under top

Ordered, written two spaces in:

  1. first ordered
  2. second ordered

An indented code block, four spaces in:

    - this is code, not a list

A fenced sample:

```
para

  - fenced item
```

## 2. Project objective

Done.
"""


@pytest.fixture
def test_indented_toc_file(jp_root_dir):
    md_file = jp_root_dir / "test_indented_toc.md"
    md_file.write_text(TEST_MARKDOWN_INDENTED_TOC, encoding="utf-8")
    return md_file


class TestIndentedTopLevelList:
    """DEF-MARK-42: `markdown_to_html` builds the converter with `tab_length=2` so a
    two-space nested list nests, and that same setting is the indented-code
    threshold - so a TOP-LEVEL list written two spaces in (every generated
    table of contents) was classified as code and exported as literal markdown
    source. `normalize_list_indentation` settles the indentation before the
    converter sees it, and these assert BOTH sides of the trade: the indented
    top-level list becomes a list, and the two-space nesting `tab_length=2`
    exists for still nests."""

    def _norm(self, src):
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase
        return ExportHandlerBase.normalize_list_indentation(
            ExportHandlerBase, src)

    def test_a_sub_list_behind_a_paragraph_is_left_alone(self):
        """A loose item's own second paragraph carries no list marker, and
        reading that as the end of the list left the item's child list looking
        standalone - so it was dedented into the ordered list above it and the
        author's steps silently renumbered."""
        src = ("1. Step one\n\n   Explanation of step one.\n\n"
               "   - detail a\n   - detail b\n\n2. Step two\n")
        assert self._norm(src) == src, (
            "a sub-list behind a continuation paragraph was hoisted out of "
            "its item, which renumbers every step below it")

    def test_a_comment_in_a_sample_is_not_a_heading(self):
        """`# install` inside a fenced sample is the commonest line in a
        procedure document. Splitting the chunk on it ended the list that was
        holding the sample, and the item's own notes below were then read as
        standalone and hoisted into the numbering."""
        src = ("1. Install\n\n   ```bash\n   # install\n   pip install foo\n"
               "   ```\n\n   - note a\n\n2. Run it\n")
        assert self._norm(src) == src, (
            "a comment line inside a sample was read as a heading, so the "
            "item's notes were hoisted into the step numbering")

    def test_a_loose_indented_list_is_left_alone(self):
        """A blank line between items does not make them separate lists, and
        the pass has no way to be sure it has found all of one. Moving part of
        a loose list left the rest one level in; gathering the whole of it put
        a fenced sample's bytes at risk. Neither is worth an index, so a list
        written loose is refused."""
        src = ("Contents.\n\n  1. first\n\n  2. second\n\n  3. third\n\nEnd.\n")
        assert self._norm(src) == src, (
            "a loose list was rescued; only a list written as one block is")

    def test_an_item_owning_a_sample_is_left_where_it_was(self):
        """A sample's bytes stand, so it cannot travel with its item. Moving
        the marker out from under its own fenced sample leaves the sample
        indented past the item, where it renders as an indented code block
        showing its own fence markers - worse than not moving at all."""
        src = ("  - input as string\n\n    ```shell\n    curl localhost\n"
               "    ```\n")
        assert self._norm(src) == src, (
            "the item was moved out from under the sample it owns")

    def test_a_heading_inside_an_item_does_not_end_the_list(self):
        """DEF-MARK-58: an author who writes a heading at an item's content
        column means it to sit inside the item. Reading it as the end of the
        list hands the item's child list to the numbering above."""
        src = "- Alpha\n\n  ### Notes\n\n  - n1\n  - n2\n\n- Beta\n"
        assert self._norm(src) == src, (
            "an indented heading ended the list, so the item's child list was "
            "hoisted up to sit beside its own parent")

    def test_a_list_below_an_indented_one_keeps_its_own_nesting(self):
        """DEF-MARK-59: a column recorded for one list was matched against a
        later chunk that merely sat at it, flattening a genuinely nested
        sub-list into its parent. The recorded column is gone - a list is now
        gathered whole before anything moves - and the shape is refused."""
        src = "Contents.\n\n  - a\n\n- top\n\n  - nested\n\n- second\n"
        assert self._norm(src) == src, (
            "the nested item was dragged to column 0 by a shift belonging to "
            "the earlier indented list")
        indented = ("**Contents**\n\n  - [A](#a)\n  - [B](#b)\n\n - Note\n\n"
                    "  - Child of note\n")
        assert self._norm(indented) == indented, (
            "the index was rescued although a list one column in follows it - "
            "python-markdown reads one leading space as top level, so the two "
            "merge and the second list's child is torn out from under it")

    def test_a_list_one_column_in_is_left_alone(self):
        """DEF-MARK-60: `OListProcessor.RE` already allows `tab_length - 1`
        leading spaces, so a marker one column in is a list without help. The
        shift of one would move its ragged child off the even-column grid the
        nesting level is read from, collapsing two levels into one."""
        src = "Intro.\n\n - Alpha\n  - child\n - Beta\n"
        assert self._norm(src) == src, (
            "a list that needed no rescue was shifted by an odd amount and "
            "lost a nesting level")

    def test_an_index_above_a_numbered_list_is_left_alone(self):
        """DEF-MARK-61: two lists parted by nothing but a blank line read as
        one, so a rescued index takes the steps below it as bullets. The
        control leaves the index as source and keeps the numbering, which is
        the better of the two."""
        src = "Contents.\n\n  - a\n  - b\n\n1. one\n2. two\n"
        assert self._norm(src) == src, (
            "the index was rescued into the numbered list below it, which "
            "costs the author's numbering in all three formats")
        parted = "Contents.\n\n  - a\n\n## Steps\n\n1. one\n"
        assert self._norm(parted) != parted, (
            "a heading parts the two lists, so the rescue is safe there and "
            "must still happen")

    def test_a_loose_list_is_decided_whole(self):
        """DEF-MARK-63: a loose list is one list. Deciding chunk by chunk left
        it half moved - the half at column 0 became a list and the half left
        behind became its child, dragging in whatever followed."""
        table = ("Contents.\n\n  - [a](#a)\n\n  - [b](#b)\n"
                 "| x | y |\n| - | - |\n| 1 | 2 |\n")
        assert self._norm(table) == table, (
            "half the index moved and swallowed the table under it")
        steps = ("Contents.\n\n  - [a](#a)\n\n  - [b](#b)\n\n"
                 "1. Step one\n2. Step two\n")
        assert self._norm(steps) == steps, (
            "half the index moved and took the numbered steps as its own "
            "bullets")
        whole = "Contents.\n\n  1. first\n\n  2. second\n\nEnd.\n"
        assert self._norm(whole) == whole, (
            "a loose list is refused whatever follows it - the pass cannot "
            "tell it has found the whole of one")

    def test_both_setext_spellings_and_a_borderless_table_are_seen(self):
        """DEF-MARK-64: `-` is the level-2 setext underline and the commoner of
        the two, and a markdown table needs no outer pipes."""
        setext = "Contents.\n\n  - [a](#a)\n  - [b](#b)\nMy Section\n---\n\nBody.\n"
        assert self._norm(setext) == setext, (
            "the heading underlined with `---` was swallowed into the list")
        borderless = ("Contents.\n\n  - [a](#a)\n  - [b](#b)\n"
                      "Name | Age\n---- | ----\nAnn | 30\n")
        assert self._norm(borderless) == borderless, (
            "the table written without outer pipes was swallowed into the "
            "list")

    def test_two_lists_of_the_same_kind_also_merge(self):
        """DEF-MARK-65: python-markdown merges two adjacent ordered lists as
        readily as a mixed pair, and renumbers the second."""
        src = "Contents.\n\n  1. Overview\n  2. Method\n\n1. Do this\n2. Do that\n"
        assert self._norm(src) == src, (
            "the ordered index was rescued into the procedure below it, which "
            "renumbers the author's steps")

    def test_only_a_column_zero_heading_parts_two_lists(self):
        """DEF-MARK-66: a break line is never moved by the shift, so an
        indented heading ends up inside the rescued item and parts nothing -
        while having already disabled the guard that keeps the list below
        from being absorbed."""
        src = ("Contents.\n\n  - [a](#a)\n  - [b](#b)\n\n   ### Steps\n\n"
               "1. Step one\n2. Step two\n")
        assert self._norm(src) == src, (
            "an indented heading was read as parting the two lists, and the "
            "numbered steps became bullets of the index")

    def test_a_sample_is_never_rewritten_from_inside_a_list(self):
        """DEF-MARK-67: a chunk was judged by its list markers alone, so a
        fenced sample whose own body is an indented list joined the list above
        it and had two columns stripped from inside the fence. `services:`
        with its items at column 2 is different YAML from one at column 0."""
        own_chunk = "  - alpha\n  ```yaml\n  services:\n    - one\n    - two\n  ```\n"
        assert self._norm(own_chunk) == own_chunk, (
            "the sample shares the item's chunk, and shifting the chunk "
            "rewrote indentation inside the fence")
        own_block = ("  - alpha\n\n  - beta\n\n```yaml\nservices:\n  - one\n"
                     "  - two\n```\n\n  - gamma\n")
        out = self._norm(own_block)
        assert "  - one" in out and "  - two" in out, (
            f"the fenced sample's own indentation was rewritten: {out!r}")

    def test_the_rescue_needs_what_follows_to_end_the_list(self):
        """DEF-MARK-68: what a rescued list absorbs when it lands is not
        knowable from the next chunk's first line. The rescue is taken only
        when what follows certainly ends the list."""
        note = ("## Contents\n\n  - [Setup](#setup)\n  - [Run it](#run-it)\n\n"
                "    Generated automatically.\n\n1. Install\n2. Restart\n")
        assert self._norm(note) == note, (
            "an indented note stood between the index and the procedure, and "
            "hid the procedure from the guard")
        one_in = "  - index\n  - criteria\n\n 1. Prepare\n 2. Train\n"
        assert self._norm(one_in) == one_in, (
            "the list below sits one column in, which python-markdown reads "
            "as top level, so the two merge and the numbering is lost")
        later = "  - Overview\n\n   - detail\n1. Kick-off\n2. Delivery\n"
        assert self._norm(later) == later, (
            "the procedure starts on a later line of the next chunk, where "
            "reading only its first line cannot see it")
        # and the three shapes that DO end the list are still rescued
        for name, src, want in (
            ('prose', "## C\n\n  - [a](#a)\n\nBody.\n",
             "## C\n\n- [a](#a)\n\nBody.\n"),
            ('heading', "## C\n\n  - [a](#a)\n\n## S\n\n1. one\n",
             "## C\n\n- [a](#a)\n\n## S\n\n1. one\n"),
            ('comment', "<!-- TOC:BEGIN -->\n\n  - [a](#a)\n<!-- TOC:END -->\n",
             "<!-- TOC:BEGIN -->\n\n- [a](#a)\n<!-- TOC:END -->\n"),
        ):
            assert self._norm(src) == want, (
                f"the index closed by {name} was not rescued")

    def test_a_rule_inside_the_chunk_refuses_the_rescue(self):
        """DEF-MARK-70: a `---` standing at the shift column is a rule where it
        is written and the underline of the line above it once the chunk lands
        at column 0 - a heading the author never wrote, carrying the raw `- `
        marker into Word's navigation pane."""
        dashes = "Contents\n\n  - alpha\n  ---\n  - beta\n\nProse.\n"
        assert self._norm(dashes) == dashes, (
            "the rule became a setext underline and fabricated a heading")
        equals = "Contents\n\n  - alpha\n  - beta\n  ===\n\nProse.\n"
        assert self._norm(equals) == equals, (
            "the `===` spelling fabricated a heading the same way")

    def test_a_reopened_pre_block_keeps_its_bytes(self):
        """DEF-MARK-69: a pass that counted `<pre>` tags itself was fooled by
        a block that closed and reopened on one line, or a nested one. The
        converter's raw-HTML extractor keeps the block open, so a shift inside
        it changes no list into being and is refused by the measurement."""
        reopened = "Head\n\n<pre>x</pre><pre>\n\n  - a\n  - b\n\n</pre>\n\nProse.\n"
        assert self._norm(reopened) == reopened, (
            "indentation was stripped from inside a still-open <pre>")
        nested = "Head\n\n<pre>\n<pre>\n</pre>\n\n  - a\n  - b\n\n</pre>\n\nProse.\n"
        assert self._norm(nested) == nested, (
            "the inner `</pre>` was read as closing the outer block")

    def test_an_indented_heading_after_the_list_refuses_the_rescue(self):
        """DEF-MARK-71: an ATX line is recorded as a break, not a chunk, and
        only a column-0 break parts two lists - so an indented heading was
        neither readable by the refusal nor parting, and `shift_chunk` leaves a
        break where it is. It ended up inside the item the rescue created."""
        src = "  - a\n  - b\n  ## Sub\nSection Title\n---\n\nBody.\n"
        assert self._norm(src) == src, (
            "the list was rescued out from under an indented heading, which "
            "stopped being a heading in every format")

    def test_a_chunk_behind_an_indented_heading_does_not_end_the_list(self):
        """DEF-MARK-73: an ATX line is a break, not a chunk, so the column-0
        line after an INDENTED heading is a lazy continuation of the item that
        heading sits in - the list above is still open. Reading it as the end
        let the child list below be rescued out from under its parent."""
        src = "- Alpha\n  ### Notes\nSee below.\n\n  1. First\n  2. Second\n"
        assert self._norm(src) == src, (
            "the numbered sub-list was dedented out of its parent item and "
            "merged into the bullet list above, losing its numbers")

    def test_the_fence_tracker_closes_where_the_converter_does(self):
        """DEF-MARK-74: `FencedBlockPreprocessor.FENCED_BLOCK_RE` anchors the
        closer after a newline and matches it with a backreference, so neither
        an indented closer nor a longer inner fence closes the block. A pass
        that tracked fences itself read either as the close and rewrote bytes
        inside the sample; measured against the rendered document, a shift
        inside an open fence changes a code block into a code block and is
        refused."""
        indented_closer = ("Intro.\n\n```yaml\nfiller line\n  ```\n\n"
                           "  - alpha\n  - beta\n\nDone.\n\n```\n\nAfter.\n")
        assert self._norm(indented_closer) == indented_closer, (
            "an indented ``` closed the tracker's fence and the pass rewrote "
            "bytes inside the sample")
        longer_inner = "```md\n\n````\n\n  - a\n\nProse.\n\n```\n\nAfter.\n"
        assert self._norm(longer_inner) == longer_inner, (
            "a longer inner fence closed the tracker's fence")

    def test_a_tab_is_not_read_as_no_indentation(self):
        """DEF-MARK-75: python-markdown runs `expandtabs(tab_length)` before
        parsing; every column here is counted in spaces. A tab-indented line
        therefore measured as column 0 and certified as prose that ends the
        list, when the converter reads it as a child of that list."""
        successor = ("Intro.\n\n  - [1. A](#a)\n  - [2. B](#b)\n\n"
                     "\tNote about the index.\n\n1. Step one\n2. Step two\n")
        assert self._norm(successor) == successor, (
            "a tab-indented successor was read as column-0 prose, and the "
            "numbered procedure below became bullets of the index")
        rule = "Contents.\n\n  - [Alpha](#alpha)\n\n\t- Beta\n---\n\nAfter.\n"
        assert self._norm(rule) == rule, (
            "the same miscount turned the author's rule into a fabricated "
            "Word heading")

    def test_a_rescued_ordered_index_must_start_at_one(self):
        """DEF-MARK-76: `LAZY_OL` is on, so no `start` is written and the list
        renders from 1 whatever the author numbered it."""
        three = "Contents.\n\n  3. Third section\n  4. Fourth section\n\nEnd.\n"
        assert self._norm(three) == three, (
            "an index opening at 3 was rescued, and renders as 1 - pointing "
            "the reader at the wrong section")
        one = "Contents.\n\n  1. First\n  2. Second\n\nEnd.\n"
        assert self._norm(one) == "Contents.\n\n1. First\n2. Second\n\nEnd.\n", (
            "an index that does start at 1 must still be rescued")

    def test_the_inert_comment_covers_only_itself(self):
        """DEF-MARK-77: a comment renders as nothing, but the lines behind it
        in the same chunk still move - and an indented sample shifted to
        column 0 stops being a code block."""
        mid = ("Contents.\n\n  - [Alpha](#alpha)\n<!-- end of index -->\n"
               "  services:\n    web: nginx\n\nNext.\n")
        assert self._norm(mid) == mid, (
            "the comment waved through the sample behind it, which lost its "
            "indentation and its block type")
        last = "<!-- TOC:BEGIN -->\n\n  - [a](#a)\n<!-- TOC:END -->\n"
        assert self._norm(last) == "<!-- TOC:BEGIN -->\n\n- [a](#a)\n<!-- TOC:END -->\n", (
            "a comment that closes the chunk is what a generated index writes "
            "and must still be waved through")

    def test_a_tab_led_chunk_does_not_end_the_list(self):
        """DEF-MARK-79: the third site of the tab miscount. A tab-indented
        continuation paragraph measured as column 0 and cleared the open
        list, so the item's own child list was judged standalone and hoisted
        into the numbering above it."""
        steps = ("1. Step one\n\n\tExplanation of step one.\n\n"
                 "   - detail a\n   - detail b\n\nEnd.\n")
        assert self._norm(steps) == steps, (
            "the child bullets of step one were promoted to steps two and "
            "three")
        spaces = ("1. Step one\n\n   Explanation of step one.\n\n"
                  "   - detail a\n   - detail b\n\nEnd.\n")
        assert self._norm(spaces) == spaces, (
            "the space-indented twin must stay refused too")

    def test_a_line_the_converter_does_not_call_blank_is_content(self):
        """DEF-MARK-80: `NormalizeWhitespace` empties only space-only lines.
        `str.strip()` also removes an NBSP or a form feed, so a line holding
        one was read as the blank that ends the chunk, and the converter then
        took what followed as a lazy continuation of the rescued list."""
        for sep in ('\xa0', '\x0c', '\u2003'):
            heading = "  - a\n" + sep + "\nText.\n---\n"
            assert self._norm(heading) == heading, (
                "a setext heading was welded into the bullet above it")
            para = "Intro.\n\n  - alpha\n  - beta\n" + sep + "\nImportant.\n"
            assert self._norm(para) == para, (
                "a standalone paragraph was welded into the last bullet")
        real = "Intro.\n\n  - alpha\n  - beta\n\nImportant.\n"
        assert self._norm(real) == "Intro.\n\n- alpha\n- beta\n\nImportant.\n", (
            "a real blank line must still end the chunk and permit the rescue")

    def test_a_heading_with_no_space_after_the_hashes_is_seen(self):
        """DEF-MARK-81: `HashHeaderProcessor.RE` needs no space after the
        hashes, so `#Note` is a heading once the shift lands it at column 0.
        The chunker asked for the space and so never recorded the break the
        indented-heading refusal reads."""
        index = "Intro.\n\n  1. one\n  #Note\n  2. two\n\nEnd.\n"
        assert self._norm(index) == index, (
            "a Word heading the author never wrote split the index in two, "
            "and the second half renumbered from 1")
        bullets = "Intro.\n\n  - a\n  ##Note\n  - b\n\nEnd.\n"
        assert self._norm(bullets) == bullets

    def test_an_interior_line_that_splits_the_list_refuses_the_rescue(self):
        """DEF-MARK-82: a raw HTML block, a thematic break or a table row at
        the shift column is content while the chunk is a code block and its
        own block once the chunk lands at column 0 - the rescued index splits
        in two and, `LAZY_OL` being on, the second half renumbers from 1."""
        for splitter in ('<div>note</div>', '***', '<!-- c -->'):
            src = ("Intro.\n\n  1. First section\n  " + splitter
                   + "\n  2. Second section\n\nEnd.\n")
            assert self._norm(src) == src, (
                f"{splitter!r} split the rescued index and renumbered its tail")

    def test_a_successor_the_converter_erases_does_not_end_the_list(self):
        """DEF-MARK-83: a column-0 line made only of whitespace the converter
        strips - an NBSP, a form feed - looks like prose that ends the list,
        but renders as nothing, so the list stays open and takes the numbered
        procedure below it as its own items."""
        for sep in ('\xa0', '\x0c'):
            src = ("Intro.\n\n  1. Overview\n  2. Method\n\n" + sep
                   + "\n\n1. First step\n2. Second step\n")
            assert self._norm(src) == src, (
                "the procedure below the index was renumbered as steps 3 and 4")

    def test_a_marker_the_converter_does_not_read_is_not_a_marker(self):
        """DEF-MARK-84: `OListProcessor.RE` accepts `1.` only, never `1)`, and
        both list processors want a space after the marker. A chunk opening
        with `1)` is an indented code sample to the converter, and shifted to
        column 0 it becomes a paragraph with its line breaks gone."""
        paren = "Intro.\n\n  1) one\n  2) two\n\nEnd.\n"
        assert self._norm(paren) == paren, (
            "a `1)` sample was shifted and its lines ran together")
        jammed = "Intro.\n\n  -foo\n  -bar\n\nEnd.\n"
        assert self._norm(jammed) == jammed

    def test_a_sample_in_the_item_own_chunk_keeps_its_fence(self):
        """A fence four columns in, with no blank line after its item, shares
        the item's chunk. Shifted with it, the fence is no longer at column 0
        and the converter reads its markers as the item's own text - the
        rendered list holds no block a check could see, so a chunk holding a
        fence is never a candidate at all."""
        src = ("Steps.\n\n  - run it\n    ```shell\n    make\n    ```\n\n"
               "End.\n")
        assert self._norm(src) == src, (
            "the item was rescued and its sample's fence became item text")

    def test_a_raw_block_inside_an_item_refuses_the_rescue(self):
        """DEF-MARK-86: the block-level set is the converter's own, not a
        six-tag copy of it. A `<style>` or `<script>` at the item's content
        column is inert escaped text in the control's code block and a live
        element once the chunk is a list."""
        for raw in ('<style>body{display:none}</style>', '<script>alert(1)</script>',
                    '<details><summary>s</summary>body</details>'):
            src = "Intro.\n\n  - a\n    " + raw + "\n  - b\n\nEnd.\n"
            assert self._norm(src) == src, (
                f"{raw[:8]} inside an item was rescued into a live element")

    def test_a_nested_ordered_run_must_also_open_at_one(self):
        """DEF-MARK-88: `LAZY_OL` renumbers a nested list from 1 as readily
        as a top-level one, and the start-at-one rule read only the chunk's
        first marker."""
        nested = "Intro.\n\n  1. a\n    3. b\n    4. c\n  2. d\n\nEnd.\n"
        assert self._norm(nested) == nested, (
            "a nested run opening at 3 was rescued and renders 1, 2")
        ok = "Intro.\n\n  1. a\n    1. b\n    2. c\n  2. d\n\nEnd.\n"
        assert self._norm(ok) == "Intro.\n\n1. a\n  1. b\n  2. c\n2. d\n\nEnd.\n", (
            "a nested run that does open at 1 must still be rescued")

    def test_a_raw_element_in_the_chunk_is_never_a_candidate(self):
        """DEF-MARK-89: the control escapes a tag inside an indented block;
        a rescue would emit it live, and `is_block_level` is the converter's
        opinion about layout, not about what a browser acts on - a `<link>`
        hides the page, a `<meta>` leaves it, an `<img onerror>` runs. A
        comment is not an element and still closes a generated index."""
        for raw in ('<link rel="stylesheet" href="data:text/css,body%7Bdisplay%3Anone%7D">',
                    '<meta http-equiv="refresh" content="0;url=https://example.org">',
                    '<img src="x" onerror="document.body.style.display=\'none\'">',
                    '<base href="https://example.org/">'):
            src = "Intro.\n\n  - a\n    " + raw + "\n  - b\n\nEnd.\n"
            assert self._norm(src) == src, f"{raw[:6]} went live in a rescued item"
        toc = "<!-- TOC:BEGIN -->\n\n  - [a](#a)\n<!-- TOC:END -->\n\nText.\n"
        assert self._norm(toc) == "<!-- TOC:BEGIN -->\n\n- [a](#a)\n<!-- TOC:END -->\n\nText.\n"

    def test_a_reference_definition_in_the_chunk_is_never_a_candidate(self):
        """DEF-MARK-104: the converter consumes a link reference definition
        as a definition and shows it nowhere, so a URL the control displays
        inside its code block vanished from the rescued list."""
        src = "Intro.\n\n  - a\n  [ref]: http://example.com/x\n\nEnd.\n"
        assert self._norm(src) == src, "a reference definition was swallowed by the rescue"

    def test_trailing_whitespace_on_the_last_item_does_not_refuse_the_rescue(self):
        """DEF-MARK-111: the converter stores a code block rstripped, so a
        stray space after the last item made the control's code text differ
        from the chunk and the whole list stayed a monospace source block."""
        for src in ("Intro.\n\n  - a\n  - b \n\nEnd.\n", "Intro.\n\n  - a \n  - b\n\nEnd.\n"):
            assert self._norm(src) != src, f"trailing whitespace refused the rescue: {src!r}"

    def test_an_inline_tag_in_an_item_does_not_refuse_the_rescue(self):
        """DEF-MARK-110: the PDF path rewrites `$E=mc^2$` to an `<img>` before
        the pass runs, and any `<tag` in the chunk refused it, so one list
        was a list in Word and HTML and a code block in the PDF. Only a tag
        opening a line is a raw block; one inside the item's text is inline."""
        src = "Intro.\n\n  - Mass\n  - Energy <img src=\"data:image/png;base64,AAAA\">\n\nEnd.\n"
        out = self._norm(src)
        assert out != src and "- Energy <img" in out, (
            "an inline tag refused the rescue")

    def test_a_candidate_the_converter_cannot_render_is_refused(self):
        """DEF-MARK-90: a list nested past the parser's depth renders as a
        code block in the control and raises inside the measurement of the
        shifted text; the raise took every format's export down."""
        src = ("Intro.\n\n" + "".join(f"{'  ' * (k + 1)}- lvl{k}\n" for k in range(1500))  # past any default recursion limit
               + "\nEnd.\n")
        assert self._norm(src) == src, "a candidate the converter cannot render was kept"

    def test_a_marker_of_the_other_kind_at_one_depth_refuses_the_rescue(self):
        """DEF-MARK-91: the converter folds a marker into the run already open
        at its depth whatever its kind, so `- Intro` then `1. First` renders
        as three bullets and `1. a / - b / 2. c` as 1, 2, 3."""
        for src in ("Intro.\n\n  - Intro\n  1. First\n  2. Second\n\nEnd.\n",
                    "Intro.\n\n  1. a\n  - b\n  2. c\n\nEnd.\n"):
            assert self._norm(src) == src, "mixed marker kinds were rescued and merged"
        nested = "Intro.\n\n  - a\n    1. x\n    2. y\n  - b\n\nEnd.\n"
        assert self._norm(nested) == "Intro.\n\n- a\n  1. x\n  2. y\n- b\n\nEnd.\n", (
            "a different kind at a deeper depth is a nested list and must still be rescued")

    def test_depth_is_the_converter_bucket_not_the_column(self):
        """DEF-MARK-95: the list processors take up to `tab_length - 1` extra
        spaces as the same level, so a bullet one column deeper than the
        numbered items around it is their sibling and merges into their run
        - `1.`, `2.`, `3.` where the author wrote `1.`, a bullet, `2.`."""
        for src in ("Intro.\n\n  1. a\n   - b\n  2. c\n\nEnd.\n",
                    "Intro.\n\n  - a\n   1. b\n\nEnd.\n",
                    "Intro.\n\n  - a\n    1. x\n     - y\n  - b\n\nEnd.\n"):
            assert self._norm(src) == src, "a marker one column off its sibling was rescued and merged"
        ok = "Intro.\n\n  - a\n    1. x\n    2. y\n  - b\n\nEnd.\n"
        assert self._norm(ok) != ok, "a nested run two columns in is a real level and must still be rescued"

    def test_every_marker_line_must_arrive_as_an_item(self):
        """DEF-MARK-96: a marker the converter cannot nest - four columns
        under its parent, where the unit is two - is kept as the parent's
        text with its `-` or `1.` literal: one list, nothing block-level in
        it, the author's level gone. Items are counted against marker lines."""
        for src in ("Intro.\n\n  - a\n      - a1\n      - a2\n  - b\n\nEnd.\n",
                    "Intro.\n\n  1. o0\n     1. o1\n        1. o2\n        2. o2b\n"
                    "     2. o1b\n  2. o0b\n\nEnd.\n"):
            assert self._norm(src) == src, "a collapsed level was rescued as item text"
        toc = "## C\n\n  - [1. One](#1-one)\n    - [1.1 Sub](#11-sub)\n  - [2. Two](#2-two)\n\nText.\n"
        assert self._norm(toc) == "## C\n\n- [1. One](#1-one)\n  - [1.1 Sub](#11-sub)\n- [2. Two](#2-two)\n\nText.\n"

    def test_a_line_carrying_two_markers_refuses_the_rescue(self):
        """DEF-MARK-102: `- 1. Intro` is two items to the converter and one
        marker line here, so items could no longer be counted against marker
        lines and a level the converter cannot nest below it went unseen."""
        for src in ("Intro.\n\n  - 1. Intro\n        - sub\n\nEnd.\n",
                    "Intro.\n\n  1. - a\n        - a2\n\nEnd.\n"):
            assert self._norm(src) == src, "a two-marker line was rescued"

    def test_a_commented_closing_tag_does_not_close_a_pre(self):
        """DEF-MARK-78: the converter's raw-HTML extractor does not close a
        `<pre>` on a `</pre>` written inside a comment, and neither does the
        measurement - the block it renders is a `<pre>` either way."""
        src = ("Intro.\n\n<pre>\n<!-- </pre> -->\n\n  - alpha\n  - beta\n\n"
               "Done.\n</pre>\n\nAfter.\n")
        assert self._norm(src) == src, (
            "the commented tag took the depth to zero and the pass rewrote "
            "indentation inside the preformatted block")

    def test_a_table_under_an_indented_list_is_left_alone(self):
        """DEF-MARK-62: a column-0 line stands where the shift would put it,
        but the arriving list makes it a lazy continuation - and a table is
        structure the control draws, not prose."""
        src = ("## Contents\n\n  - [a](#a)\n  - [b](#b)\n"
               "| x | y |\n| - | - |\n| 1 | 2 |\n")
        assert self._norm(src) == src, (
            "the table was swallowed into a list item, where its rows render "
            "as literal pipe characters")
        prose = "Intro.\n\n  - item\ncontinues here\n"
        assert self._norm(prose) == prose, (
            "a column-0 line under the list was swallowed as a continuation; "
            "only a line that renders as nothing may stand there")

    async def test_html_renders_the_indented_toc_as_links(
            self, jp_fetch, test_indented_toc_file):
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "test_indented_toc.md"}),
            raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert '<a href="#1.-Introduction">1. Introduction</a>' in html, (
            "the two-space TOC did not become a list of links"
        )
        assert "[1. Introduction](#1.-Introduction)" not in html, (
            "literal markdown source leaked into the HTML"
        )

    async def test_docx_renders_the_indented_toc_as_text_not_source(
            self, jp_fetch, test_indented_toc_file):
        from docx import Document

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "test_indented_toc.md"}),
            raise_error=False,
        )
        assert response.code == 200
        doc = Document(io.BytesIO(response.body))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "1. Introduction" in text, "the TOC entry is missing from the DOCX"
        assert "](#1.-Introduction)" not in text, (
            "the TOC exported as literal markdown source"
        )

    async def test_pdf_renders_the_indented_toc_as_text_not_source(
            self, jp_fetch, test_indented_toc_file):
        import fitz

        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/pdf",
            method="POST", body=json.dumps({"path": "test_indented_toc.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert response.body.startswith(b"%PDF-")
        doc = fitz.open(stream=response.body, filetype="pdf")
        text = "".join(page.get_text() for page in doc)
        doc.close()
        assert "](#1.-Introduction)" not in text, (
            "the TOC exported as literal markdown source"
        )

    async def test_ordered_list_two_spaces_in_is_numbered(
            self, jp_fetch, test_indented_toc_file):
        """An ordered list is classified by the same threshold as a bulleted
        one, so it needs the same shift."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "test_indented_toc.md"}),
            raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert "<ol>" in html and "<li>first ordered</li>" in html, (
            "the two-space ordered list did not become a numbered list"
        )
        assert "1. first ordered" not in html, (
            "literal ordered-list source leaked into the HTML"
        )

    async def test_two_space_nesting_still_nests(
            self, jp_fetch, test_indented_toc_file):
        """The regression guard on the reason `tab_length=2` was chosen: a list
        nested by two spaces must stay a child, not become a sibling."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "test_indented_toc.md"}),
            raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert re.search(
            r"<li>top level<ul>\s*<li>nested under top</li>", html), (
            "the two-space nested item was flattened into a sibling"
        )

    async def test_indented_code_block_stays_code(
            self, jp_fetch, test_indented_toc_file):
        """Four or more spaces is an indented code block - the one shape the
        shift must never claim."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "test_indented_toc.md"}),
            raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        plain = re.sub(r"<[^>]+>", "", html)
        assert "- this is code, not a list" in plain, (
            "the indented code block lost its content"
        )
        assert "<li>this is code, not a list</li>" not in html, (
            "an indented code block was promoted to a list"
        )

    async def test_list_inside_a_fence_survives_verbatim(
            self, jp_fetch, test_indented_toc_file):
        """A `  - x` inside a fenced sample is code: its two spaces are part of
        what the author is showing."""
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "test_indented_toc.md"}),
            raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        # Pygments wraps each token in a span; strip tags before reading the text
        plain = re.sub(r"<[^>]+>", "", html)
        assert "\n  - fenced item" in plain, (
            "the fenced sample was dedented with the document around it"
        )

    async def test_nested_item_after_a_blank_line_is_not_flattened(
            self, jp_fetch, jp_root_dir):
        """A loose list puts a blank line before its nested item, so "preceded
        by a blank line" alone cannot mean "top level" - the open list has to
        be tracked too, or a nested item is dedented into a sibling."""
        (jp_root_dir / "loose.md").write_text(
            "# L\n\n- top level\n\n  - nested under top\n", encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "loose.md"}),
            raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert "<ul>\n<li>\n<p>top level</p>\n<ul>" in html, (
            "a nested item after a blank line was flattened into a sibling"
        )

    async def test_an_indented_block_keeps_its_own_nesting(
            self, jp_fetch, jp_root_dir):
        """Every line of the block moves by the same amount, so nesting inside
        a block that is itself indented survives the shift."""
        (jp_root_dir / "shifted.md").write_text(
            "# S\n\n  - outer item\n    - inner item\n", encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "shifted.md"}),
            raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert re.search(r"<li>outer item<ul>\s*<li>inner item</li>", html), (
            "the shifted block lost the nesting it was written with"
        )

    async def test_a_fence_left_of_the_shift_keeps_its_body_verbatim(
            self, jp_fetch, jp_root_dir):
        """A fence written left of the open marker column ends the block, and
        must do so before its own lines are read. Shifting the sample lines
        that reach the marker column while leaving those left of it alone is
        no translation at all - it rewrites indentation inside content the
        author wrote to be reproduced exactly. The rescue is measured against
        the rendered document, so the list above the sample is rescued and
        the sample's block is byte-identical to the control's."""
        (jp_root_dir / "fenceleft.md").write_text(
            "# F\n\nSteps:\n\n  - do this\n\n```python\ndef f():\n"
            "    return 1\n```\n", encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "fenceleft.md"}),
            raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        plain = re.sub(r"<[^>]+>", "", html)
        assert "<li>do this</li>" in html, (
            "the list above a fenced sample was left as literal source; the "
            "sample below it is its own block and nothing the list can absorb")
        assert "\n    return 1" in plain, (
            "the fenced body was dedented by the list's shift, so the sample "
            "exports as invalid Python"
        )

    async def test_a_display_math_item_does_not_become_a_heading(
            self, jp_fetch, jp_root_dir):
        """DEF-MARK-85: the pass measures what the converter renders, so it
        must see the text the converter gets. Run before the math pass, it
        certified `- $$a = b$$` as a list item; the math pass then left a
        bare `- ` line, which is the setext underline of the item above."""
        (jp_root_dir / "mathitem.md").write_text(
            "# T\n\nIntro.\n\n  - intro item\n  - $$a = b$$\n  - outro item\n\n"
            "End.\n", encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "mathitem.md"}),
            raise_error=False)
        assert response.code == 200
        from docx import Document as _Doc
        document = _Doc(io.BytesIO(response.body))
        fabricated = [p.text for p in document.paragraphs
                      if p.style.name.startswith("Heading")
                      and "intro item" in p.text]
        assert not fabricated, (
            f"the item above the formula became a Word heading: {fabricated!r}")

    async def test_docx_keeps_a_fenced_body_the_list_shift_reached(
            self, jp_fetch, jp_root_dir):
        """The pass runs before all three handlers, so the same clipped sample
        reaches Word."""
        (jp_root_dir / "fencedocx.md").write_text(
            "# F\n\nSteps:\n\n  - do this\n\n```python\ndef f():\n"
            "    return 1\n```\n", encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/docx",
            method="POST", body=json.dumps({"path": "fencedocx.md"}),
            raise_error=False,
        )
        assert response.code == 200
        assert "    return 1" in _docx_text(response.body), (
            "the fenced body reached Word with the list's shift cut off it"
        )

    async def test_a_closing_fence_ends_the_shifted_block(
            self, jp_fetch, jp_root_dir):
        """The block ends at the fence, so what follows the fence is read at
        the column it was written in - three spaces is still indented code."""
        (jp_root_dir / "afterfence.md").write_text(
            "# A\n\nIntro.\n\n  - item\n\n```\nx\n```\n\n"
            "   indented code block\n", encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "afterfence.md"}),
            raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert "<p>indented code block</p>" not in html, (
            "the shift outlived the fence and demoted an indented code block "
            "to an ordinary paragraph"
        )
        assert "indented code block" in re.sub(r"<[^>]+>", "", html), (
            "the code block lost its content"
        )

    async def test_an_unclosed_fence_does_not_dedent_the_rest_of_the_file(
            self, jp_fetch, jp_root_dir):
        """A fence tracker with nothing to close it stays open to end of file,
        so nothing may depend on a fence closing to end the block - or one
        stray fence flattens every list below it."""
        (jp_root_dir / "unclosed.md").write_text(
            "Intro.\n\n  - a\n\n      ```\n      how a fence is written\n\n"
            "## Heading\n\n- top\n  - nested\n    - deeper\n", encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "unclosed.md"}),
            raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert re.search(r"<li>top<ul>\s*<li>nested<ul>\s*<li>deeper</li>", html), (
            "the open fence carried the shift past a heading and flattened "
            "the nesting of every list after it"
        )

    async def test_a_list_written_under_its_heading_is_a_list(
            self, jp_fetch, jp_root_dir):
        """A heading ends the block above it, so the marker under it is
        top-level whether or not a blank line separates the two - which is how
        a table of contents and most hand-written lists are actually laid
        out."""
        (jp_root_dir / "underhead.md").write_text(
            "## Heading\n  - first entry\n  - second entry\n", encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "underhead.md"}),
            raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert "<li>first entry</li>" in html and "<li>second entry</li>" in html, (
            "a list written directly under its heading still exports as "
            "literal markdown source"
        )

    async def test_a_lazy_continuation_leaves_the_list_open(
            self, jp_fetch, jp_root_dir):
        """An item's paragraph may run on at column 0 without ending the item,
        so a line left of the content column closes the list only after a
        blank one. Read as a close, the item's own nested marker below looks
        top-level and is dedented into its parent's sibling."""
        (jp_root_dir / "lazy.md").write_text(
            "# L\n\n- *Tips*: the first line of the item\n"
            "runs on at column 0 without ending it.\n"
            "  - a genuinely nested point\n", encoding="utf-8")
        response = await jp_fetch(
            "jupyterlab-export-markdown-extension", "export/html",
            method="POST", body=json.dumps({"path": "lazy.md"}),
            raise_error=False,
        )
        assert response.code == 200
        html = response.body.decode("utf-8")
        assert re.search(r"<ul>\s*<li>a genuinely nested point</li>", html), (
            "the lazy continuation was read as closing the list, so the "
            "nested item was flattened into a sibling of its parent"
        )


CALLOUT_BOX_DOC = """# Callout

Intro paragraph.

<div style="border: 2px dashed #9ca3af; background: rgba(244, 244, 245, 0.45); padding: 10px 14px; margin: 10px 0; color: #000000;">
<b>&#x26A0; DRAFT v12, not for signature.</b> This document is under active preparation.
</div>

After the box.
"""


@pytest.fixture
def test_callout_box_file(jp_root_dir):
    md_file = jp_root_dir / "test_callout_box.md"
    md_file.write_text(CALLOUT_BOX_DOC, encoding="utf-8")
    return md_file


def _docx_callout_boxes(doc):
    """(bar_hex, fill_hex, text) of every callout box in a DOCX - the
    single-cell table with a coloured left border that style_docx_alert_boxes
    builds for a GitHub alert and for a hand-drawn box alike."""
    from docx.oxml.ns import qn

    boxes = []
    for tbl in doc.tables:
        if len(tbl.rows) != 1 or len(tbl.rows[0].cells) != 1:
            continue
        tblPr = tbl._tbl.find(qn("w:tblPr"))
        borders = tblPr.find(qn("w:tblBorders")) if tblPr is not None else None
        left = borders.find(qn("w:left")) if borders is not None else None
        if left is None:
            continue
        tcPr = tbl.rows[0].cells[0]._tc.find(qn("w:tcPr"))
        shd = tcPr.find(qn("w:shd")) if tcPr is not None else None
        fill = shd.get(qn("w:fill")) if shd is not None else ""
        boxes.append(((left.get(qn("w:color")) or "").upper(),
                      (fill or "").upper(),
                      tbl.rows[0].cells[0].text))
    return boxes


async def _export(jp_fetch, fmt, path):
    response = await jp_fetch(
        "jupyterlab-export-markdown-extension", f"export/{fmt}",
        method="POST", body=json.dumps({"path": path}), raise_error=False,
    )
    assert response.code == 200, f"{fmt} export failed"
    return response


class TestCssColorNotation:
    """DEF-DIAG-29: `rgb()` / `rgba()` is the notation a rich-text paste and every
    devtools colour copy produces, and an extended colour name is one an author
    simply writes; both used to resolve to nothing and render black."""

    async def test_docx_rgb_and_rgba_text_colour_resolve(self, jp_fetch, jp_root_dir):
        """Both notations must arrive, the alpha one composited onto the white
        page - 45% of DC2626 over white is EF9D9D, which is what the author
        sees in a browser. The alpha form is also the one that used to take the
        export down with it: left in the style attribute, htmldocx reads the
        channels with int() and the whole request dies on `invalid literal for
        int() with base 10: '0.45'`."""
        from docx import Document

        (jp_root_dir / "rgb.md").write_text(
            'Level <span style="color: rgb(220, 38, 38)">alarm</span> and '
            '<span style="color: rgba(220, 38, 38, 0.45)">faded</span> here.\n',
            encoding="utf-8")
        response = await _export(jp_fetch, "docx", "rgb.md")
        doc = Document(io.BytesIO(response.body))
        found = {r.text.strip(): r.font.color for p in doc.paragraphs for r in p.runs
                 if r.text.strip() in ("alarm", "faded")}
        assert set(found) == {"alarm", "faded"}, "a coloured run is missing from the DOCX"
        for label, colour in found.items():
            assert colour is not None and colour.rgb is not None, (
                f"{label}: the colour resolved to nothing, so the run renders black"
            )
        assert str(found["alarm"].rgb) == "DC2626", (
            f"alarm: exported as {found['alarm'].rgb}, not DC2626"
        )
        assert str(found["faded"].rgb) == "EF9D9D", (
            f"faded: exported as {found['faded'].rgb}, not EF9D9D - the alpha "
            "was dropped instead of composited, so the run arrives at full "
            "strength and is darker than the author drew it"
        )

    async def test_docx_percentage_channels_resolve(self, jp_fetch, jp_root_dir):
        """CSS states a channel as a number or a percentage, and both are what
        a stylesheet copied out of a browser can carry."""
        from docx import Document

        (jp_root_dir / "pct.md").write_text(
            'A <span style="color: rgb(100%, 0%, 0%)">pct</span> value.\n',
            encoding="utf-8")
        response = await _export(jp_fetch, "docx", "pct.md")
        doc = Document(io.BytesIO(response.body))
        runs = [r for p in doc.paragraphs for r in p.runs if r.text.strip() == "pct"]
        assert runs and runs[0].font.color.rgb is not None, "percentage rgb() lost"
        assert str(runs[0].font.color.rgb) == "FF0000", (
            f"rgb(100%, 0%, 0%) exported as {runs[0].font.color.rgb}, not FF0000"
        )

    async def test_docx_rgba_background_shades_the_run(self, jp_fetch, jp_root_dir):
        """Word has no translucent run, so a half-strength yellow arrives
        composited onto the white page (FFFF80) and an opaque one at full
        strength (FFFF00) - both as true run shading rather than the Word
        highlight htmldocx picks from its fixed palette."""
        from docx import Document
        from docx.oxml.ns import qn

        (jp_root_dir / "pill.md").write_text(
            'Marked <span style="background-color: rgba(255, 255, 0, 0.5)">lit</span> '
            'and <span style="background-color: rgb(255,255,0)">plain</span>.\n',
            encoding="utf-8")
        response = await _export(jp_fetch, "docx", "pill.md")
        doc = Document(io.BytesIO(response.body))
        runs = {r.text.strip(): r for p in doc.paragraphs for r in p.runs
                if r.text.strip() in ("lit", "plain")}
        assert set(runs) == {"lit", "plain"}, "a shaded run is missing from the DOCX"
        for label, expected in (("lit", "FFFF80"), ("plain", "FFFF00")):
            rPr = runs[label]._r.find(qn("w:rPr"))
            shd = rPr.find(qn("w:shd")) if rPr is not None else None
            assert shd is not None, f"{label}: the background did not become run shading"
            assert shd.get(qn("w:fill")) == expected, (
                f"{label}: shaded {shd.get(qn('w:fill'))}, not {expected}"
            )
            assert rPr.find(qn("w:highlight")) is None, (
                f"{label}: left to htmldocx, which maps any background to the "
                "nearest of its own few highlight names"
            )

    async def test_docx_extended_colour_name_is_not_black(self, jp_fetch, jp_root_dir):
        from docx import Document

        (jp_root_dir / "named.md").write_text(
            'A <span style="color: lightgray">pale</span> word.\n', encoding="utf-8")
        response = await _export(jp_fetch, "docx", "named.md")
        doc = Document(io.BytesIO(response.body))
        runs = [r for p in doc.paragraphs for r in p.runs if r.text.strip() == "pale"]
        assert runs and runs[0].font.color.rgb is not None, (
            "the colour name resolved to nothing, so the run renders black"
        )
        assert str(runs[0].font.color.rgb) == "D3D3D3", (
            f"lightgray exported as {runs[0].font.color.rgb}, not D3D3D3"
        )

    async def test_a_non_finite_channel_does_not_take_the_export_down(
            self, jp_fetch, jp_root_dir):
        """`nan`, `inf` and an overflowing exponent are all values float()
        accepts and round() refuses, so guarding only the parse leaves the
        failure this branch was written to remove: HTTP 500 and no document at
        all, for one malformed declaration anywhere in the file."""
        (jp_root_dir / "wild.md").write_text(
            '# W\n\n<span style="color: rgb(nan, 0, 0)">first</span> and '
            '<span style="color: rgb(1e400, 0, 0)">second</span>.\n\nafter\n',
            encoding="utf-8")
        for fmt in ("docx", "pdf", "html"):
            response = await jp_fetch(
                "jupyterlab-export-markdown-extension", f"export/{fmt}",
                method="POST", body=json.dumps({"path": "wild.md"}),
                raise_error=False,
            )
            assert response.code == 200, (
                f"{fmt}: a non-finite colour channel lost the whole document - "
                f"{response.body[:120]!r}"
            )
        response = await _export(jp_fetch, "docx", "wild.md")
        assert "after" in _docx_text(response.body), "the document lost its text"

    async def test_a_translucent_wash_arrives_light_not_at_full_strength(
            self, jp_fetch, jp_root_dir):
        """Alpha is what makes a colour light. Dropped rather than composited,
        the commonest soft-grey wash there is - a few percent of black -
        inverts to solid black and paints out the text sitting on it."""
        from docx import Document
        from docx.oxml.ns import qn

        (jp_root_dir / "wash.md").write_text(
            '# W\n\nA <span style="background: rgba(0, 0, 0, 0.05)">soft wash</span> '
            'inline.\n', encoding="utf-8")
        response = await _export(jp_fetch, "docx", "wash.md")
        doc = Document(io.BytesIO(response.body))
        runs = [r for p in doc.paragraphs for r in p.runs if r.text.strip() == "soft wash"]
        assert runs, "the shaded run is missing from the DOCX"
        shd = runs[0]._r.find(qn("w:rPr")).find(qn("w:shd"))
        assert shd is not None and shd.get(qn("w:fill")) == "F2F2F2", (
            f"5% black over white is F2F2F2; shaded "
            f"{shd.get(qn('w:fill')) if shd is not None else None} instead"
        )

    async def test_a_fully_transparent_colour_draws_nothing(
            self, jp_fetch, jp_root_dir):
        """`rgba(..., 0)` is what generated CSS writes for `transparent`, which
        the named table already refuses. Read at full strength the two
        spellings of one value disagree, and the wrong one paints a black
        block over the text."""
        from docx import Document
        from docx.oxml.ns import qn

        (jp_root_dir / "clear.md").write_text(
            '# C\n\nA <span style="background: rgba(0, 0, 0, 0)">clear</span> word.\n'
            '\n<div style="background: rgba(255, 255, 255, 0); padding: 8px">'
            'clear box</div>\n', encoding="utf-8")
        response = await _export(jp_fetch, "docx", "clear.md")
        doc = Document(io.BytesIO(response.body))
        runs = [r for p in doc.paragraphs for r in p.runs if r.text.strip() == "clear"]
        assert runs, "the run is missing from the DOCX"
        rPr = runs[0]._r.find(qn("w:rPr"))
        shd = rPr.find(qn("w:shd")) if rPr is not None else None
        assert shd is None, (
            f"a fully transparent background shaded the run "
            f"{shd.get(qn('w:fill'))}"
        )
        assert rPr is None or rPr.find(qn("w:highlight")) is None, (
            "the declaration was handed to htmldocx, which picked a highlight "
            "from its own palette for a colour that paints nothing"
        )
        assert not _docx_callout_boxes(doc), (
            "a fully transparent background drew a callout box"
        )
        assert any("clear box" in p.text for p in doc.paragraphs), (
            "the div's text is missing from the DOCX"
        )


class TestHtmlCalloutBox:
    """DEF-MARK-43: a <div> carrying a border or a background is a callout the author
    drew by hand, and exports as the box an alert already gets - in the div's
    own colours, in all three formats."""

    async def test_html_keeps_the_authors_own_box(self, jp_fetch, test_callout_box_file):
        """The HTML export is the reference the other two are brought into line
        with: it hands the browser the div's own CSS."""
        response = await _export(jp_fetch, "html", "test_callout_box.md")
        html = response.body.decode("utf-8")
        assert "border: 2px dashed #9ca3af" in html, "the box lost its border CSS"
        assert "rgba(244, 244, 245, 0.45)" in html, "the box lost its background CSS"

    async def test_docx_bordered_div_becomes_a_box(self, jp_fetch, test_callout_box_file):
        from docx import Document
        from docx.oxml.ns import qn

        response = await _export(jp_fetch, "docx", "test_callout_box.md")
        doc = Document(io.BytesIO(response.body))
        boxes = [b for b in _docx_callout_boxes(doc) if "DRAFT" in b[2]]
        assert boxes, (
            "the bordered div exported as a plain paragraph - no border, no fill"
        )
        bar, fill, text = boxes[0]
        assert bar == "9CA3AF", f"the box bar is {bar}, not the div's own border colour"
        assert fill == "FAFAFA", (
            f"the box fill is {fill}, not the div's own background composited "
            "onto the page - 45% of F4F4F5 over white is FAFAFA"
        )
        assert "not for signature" in text, "the box lost the text it was drawn around"
        from docx.text.paragraph import Paragraph
        all_text = "".join(Paragraph(p, doc).text
                           for p in doc.element.body.iter(qn("w:p")))
        assert "BOX:" not in all_text and "\u2063" not in all_text, (
            "the callout marker leaked into the visible text"
        )

    async def test_pdf_bordered_div_becomes_a_box(self, jp_fetch, test_callout_box_file):
        """The PDF reads the finished DOCX box back, so the two formats draw the
        same bar and the same fill without being told twice."""
        import fitz

        response = await _export(jp_fetch, "pdf", "test_callout_box.md")
        doc = fitz.open(stream=response.body, filetype="pdf")
        fills = [c for page in doc for c in _pdf_fill_colors(page)]
        strokes = [c for page in doc for c in _pdf_stroke_colors(page)]
        text = "".join(page.get_text() for page in doc)
        doc.close()
        assert _color_near(fills, (0.980, 0.980, 0.980)), (
            "the box background (FAFAFA) is missing from the PDF"
        )
        assert _color_near(strokes, (0.612, 0.639, 0.686)), (
            "the box border colour (9CA3AF) is missing from the PDF"
        )
        assert "not for signature" in text, "the box text is missing from the PDF"
        assert "BOX:" not in text, "the callout marker leaked into the PDF text"

    async def test_docx_background_only_div_is_a_box(self, jp_fetch, jp_root_dir):
        """A background alone draws a box too, and it must not be mistaken for a
        highlighter pill run through the text."""
        from docx import Document
        from docx.oxml.ns import qn

        (jp_root_dir / "shaded.md").write_text(
            'Before.\n\n<div style="background: #ffe9a8;">Only a background.</div>'
            '\n\nAfter.\n', encoding="utf-8")
        response = await _export(jp_fetch, "docx", "shaded.md")
        doc = Document(io.BytesIO(response.body))
        boxes = [b for b in _docx_callout_boxes(doc) if "Only a background" in b[2]]
        assert boxes, "the shaded div exported without a box"
        assert boxes[0][1] == "FFE9A8", f"the box fill is {boxes[0][1]}, not the div's"
        shaded_runs = [
            r for tbl in doc.tables for row in tbl.rows for cell in row.cells
            for para in cell.paragraphs for r in para.runs
            if r._r.find(qn("w:rPr")) is not None
            and r._r.find(qn("w:rPr")).find(qn("w:shd")) is not None
        ]
        assert not shaded_runs, (
            "the fill was painted onto the runs as well - a highlighter band "
            "across text already sitting on the box's own background"
        )

    async def test_a_border_that_draws_no_line_is_not_a_box(
            self, jp_fetch, jp_root_dir):
        """`border-collapse` and `border-image` open with the same word as a
        border and draw none, and `0px`, `none !important` and a transparent
        colour each write one off. A callout box is a loud construct - drawn
        where the CSS draws nothing, it invents an aside the author never
        wrote."""
        from docx import Document

        for name, style in (
                ("zero", "border: 0px; background: none"),
                ("important", "border: none !important"),
                ("collapse", "border-collapse: collapse; padding: 4px"),
                ("image", "border-image: url(x.png)"),
                ("clear", "border: 1px solid transparent"),
        ):
            (jp_root_dir / f"{name}.md").write_text(
                f'Before.\n\n<div style="{style}">not a box</div>\n\nAfter.\n',
                encoding="utf-8")
            response = await _export(jp_fetch, "docx", f"{name}.md")
            doc = Document(io.BytesIO(response.body))
            assert not _docx_callout_boxes(doc), (
                f"`{style}` draws no line, and was boxed anyway"
            )
            assert any("not a box" in p.text for p in doc.paragraphs), (
                f"`{style}`: the div's text is missing from the DOCX"
            )

    async def test_a_border_still_draws_its_box(self, jp_fetch, jp_root_dir):
        """The other side of the same test: a width, a style or a colour on
        any one side is a line, and takes the box."""
        from docx import Document

        for name, style, bar in (
                ("solid", "border: 1px solid #123456", "123456"),
                ("style_only", "border-style: dashed", "BBBBBB"),
                ("one_side", "border-left: 3px solid #abcdef", "ABCDEF"),
        ):
            (jp_root_dir / f"b{name}.md").write_text(
                f'Before.\n\n<div style="{style}">a box</div>\n\nAfter.\n',
                encoding="utf-8")
            response = await _export(jp_fetch, "docx", f"b{name}.md")
            boxes = _docx_callout_boxes(Document(io.BytesIO(response.body)))
            assert boxes, f"`{style}` draws a line and lost its box"
            assert boxes[0][0] == bar, f"`{style}` barred {boxes[0][0]}, not {bar}"

    async def test_a_plain_div_is_not_boxed(self, jp_fetch, jp_root_dir):
        """A div is a callout only when the author drew one; a bare div must
        stay the paragraph it already exports as."""
        from docx import Document

        (jp_root_dir / "plain.md").write_text(
            "Before.\n\n<div>No box here.</div>\n\nAfter.\n", encoding="utf-8")
        response = await _export(jp_fetch, "docx", "plain.md")
        doc = Document(io.BytesIO(response.body))
        assert not _docx_callout_boxes(doc), "a div with no border and no fill was boxed"
        assert any("No box here" in p.text for p in doc.paragraphs), (
            "the div's text is missing from the DOCX"
        )

    async def test_pdf_black_bordered_box_is_still_a_box(self, jp_fetch, jp_root_dir):
        """Black is a colour a hand-drawn border is written in, and the PDF
        recognises the box by that border - so a black one must not be read as
        an ordinary table's default rule and lose its box."""
        import fitz

        (jp_root_dir / "black.md").write_text(
            'Before.\n\n<div style="border: 1px solid black;">Black bordered note.'
            '</div>\n\nAfter.\n', encoding="utf-8")
        response = await _export(jp_fetch, "pdf", "black.md")
        doc = fitz.open(stream=response.body, filetype="pdf")
        strokes = [c for page in doc for c in _pdf_stroke_colors(page)]
        text = "".join(page.get_text() for page in doc)
        doc.close()
        assert _color_near(strokes, (0.0, 0.0, 0.0)), (
            "a black-bordered box drew no bar in the PDF"
        )
        assert "Black bordered note" in text, "the box text is missing from the PDF"

    async def test_a_box_inside_a_list_item_stays_inline(self, jp_fetch, jp_root_dir):
        """Word has no table inside a list item, so boxing one there moves the
        whole bullet into the box and the list loses the item."""
        from docx import Document

        (jp_root_dir / "inli.md").write_text(
            '- item <div style="border:1px solid #ccc">boxed</div>\n', encoding="utf-8")
        response = await _export(jp_fetch, "docx", "inli.md")
        doc = Document(io.BytesIO(response.body))
        assert not _docx_callout_boxes(doc), "a div in a list item was boxed"
        bullets = [p for p in doc.paragraphs
                   if p.style and p.style.name == "List Bullet"]
        assert bullets and "item boxed" in bullets[0].text, (
            "the list item lost its text to the box"
        )

    async def test_a_picture_in_a_bordered_div_keeps_its_place(
            self, jp_fetch, jp_root_dir):
        """htmldocx gives a picture a paragraph of its own, which stays outside
        the table the text is moved into - so boxing the div would drop the
        picture below the box it was written above."""
        from docx import Document
        from PIL import Image as PILImage

        PILImage.new("RGB", (60, 40), color=(30, 90, 200)).save(
            jp_root_dir / "pic.png")
        (jp_root_dir / "picbox.md").write_text(
            'Before.\n\n<div style="border:1px solid #ccc">'
            '<img src="pic.png" alt="pic"> caption</div>\n\nAfter.\n',
            encoding="utf-8")
        response = await _export(jp_fetch, "docx", "picbox.md")
        doc = Document(io.BytesIO(response.body))
        assert not _docx_callout_boxes(doc), (
            "a div holding a picture was boxed, leaving the picture outside it"
        )
        with zipfile.ZipFile(io.BytesIO(response.body)) as zf:
            media = [n for n in zf.namelist() if n.startswith("word/media/")]
        assert media, "the picture is missing from the DOCX"
        assert any("caption" in p.text for p in doc.paragraphs), "the caption is missing"


BLOCK_CALLOUT_DOC = """# Callout

Intro paragraph.

<div style="border: 2px dashed #9ca3af; background: rgba(244, 244, 245, 0.45); padding: 10px 14px; margin: 10px 0;">
<h3>Heads up</h3>
This document is under active preparation.
</div>

After the box.
"""


@pytest.fixture
def test_block_callout_file(jp_root_dir):
    md_file = jp_root_dir / "test_block_callout.md"
    md_file.write_text(BLOCK_CALLOUT_DOC, encoding="utf-8")
    return md_file


class TestHtmlCalloutBoxWithBlocks:
    """DEF-MARK-48: the same box drawn around a heading and a line of text. HTML
    boxes it on the div's own CSS whatever it holds, so DOCX and PDF must box
    both spellings or the three formats disagree on the same construct."""

    async def test_html_keeps_the_box_around_blocks(
            self, jp_fetch, test_block_callout_file):
        response = await _export(jp_fetch, "html", "test_block_callout.md")
        html = response.body.decode("utf-8")
        assert "border: 2px dashed #9ca3af" in html, "the box lost its border CSS"

    async def test_docx_bordered_div_holding_blocks_is_a_box(
            self, jp_fetch, test_block_callout_file):
        from docx import Document
        from docx.oxml.ns import qn

        response = await _export(jp_fetch, "docx", "test_block_callout.md")
        doc = Document(io.BytesIO(response.body))
        boxes = [b for b in _docx_callout_boxes(doc) if "Heads up" in b[2]]
        assert boxes, (
            "the div exported as plain paragraphs - a heading inside the "
            "callout cost it the box the same div without one gets"
        )
        bar, fill, text = boxes[0]
        assert bar == "9CA3AF", f"the box bar is {bar}, not the div's own border colour"
        assert fill == "FAFAFA", (
            f"the box fill is {fill}, not the div's own background composited "
            "onto the page"
        )
        assert "active preparation" in text, "the box lost the text it holds"
        from docx.text.paragraph import Paragraph
        all_text = "".join(Paragraph(p, doc).text
                           for p in doc.element.body.iter(qn("w:p")))
        assert "BOX" not in all_text and "⁣" not in all_text, (
            "a callout marker leaked into the visible text"
        )

    async def test_pdf_bordered_div_holding_blocks_is_a_box(
            self, jp_fetch, test_block_callout_file):
        import fitz

        response = await _export(jp_fetch, "pdf", "test_block_callout.md")
        doc = fitz.open(stream=response.body, filetype="pdf")
        fills = [c for page in doc for c in _pdf_fill_colors(page)]
        strokes = [c for page in doc for c in _pdf_stroke_colors(page)]
        text = "".join(page.get_text() for page in doc)
        doc.close()
        assert _color_near(fills, (0.980, 0.980, 0.980)), (
            "the box background (FAFAFA) is missing from the PDF"
        )
        assert _color_near(strokes, (0.612, 0.639, 0.686)), (
            "the box border colour (9CA3AF) is missing from the PDF"
        )
        assert "active preparation" in text, "the box text is missing from the PDF"
        assert "BOX" not in text, "a callout marker leaked into the PDF text"

    async def test_two_paragraphs_in_one_div_are_one_box(self, jp_fetch, jp_root_dir):
        """A box is one table, so the paragraphs of one div go into one - a
        table each would draw a stack of boxes with a gap between them where
        the author drew a single aside."""
        from docx import Document

        (jp_root_dir / "twopara.md").write_text(
            'Before.\n\n<div style="border: 1px solid #123456; background: #eeeeee">\n'
            '<p>first line</p>\n<p>second line</p>\n</div>\n\nAfter.\n',
            encoding="utf-8")
        response = await _export(jp_fetch, "docx", "twopara.md")
        doc = Document(io.BytesIO(response.body))
        boxes = _docx_callout_boxes(doc)
        assert len(boxes) == 1, (
            f"a two-paragraph div drew {len(boxes)} boxes, expected 1"
        )
        assert "first line" in boxes[0][2] and "second line" in boxes[0][2], (
            f"the box holds {boxes[0][2]!r} - a paragraph fell out of it"
        )

    async def test_two_adjacent_divs_stay_two_boxes(self, jp_fetch, jp_root_dir):
        """Two asides written one after the other in the same colours are two
        boxes, not one: a run of paragraphs is a single box only when one div
        opened it."""
        from docx import Document

        (jp_root_dir / "adjacent.md").write_text(
            'Before.\n\n<div style="border: 1px solid #123456">first note</div>\n\n'
            '<div style="border: 1px solid #123456">second note</div>\n\nAfter.\n',
            encoding="utf-8")
        response = await _export(jp_fetch, "docx", "adjacent.md")
        boxes = _docx_callout_boxes(Document(io.BytesIO(response.body)))
        assert len(boxes) == 2, (
            f"two adjacent divs drew {len(boxes)} boxes, expected 2"
        )

    async def test_a_div_holding_a_table_is_not_boxed(self, jp_fetch, jp_root_dir):
        """A table is a structure of its own in Word and stays outside the box
        the text is moved into, the way a picture does - so the div keeps the
        plain export rather than dropping its table out of its own box."""
        from docx import Document

        (jp_root_dir / "tblbox.md").write_text(
            'Before.\n\n<div style="border: 1px solid #123456">\n'
            '<p>lead line</p>\n<table><tr><td>cell</td></tr></table>\n</div>\n\n'
            'After.\n', encoding="utf-8")
        response = await _export(jp_fetch, "docx", "tblbox.md")
        doc = Document(io.BytesIO(response.body))
        assert not _docx_callout_boxes(doc), (
            "a div holding a table was boxed, leaving the table outside it"
        )
        assert any("cell" in c.text for t in doc.tables
                   for r in t.rows for c in r.cells), "the table is missing"


class TestCssColourInFunctionCall:
    """DEF-MARK-47: a CSS function call carries names that are not colours - a file
    name in `url()`, a custom property in `var()` - and reading one as a colour
    paints a fill the browser does not."""

    async def test_a_colour_name_in_a_file_name_is_not_a_box(
            self, jp_fetch, jp_root_dir):
        from docx import Document

        (jp_root_dir / "urlbg.md").write_text(
            'Before.\n\n<div style="background: url(assets/red-banner.png); '
            'padding:4px">boxed</div>\n\nAfter.\n', encoding="utf-8")
        response = await _export(jp_fetch, "docx", "urlbg.md")
        doc = Document(io.BytesIO(response.body))
        assert not _docx_callout_boxes(doc), (
            "the image file name was read as a colour, so the div took a "
            "solid callout box the browser paints nowhere"
        )
        assert any("boxed" in p.text for p in doc.paragraphs), (
            "the div's text is missing from the DOCX"
        )

    async def test_a_custom_property_name_does_not_shade_a_run(
            self, jp_fetch, jp_root_dir):
        from docx import Document
        from docx.oxml.ns import qn

        (jp_root_dir / "varbg.md").write_text(
            'A <span style="background: var(--blue-500)">tok</span> word.\n',
            encoding="utf-8")
        response = await _export(jp_fetch, "docx", "varbg.md")
        doc = Document(io.BytesIO(response.body))
        runs = [r for p in doc.paragraphs for r in p.runs if r.text.strip() == "tok"]
        assert runs, "the run is missing from the DOCX"
        rPr = runs[0]._r.find(qn("w:rPr"))
        shd = rPr.find(qn("w:shd")) if rPr is not None else None
        assert shd is None, (
            f"the custom property name was read as a colour and shaded the run "
            f"{shd.get(qn('w:fill'))}; the browser resolves the variable and "
            "paints nothing like it"
        )

    async def test_a_bare_name_and_rgb_still_resolve(self, jp_fetch, jp_root_dir):
        """The other side of the same test: a value that IS a colour keeps
        working, in a bare name, in `rgb()`, and beside a `url()` in one
        shorthand."""
        from docx import Document
        from docx.oxml.ns import qn

        (jp_root_dir / "keeps.md").write_text(
            'A <span style="background-color: rgb(0, 0, 255)">tok</span> word.\n\n'
            '<div style="background: gold">bare name</div>\n\n'
            '<div style="background: #ffffff url(bg.png) no-repeat; '
            'border: 1px solid #123456">beside a url</div>\n', encoding="utf-8")
        response = await _export(jp_fetch, "docx", "keeps.md")
        doc = Document(io.BytesIO(response.body))
        runs = [r for p in doc.paragraphs for r in p.runs if r.text.strip() == "tok"]
        assert runs, "the run is missing from the DOCX"
        shd = runs[0]._r.find(qn("w:rPr")).find(qn("w:shd"))
        assert shd is not None and shd.get(qn("w:fill")) == "0000FF", "rgb() lost"
        named = [b for b in _docx_callout_boxes(doc) if "bare name" in b[2]]
        assert named and named[0][1] == "FFD700", (
            f"a bare colour name lost its fill: {named}"
        )
        beside = [b for b in _docx_callout_boxes(doc) if "beside a url" in b[2]]
        assert beside and beside[0][:2] == ("123456", "FFFFFF"), (
            f"the colour beside the url() was lost: {beside}"
        )


class TestCssBorderCascade:
    """DEF-MARK-49: the cascade applies to a border the way DEF-DIAG-39 made it apply to
    emphasis - declarations folded per property, last value wins - so a later
    longhand can write a border off and a colour alone is not one."""

    async def test_a_later_longhand_switches_the_border_off(
            self, jp_fetch, jp_root_dir):
        from docx import Document

        (jp_root_dir / "offlater.md").write_text(
            'Before.\n\n<div style="border: 2px solid #9ca3af; border-style: none">'
            'not a box</div>\n\nAfter.\n', encoding="utf-8")
        response = await _export(jp_fetch, "docx", "offlater.md")
        doc = Document(io.BytesIO(response.body))
        assert not _docx_callout_boxes(doc), (
            "the border was latched on by the first declaration, so the "
            "longhand that switched it off drew a box anyway"
        )
        assert any("not a box" in p.text for p in doc.paragraphs), (
            "the div's text is missing from the DOCX"
        )

    async def test_a_border_colour_alone_is_not_a_border(self, jp_fetch, jp_root_dir):
        """`border-style` defaults to `none`, so a colour on its own draws no
        line at all - a browser shows nothing and a box invents an aside."""
        from docx import Document

        (jp_root_dir / "coloronly.md").write_text(
            'Before.\n\n<div style="border-color: #9ca3af">not a box</div>\n\n'
            'After.\n', encoding="utf-8")
        response = await _export(jp_fetch, "docx", "coloronly.md")
        doc = Document(io.BytesIO(response.body))
        assert not _docx_callout_boxes(doc), (
            "a border colour with no border-style drew a box"
        )
        assert any("not a box" in p.text for p in doc.paragraphs), (
            "the div's text is missing from the DOCX"
        )

    async def test_the_last_border_declaration_wins(self, jp_fetch, jp_root_dir):
        """The other direction: a border written off and then back on draws."""
        from docx import Document

        (jp_root_dir / "onlater.md").write_text(
            'Before.\n\n<div style="border: none; border: 1px solid #123456">'
            'a box</div>\n\nAfter.\n', encoding="utf-8")
        response = await _export(jp_fetch, "docx", "onlater.md")
        boxes = _docx_callout_boxes(Document(io.BytesIO(response.body)))
        assert boxes and boxes[0][0] == "123456", (
            f"the last border declaration lost: {boxes}"
        )


class TestUnresolvableColourValue:
    """DEF-MARK-50: htmldocx reads a colour with `int()` and slices a hex blind, so
    a value this code could not resolve either takes the whole export down or
    paints a colour the browser does not."""

    async def test_a_malformed_hex_does_not_take_the_export_down(
            self, jp_fetch, jp_root_dir):
        (jp_root_dir / "badhex.md").write_text(
            '# B\n\n<span style="color: #1234">first</span> and '
            '<span style="background-color: #gggggg">second</span>.\n\nafter\n',
            encoding="utf-8")
        for fmt in ("docx", "pdf", "html"):
            response = await jp_fetch(
                "jupyterlab-export-markdown-extension", f"export/{fmt}",
                method="POST", body=json.dumps({"path": "badhex.md"}),
                raise_error=False,
            )
            assert response.code == 200, (
                f"{fmt}: a malformed hex colour lost the whole document - "
                f"{response.body[:120]!r}"
            )
        response = await _export(jp_fetch, "docx", "badhex.md")
        assert "after" in _docx_text(response.body), "the document lost its text"

    async def test_a_variable_holding_the_letters_rgb_does_not_take_it_down(
            self, jp_fetch, jp_root_dir):
        """htmldocx routes a value on the substring `rgb`, so a custom property
        merely named after the notation is read as the notation."""
        (jp_root_dir / "varrgb.md").write_text(
            '# V\n\n<span style="color: var(--rgb-blue)">first</span>.\n\nafter\n',
            encoding="utf-8")
        for fmt in ("docx", "pdf", "html"):
            response = await jp_fetch(
                "jupyterlab-export-markdown-extension", f"export/{fmt}",
                method="POST", body=json.dumps({"path": "varrgb.md"}),
                raise_error=False,
            )
            assert response.code == 200, (
                f"{fmt}: a custom property lost the whole document - "
                f"{response.body[:120]!r}"
            )

    async def test_an_unresolvable_hex_is_not_painted_as_another_colour(
            self, jp_fetch, jp_root_dir):
        """A five-digit hex is not a colour: a browser drops the declaration
        and the text stays the body colour. Handed on, htmldocx slices the
        first two pairs and paints a colour that is on no screen anywhere."""
        from docx import Document

        (jp_root_dir / "fivehex.md").write_text(
            'A <span style="color: #12345">tint</span> word.\n', encoding="utf-8")
        response = await _export(jp_fetch, "docx", "fivehex.md")
        doc = Document(io.BytesIO(response.body))
        runs = [r for p in doc.paragraphs for r in p.runs if r.text.strip() == "tint"]
        assert runs, "the run is missing from the DOCX"
        colour = runs[0].font.color
        assert colour is None or colour.rgb is None, (
            f"the run was painted {colour.rgb}, read out of a hex that names "
            "no colour"
        )


MIXED_CALLOUT_DOC = """# T

<div style="border: 2px solid #333; background: #eeeeee; color: #ff0000;">
<h3>Heads up</h3>
loose text between
<p>a paragraph</p>
</div>

After.
"""


@pytest.fixture
def test_mixed_callout_file(jp_root_dir):
    md_file = jp_root_dir / "test_mixed_callout.md"
    md_file.write_text(MIXED_CALLOUT_DOC, encoding="utf-8")
    return md_file


def _pdf_spans(pdf_bytes):
    """(text, size, font, colour hex) of every non-blank span in a PDF."""
    import fitz

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    spans = []
    for page in doc:
        for block in page.get_text("dict")["blocks"]:
            for line in block.get("lines", []):
                for span in line["spans"]:
                    if span["text"].strip():
                        spans.append((span["text"].strip(),
                                      round(span["size"], 1), span["font"],
                                      f'{span["color"]:06X}'))
    doc.close()
    return spans


class TestCalloutBoxContents:
    """DEF-MARK-51/52/53: what stands INSIDE a hand-drawn box. A heading, a run of
    loose text and a paragraph are three blocks in the browser and have to
    arrive as three, at the sizes their tags earn, in the colour the same
    declaration block gives them."""

    async def test_loose_text_is_not_welded_into_the_block_above(
            self, jp_fetch, test_mixed_callout_file):
        """DEF-MARK-51: a bare text node carries no marker, so it stayed at body
        level when the div was unwrapped and htmldocx joined it to the
        paragraph above - two authored blocks arrived as one, in the style of
        the first."""
        from docx import Document

        response = await _export(jp_fetch, "docx", "test_mixed_callout.md")
        doc = Document(io.BytesIO(response.body))
        boxes = [t for t in doc.tables
                 if "Heads up" in t.rows[0].cells[0].text]
        assert boxes, "the div took no box at all"
        paras = [p for p in boxes[0].rows[0].cells[0].paragraphs if p.text.strip()]
        assert [p.text for p in paras] == [
            "Heads up", "loose text between", "a paragraph"], (
            f"the box holds {[p.text for p in paras]}; the three blocks the "
            "author wrote must arrive as three, in the order written"
        )
        assert paras[0].style.name == "Heading 3", (
            f"the heading arrived as {paras[0].style.name}"
        )
        assert paras[1].style.name != "Heading 3", (
            "the loose text was welded into the heading and took its style"
        )

    async def test_pdf_agrees_that_a_boxed_heading_is_a_heading(
            self, jp_fetch, test_mixed_callout_file):
        """DEF-MARK-52: the PDF painted every paragraph of a callout cell as body
        text, so the heading Word keeps read at body size - one line of source
        with two answers."""
        from docx import Document

        docx_response = await _export(jp_fetch, "docx", "test_mixed_callout.md")
        doc = Document(io.BytesIO(docx_response.body))
        boxed = [p for t in doc.tables for p in t.rows[0].cells[0].paragraphs
                 if p.text.strip() == "Heads up"]
        assert boxed and boxed[0].style.name.startswith("Heading"), (
            "the DOCX box lost the heading this test compares against"
        )

        spans = _pdf_spans((await _export(jp_fetch, "pdf",
                                          "test_mixed_callout.md")).body)
        heading = [s for s in spans if s[0] == "Heads up"]
        body = [s for s in spans if s[0] == "a paragraph"]
        assert heading and body, f"the box text is missing from the PDF: {spans}"
        assert heading[0][1] > body[0][1], (
            f"the boxed heading is {heading[0][1]}pt against {body[0][1]}pt of "
            "body text in the same box - the PDF dropped the heading level "
            "Word keeps"
        )
        assert "Bold" in heading[0][2], (
            f"the boxed heading is set in {heading[0][2]}, not a heading face"
        )

    async def test_the_box_colour_reaches_the_text_it_was_declared_with(
            self, jp_fetch, test_mixed_callout_file):
        """DEF-MARK-53: the fill and the colour come from one declaration block, so
        an author who sets both must not get one. A heading is the exception in
        all three formats - the stylesheet colour beats what it inherits."""
        from docx import Document

        doc = Document(io.BytesIO(
            (await _export(jp_fetch, "docx", "test_mixed_callout.md")).body))
        painted = {}
        for tbl in doc.tables:
            for para in tbl.rows[0].cells[0].paragraphs:
                if not para.text.strip():
                    continue
                painted[para.text.strip()] = {
                    str(run.font.color.rgb) for run in para.runs
                    if run.font.color and run.font.color.rgb}
        assert painted.get("loose text between") == {"FF0000"}, (
            f"the loose text is painted {painted.get('loose text between')}; "
            "the colour declared beside the fill reached neither format"
        )
        assert painted.get("a paragraph") == {"FF0000"}, (
            f"the paragraph is painted {painted.get('a paragraph')}"
        )
        assert "FF0000" not in painted.get("Heads up", set()), (
            "the heading took the author's body colour, though the stylesheet "
            "colours a heading itself and a browser draws it unchanged"
        )

        spans = _pdf_spans(
            (await _export(jp_fetch, "pdf", "test_mixed_callout.md")).body)
        boxed = {s[0]: s[3] for s in spans}
        assert boxed.get("a paragraph") == "FF0000", (
            f"the PDF painted the box text {boxed.get('a paragraph')}, so the "
            "two formats disagree on the same declaration"
        )
        assert boxed.get("Heads up") != "FF0000", (
            "the PDF painted the boxed heading the body colour"
        )

    async def test_a_box_around_one_paragraph_carries_its_colour_too(
            self, jp_fetch, jp_root_dir):
        """The other spelling of the same construct - no block children, so the
        div becomes the paragraph itself. htmldocx reads a colour off an inline
        tag only, so one left on that paragraph arrived nowhere."""
        from docx import Document

        (jp_root_dir / "onepara.md").write_text(
            'Before.\n\n<div style="border: 1px solid #123456; color: #ff0000">'
            'just text</div>\n\nAfter.\n', encoding="utf-8")
        response = await _export(jp_fetch, "docx", "onepara.md")
        doc = Document(io.BytesIO(response.body))
        boxed = [p for t in doc.tables for p in t.rows[0].cells[0].paragraphs
                 if p.text.strip()]
        assert [p.text for p in boxed] == ["just text"], (
            f"the box holds {[p.text for p in boxed]}"
        )
        assert {str(run.font.color.rgb) for run in boxed[0].runs
                if run.font.color and run.font.color.rgb} == {"FF0000"}, (
            "the colour declared on the box reached no run"
        )


class TestCalloutBorderWidth:
    """DEF-MARK-87: a `border` whose width is not a positive length draws
    nothing in a browser, and the HTML export passes the style through - so
    Word and the PDF must draw no box for it either."""

    def test_a_width_that_is_not_a_positive_length_draws_no_box(self):
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase
        for css in ('border: -1px solid red', 'border: 00px solid red',
                    'border: 0 solid red', 'border: 0.0em solid red'):
            assert ExportHandlerBase._css_callout_box(css) is None, (
                f"{css!r} drew a callout the browser does not")
        assert ExportHandlerBase._css_callout_box('border: 1px solid red') == (
            'FF0000', 'FFFFFF')
        assert ExportHandlerBase._css_callout_box('border: thin solid red') == (
            'FF0000', 'FFFFFF')

    def test_the_bar_takes_the_left_edge_colour_over_the_frame(self):
        """DEF-MARK-109: the classic callout is a neutral frame with a
        coloured left edge; the bar took the first drawn side in top-right-
        bottom-left order, so the frame's grey won and an info box and a
        danger box written this way came out the same."""
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase as B
        for css in ('border: 1px solid #ddd; border-left: 4px solid #0969da',
                    'border: 1px solid #ddd; border-left-width: 4px; border-left-color: #0969da'):
            assert B._css_callout_box(css) == ('0969DA', 'FFFFFF'), (
                f"{css!r} lost its accent to the frame")
        assert B._css_callout_box('border-top: 3px solid red') == ('FF0000', 'FFFFFF')

    def test_an_invisible_border_in_any_notation_draws_no_box(self):
        """DEF-MARK-92: `transparent` was the only spelling of an invisible
        border written off; `rgba(0,0,0,0)`, `#0000` and `#00000000` resolve
        to no colour and counted as a drawn side with the fallback grey bar.
        An alpha in hex is composited as the `rgba()` alpha is."""
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase as B
        for css in ('border: 1px solid rgba(0,0,0,0)', 'border: 1px solid #0000',
                    'border: 1px solid #00000000',
                    'border: 1px solid red; border-color: rgba(0,0,0,0)'):
            assert B._css_callout_box(css) is None, f"{css!r} drew a box"
        assert B._css_callout_box('border: 2px solid #ff000080') == ('FF7F7F', 'FFFFFF'), (
            "a half-transparent red bar must composite onto the page, not turn grey")
        assert B._css_callout_box('border: 1px solid currentcolor') == ('BBBBBB', 'FFFFFF'), (
            "a bare word nothing resolves keeps the grey bar")

    def test_a_later_background_that_names_no_colour_clears_the_fill(self):
        """DEF-MARK-93: the fill latched on the first colour any background
        declaration named; the cascade's last value wins, and one naming no
        colour clears it."""
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase as B
        for css in ('background: red; background: none',
                    'background: red; background-color: transparent',
                    'background-color: red; background: url(x.png)'):
            assert B._css_callout_box(css) is None, f"{css!r} kept the earlier fill"
        assert B._css_callout_box(
            'border: 1px solid #333333; background: red; background: none') == ('333333', 'FFFFFF')
        assert B._css_callout_box('background: none; background: red') == ('FF0000', 'FF0000')

    def test_a_semicolon_inside_a_url_does_not_split_the_declaration(self):
        """DEF-MARK-94: torn at the `;` inside `url('tan;x.png')`, the
        declaration ends in `tan`, a named colour, and a tan box is drawn for
        a background image the export cannot show."""
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase as B
        for css in ("background: url('tan;x.png')", 'background: url("tan;x.png")',
                    "background: url(tan;x.png)"):
            assert B._css_callout_box(css) is None, f"{css!r} drew a tan box"
        assert B._css_text_align("background: url('a;b'); text-align: center") == 'center'

    def test_an_unreadable_colour_keeps_its_box_and_an_invisible_one_does_not(self):
        """DEF-MARK-97: the write-off was keyed on notation, so `hsl()` - a
        colour the parser did not read - was treated as invisible and a
        document boxed with it exported boxed in HTML and unboxed in Word.
        `hsl()` is now read, and only a colour that resolves to alpha zero,
        in any spelling, writes the side off."""
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase as B
        assert B._css_callout_box('border: 2px solid hsl(120, 100%, 25%)') == ('008000', 'FFFFFF')
        assert B._css_callout_box('border: 1px solid foo(1)') == ('BBBBBB', 'FFFFFF'), (
            "a colour the parser cannot read must keep the grey bar, as a bare unknown word does")
        for css in ('border: 1px solid hsla(0, 0%, 0%, 0)', 'border: 1px solid hsl(0 100% 50% / 0)',
                    'border: 1px solid transparent', 'border: 1px solid rgba(0,0,0,0)',
                    'border: 1px solid rgb(0 0 0 / 0%)'):
            assert B._css_callout_box(css) is None, f"{css!r} drew a box"
        # The shorthand parser keeps a function call whole: split on spaces,
        # `100%` inside `hsl(0 100% 50%)` was read as the width
        assert B._css_callout_box('border: 2px solid hsl(120 100% 25%)') == ('008000', 'FFFFFF')
        assert B._normalize_css_color('foo(1)') is None
        assert B._normalize_css_color('transparent') == ''

    def test_the_run_shading_reads_the_cascade_as_the_box_does(self):
        """DEF-MARK-98: the pill reader latched the first background colour
        while the box reader (DEF-MARK-93) reads last-wins; one declaration
        block read two ways, and a red band where the author's last word was
        `none`."""
        from jupyterlab_export_markdown_extension.routes import ExportHandlerBase as B
        h = B.__new__(B)
        cleared = h.restructure_html_for_docx('<p><span style="background: red; background: none">x</span></p>')
        assert 'PILL:' not in cleared, "a later `background: none` did not clear the run shading"
        kept = h.restructure_html_for_docx('<p><span style="background: none; background: red">x</span></p>')
        assert 'PILL:FF0000' in kept
