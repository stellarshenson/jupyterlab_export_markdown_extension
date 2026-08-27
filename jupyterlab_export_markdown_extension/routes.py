"""
Route handlers for the markdown export extension.

Provides API endpoints for exporting markdown files to PDF, DOCX, and HTML formats
using pure Python libraries (no system dependencies).
"""

from __future__ import annotations

import asyncio
import json
import os
import base64
import ipaddress
import re
import io
import socket
import tempfile
import time
import unicodedata
from html import unescape
from pathlib import Path
from urllib.parse import quote, unquote, urlparse

from jupyter_server.base.handlers import APIHandler
from jupyter_server.utils import url_path_join
import tornado.web

# Reportlab imports for DOCX-to-PDF conversion
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage,
        Preformatted, XPreformatted
    )
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


class _MermaidBlock:
    """One fenced mermaid block, shaped like the `re.Match` it replaced so
    callers read `.start()`, `.end()`, `.group(0)` and `.group(1)` as before."""

    __slots__ = ('_content', '_start', '_end', '_source')

    def __init__(self, content: str, start: int, end: int, source: str):
        self._content, self._start, self._end, self._source = (
            content, start, end, source)

    def start(self) -> int:
        return self._start

    def end(self) -> int:
        return self._end

    def group(self, n: int = 0) -> str:
        if n == 0:
            return self._content[self._start:self._end]
        if n == 1:
            return self._source
        raise IndexError('no such group')  # never answer a question not asked


class ChromiumUnavailableError(RuntimeError):
    """Raised when Playwright cannot launch Chromium (binary or sys-libs missing).

    Carries the original error message so the handler can return install
    guidance to the frontend.
    """


#: The one command that fixes a missing Chromium, quoted rather than reworded
#: wherever this module says it. `cli.py` and `src/index.ts` restate the same
#: string because neither can import from here - the installer CLI must run
#: without this module's heavy imports, and the frontend is another language -
#: so those two are the copies to keep in step if the entry point is renamed.
CHROMIUM_INSTALL_COMMAND = 'jupyterlab-export-markdown-extension install'

#: Why a diagram kept its source instead of becoming a picture, and what the
#: caller can do about it. Both the renderer (which detects the failures) and
#: the handler (which reports them) name codes from this one table, so a code
#: cannot be emitted that has no message. The message IS the whole remedy - a
#: caller reading `X-Export-Warnings` off a binary response has nowhere else
#: to look.
MERMAID_WARNINGS = {
    'chromium-unavailable': (
        'Mermaid diagrams were not rendered because Chromium could not '
        f'launch, so their source was kept. Run: {CHROMIUM_INSTALL_COMMAND}'
    ),
    'bundle-missing': (
        'Mermaid diagrams were not rendered because the bundled mermaid '
        'renderer is missing or could not be loaded, so their source was '
        'kept. Reinstall jupyterlab_export_markdown_extension'
    ),
    'syntax-error': (
        'Mermaid could not draw these diagrams, so their source was kept. '
        'Check the diagram syntax'
    ),
    'layout-unsupported': (
        'These diagrams ask for a layout engine the server does not carry '
        '(mermaid-layout-elk is not bundled), so their source was kept. They '
        'still render in JupyterLab; export them from the UI, or use the '
        'default layout'
    ),
    'render-timeout': (
        'Mermaid did not finish drawing this diagram in time, so its source '
        'was kept. Simplify it or split it up'
    ),
    'skipped': (
        'Rendering stopped at the diagram that did not finish, so these were '
        'never attempted and kept their source. Fix the one reported under '
        'render-timeout'
    ),
    'budget-exhausted': (
        'Rendering ran past the time this export allows for diagrams, so '
        'these were never attempted and kept their source'
    ),
    'rasterize-failed': (
        'These diagrams were drawn but could not be converted to an image, '
        'so their source was kept'
    ),
    'render-failed': (
        'The mermaid renderer failed, so these diagrams kept their source. '
        'Check the Jupyter server log'
    ),
}


#: A line break as an author may write it - any case, any attributes, closed or
#: not. `\b` would accept a custom element like `<br-spacer>`, whose break is
#: not a break at all.
BREAK_TAG_RE = re.compile(r'<br(?:\s[^>]*?)?/?>', re.IGNORECASE)

#: The same tag, at the very end of a string. Used where a break is about to be
#: appended and one the author already wrote has to count towards it.
TRAILING_BREAK_RE = re.compile(BREAK_TAG_RE.pattern + r'\s*\Z', re.IGNORECASE)

#: Built once - the rule is stateless, and `markdown_to_html` runs per export
_MANUAL_BREAK_RULE = None


def manual_break_aware_nl2br():
    """The `nl2br` rule, minus the break it would add after a manual one.

    `nl2br` turns every newline into a break. A line the author ended with
    `<br>` - the idiom for a question above its answer - therefore gets two,
    and the resulting blank line makes the gap inside the pair match the gap
    between pairs, so the grouping reads backwards. The author asked for one
    break and should get one.

    Deciding that needs to happen where provenance still exists, which is the
    inline stage. By the time the HTML is serialized a break the author typed
    and a break `nl2br` generated are the same six characters, and so is the
    one core Markdown generates for the two-trailing-spaces idiom - collapsing
    on shape there deletes authored breaks. Here the author's tag is still a
    stashed node, and the two-space break is consumed by the `linebreak`
    pattern (priority 100) long before this one (priority 5) is reached, so
    `text<br>` + two spaces keeps both of its breaks.

    Everything the resolution touches is guarded: if Markdown's internals move,
    this cannot tell a manual break from a generated one and falls back to
    emitting the break, which is exactly `nl2br`'s own behaviour.
    """
    global _MANUAL_BREAK_RULE
    if _MANUAL_BREAK_RULE is not None:
        return _MANUAL_BREAK_RULE

    from markdown.extensions import Extension
    from markdown.inlinepatterns import InlineProcessor
    from markdown import util
    from xml.etree import ElementTree

    # The author's tag as it sits in the text at this point: an inline
    # placeholder, at the very end of what precedes the newline
    trailing_node = re.compile(
        re.escape(util.INLINE_PLACEHOLDER).replace(re.escape('%s'), r'(\d+)')
        + r'\Z')

    class NewlineToBreak(InlineProcessor):
        def handleMatch(self, m, data):
            if self._after_manual_break(data[:m.start(0)]):
                return None, None, None
            return ElementTree.Element('br'), m.start(0), m.end(0)

        def _after_manual_break(self, before):
            """True when a break tag the author typed ends ``before``.

            Trailing blanks are skipped: one space after `<br>` is invisible
            in an editor and must not decide how the document renders. Two
            spaces never reach here - `linebreak` has already claimed them. A
            tab cannot either - `expandtabs` runs before inline processing.

            A break nested inside inline markup (`**Q<br>**`) is NOT detected:
            the trailing node is then the emphasis Element, not the break, and
            the pair renders with a blank line. Known limitation, pinned by
            `test_break_inside_emphasis_is_a_known_limitation`.
            """
            node = trailing_node.search(before.rstrip(' \xa0'))
            if node is None:
                return False
            try:
                inline = self.md.treeprocessors['inline'].stashed_nodes
                stashed = inline.get(node.group(1))
                if not isinstance(stashed, str):
                    return False       # a real element, so not raw HTML
                held = util.HTML_PLACEHOLDER_RE.fullmatch(stashed)
                if held is None:
                    return False
                raw = self.md.htmlStash.rawHtmlBlocks[int(held.group(1))]
            except (AttributeError, IndexError, KeyError, TypeError):
                return False           # cannot tell - behave like plain nl2br
            return isinstance(raw, str) and bool(BREAK_TAG_RE.fullmatch(raw.strip()))

    class ManualBreakAwareNl2Br(Extension):
        def extendMarkdown(self, md):
            # Same slot and priority nl2br itself uses
            md.inlinePatterns.register(NewlineToBreak(r'\n', md), 'nl', 5)

    _MANUAL_BREAK_RULE = ManualBreakAwareNl2Br()
    return _MANUAL_BREAK_RULE


class PlaywrightSvgRenderer:
    """Owns a headless Chromium: rasterizes SVG to PNG, and draws mermaid.

    Reuses one browser process across multiple render() calls within an
    `async with` block. Uses a real browser engine, so CSS classes,
    @font-face, gradients, filters, and @media (prefers-color-scheme)
    all behave the way they do in the user's browser.

    The output PNG is `width` pixels wide; height follows the SVG's
    viewBox aspect ratio. Chromium anti-aliases natively, so no
    supersampling is needed - it renders straight to the target size.
    Pass supersample > 1 only if a specific SVG needs extra smoothing.

    `render_mermaid` uses the same browser to turn mermaid source into a
    diagram, so the class also carries the bundle path, the mermaid options
    and the render budget.
    """

    def __init__(self, color_scheme: str = 'light', offline: bool = False):
        if color_scheme not in ('light', 'dark'):
            color_scheme = 'light'
        self.color_scheme = color_scheme
        self.offline = offline
        self._pw = None
        self._browser = None

    async def __aenter__(self):
        try:
            from playwright.async_api import async_playwright
        except ImportError as e:
            raise ChromiumUnavailableError(
                f'playwright not installed: {e}'
            ) from e
        try:
            self._pw = await async_playwright().start()
            # An offline browser cannot resolve a name at all. Route
            # interception stops requests, but Chromium resolves a
            # `<link rel="dns-prefetch">` without emitting one, and a mermaid
            # label can carry that tag - a DNS leak out of the server from
            # someone else's markdown. Used for content this server generated.
            self._browser = await self._pw.chromium.launch(
                headless=True,
                args=(['--host-resolver-rules=MAP * ~NOTFOUND',
                       '--disable-features=NetworkPrediction']
                      if self.offline else []),
            )
        except Exception as e:
            # Most common failure is missing system libs (libnspr4, libnss3
            # etc.) or Chromium binary not downloaded yet.
            await self._safe_stop()
            raise ChromiumUnavailableError(str(e)) from e
        return self

    async def __aexit__(self, exc_type, exc, tb):
        await self._safe_stop()

    async def _safe_stop(self):
        try:
            if self._browser is not None:
                # Bounded: the one path that gets here with a renderer still
                # spinning is a mermaid layout timeout, and an unbounded close
                # would hand the request back to that same wedge.
                await asyncio.wait_for(self._browser.close(),
                                       timeout=self.CONTEXT_CLOSE_TIMEOUT_S)
        except Exception:
            pass
        try:
            if self._pw is not None:
                await asyncio.wait_for(self._pw.stop(),
                                       timeout=self.CONTEXT_CLOSE_TIMEOUT_S)
        except Exception:
            pass
        self._browser = None
        self._pw = None

    # Chromium refuses a viewport, and silently fails a texture, past roughly
    # 16384px. Held a little under it for the compositor's own overhead.
    MAX_RASTER_PX = 16000

    #: The mermaid UMD bundle, copied out of node_modules into the wheel by
    #: `jlpm vendor:mermaid`. Loading it from disk keeps a server-side render
    #: offline - no CDN, no network call from the export path.
    MERMAID_JS_PATH = Path(__file__).parent / 'vendor' / 'mermaid.min.js'

    #: What THIS extension's frontend initialises mermaid with before it
    #: captures a diagram (`renderMermaidInTheme` in `src/index.ts`), not what
    #: `@jupyterlab/mermaid` defaults to. The two differ: the frontend sets
    #: `securityLevel: 'loose'`, and under `strict` mermaid turns off HTML
    #: labels, so the same diagram would come out of the API visibly unlike
    #: the one the UI produces - the divergence this pass exists to remove.
    #: Loose costs nothing here because the render has no network (see
    #: `render_mermaid`) and the SVG is never scripted, only screenshotted.
    #: `theme` is per-render.
    MERMAID_INIT_OPTIONS = {
        'startOnLoad': False,
        'securityLevel': 'loose',
        # Not set by `renderMermaidInTheme` itself - `initialize` merges onto
        # the config `@jupyterlab/mermaid`'s manager already applied, so the
        # browser keeps these ceilings. Without them the server falls back to
        # mermaid's own much lower defaults and reports a large generated
        # diagram as a syntax error that JupyterLab renders perfectly well.
        'maxTextSize': 100000,
        'maxEdges': 100000,
    }

    #: Seconds a single diagram gets to lay itself out. Mermaid runs in the
    #: page's one JS thread, so without this a graph whose layout does not
    #: converge holds the export request open with nothing to return.
    MERMAID_RENDER_TIMEOUT_S = 30

    #: Seconds all of a document's diagrams get between them. The per-diagram
    #: timeout does not bound the request: 200 diagrams that each take 4s hold
    #: it open for 13 minutes, long after the client gave up.
    MERMAID_TOTAL_BUDGET_S = 180

    #: Seconds to wait on closing the page. The one path that reaches here
    #: with a wedged renderer is the layout timeout, which is exactly where an
    #: unbounded close would hang.
    CONTEXT_CLOSE_TIMEOUT_S = 15

    @classmethod
    def mermaid_bundle_available(cls) -> bool:
        """Whether this installation can render mermaid at all. Lets a caller
        decide without launching a browser, and without knowing where the
        bundle lives."""
        return cls.MERMAID_JS_PATH.exists()

    @staticmethod
    async def _block_network(ctx) -> None:
        """Refuse every request the page makes.

        A mermaid diagram is self-contained - the bundle is injected from
        disk, the styles are inlined - so nothing here legitimately needs the
        network. Left open, a diagram is an SSRF primitive: mermaid's HTML
        labels survive into a `<foreignObject>`, so a label of
        `A["<img src='http://169.254.169.254/...'>"]` in someone else's
        markdown makes THIS server fetch that URL. Measured before the block:
        three outbound requests from one exported diagram.
        """
        await ctx.route('**/*', lambda route: route.abort())

    #: Mermaid's own words when a diagram asks for a layout engine that was
    #: never registered. Told apart from a syntax error because the remedy is
    #: the opposite: the diagram is fine, this server just cannot draw it.
    _MERMAID_MISSING_LAYOUT_RE = re.compile(
        r'layout\s+(?:algorithm\s+)?\S+\s+is\s+not\s+registered'
        r'|no\s+layout\s+loader|registerLayoutLoaders'
        r'|unknown\s+layout', re.IGNORECASE)

    async def render_mermaid(self, sources: list[str], *, theme: str,
                             png_width: int | None
                             ) -> list[tuple[str | bytes | None, str | None]]:
        """Render mermaid sources, one browser page for all of them.

        Returns one `(result, reason)` pair per source, in order, exactly one
        of them set. `result` is the SVG text, or the rasterized PNG bytes
        when `png_width` is given - rasterizing here rather than leaving it to
        `extract_data_uri_images` reuses this browser, so a document whose only
        images are diagrams starts Chromium once instead of twice. (One that
        also carries the frontend's own SVGs still pays for the second.)
        `reason` is a code from `MERMAID_WARNINGS` naming why the caller must
        keep that diagram's source instead.

        Never raises for a diagram - a failure is always a reason code, and
        even a browser that dies mid-run keeps the diagrams already drawn.
        """
        out: list[tuple[str | bytes | None, str | None]] = []

        def fill_rest(reason: str) -> None:
            out.extend([(None, reason)] * (len(sources) - len(out)))

        ctx = await self._browser.new_context(color_scheme=self.color_scheme)
        try:
            await self._block_network(ctx)
            page = await ctx.new_page()
            await page.set_content(
                "<!DOCTYPE html><html><head><meta charset='utf-8'>"
                "</head><body></body></html>"
            )
            await page.add_script_tag(path=str(self.MERMAID_JS_PATH))
            # A bundle that loaded but did not define the global is a bundle
            # of the wrong kind (an ESM build injected as a classic script).
            # Say so, rather than blame every diagram in the document.
            if await page.evaluate("() => typeof mermaid") == 'undefined':
                fill_rest('bundle-missing')
                return out
            await page.evaluate(
                '(opts) => mermaid.initialize(opts)',
                {**self.MERMAID_INIT_OPTIONS, 'theme': theme},
            )

            deadline = time.monotonic() + self.MERMAID_TOTAL_BUDGET_S
            for i, source in enumerate(sources):
                if time.monotonic() >= deadline:
                    fill_rest('budget-exhausted')
                    break
                # A diagram gets its full timeout only if that much budget is
                # left; nearer the deadline the window shrinks. Remember which
                # limit applied, so a diagram that draws fine but ran out the
                # document's total budget is not blamed for being slow.
                remaining = deadline - time.monotonic()
                budget_limited = remaining < self.MERMAID_RENDER_TIMEOUT_S
                try:
                    svg = await asyncio.wait_for(page.evaluate(
                        """async ({src, id}) => {
                            const {svg} = await mermaid.render(id, src);
                            return svg;
                        }""",
                        {'src': source, 'id': f'export-mermaid-{i}'},
                    ), timeout=max(1.0, min(self.MERMAID_RENDER_TIMEOUT_S,
                                            remaining)))
                except asyncio.TimeoutError:
                    if budget_limited:
                        # The document's total budget ran out, not this diagram
                        # - it never got a full window. Report the true cause.
                        fill_rest('budget-exhausted')
                        break
                    # That layout still holds the page's only JS thread, so
                    # every diagram behind it would time out in turn. Give them
                    # all their source back rather than spend the timeout again
                    # and again on a page that is not coming back - and say
                    # which one stopped the run, since the rest are innocent.
                    out.append((None, 'render-timeout'))
                    fill_rest('skipped')
                    break
                except Exception as e:
                    if page.is_closed() or not self._browser.is_connected():
                        # The page is gone; every later evaluate would fail the
                        # same way and each would be reported as the author's
                        # syntax error. Stop, and keep what did render. Decided
                        # from page STATE, not the message text - mermaid echoes
                        # the offending line, so a node label reading "Target
                        # page" or "Connection closed" beside a real syntax
                        # error must not be read as the browser dying.
                        fill_rest('render-failed')
                        break
                    out.append((None, 'layout-unsupported'
                                if self._MERMAID_MISSING_LAYOUT_RE.search(str(e))
                                else 'syntax-error'))
                    continue

                if not svg:
                    out.append((None, 'syntax-error'))
                elif png_width is None:
                    out.append((svg, None))
                else:
                    try:
                        out.append((await asyncio.wait_for(
                            self.render(svg.encode('utf-8'), width=png_width,
                                        block_network=True),
                            timeout=max(1.0, deadline - time.monotonic()),
                        ), None))
                    except Exception:
                        # Including a timeout: the loop samples the budget only
                        # at the top, so a single hung screenshot would hold the
                        # request open past every other bound.
                        out.append((None, 'rasterize-failed'))
            return out
        finally:
            # Cleanup must not throw away the diagrams that rendered: this
            # runs after `return out`, so an exception here would discard the
            # whole list and degrade a good export to source blocks.
            try:
                await asyncio.wait_for(ctx.close(),
                                       timeout=self.CONTEXT_CLOSE_TIMEOUT_S)
            except Exception:
                pass

    @staticmethod
    def _viewbox_dims(svg_text: str) -> tuple[float, float]:
        m = re.search(r'viewBox\s*=\s*"([^"]+)"', svg_text)
        if m:
            parts = re.split(r'[\s,]+', m.group(1).strip())
            if len(parts) >= 4:
                try:
                    return float(parts[2]), float(parts[3])
                except ValueError:
                    pass
        wm = re.search(r'\bwidth\s*=\s*"([0-9.]+)', svg_text)
        hm = re.search(r'\bheight\s*=\s*"([0-9.]+)', svg_text)
        return (
            float(wm.group(1)) if wm else 800.0,
            float(hm.group(1)) if hm else 600.0,
        )

    async def render(self, svg_bytes: bytes, *,
                     width: int = 1920, supersample: int = 1,
                     block_network: bool = False) -> bytes:
        svg_text = svg_bytes.decode('utf-8', errors='replace')

        # Drop any XML declaration / DOCTYPE prologue before the <svg> root.
        # The SVG is embedded into an HTML document and parsed by Chromium's
        # HTML parser, which does not understand a DOCTYPE internal subset
        # (`<!DOCTYPE svg [ <!ENTITY ...> ]>`). It ends the DOCTYPE at the
        # first `>`, spilling the rest of the subset - including the closing
        # `]>` - as a visible text node in the top-left corner. JupyterLab's
        # mermaid renderer prepends exactly such a prologue, so this leak
        # showed up on every exported mermaid diagram. Cutting to the first
        # `<svg` removes it; Chromium renders the element fine without it.
        svg_start = svg_text.find('<svg')
        if svg_start > 0:
            svg_text = svg_text[svg_start:]

        # Many mermaid outputs declare a viewBox far larger than the drawn
        # shape, rasterizing the diagram tiny inside a mostly-empty canvas.
        # Load the SVG, measure the content bounding box with getBBox(), and -
        # ONLY when the declared viewBox has real excess whitespace - tighten
        # the viewBox to it. A viewBox that already frames the content is left
        # alone: getBBox reports geometry only (no stroke, markers or filters),
        # so cropping a well-framed diagram would clip node borders and
        # arrowheads that its own padding is there to hold.
        declared_w, declared_h = self._viewbox_dims(svg_text)
        probe = max(1, int(width))

        ctx = await self._browser.new_context(
            color_scheme=self.color_scheme,
            viewport={'width': probe, 'height': probe},
        )
        try:
            if block_network:
                # Only for an SVG this server generated from document text -
                # see `_block_network`. An SVG the user supplied keeps the
                # browser's own behaviour, web fonts and all, as it always had.
                await self._block_network(ctx)
            page = await ctx.new_page()
            probe_html = (
                "<!DOCTYPE html><html><head><meta charset='utf-8'><style>"
                "*{margin:0;padding:0}"
                "html,body{background:transparent;overflow:hidden}"
                "svg{display:block}"
                "</style></head><body>"
                f"{svg_text}"
                "</body></html>"
            )
            await page.set_content(probe_html, wait_until='load')

            # Real content bounds in the SVG's own (viewBox) coordinate space.
            bbox = await page.evaluate(
                """() => {
                    const svg = document.querySelector('svg');
                    if (!svg) return null;
                    try {
                        const b = svg.getBBox();
                        if (!b || b.width <= 0 || b.height <= 0) return null;
                        return {x: b.x, y: b.y, width: b.width, height: b.height};
                    } catch (e) { return null; }
                }"""
            )

            tighten = None  # (x, y, w, h) when we crop the viewBox, else keep it
            if bbox:
                fill_w = bbox['width'] / declared_w if declared_w > 0 else 0.0
                fill_h = bbox['height'] / declared_h if declared_h > 0 else 0.0
                if min(fill_w, fill_h) < 0.8:  # >20% empty in some dimension
                    # One pad, taken from the SMALLER extent, on both axes. The
                    # crop is scaled uniformly to the page, so an equal pad here
                    # is an equal printed margin all round. Taking it from the
                    # larger extent swamps the short axis - 4% of an 844pt width
                    # is 34pt, which nearly doubles a 70pt-tall flowchart - and
                    # padding each axis by its own extent instead trades that for
                    # a diagram whose side gutters are five times its top ones,
                    # and stretches an already extreme aspect ratio further.
                    pad = max(6.0, 0.04 * min(bbox['width'], bbox['height']))
                    tighten = (bbox['x'] - pad, bbox['y'] - pad,
                               bbox['width'] + 2 * pad, bbox['height'] + 2 * pad)

            # Size from the box we will render: the tightened box if cropping,
            # otherwise the declared viewBox (which we leave in place).
            vb_w, vb_h = (tighten[2], tighten[3]) if tighten else (declared_w, declared_h)
            nominal_w = probe
            nominal_h = max(1, round(vb_h * nominal_w / vb_w)) if vb_w > 0 else nominal_w
            target_w = max(1, nominal_w * supersample)
            target_h = max(1, nominal_h * supersample)
            # A long single-column flowchart is many times taller than it is
            # wide, and at the configured export width its raster passes the
            # ~16384px Chromium caps on a viewport and a texture. The screenshot
            # then throws and the diagram is dropped from the export entirely,
            # so scale the whole raster down instead - the aspect ratio and the
            # page fit are unaffected, only the pixel density.
            longest = max(target_w, target_h)
            if longest > self.MAX_RASTER_PX:
                shrink = self.MAX_RASTER_PX / longest
                target_w = max(1, int(target_w * shrink))
                target_h = max(1, int(target_h * shrink))

            # Rewrite the viewBox only when cropping; otherwise keep the SVG's
            # own (honouring a non-zero origin). Size the element either way.
            await page.evaluate(
                """({vb, w, h}) => {
                    const svg = document.querySelector('svg');
                    if (!svg) return;
                    if (vb) svg.setAttribute('viewBox', vb);
                    svg.setAttribute('preserveAspectRatio', 'xMidYMid meet');
                    // Mermaid stamps an inline `max-width: <natural>px` on the
                    // root. It caps the element at its natural size, so the
                    // width set below is ignored and the diagram paints small
                    // inside the target canvas - the rest rasterizing as
                    // whitespace. Clear both caps before sizing.
                    svg.style.maxWidth = 'none';
                    svg.style.maxHeight = 'none';
                    svg.setAttribute('width', w);
                    svg.setAttribute('height', h);
                    svg.style.width = w + 'px';
                    svg.style.height = h + 'px';
                    for (const el of [document.documentElement, document.body]) {
                        el.style.width = w + 'px';
                        el.style.height = h + 'px';
                    }
                }""",
                {'vb': (f'{tighten[0]} {tighten[1]} {tighten[2]} {tighten[3]}'
                        if tighten else None),
                 'w': target_w, 'h': target_h},
            )
            await page.set_viewport_size({'width': target_w, 'height': target_h})

            png_bytes = await page.screenshot(
                type='png',
                omit_background=True,
                clip={'x': 0, 'y': 0, 'width': target_w, 'height': target_h},
            )
        finally:
            await ctx.close()

        if supersample > 1:
            from PIL import Image as _PILImage
            img = _PILImage.open(io.BytesIO(png_bytes))
            final = img.resize((nominal_w, nominal_h), _PILImage.LANCZOS)
            buf = io.BytesIO()
            final.save(buf, 'PNG')
            png_bytes = buf.getvalue()

        return png_bytes


class ExportHandlerBase(APIHandler):
    """Base class for export handlers with common functionality."""

    #: w:tblCellMar of the 'Light List Accent 1' style applied to every content
    #: table, 108 twips each side. Column widths are outer widths, so this much
    #: of each is margin rather than content.
    DOCX_CELL_MARGIN_TWIPS = 216

    #: Padding SimpleDocTemplate puts inside its frame on every side, so a
    #: flowable has this much less room than the page margins suggest.
    PDF_FRAME_PADDING = 6

    #: Page margin of the exported PDF, every side.
    PDF_PAGE_MARGIN = 36

    #: Left and right indent of a PDF code block, in points. Subtracted from
    #: the frame before a line is measured, so a wrapped line stops where the
    #: block's own background does.
    PDF_CODE_INDENT = 6

    #: reportlab's default LEFT/RIGHTPADDING inside a table cell. Distinct from
    #: PDF_FRAME_PADDING (which happens to share the value): this is the gap
    #: around cell text, that is the gap inside the page frame.
    PDF_TABLE_CELL_PADDING = 6

    #: Base body size in points for each `exportFontSize` setting. Every other
    #: size in every format is a proportion of this one, so the whole document
    #: scales together rather than only its paragraphs.
    EXPORT_FONT_SIZES = {'small': 10.0, 'medium': 12.0, 'large': 13.0}
    DEFAULT_FONT_SIZE_PT = EXPORT_FONT_SIZES['medium']

    #: Body size the DOCX template is built around: its named styles carry
    #: explicit sizes proportioned against this, and so does the column-width
    #: estimate in `fit_docx_table_to_page`.
    DOCX_TEMPLATE_BASE_PT = 11.0

    #: Every size the PDF draws, as a multiple of the base body size, with the
    #: leading each one needs. Sized against the DOCX template's proportions so
    #: the two formats render the same document at the same scale.
    PDF_TYPE_SCALE = {
        'body': (1.0, 1.2),
        'heading1': (1.4, 1.8),
        'heading2': (1.2, 1.5),
        'heading3': (1.1, 1.4),
        # H4-H6 carry no size of their own in the DOCX template either - they
        # are told apart by weight, slant and colour, not by size
        'heading4': (1.0, 1.3),
        'heading5': (1.0, 1.3),
        'heading6': (1.0, 1.3),
        'table': (0.9, 1.1),
        'code': (0.8, 1.0),
    }

    #: Face the PDF gives each heading below level 3, read off the DOCX
    #: template python-docx builds from: (bold, italic, colour). Without them
    #: every level under 3 shared the Heading 3 style, so a sub-subsection read
    #: as a sibling of its parent - and the PDF disagreed with its own DOCX.
    PDF_MINOR_HEADING_FACES = {
        4: (True, True, '#4F81BD'),
        5: (False, False, '#243F60'),
        6: (False, True, '#243F60'),
    }

    @classmethod
    def font_size_pt(cls, name) -> float:
        """Base body size in points for a setting value.

        Anything unrecognised - an older client that does not send the
        setting, a hand-edited value, a value of the wrong type entirely -
        falls back to the default rather than failing an export over a
        cosmetic choice. An explicit number is accepted but clamped to a range
        that still produces a readable document.
        """
        if isinstance(name, bool) or not isinstance(name, (int, float, str)):
            return cls.DEFAULT_FONT_SIZE_PT   # a list or dict is not a size
        if isinstance(name, str):
            return cls.EXPORT_FONT_SIZES.get(name, cls.DEFAULT_FONT_SIZE_PT)
        # An explicit number is honoured, but clamped: 0 gives zero-height
        # flowables and a huge value one glyph per page, neither of which is
        # a document
        return min(max(float(name), 6.0), 32.0)

    @classmethod
    def pdf_type(cls, role: str, base_pt: float) -> dict:
        """`fontSize`/`leading` kwargs for one role at a given base size."""
        size_ratio, leading_ratio = cls.PDF_TYPE_SCALE[role]
        return {'fontSize': base_pt * size_ratio,
                'leading': base_pt * leading_ratio}

    def apply_docx_font_size(self, document, base_pt: float) -> None:
        """Scale a document's type so its body text is ``base_pt``.

        Word resolves an unset size through the style chain, so setting Normal
        moves body text, tables and lists together. The styles that carry an
        explicit size - headings, caption, title - are scaled by the same
        factor instead of being overwritten, which keeps the proportions the
        template was designed with.
        """
        from docx.shared import Pt

        scale = base_pt / self.DOCX_TEMPLATE_BASE_PT
        if scale != 1.0:
            for style in document.styles:
                try:
                    size = style.font.size
                except (AttributeError, NotImplementedError):
                    continue  # table and numbering styles carry no font
                if size is not None:
                    style.font.size = Pt(round(size.pt * scale, 1))
        document.styles['Normal'].font.size = Pt(base_pt)

    #: Bounds of `svgPixelWidth`, mirroring schema/plugin.json. A width is a
    #: raster size: zero produces nothing, and an absurd one a canvas Chromium
    #: refuses.
    SVG_PIXEL_WIDTH_RANGE = (400, 4096)
    DEFAULT_SVG_PIXEL_WIDTH = 1920

    #: Bounds of `mathPixelWidth`, mirroring schema/plugin.json.
    MATH_PIXEL_WIDTH_RANGE = (200, 3000)
    DEFAULT_MATH_PIXEL_WIDTH = 800

    @classmethod
    def math_pixel_width(cls, value) -> int:
        """Math rasterization width for a request value.

        Same hazard as `svg_pixel_width`: `"mathPixelWidth": null` reaches
        the renderer as None, where the DPI arithmetic raises, the raise is
        swallowed, and the PDF ships its equations as literal `$x^2$` text.
        """
        return cls._pixel_width(value, cls.DEFAULT_MATH_PIXEL_WIDTH,
                                cls.MATH_PIXEL_WIDTH_RANGE)

    @classmethod
    def svg_pixel_width(cls, value) -> int:
        """Rasterization width for a request value.

        `data.get(key, default)` only defaults a MISSING key, so a client that
        sends `"svgPixelWidth": null` - or a string, or a bool - lands here
        with something that is not a width. It must not reach the renderer:
        `None` there silently selects SVG output, which Word cannot display,
        and no failure is raised to warn anyone.
        """
        return cls._pixel_width(value, cls.DEFAULT_SVG_PIXEL_WIDTH,
                                cls.SVG_PIXEL_WIDTH_RANGE)

    @staticmethod
    def _pixel_width(value, default: int, bounds: tuple[int, int]) -> int:
        if isinstance(value, bool) or not isinstance(value, (int, float, str)):
            return default
        try:
            width = int(float(value))
        except (TypeError, ValueError, OverflowError):
            # OverflowError is not a ValueError: `Infinity` survives
            # json.loads, and "1e400" survives the isinstance gate.
            return default
        low, high = bounds
        return min(max(width, low), high)

    def get_absolute_path(self, relative_path: str) -> Path:
        """Convert a relative path to an absolute path within the server root."""
        root_dir = self.contents_manager.root_dir
        return Path(root_dir) / relative_path

    @staticmethod
    def _ip_is_blocked(ip_str: str) -> bool:
        """True for any non-public address - the ranges an SSRF would target
        (cloud metadata 169.254.169.254, localhost, intranet, CGNAT). Allowlist
        via ``is_global`` rather than enumerating private ranges, so special-use
        blocks we didn't think of stay blocked.
        """
        try:
            ip = ipaddress.ip_address(ip_str)
        except ValueError:
            return True
        return not ip.is_global

    @classmethod
    def _host_is_blocked(cls, host: str | None) -> bool:
        """Resolve a hostname and block it if any resolved IP is non-public.
        Unresolvable hosts are blocked (fail closed)."""
        if not host:
            return True
        try:
            infos = socket.getaddrinfo(host, None)
        except OSError:
            return True
        return any(cls._ip_is_blocked(info[4][0]) for info in infos)

    def read_markdown_file(self, path: Path) -> str:
        """Read and return the contents of a markdown file."""
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()

    #: Any fenced block's opening or closing line: its indentation, the run of
    #: backticks or tildes, then the info string. Three backticks is the
    #: CommonMark minimum, so a ```mermaid block is one of these too - which is
    #: the point: whether it is a diagram or an example depends on whether a
    #: fence is already open.
    #:
    #: The prefix is captured, not bounded, and allows `>`. CommonMark measures
    #: a fence's indent from its container's content column, so a block inside
    #: a list item sits four or more spaces in and a quoted one behind a `>` -
    #: both still fences. The regex this replaced was position-blind and
    #: JupyterLab renders both, so refusing them here would drop a diagram the
    #: browser had already counted and put every later picture on the wrong
    #: fence.
    _FENCE_RE = re.compile(r'^([ \t>]*)(`{3,}|~{3,})(.*)$')

    #: The exact character set JS `String.prototype.trim()` removes - ECMAScript
    #: WhiteSpace plus LineTerminator. marked (JS) is the authority for what an
    #: info string trims to, and Python's `str.strip()` differs on a few code
    #: points: it keeps U+FEFF (a BOM after `mermaid` would count as a diagram
    #: in the browser but not here) and strips control chars JS does not. Using
    #: this set with `str.strip()` makes the trim byte-equivalent to marked's.
    _JS_TRIM = ''.join(chr(c) for c in (
        0x09, 0x0a, 0x0b, 0x0c, 0x0d, 0x20, 0xa0, 0x1680,
        0x2000, 0x2001, 0x2002, 0x2003, 0x2004, 0x2005, 0x2006, 0x2007,
        0x2008, 0x2009, 0x200a, 0x2028, 0x2029, 0x202f, 0x205f, 0x3000,
        0xfeff,
    ))

    @classmethod
    def iter_mermaid_blocks(cls, content: str):
        """Every mermaid block the document actually means as a diagram.

        Yields a match-like object per block, in document order, with the
        source in group 1 - so a caller reads it exactly as it read the
        `re.finditer` this replaced.

        A ```mermaid opened while another fence is already open is
        documentation showing the syntax, not a diagram: nobody wants their
        code sample replaced by a picture of itself. Only a same-character
        fence at least as long as the opener AND carrying no info string
        closes a block (CommonMark), so ```mermaid can sit quoted inside a
        plain ```text block as easily as inside a ```` ```` ```` one.
        """
        lines = content.split('\n')
        offsets, at = [], 0
        for line in lines:
            offsets.append(at)
            at += len(line) + 1

        char: str | None = None      # the open fence, if any
        length = 0
        opened_mermaid = False       # ...and whether it is a diagram
        quoted_block = False         # ...and whether a blockquote holds it
        open_q = 0                   # ...and at what blockquote depth
        list_col = 0                 # content column of the open list item
        list_q = 0                   # ...at what blockquote depth it lives
        body_start = 0
        block_start = 0
        open_spaces = 0

        for i, line in enumerate(lines):
            # Everything about where a fence sits is measured in the coordinate
            # marked re-lexes a line in: blockquote markers stripped, tabs
            # expanded. A fence is a diagram only where marked draws one, and
            # the browser pairs diagrams with fences by position, so counting
            # one marked skips - or skipping one it draws - puts every later
            # picture on the wrong fence.
            q, text = cls._quote_stripped(line)
            spaces = len(text) - len(text.lstrip(' '))

            if char is None:
                # Track the innermost open list item's content column. Inside an
                # item the indentation is the item's, and the fence starts at
                # that column; four spaces PAST it is an indented code block,
                # exactly as at the top level.
                col = cls._list_content_col(text)
                if col is not None:
                    list_col, list_q = col, q
                elif text.strip(cls._JS_TRIM) and (q != list_q
                                                   or spaces < list_col):
                    list_col = 0

            # A blockquote ends at the first line that is not quoted, and
            # CommonMark closes any fence still open inside it. Checked before
            # the fence branch so a line that both ends the quote AND is itself
            # a fence (```python right after a quoted ```mermaid) closes the
            # block first instead of running on and eating that next fence.
            if char is not None and quoted_block and q < open_q:
                if opened_mermaid:
                    # marked renders a container-closed block, so the browser
                    # counted it; dropping it here shifts every later picture.
                    end = offsets[i] - 1
                    yield _MermaidBlock(content, block_start, end,
                                        cls._body_source(
                                            content[body_start:end],
                                            quoted_block, open_spaces))
                char, opened_mermaid, quoted_block = None, False, False

            fence = cls._FENCE_RE.match(line)
            if not fence:
                continue
            indent, marker = fence.group(1), fence.group(2)
            info = fence.group(3).strip(cls._JS_TRIM)
            if char is None:
                base = list_col if (list_col and q == list_q) else 0
                if spaces - base > 3:
                    continue  # indented code, not a fence opener
                char, length = marker[0], len(marker)
                open_spaces = spaces
                open_q = q
                # marked decides what a DIAGRAM is: `languages.includes(lang)`
                # is an exact, case-sensitive match on the trimmed info string,
                # and marked's fence rule treats ~~~ exactly like ```.
                # Lowercasing made a ```MERMAID sample a picture of itself, and
                # refusing ~~~ dropped a diagram the browser had counted.
                # Residue: a ~~~ block left un-rendered is not covered by the
                # ``` fence protection in the math and alert passes (DEF-DIAG-19).
                opened_mermaid = info == 'mermaid'
                quoted_block = q > 0
                # Start at the backticks, not at the start of the line: the
                # indentation belongs to whatever container holds the block,
                # and swallowing it drops the image out of that container.
                block_start = offsets[i] + len(indent)
                body_start = offsets[i] + len(line) + 1
            elif (marker[0] == char and len(marker) >= length and not info
                  and q == open_q and spaces <= open_spaces + 3):
                # A closer has to live in the opener's container: a quoted ```
                # inside an unquoted fence, or one indented far enough to be
                # code, is body text. Both are how a document that quotes fenced
                # markdown reads, and closing early truncates the source and
                # shifts every later block.
                if opened_mermaid:
                    yield _MermaidBlock(content, block_start,
                                        offsets[i] + len(line),
                                        cls._body_source(
                                            content[body_start:offsets[i]],
                                            quoted_block, open_spaces))
                char = None
                opened_mermaid = False
                quoted_block = False

        if opened_mermaid:
            # Same rule at the end of the document: CommonMark closes the fence
            # there and marked renders the block, so it is a diagram.
            end = offsets[-1] + len(lines[-1])
            yield _MermaidBlock(content, block_start, end,
                                cls._body_source(content[body_start:end],
                                                 quoted_block, open_spaces))

    #: A blockquote marker at the head of a line inside a quoted block. The
    #: markers are the quote's syntax, not the diagram's - markdown strips
    #: them before the code ever reaches a renderer, and mermaid cannot parse
    #: `> flowchart LR`.
    _QUOTE_MARKER_RE = re.compile(r'^[ \t]*>[ ]?', re.MULTILINE)

    #: A blockquote marker at the head of a line: up to three leading spaces,
    #: a `>`, and one optional space that marked strips with it.
    _QUOTE_PREFIX_RE = re.compile(r'^ {0,3}>[ ]?')

    #: A list item marker and the spaces after it. `\r` is in the trailing
    #: class because `_quote_stripped` keeps a CRLF line's `\r`, so a bare `-`
    #: marker must still read as a list item.
    _LIST_MARKER_RE = re.compile(r'^( {0,3})([-*+]|\d{1,9}[.)])([ \t\r]*)')

    @classmethod
    def _quote_stripped(cls, line: str) -> tuple[int, str]:
        """The line in the coordinate marked re-lexes it in.

        Blockquote markers removed (with the one space each carries) and tabs
        expanded, so a fence's indentation is measured against its container's
        content column, not the raw start of the line. Returns the blockquote
        depth and the stripped text.
        """
        # Expand tabs first: a tab right after `>` advances to the next tab
        # stop, and the blockquote marker consumes one column of it, so the
        # rest counts as the fence's indent. marked draws `>\t```mermaid`.
        line = line.expandtabs(4)
        q = 0
        m = cls._QUOTE_PREFIX_RE.match(line)
        while m:
            q += 1
            line = line[m.end():]
            m = cls._QUOTE_PREFIX_RE.match(line)
        return q, line

    @classmethod
    def _list_content_col(cls, text: str) -> int | None:
        """The column a list item's content begins at, or None if `text` (already
        quote-stripped) is not a list item.

        CommonMark: the content column is the marker end plus the spaces after
        it, except a bare marker or five-plus spaces (indented code inside the
        item) counts as one - which is the column a nested fence indents into.
        """
        m = cls._LIST_MARKER_RE.match(text)
        if not m:
            return None
        after = m.group(3).replace('\r', '')
        if m.end() != len(text) and not after:
            return None  # marker jammed against content, e.g. `-->`
        base = len(m.group(1)) + len(m.group(2))
        n = len(after)
        if m.end() == len(text) or n == 0 or n >= 5:
            return base + 1
        return base + n

    @classmethod
    def _body_source(cls, source: str, quoted: bool, indent: int) -> str:
        """The source as the diagram author wrote it.

        A nested quote carries a marker per level, so strip while they match;
        and a CRLF document keeps its `\\r` through the slice, which mermaid
        reads as part of the last token on every line.

        A block inside a list item is indented, and that indent is the item's,
        not the diagram's - marked hands mermaid the body without it, so this
        removes the same columns the opening fence carried. Leaving them on is
        what makes a diagram render from the preview but not from the API.
        """
        source = source.replace('\r\n', '\n')
        # `indent` is already the container-relative column (the quote's own
        # space was stripped when it was measured), so the body is dedented by
        # exactly that after its own quote markers come off.
        while quoted and cls._QUOTE_MARKER_RE.search(source):
            source = cls._QUOTE_MARKER_RE.sub('', source)
        out = []
        for line in source.split('\n'):
            # A block closed by its container rather than by a fence ends mid
            # line-ending, so the last line can still carry a lone `\r`.
            if line.endswith('\r'):
                line = line[:-1]
            cut = col = 0
            while cut < len(line) and col < indent:
                if line[cut] == ' ':
                    col += 1
                elif line[cut] == '\t':
                    col += 4
                else:
                    break
                cut += 1
            out.append(line[cut:])
        return '\n'.join(out)

    @classmethod
    def count_mermaid_blocks(cls, content: str) -> int:
        return sum(1 for _ in cls.iter_mermaid_blocks(content))

    @classmethod
    def sub_mermaid_blocks(cls, content: str, replace) -> str:
        """`re.sub` over exactly the blocks `iter_mermaid_blocks` yields, so
        the substitution and the collection can never see different sets."""
        out: list[str] = []
        last = 0
        for match in cls.iter_mermaid_blocks(content):
            out.append(content[last:match.start()])
            out.append(replace(match))
            last = match.end()
        out.append(content[last:])
        return ''.join(out)

    @staticmethod
    def color_scheme_for(theme: str) -> str:
        """The browser colour scheme an export theme resolves to. Mermaid and
        `prefers-color-scheme` have light and dark and no third option, so the
        follow-the-viewer settings - 'auto' for the DOCX theme, 'system' for
        the HTML one - render light."""
        return 'dark' if theme == 'dark' else 'light'

    def replace_mermaid_with_images(self, content: str, mermaid_diagrams: list,
                                      use_png: bool = False) -> tuple[str, list[int]]:
        """
        Replace mermaid code blocks with pre-rendered images from the frontend.

        Args:
            content: Markdown content with mermaid code blocks
            mermaid_diagrams: List of dicts with 'index' and 'svg' (base64 data URI)
            use_png: If True, convert SVG to PNG (for DOCX compatibility)

        Returns:
            The content with each supplied diagram substituted, and the
            document positions of the blocks left behind. Those positions are
            what a warning about an un-rendered diagram must quote: numbering
            the leftovers 0, 1, 2 would send the reader to the wrong diagram
            whenever the frontend supplied some but not all of them.
        """
        if not mermaid_diagrams:
            return content, list(range(self.count_mermaid_blocks(content)))

        # Create lookup dicts by index for both SVG and PNG
        diagrams_by_index = {}
        for d in mermaid_diagrams:
            diagrams_by_index[d['index']] = {
                'svg': d.get('svg', ''),
                'png': d.get('png', '')
            }

        current_index = [0]  # Use list to allow mutation in nested function
        left_behind: list[int] = []

        def replace_mermaid(match):
            idx = current_index[0]
            current_index[0] += 1

            if idx in diagrams_by_index:
                diagram = diagrams_by_index[idx]
                svg_data_uri = diagram['svg']
                png_data_uri = diagram['png']

                if use_png and png_data_uri:
                    # Prefer PNG from frontend (client-side Canvas conversion)
                    return f'![Mermaid Diagram]({png_data_uri})'
                # Otherwise emit SVG; extract_data_uri_images() converts to
                # PNG server-side via Playwright when convert_svg=True.
                return f'![Mermaid Diagram]({svg_data_uri})'

            # No pre-rendered diagram available, keep original
            left_behind.append(idx)
            return match.group(0)

        content = self.sub_mermaid_blocks(content, replace_mermaid)
        return content, left_behind

    #: Diagram indices listed per warning before the header just reports the
    #: count. A generated document can carry hundreds of diagrams, and this is
    #: what keeps the header bounded. At most six codes can co-occur (the
    #: document-wide ones exclude the per-diagram ones), so the header stays
    #: under 1.7KB - inside nginx's 4KB default `proxy_buffer_size`, which has
    #: to hold the whole response header block, not just this one.
    MAX_REPORTED_DIAGRAMS = 10


    async def render_mermaid_server_side(
        self, content: str, *, color_scheme: str, png_width: int | None,
        diagram_indices: list[int],
    ) -> tuple[str, list[dict]]:
        """Render any mermaid block the frontend did not, and inline it.

        Mermaid is a browser library, so the extension renders diagrams in the
        page and posts them as `mermaidDiagrams`. A caller driving the REST API
        directly - a script, curl, a scheduled job - has no browser, so nothing
        arrives and every diagram used to export as its own source code. Render
        those blocks with the vendored mermaid bundle in the same headless
        Chromium that rasterizes SVGs: `png_width` gives a PNG at that width
        for DOCX and PDF, `None` an SVG for HTML - exactly what the frontend
        posts for each.

        Runs after `replace_mermaid_with_images`, so a diagram the frontend
        supplied is already gone and only what it left is rendered - the UI
        export is untouched, and a document with no mermaid never starts a
        browser. `diagram_indices` is that method's account of which document
        positions the blocks still here occupy, so a warning quotes the number
        the reader would count to in their own file.

        ORDERING RULE: both mermaid passes run on the document as read, before
        any pass that rewrites the source. The frontend counted its diagrams in
        the file the author wrote and pairing is by position, so a pass that
        adds or removes a fence first shifts every later picture onto the wrong
        one. `preprocess_github_alerts` does exactly that - it folds an alert
        body onto a single line, which erases a fence inside a `> [!NOTE]`
        completely, so a document with one diagram in a note and one after it
        exported the note's picture in the later diagram's place.

        Returns the content and a warning per reason a diagram kept its source.
        Nothing here fails an export: an export missing a picture still beats
        no export at all, so every failure degrades to the source block and is
        reported instead.
        """
        sources = [m.group(1) for m in self.iter_mermaid_blocks(content)]
        if not sources:
            return content, []

        # Checked before the browser starts, so a broken install costs no
        # Chromium launch to discover.
        if not PlaywrightSvgRenderer.mermaid_bundle_available():
            # Logged as well as reported: a mis-packaged wheel degrades every
            # export in the same silent way a caller who never looks at the
            # header cannot tell from the feature not existing.
            self.log.warning('mermaid: %s', MERMAID_WARNINGS['bundle-missing'])
            return content, self.group_mermaid_warnings(
                ['bundle-missing'] * len(sources), diagram_indices)

        # One parameter carries the whole PNG-vs-SVG decision, so what the
        # renderer produces and how the data URI labels it cannot disagree.
        mime = 'image/png' if png_width is not None else 'image/svg+xml'

        try:
            async with PlaywrightSvgRenderer(color_scheme=color_scheme,
                                             offline=True) as renderer:
                results = await renderer.render_mermaid(
                    sources,
                    theme='dark' if color_scheme == 'dark' else 'default',
                    png_width=png_width,
                )
        except ChromiumUnavailableError:
            self.log.warning('mermaid: %s',
                             MERMAID_WARNINGS['chromium-unavailable'])
            return content, self.group_mermaid_warnings(
                ['chromium-unavailable'] * len(sources), diagram_indices)
        except Exception:
            # Anything else the browser session can throw - the page failing to
            # load, the bundle deleted under us, the context dying. Per-diagram
            # faults are already handled inside render_mermaid; this is the
            # whole session going down, and it must still not fail the export.
            self.log.exception('mermaid: the render session failed')
            return content, self.group_mermaid_warnings(
                ['render-failed'] * len(sources), diagram_indices)

        rendered = iter(results)
        reasons: list[str | None] = []

        def replace(match):
            image, reason = next(rendered, (None, 'render-failed'))
            try:
                if image is None:
                    return match.group(0)  # keep the source, reported below
                raw = image if isinstance(image, bytes) else image.encode('utf-8')
                data_uri = (f'data:{mime};base64,'
                            + base64.b64encode(raw).decode('ascii'))
                return f'![Mermaid Diagram]({data_uri})'
            except Exception:
                # The only unguarded step left in the pass, and it must not be
                # the one that fails an export - a lone surrogate in a label
                # would take UTF-8 encoding down with it.
                self.log.exception('mermaid: could not inline a rendered diagram')
                reason = 'render-failed'
                return match.group(0)
            finally:
                reasons.append(reason)

        content = self.sub_mermaid_blocks(content, replace)
        return content, self.group_mermaid_warnings(reasons, diagram_indices)

    def group_mermaid_warnings(self, reasons: list[str | None],
                               diagram_indices: list[int]) -> list[dict]:
        """Group per-diagram failure reasons into one warning each, in the
        order the diagrams appear.

        `reasons[k]` is the k-th block this pass saw; `diagram_indices[k]` is
        where that block sits among the document's own diagrams, and that is
        the number reported. `count` is always the true total; `diagrams`
        lists the first `MAX_REPORTED_DIAGRAMS`, so a document with hundreds
        of broken diagrams cannot grow the header past what an HTTP header
        block will carry.
        """
        order: list[str] = []
        by_code: dict[str, list[int]] = {}
        for position, reason in enumerate(reasons):
            if reason is None:
                continue
            if reason not in by_code:
                by_code[reason] = []
                order.append(reason)
            by_code[reason].append(
                diagram_indices[position] if position < len(diagram_indices)
                else position)
        return [{
            'code': code,
            'count': len(by_code[code]),
            'diagrams': by_code[code][:self.MAX_REPORTED_DIAGRAMS],
            'message': MERMAID_WARNINGS[code],
        } for code in order]

    def set_export_warnings(self, warnings: list[dict]) -> None:
        """Report what went wrong on an export that still succeeded.

        The body is a document, so a warning has nowhere to go but a header.
        Absent when everything rendered.

        Nothing in `src/` reads it: asked how warnings should reach a caller,
        the answer was the response header alone rather than a dialog. It is
        an API channel by decision - a UI export reaches this path only when
        the mermaid manager token is missing, since otherwise the browser
        renders every diagram and there is nothing left to warn about.
        """
        if not warnings:
            return
        self.set_header('X-Export-Warnings', json.dumps(warnings))
        # Without this a cross-origin `fetch()` is handed the document with the
        # warnings stripped - the one channel, silently gone. `add_header`,
        # because something else may already expose one.
        self.add_header('Access-Control-Expose-Headers', 'X-Export-Warnings')

    def set_attachment_filename(self, filename: str) -> None:
        """Name the download, including when the name is not ASCII.

        Tornado refuses any header value outside latin-1, so interpolating a
        name straight into `Content-Disposition` raises `Unsafe header value`
        before a byte of the document is written - the export fails outright
        rather than downloading under an ugly name. Polish, Czech and Greek
        source files all hit this; `zniesławienie-....docx` is the case that
        found it.

        RFC 6266 is the way out: send both an ASCII `filename` for anything
        that predates the standard and `filename*`, which carries the real
        name percent-encoded. Every current browser prefers `filename*`.
        """
        # NFKD splits the accented Latin letters into base + combining mark,
        # which 'ignore' then drops to the base. It leaves stroked letters
        # alone (they have no decomposition), so map those first or 'ł' would
        # vanish rather than degrade to 'l'.
        stroked = str.maketrans({'ł': 'l', 'Ł': 'L', 'đ': 'd', 'Đ': 'D',
                                 'ø': 'o', 'Ø': 'O'})
        stem, _, suffix = filename.rpartition('.')
        ascii_stem = unicodedata.normalize('NFKD', (stem or filename).translate(stroked))
        ascii_stem = ascii_stem.encode('ascii', 'ignore').decode('ascii')
        # Anything left that could close the quoted string or inject a header
        ascii_stem = re.sub(r'[^A-Za-z0-9._-]', '_', ascii_stem).strip('_')
        # A name in a non-Latin script leaves nothing behind, and ".docx" alone
        # is a hidden file rather than a document
        fallback = f'{ascii_stem or "export"}.{suffix}' if suffix else (ascii_stem or 'export')

        self.set_header(
            'Content-Disposition',
            'attachment; filename="{}"; filename*=UTF-8\'\'{}'.format(
                fallback, quote(filename, safe='')
            ),
        )

    # Inline images carrying an explicit display size at or below this many
    # CSS px are treated as small badges/pills, not full-width diagrams.
    _BADGE_MAX_PX = 200
    # Pixel-density multiplier applied to a badge's display size when
    # rasterizing, so it stays crisp when zoomed in Word.
    _BADGE_RENDER_SCALE = 4
    # Hard ceiling on a badge's rasterized width (px) - guards against a crafted
    # SVG viewBox aspect ratio driving the Playwright canvas to an absurd size.
    _BADGE_RENDER_MAX_PX = 4096

    @classmethod
    def _badge_render_spec(cls, full_tag: str, vb_w: float, vb_h: float):
        """Render width and DPI for a small, explicitly-sized inline image.

        Reads height/width from the tag's ``style`` (``max-height`` preferred
        over ``height``) or its ``width``/``height`` attributes. When the
        resulting display size is small, returns ``(render_width_px, dpi)`` so
        the PNG embeds at that physical size with ``_BADGE_RENDER_SCALE``x pixel
        density. Returns ``None`` for unsized or large images, which then
        rasterize at the full diagram width.
        """
        num = r'(\d+(?:\.\d+)?)'  # strict number - no ValueError on "1.2.3px"
        max_h = max_w = None
        style_m = re.search(r'style=["\']([^"\']*)["\']', full_tag, re.IGNORECASE)
        if style_m:
            s = style_m.group(1)
            # max-height wins over plain height (CSS clamps height to max-height);
            # the leading delimiter stops `line-height` / `stroke-width` matching.
            mh = (re.search(r'(?:^|[;\s])max-height\s*:\s*' + num + r'\s*px', s, re.IGNORECASE)
                  or re.search(r'(?:^|[;\s])height\s*:\s*' + num + r'\s*px', s, re.IGNORECASE))
            mw = (re.search(r'(?:^|[;\s])max-width\s*:\s*' + num + r'\s*px', s, re.IGNORECASE)
                  or re.search(r'(?:^|[;\s])width\s*:\s*' + num + r'\s*px', s, re.IGNORECASE))
            if mh:
                max_h = float(mh.group(1))
            if mw:
                max_w = float(mw.group(1))
        # HTML width/height attributes. Tokenise attributes so a `width=`/
        # `height=` sitting inside another attribute's quoted value (e.g.
        # alt="x width=5") or a `data-width`/`data-height` is never read as
        # geometry. Bare number, optional px; %, em etc. are rejected.
        if max_h is None or max_w is None:
            for name, dq, sq, bare in re.findall(
                r'([-\w]+)\s*=\s*(?:"([^"]*)"|\'([^\']*)\'|([^\s>]+))',
                full_tag, re.IGNORECASE,
            ):
                key = name.lower()
                if key not in ('width', 'height'):
                    continue
                vm = re.fullmatch(num + r'(?:px)?', (dq or sq or bare).strip(), re.IGNORECASE)
                if not vm:
                    continue
                if key == 'height' and max_h is None:
                    max_h = float(vm.group(1))
                elif key == 'width' and max_w is None:
                    max_w = float(vm.group(1))
        if max_h is None and max_w is None:
            return None
        if vb_w <= 0 or vb_h <= 0:
            return None
        # Effective rendered height = smallest height the constraints permit
        # (a max-width can shrink a tall image below its max-height). Qualify on
        # height so wide-but-short pills (max-width:1000px) still count as badges.
        aspect = vb_w / vb_h
        eff_h = max_h
        if max_w is not None:
            h_from_w = max_w / aspect
            eff_h = h_from_w if eff_h is None else min(eff_h, h_from_w)
        if eff_h > cls._BADGE_MAX_PX:
            return None
        disp_w = eff_h * aspect
        render_w = min(cls._BADGE_RENDER_MAX_PX,
                       max(1, round(disp_w * cls._BADGE_RENDER_SCALE)))
        return render_w, 96 * cls._BADGE_RENDER_SCALE

    async def extract_data_uri_images(self, html: str, temp_dir: str,
                                      convert_svg: bool = False,
                                      svg_pixel_width: int = 1920,
                                      color_scheme: str = 'light') -> str:
        """
        Extract data URI images to temp files for htmldocx compatibility.

        When convert_svg=True, SVG images are rasterized to PNG via Playwright
        Chromium so CSS (including @media (prefers-color-scheme)), web fonts,
        filters and tspan layout match a real browser. One Chromium instance
        is reused across all SVGs in this call.

        Raises ChromiumUnavailableError if convert_svg=True and Chromium
        cannot launch (binary or system libs missing).

        Args:
            html: HTML content with data URI images
            temp_dir: Directory to store temp image files
            convert_svg: If True, convert SVG images to PNG via Playwright
            svg_pixel_width: Target pixel width for each rasterized SVG;
                height follows the SVG's viewBox aspect ratio
            color_scheme: 'light' or 'dark'; passed to the browser context so
                @media (prefers-color-scheme) rules in the SVG resolve correctly
        """
        img_pattern = r'<img\s+[^>]*src=["\']data:image/([^;]+);base64,([^"\']+)["\'][^>]*>'
        ext_map = {
            'png': '.png', 'jpeg': '.jpg', 'jpg': '.jpg',
            'gif': '.gif', 'svg+xml': '.svg',
        }

        matches = list(re.finditer(img_pattern, html))
        if not matches:
            return html

        svg_indices = [
            i for i, m in enumerate(matches)
            if convert_svg and m.group(1) == 'svg+xml'
        ]
        rendered: dict[int, bytes | None] = {}
        # Per-image target DPI: small inline badges get a high DPI so their
        # rasterized PNG embeds at the badge's physical size (not page width).
        target_dpi: dict[int, int] = {}

        if svg_indices:
            async with PlaywrightSvgRenderer(color_scheme=color_scheme) as renderer:
                for i in svg_indices:
                    try:
                        svg_bytes = base64.b64decode(matches[i].group(2))
                        svg_text = svg_bytes.decode('utf-8', errors='replace')
                        vb_w, vb_h = PlaywrightSvgRenderer._viewbox_dims(svg_text)
                        spec = self._badge_render_spec(
                            matches[i].group(0), vb_w, vb_h
                        )
                        render_w = spec[0] if spec is not None else svg_pixel_width
                        rendered[i] = await renderer.render(
                            svg_bytes, width=render_w
                        )
                        # Commit the badge DPI only after a successful render so
                        # a later failure can't leave a stale badge flag set.
                        if spec is not None:
                            target_dpi[i] = spec[1]
                    except ChromiumUnavailableError:
                        # Bubble up; handler converts to a typed HTTP error
                        # so the frontend can show its install-required popup.
                        raise
                    except Exception:
                        rendered[i] = None  # render failed; keep original SVG

        out: list[str] = []
        last = 0
        for i, m in enumerate(matches):
            out.append(html[last:m.start()])
            last = m.end()

            img_type = m.group(1)
            base64_data = m.group(2)
            full_tag = m.group(0)
            ext = ext_map.get(img_type, '.png')

            try:
                if i in rendered:
                    if rendered[i] is None:
                        out.append(full_tag)
                        continue
                    img_bytes = rendered[i]
                    ext = '.png'
                else:
                    img_bytes = base64.b64decode(base64_data)

                import hashlib
                hash_id = hashlib.md5(img_bytes).hexdigest()[:8]
                filepath = os.path.join(temp_dir, f'img_{hash_id}{ext}')
                with open(filepath, 'wb') as f:
                    f.write(img_bytes)

                # Normalize DPI for consistent sizing in DOCX. Embedded DPI
                # metadata drives python-docx's native-size computation; without
                # it the same pixel dimensions render at wildly different sizes.
                # Small inline badges get a high DPI (target_dpi[i]) so their PNG
                # embeds at the badge's physical size rather than at page width.
                dpi = target_dpi.get(i, 96)
                if ext in ('.jpg', '.jpeg', '.png'):
                    try:
                        from PIL import Image
                        img = Image.open(filepath)
                        img.save(filepath, dpi=(dpi, dpi))
                    except Exception:
                        pass

                if i in target_dpi:
                    # Badge: rely on native DPI size; drop width/height/style so
                    # nothing scales it back up to the cell width. Keep alt only.
                    alt = re.search(r'(alt=["\'][^"\']*["\'])', full_tag)
                    out.append(
                        f'<img src="{filepath}" {alt.group(1)}>' if alt
                        else f'<img src="{filepath}">'
                    )
                else:
                    other_attrs = re.findall(
                        r'((?:alt|style|width|height)=["\'][^"\']*["\'])', full_tag
                    )
                    attrs_str = ' '.join(other_attrs)
                    if attrs_str:
                        out.append(f'<img src="{filepath}" {attrs_str}>')
                    else:
                        out.append(f'<img src="{filepath}">')
            except Exception:
                out.append(full_tag)

        out.append(html[last:])
        return ''.join(out)

    def embed_images_as_base64(self, content: str, markdown_dir: Path) -> str:
        """
        Replace local and remote image references with base64-encoded data URIs.

        Handles both Markdown image syntax ``![alt](path)`` and raw HTML
        ``<img src="path">`` tags (the latter common inside Markdown tables,
        where inline badges/pills are written as HTML).

        - Local paths are read from disk relative to ``markdown_dir`` but must
          stay within the Jupyter server root (the user's workspace); paths that
          escape it (absolute, ``../`` traversal) are refused.
        - ``http://`` / ``https://`` URLs are fetched (10 s timeout, 5 MB cap)
          so badges and other remote images embed in DOCX/PDF, but only when the
          host resolves to a public IP (SSRF guard). Failures fall back to the
          original reference silently.
        - ``data:`` URIs are passed through.
        """
        img_pattern = r'!\[([^\]]*)\]\(([^)"\s]+)(?:\s+"[^"]*")?\)'
        mime_types = {
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.gif': 'image/gif',
            '.svg': 'image/svg+xml',
            '.webp': 'image/webp',
        }
        remote_cache: dict[str, str | None] = {}
        # Containment boundary for local reads - the server root the user is
        # already allowed to browse.
        try:
            root_dir = Path(self.contents_manager.root_dir).resolve()
        except Exception:
            root_dir = None

        def fetch_remote(url: str) -> str | None:
            if url in remote_cache:
                return remote_cache[url]
            try:
                import http.client
                import urllib.request as _urlreq
                from urllib.request import (
                    Request, build_opener, HTTPRedirectHandler,
                    HTTPHandler, HTTPSHandler,
                )
                parsed = urlparse(url)
                # SSRF guard: only public http(s) hosts. Blocks localhost,
                # intranet and the 169.254.169.254 cloud-metadata endpoint.
                if parsed.scheme not in ('http', 'https') or \
                        self._host_is_blocked(parsed.hostname):
                    remote_cache[url] = None
                    return None

                # When an HTTP proxy applies, the socket connects to the proxy
                # (often a private IP), so the peer-IP check can't run - the host
                # pre-check above plus the proxy's own egress policy guard it.
                try:
                    proxied = bool(_urlreq.getproxies().get(parsed.scheme)) and \
                        not _urlreq.proxy_bypass(parsed.hostname or '')
                except Exception:
                    proxied = False

                host_blocked = self._host_is_blocked
                ip_blocked = self._ip_is_blocked

                def _check_peer(conn):
                    # Authoritative SSRF check: validate the IP actually connected
                    # to. Closes the DNS-rebinding/TOCTOU gap between the host
                    # pre-check and urllib's own connect-time resolution.
                    ip = conn.sock.getpeername()[0]
                    if ip_blocked(ip):
                        conn.close()
                        raise OSError(f'blocked non-public address: {ip}')

                class _GHTTP(http.client.HTTPConnection):
                    def connect(self):
                        super().connect()
                        _check_peer(self)

                class _GHTTPS(http.client.HTTPSConnection):
                    def connect(self):
                        super().connect()
                        _check_peer(self)

                class _GHTTPHandler(HTTPHandler):
                    def http_open(self, req):
                        return self.do_open(_GHTTP, req)

                class _GHTTPSHandler(HTTPSHandler):
                    def https_open(self, req):
                        return self.do_open(
                            _GHTTPS, req,
                            context=getattr(self, '_context', None),
                            check_hostname=getattr(self, '_check_hostname', None),
                        )

                class _GuardedRedirect(HTTPRedirectHandler):
                    def redirect_request(self, req, fp, code, msg, hdrs, newurl):
                        p = urlparse(newurl)
                        if p.scheme not in ('http', 'https') or host_blocked(p.hostname):
                            return None  # refuse redirect to a non-public host
                        return super().redirect_request(
                            req, fp, code, msg, hdrs, newurl)

                if proxied:
                    # Default handlers route through the proxy; keep the redirect
                    # host guard but skip the peer-IP check (peer is the proxy).
                    opener = build_opener(_GuardedRedirect())
                else:
                    opener = build_opener(
                        _GHTTPHandler(), _GHTTPSHandler(), _GuardedRedirect()
                    )
                req = Request(
                    url,
                    headers={'User-Agent': 'jupyterlab-export-markdown-extension'},
                )
                with opener.open(req, timeout=10) as resp:
                    raw = resp.read(5 * 1024 * 1024 + 1)
                    if len(raw) > 5 * 1024 * 1024:
                        remote_cache[url] = None
                        return None
                    ctype = (resp.headers.get('Content-Type') or '').split(';', 1)[0].strip()
                if not ctype:
                    ext = os.path.splitext(url)[1].lower()
                    ctype = mime_types.get(ext, 'application/octet-stream')
                b64 = base64.b64encode(raw).decode('utf-8')
                data_uri = f'data:{ctype};base64,{b64}'
                remote_cache[url] = data_uri
                return data_uri
            except Exception:
                remote_cache[url] = None
                return None

        def local_data_uri(img_path: str) -> str | None:
            try:
                # URL-decode the path (handles %20 for spaces, etc.). resolve()
                # can raise (NUL byte, symlink loop) - treat as unresolvable.
                full_path = (markdown_dir / unquote(img_path)).resolve()
                # Containment: stay within the server root. Refuses absolute
                # paths and ../ traversal that escape the workspace (e.g.
                # /etc/passwd), while still allowing legit ../sibling refs inside
                # the root. Fails closed: no known root -> refuse the read.
                if root_dir is None or not full_path.is_relative_to(root_dir):
                    return None
                if not full_path.exists():
                    return None
                with open(full_path, 'rb') as img_file:
                    img_data = base64.b64encode(img_file.read()).decode('utf-8')
                ext = full_path.suffix.lower()
                mime_type = mime_types.get(ext, 'application/octet-stream')
                return f'data:{mime_type};base64,{img_data}'
            except Exception:
                return None

        def resolve_src(src: str) -> str | None:
            """Return a base64 data URI for ``src``, or None to leave it as-is."""
            low = src.lower()
            if low.startswith('data:'):
                return None
            if low.startswith(('http://', 'https://')):
                return fetch_remote(src)
            return local_data_uri(src)

        def replace_image(match):
            alt_text = match.group(1)
            data_uri = resolve_src(match.group(2))
            if data_uri is None:
                return match.group(0)
            return f'![{alt_text}]({data_uri})'

        content = re.sub(img_pattern, replace_image, content)

        # Raw HTML <img> tags (inline badges/pills inside Markdown tables).
        # The (?<![-\w:]) guard keeps `data-src`/`lowsrc`/`xlink:src` from
        # matching as `src`.
        src_attr_re = re.compile(r'(?<![-\w:])(src\s*=\s*)(["\'])(.*?)\2', re.IGNORECASE)

        def replace_html_img(tag_match):
            tag = tag_match.group(0)

            def sub_src(m):
                data_uri = resolve_src(m.group(3))
                if data_uri is None:
                    return m.group(0)
                return f'{m.group(1)}"{data_uri}"'

            return src_attr_re.sub(sub_src, tag, count=1)

        # Quote-aware tag match so a `>` inside an attribute value (e.g.
        # alt="a > b") doesn't truncate the tag.
        img_tag_re = r'<img\b(?:[^>"\']|"[^"]*"|\'[^\']*\')*>'
        return re.sub(img_tag_re, replace_html_img, content, flags=re.IGNORECASE)

    def get_pygments_css(self, dark: bool = False) -> str:
        """Get Pygments CSS for syntax highlighting.

        Args:
            dark: If True, return dark theme (monokai) CSS for use inside
                  @media (prefers-color-scheme: dark) block.
        """
        try:
            from pygments.formatters import HtmlFormatter
            theme = 'monokai' if dark else 'default'
            return HtmlFormatter(style=theme).get_style_defs('.codehilite')
        except ImportError:
            return ''

    def highlight_code_blocks(self, content: str, use_inline_styles: bool = False) -> str:
        """Highlight code blocks using Pygments.

        Args:
            content: Markdown content with fenced code blocks
            use_inline_styles: If True, use inline styles (for DOCX). If False, use CSS classes (for HTML).

        Returns:
            Content with code blocks replaced by highlighted HTML
        """
        try:
            from pygments import highlight
            from pygments.lexers import get_lexer_by_name, guess_lexer, TextLexer
            from pygments.formatters import HtmlFormatter
        except ImportError:
            return content

        # Pattern to match fenced code blocks with optional language
        code_pattern = r'```(\w*)\n(.*?)```'

        def highlight_match(match):
            lang = match.group(1).strip().lower() or 'text'
            code = match.group(2)

            try:
                lexer = get_lexer_by_name(lang)
            except Exception:
                try:
                    lexer = guess_lexer(code)
                except Exception:
                    lexer = TextLexer()

            formatter = HtmlFormatter(
                noclasses=use_inline_styles,
                nowrap=False,
                cssclass='codehilite'
            )
            highlighted = highlight(code, lexer, formatter)

            # Fix color format for htmldocx (needs # prefix, but only 6-char hex)
            if use_inline_styles:
                # Convert 3-char hex to 6-char hex (e.g., #BBB -> #BBBBBB, #00F -> #0000FF)
                highlighted = re.sub(
                    r'color:\s*#([0-9A-Fa-f])([0-9A-Fa-f])([0-9A-Fa-f])(?![0-9A-Fa-f])',
                    lambda m: f'color: #{m.group(1)*2}{m.group(2)*2}{m.group(3)*2}',
                    highlighted
                )
                highlighted = re.sub(
                    r'background:\s*#([0-9A-Fa-f])([0-9A-Fa-f])([0-9A-Fa-f])(?![0-9A-Fa-f])',
                    lambda m: f'background: #{m.group(1)*2}{m.group(2)*2}{m.group(3)*2}',
                    highlighted
                )

            return highlighted

        return re.sub(code_pattern, highlight_match, content, flags=re.DOTALL)

    def extract_code_blocks(self, content: str) -> tuple:
        """Extract code blocks from markdown for PDF rendering.

        Returns:
            Tuple of (modified_content, code_blocks_list)
            code_blocks_list contains dicts with 'lang' and 'code' keys
        """
        code_blocks = []
        code_pattern = r'```(\w*)\n(.*?)```'

        def extract_match(match):
            lang = match.group(1).strip().lower() or 'text'
            code = match.group(2).rstrip('\n')
            idx = len(code_blocks)
            code_blocks.append({'lang': lang, 'code': code})
            # Replace with placeholder
            return f'[[CODE_BLOCK_{idx}]]'

        modified = re.sub(code_pattern, extract_match, content, flags=re.DOTALL)
        return modified, code_blocks

    @classmethod
    def code_columns(cls, avail_width, font_name: str,
                     font_size: float) -> int:
        """Characters a code line may carry before it has to wrap.

        `XPreformatted` draws every source line as exactly one line whatever
        its width - it never wraps - so an over-long line is laid off the page
        and the glyphs past the page edge are not drawn at all. The code font
        is fixed-width, so for the characters it carries a column count is an
        exact width measure and the line can be split before it is ever handed
        to reportlab. A character the code font lacks (an emoji, a CJK glyph)
        renders through a wider substitute, so a line dense with those can
        still overrun - rare in code, and still far better than the whole line
        running off the page. A tab counts as one column, which is also how
        reportlab draws it here (it does not expand tab stops).

        Returns 0 when there is nothing to measure against, which the callers
        read as "do not wrap" - the old behaviour.
        """
        if not avail_width:
            return 0
        char_width = pdfmetrics.stringWidth('M', font_name, font_size)
        if char_width <= 0:
            return 0
        room = avail_width - 2 * cls.PDF_CODE_INDENT
        # A floor, so a pathological width can never wrap to one column a line
        return max(20, int(room / char_width))

    def highlight_code_for_pdf(self, code: str, lang: str,
                               base_pt: float = None,
                               avail_width: float = None) -> list:
        """Highlight code for PDF using Pygments and return reportlab flowables.

        Args:
            avail_width: Width the flowable will be laid out in, in points.
                Given, a line too wide for it is wrapped; omitted, nothing is
                wrapped - there is no page to wrap against.

        Returns list of reportlab flowables (Preformatted paragraphs)
        """
        code_base = self.DEFAULT_FONT_SIZE_PT if base_pt is None else base_pt
        try:
            from pygments import highlight
            from pygments.lexers import get_lexer_by_name, guess_lexer, TextLexer
            from pygments.formatters import get_formatter_by_name
        except ImportError:
            # Fallback: return plain preformatted text
            if REPORTLAB_AVAILABLE:
                from reportlab.lib.styles import getSampleStyleSheet
                styles = getSampleStyleSheet()
                code_style = ParagraphStyle(
                    'CodeBlock',
                    parent=styles['Code'],
                    fontName='Courier',
                    **self.pdf_type('code', code_base),
                    backColor=colors.HexColor('#f8f8f8'),
                    leftIndent=self.PDF_CODE_INDENT,
                    rightIndent=self.PDF_CODE_INDENT,
                    spaceBefore=6,
                    spaceAfter=6
                )
                return [Preformatted(code, code_style, maxLineLength=(
                    self.code_columns(avail_width, 'Courier',
                                      code_style.fontSize) or None))]
            return []

        try:
            lexer = get_lexer_by_name(lang)
        except Exception:
            try:
                lexer = guess_lexer(code)
            except Exception:
                lexer = TextLexer()

        # Use a monospace font for code
        font_name = 'Courier'
        if REPORTLAB_AVAILABLE and 'UnicodeSans' in pdfmetrics.getRegisteredFontNames():
            # Check if we have a monospace Unicode font
            pass  # Keep Courier for code

        code_style = ParagraphStyle(
            'CodeBlock',
            fontName=font_name,
            **self.pdf_type('code', code_base),
            backColor=colors.HexColor('#f8f8f8'),
            leftIndent=self.PDF_CODE_INDENT,
            rightIndent=self.PDF_CODE_INDENT,
            spaceBefore=6,
            spaceAfter=6,
            wordWrap='CJK'  # Allow wrapping on any character
        )

        # Convert Pygments tokens to reportlab XML markup
        from pygments.token import Token, Keyword, Name, Comment, String, Number, Operator

        # Color mapping for token types
        token_colors = {
            Token.Keyword: '#008000',
            Token.Keyword.Constant: '#008000',
            Token.Keyword.Declaration: '#008000',
            Token.Keyword.Namespace: '#008000',
            Token.Keyword.Type: '#B00040',
            Token.Name.Function: '#0000FF',
            Token.Name.Class: '#0000FF',
            Token.Name.Decorator: '#AA22FF',
            Token.Name.Builtin: '#008000',
            Token.Name.Tag: '#008000',
            Token.Comment: '#3D7B7B',
            Token.Comment.Single: '#3D7B7B',
            Token.Comment.Multiline: '#3D7B7B',
            Token.String: '#BA2121',
            Token.String.Doc: '#BA2121',
            Token.String.Escape: '#AA5D1F',
            Token.Number: '#666666',
            Token.Number.Integer: '#666666',
            Token.Number.Float: '#666666',
            Token.Operator: '#666666',
            Token.Operator.Word: '#AA22FF',
        }

        # Collect each source line as (colour, text) segments, keeping the text
        # raw: escaping first would count `&amp;` as five columns when the
        # reader sees one, and the wrap below measures in real characters
        source_lines = []
        current_line = []

        for ttype, value in lexer.get_tokens(code):
            # Find color for token type
            color = None
            for token_type, token_color in token_colors.items():
                if ttype in token_type or ttype is token_type:
                    color = token_color
                    break

            # Split by newlines to handle multiline tokens
            parts = value.split('\n')
            for i, part in enumerate(parts):
                if part:
                    current_line.append((color, part))
                if i < len(parts) - 1:  # Not the last part, meaning there was a newline
                    source_lines.append(current_line)
                    current_line = []

        # Add remaining content
        if current_line:
            source_lines.append(current_line)

        columns = self.code_columns(avail_width, font_name,
                                    code_style.fontSize)

        def wrap(segments):
            """Split one source line into lines of at most `columns` chars."""
            if not columns:
                return [segments]
            lines, line, used = [], [], 0
            for color, text in segments:
                while text:
                    if used >= columns:
                        lines.append(line)
                        line, used = [], 0
                    head, text = text[:columns - used], text[columns - used:]
                    line.append((color, head))
                    used += len(head)
            lines.append(line)
            return lines

        def render(segments):
            out = []
            for color, text in segments:
                # Escape XML special characters
                text = (text.replace('&', '&amp;').replace('<', '&lt;')
                        .replace('>', '&gt;'))
                out.append(f'<font color="{color}">{text}</font>' if color
                           else text)
            return ''.join(out)

        # Join with newlines - XPreformatted preserves whitespace like Preformatted
        formatted_code = '\n'.join(render(segments)
                                   for source in source_lines
                                   for segments in wrap(source))

        return [XPreformatted(formatted_code, code_style)]

    # Alert color scheme (GitHub-aligned)
    ALERT_COLORS = {
        'NOTE':      {'border': '0969DA', 'shading': 'EDF5FD'},  # Blue
        'TIP':       {'border': '1A7F37', 'shading': 'EDFBF2'},  # Green
        'IMPORTANT': {'border': '8250DF', 'shading': 'F4EDFF'},  # Purple
        'WARNING':   {'border': '9A6700', 'shading': 'FEF9E7'},  # Amber
        'CAUTION':   {'border': 'CF222E', 'shading': 'FEF0F0'},  # Red
    }

    #: Task-list checkbox glyphs - Word's own checkbox defaults (U+2612 ballot
    #: box with X for done, U+2610 empty box), which DejaVu (PDF) and every
    #: browser (HTML) also carry. In DOCX they are tagged MS Gothic - the font
    #: Word draws its native checkboxes in - by style_docx_symbol_runs().
    TASK_CHECKBOX_DONE = '☒'
    TASK_CHECKBOX_OPEN = '☐'

    #: Font DOCX gives a symbol glyph. Word ships it on every Windows since 7,
    #: it covers each range below, and it is the face Word's own substitution
    #: usually lands on - naming it just makes the choice the same everywhere.
    DOCX_SYMBOL_FONT = 'Segoe UI Symbol'

    #: Font DOCX gives the task checkboxes - the one Word draws its own in, so
    #: they render solid beside Word's rather than as a thin outline.
    DOCX_CHECKBOX_FONT = 'MS Gothic'

    #: Ranges Cambria - the DOCX body face - has no glyphs in. A character here
    #: is left to Word's per-machine font substitution otherwise, which is what
    #: turns a star into a box on a machine with a thin font set. The arrow
    #: block starts at U+2194 on purpose: Cambria does carry U+2190-2193, so the
    #: everyday `->` keeps the body typeface instead of switching mid-sentence.
    DOCX_SYMBOL_RANGES = (
        (0x2194, 0x21FF),  # Arrows, less the four Cambria has
        (0x2500, 0x257F),  # Box Drawing
        (0x2580, 0x259F),  # Block Elements
        (0x25A0, 0x25FC),  # Geometric Shapes, less the two default emoji
        (0x25FF, 0x25FF),  # U+25FD/U+25FE are Emoji_Presentation=Yes
    )

    #: Symbols named one at a time because their blocks cannot be taken whole.
    #: Miscellaneous Symbols, Dingbats and Miscellaneous Symbols and Arrows
    #: interleave text symbols with emoji, and Word's own fallback draws an
    #: emoji in colour from Segoe UI Emoji - better than anything named here.
    #: So only the text-presentation symbols this project actually renders are
    #: listed: evidence stars and the check and cross marks its documents use.
    DOCX_SYMBOL_CHARS = frozenset('★☆✓✔✗✘')

    #: Codepoints that modify the character before them - the emoji and text
    #: presentation selectors, and the enclosing keycap. They carry no glyph of
    #: their own and must stay in the run their base character is in.
    DOCX_VARIATION_SELECTORS = '\ufe0e\ufe0f\u20e3'

    #: The selector that asks for emoji presentation. A character carrying it
    #: is left unnamed so Word's colour emoji fallback still applies.
    DOCX_EMOJI_SELECTOR = '\ufe0f'

    def preprocess_task_lists(self, content: str) -> str:
        """Render GitHub task-list markers as Unicode checkbox glyphs.

        The markdown converter carries no task-list extension, so ``- [x]`` /
        ``- [ ]`` would pass through as literal ``[x]`` / ``[ ]`` text. Swapping
        the marker for a ballot glyph (done U+2612, open U+2610) renders a
        checkbox in every export format, since the glyph flows through as plain
        list text - HTML, DOCX and PDF all carry it in their body font.

        Fenced code blocks are skipped so a ``- [ ]`` written inside a code
        sample is not rewritten. Only a real list item (``[x]`` / ``[ ]`` / a
        space) matches, so ``- [link](url)`` is left alone.
        """
        task_re = re.compile(r'^(\s*[-*+]\s+)\[([ xX])\](?:\s+|$)')
        fence_re = re.compile(r'^\s*(`{3,}|~{3,})')

        def replace(match):
            box = (self.TASK_CHECKBOX_DONE if match.group(2) in 'xX'
                   else self.TASK_CHECKBOX_OPEN)
            return f'{match.group(1)}{box} '

        # Track the fence character and length so a ``` inside a ~~~ block (or a
        # shorter ``` inside a ```` block) is treated as content. Per CommonMark
        # a block closes only on a same-char fence at least as long as the
        # opener AND carrying no info string (a ```lang line is a nested opener,
        # never a closer), while the opening fence itself may carry one.
        out, fence_char, fence_len = [], None, 0
        for line in content.split('\n'):
            fence = fence_re.match(line)
            if fence:
                marker = fence.group(1)
                bare = not line[fence.end():].strip()
                if fence_char is None:
                    fence_char, fence_len = marker[0], len(marker)
                elif marker[0] == fence_char and len(marker) >= fence_len and bare:
                    fence_char, fence_len = None, 0
            elif fence_char is None:
                line = task_re.sub(replace, line)
            out.append(line)
        return '\n'.join(out)

    # What python-markdown's `OListProcessor.RE` and `UListProcessor.RE`
    # accept once their leading-space bound is lifted - the column is measured
    # here. `1)` is no marker to the converter, nor is a marker with nothing
    # after it, and a tab after the marker is the spaces `expandtabs` makes.
    _CONVERTER_ITEM_RE = re.compile(r'^ *(\d+\.|[*+-])[ \t]+')

    def normalize_list_indentation(self, content: str) -> str:
        """Shift a top-level list written two or three spaces in back to column 0.

        `markdown_to_html` builds the converter with ``tab_length=2`` so a list
        nested by two spaces nests, and ``tab_length`` is also the
        indented-code threshold: a TOP-LEVEL list whose marker sits two spaces
        in is classified as a code block and exported as literal markdown
        source (DEF-MARK-42). Raising it to 4 only trades the defect for its mirror
        image - the two-space nested list flattens into siblings - so neither
        value alone is right, and the indentation is settled here instead,
        before the converter ever sees it. Teaching python-markdown the two
        thresholds separately was measured and abandoned: its list processors
        build their own regexes from ``tab_length`` and `get_items` indexes on
        them, so widening the marker regex raises IndexError from inside the
        parser (recorded under DEF-MARK-42).

        The promise is measured, not predicted. A candidate is a chunk - the
        run of lines between blank lines - whose first line is a marker two or
        three columns in. The document is rendered with the chunk shifted and
        without, and the shift is kept only when exactly ONE top-level block
        differs, and that block turned from an indented code block into a
        list. Everything the converter does with the shifted lines - a tab it
        expands, a line it will not call blank, a `#Note` it reads as a
        heading, a `<div>` or a rule that splits the list, a successor it
        would swallow, a list already open above - shows up as a second
        changed block, or as the wrong kind of block, and refuses the rescue.
        Eleven review rounds each found one more such construct while this
        pass tried to enumerate them; asking the converter closes the class.

        Inside the one block that changes, two more things must hold. The
        code block must have been made of exactly this chunk's lines: a loose
        list, or an item with its own indented sample below it, renders as one
        code block spanning several chunks, and shifting one chunk of it nests
        the rest under the moved item. And the list must hold nothing that
        needs a block of its own - a `===` under the last item becomes a
        heading inside it, a sample under an item a code block inside it.

        Two things the rendering cannot show are checked by hand: `LAZY_OL`
        writes no `start`, so an ordered chunk that does not open at `1.`
        would silently renumber and is refused; and a chunk holding a fence
        or a raw HTML element is never a candidate, so a sample keeps its
        bytes and a tag the control escapes is never emitted live.
        """
        import markdown
        from bs4 import BeautifulSoup

        # A fence, or a raw HTML element: the control escapes a tag inside
        # an indented block and the rescue would emit it live - a `<link>`
        # that hides the page, a `<meta>` that leaves it. A comment (`<!--`)
        # renders as nothing either way and is not an element. A link
        # reference definition is consumed by the converter as a definition
        # and shown nowhere, so a URL the control displays would vanish.
        verbatim_re = re.compile(
            r'^\s*(`{3,}|~{3,})|<[A-Za-z/]|^\s*\[[^\]]*\]:')
        # `HashHeaderProcessor.RE`: hashes at column 0, nothing else required.
        # A heading is its own block, so it ends the chunk above it
        atx_re = re.compile(r'^#{1,6}')

        def blank(ln):
            # `NormalizeWhitespace` expands tabs and then empties only
            # SPACE-only lines; a line holding an NBSP is content to it
            return not ln.strip(' \t\r')

        md = markdown.Markdown(
            extensions=['tables', 'fenced_code', 'codehilite', 'toc'],
            tab_length=2)

        def blocks(lines):
            md.reset()
            soup = BeautifulSoup(md.convert('\n'.join(lines)), 'html.parser')
            return [str(node) for node in soup.contents if node.name]

        def is_code(block):
            return block.startswith('<pre') \
                or block.startswith('<div class="codehilite">')

        def is_list(block):
            return block.startswith('<ul') or block.startswith('<ol')

        def own_block(block):
            # The converter's block-level set, not a copy of it: a rescued
            # list holds items and nested lists and nothing else - a `<p>`
            # here means a raw block sat inside an item
            return any(md.is_block_level(tag.name) and tag.name not in ('ul', 'ol', 'li')
                       for tag in BeautifulSoup(block, 'html.parser').find_all(True))

        def code_text(block):
            return BeautifulSoup(block, 'html.parser').get_text().rstrip('\n')

        def detab(ln):
            # `CodeBlockProcessor` strips one `tab_length` of indentation; a
            # chunk line with less than that was never part of the code block
            # - a comment closing a generated index stands at column 0 and
            # renders as nothing on either side of the shift
            return ln[2:]

        lines = content.split('\n')
        chunks, chunk = [], []
        for i, line in enumerate(lines):
            if blank(line) or atx_re.match(line):
                if chunk:
                    chunks.append(chunk)
                chunk = []
            else:
                chunk.append(i)
        if chunk:
            chunks.append(chunk)

        before = None
        for chunk in chunks:
            first = lines[chunk[0]]
            indent = len(first) - len(first.lstrip(' '))
            m = self._CONVERTER_ITEM_RE.match(first)
            if not m or not 2 <= indent <= 3:
                continue
            # The outermost marker sets the column, not merely the first: a
            # list written raggedly still has one top level
            marker_cols = [len(lines[i]) - len(lines[i].lstrip(' '))
                           for i in chunk if self._CONVERTER_ITEM_RE.match(lines[i])]
            shift = min(marker_cols)
            if shift < 2:
                continue
            # Numbering the rendering cannot show. `LAZY_OL` writes no
            # `start`, so an ordered run that does not open at `1.` renders
            # from 1 and the author's numbers are lost - at any depth: a
            # nested run opening at `3.` is a new list too. And the converter
            # folds a marker of the other kind into the run already open at
            # its depth, so `- Intro` then `1. First` is three bullets. Depth
            # is the converter's bucket, not the column: its list processors
            # take up to `tab_length - 1` extra spaces as the same level, so a
            # marker one column deeper than its sibling is that sibling's run.
            open_runs, renumbered = {}, False
            for i in chunk:
                item = self._CONVERTER_ITEM_RE.match(lines[i])
                if not item:
                    continue
                if self._CONVERTER_ITEM_RE.match(lines[i][item.end():]):
                    # `- 1. Intro` is two items to the converter, one marker
                    # line here: items could no longer be counted against
                    # marker lines, and a level below it is lost unseen
                    renumbered = True
                    break
                depth = (len(lines[i]) - len(lines[i].lstrip(' ')) - shift) // 2
                kind = 'ordered' if item.group(1)[0].isdigit() else 'bullet'
                open_runs = {d: k for d, k in open_runs.items() if d <= depth}
                if depth in open_runs and open_runs[depth] != kind:
                    renumbered = True
                    break
                if kind == 'ordered' and depth not in open_runs \
                        and item.group(1)[:-1] != '1':
                    renumbered = True
                    break
                open_runs[depth] = kind
            if renumbered:
                continue
            if any(verbatim_re.search(lines[i]) for i in chunk):
                continue
            candidate = list(lines)
            for i in chunk:
                if len(lines[i]) - len(lines[i].lstrip(' ')) >= shift:
                    candidate[i] = lines[i][shift:]
            try:
                if before is None:
                    before = blocks(lines)
                after = blocks(candidate)
            except RecursionError:
                # A list nested past the parser's depth: the control renders
                # it (as a code block) where the shifted text cannot be
                # rendered at all, so the shift is refused, not raised
                continue
            if len(after) != len(before):
                continue
            changed = [(b, a) for b, a in zip(before, after) if b != a]
            if len(changed) != 1:
                continue
            was, now = changed[0]
            # A marker the converter cannot nest - four columns under its
            # parent, where the nesting unit is two - is not dropped but kept
            # as the parent's text, `- a1` literal inside the item: one list,
            # nothing block-level in it, and the author's level gone
            if is_code(was) and is_list(now) \
                    and now.count('<li') == len(marker_cols) \
                    and not own_block(now) \
                    and code_text(was) == '\n'.join(
                        detab(lines[i]) for i in chunk
                        if lines[i].startswith('  ')):
                lines, before = candidate, after
        return '\n'.join(lines)

    def preprocess_github_alerts(self, content: str, show_labels: bool = False) -> str:
        """Convert GitHub-style alerts to paragraphs with markers.

        Supports: NOTE, TIP, IMPORTANT, WARNING, CAUTION.
        Zero-width space markers (\u200b) around the type name allow
        post-processing in DOCX to apply colored styling.
        Preserves <br> tags and markdown links/formatting within alert content.
        When show_labels is False, the alert type label is hidden from output.

        The marker has to sit in one paragraph for the DOCX and HTML passes to
        find it, so the body's own structure is carried by explicit breaks
        instead: a source newline becomes one break, exactly as `nl2br` gives
        body text, and the bare `>` line that separates two paragraphs of an
        alert becomes two. Requiring `"> "` on every continuation line was
        what ended the match at that bare `>`, leaving the rest of the alert
        behind as a plain blockquote - one alert rendered as two boxes.
        """
        alert_pattern = (r'> \[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\] *\n'
                         r'((?:>.*\n?)*)')

        def add_breaks(text, count):
            """Append `count` breaks, counting one the author already wrote."""
            if TRAILING_BREAK_RE.search(text):
                count -= 1
            return text + '<br>' * max(0, count)

        def replace_alert(match):
            alert_type = match.group(1).upper()
            # Strip the blockquote prefix; an emptied line is the bare `>` that
            # separates two paragraphs of the alert
            alert_content, owed = '', 0
            for line in match.group(2).split('\n'):
                line = line.lstrip('> ').strip()
                if not line:
                    if alert_content:
                        owed = 2
                    continue
                if alert_content:
                    alert_content = add_breaks(alert_content, max(owed, 1))
                alert_content += line
                owed = 0

            # Zero-width space markers for DOCX post-processing
            marker = f'\u200b{alert_type}\u200b'
            if show_labels:
                return f'\n\n**{marker}:** {alert_content}\n\n'
            else:
                return f'\n\n{marker}{alert_content}\n\n'

        return re.sub(alert_pattern, replace_alert, content)

    def clean_alert_markers_from_html(self, html: str, show_labels: bool = False) -> str:
        """Remove zero-width space markers from HTML output.

        When show_labels is True, only strip the zero-width markers, keeping
        the alert type text visible. When False, strip both markers and the
        alert type text between them.
        """
        if show_labels:
            html = html.replace('\u200b', '')
        else:
            for alert_type in self.ALERT_COLORS:
                html = html.replace(f'\u200b{alert_type}\u200b', '')
            html = html.replace('\u200b', '')
        return html

    def drop_empty_table_headers(self, html: str) -> str:
        """Remove a table's header row when every header cell is empty.

        Markdown requires a header row, so a borderless image/layout grid is
        written with an empty one (`|  |  |  |`) and the converter emits a
        `<thead>` of blank `<th>`. DOCX and PDF drop that row; HTML must drop it
        too, or the same document renders with an extra blank banded strip in
        one format and not the others.

        Matched with a regex rather than a soup round-trip: this runs on the
        finished standalone document, and reparsing it would risk disturbing the
        embedded KaTeX/Pygments markup for a change this local.
        """
        if '<thead' not in html:
            return html

        def strip_tags(text):
            return unescape(re.sub(r'<[^>]+>', '', text)).strip()

        def replace(match):
            block = match.group(0)
            # Stripping tags erases an <img> along with them, so a header row
            # of pictures would read as blank - the DOCX and PDF paths keep
            # such a row for the same reason (image-on-top/caption-below grid)
            if re.search(r'<(img|svg|picture|video|object|iframe|input)\b',
                         block, re.I):
                return block
            cells = re.findall(r'<t[hd][^>]*>(.*?)</t[hd]>', block, re.S)
            if cells and not any(strip_tags(c) for c in cells):
                return ''
            return block

        # Only when a body row survives the delete, as on the other two paths.
        # The tag pattern admits attributes so it cannot silently miss a
        # `<thead class="...">` the guard above already let through.
        if not re.search(r'</thead>\s*<tbody[^>]*>\s*<tr', html, re.I | re.S):
            return html
        return re.sub(r'<thead\b[^>]*>.*?</thead\s*>\s*', replace, html,
                      flags=re.S | re.I)

    def wrap_html_tables(self, html: str) -> str:
        """Put every table in a horizontal scroll box.

        Cell wrapping cannot shrink a table below one character per column, so
        a table with very many columns would push the whole document sideways.
        The box has to be a wrapper element: `overflow-x` on the table itself
        only takes effect with `display: block`, which turns the grid into a
        shrink-to-fit anonymous box and voids its width.

        Only outermost tables are wrapped, and a table's attributes are kept:
        a non-greedy regex would close the wrapper at a nested table's end tag,
        and browser error recovery then stretches the box over the rest of the
        document, scrolling every element that follows. Comments are matched
        and skipped so a `<table>` written inside `<!-- ... -->` (markdown
        passes comments through verbatim) cannot open a wrapper that never
        closes.

        HTML export only - the DOCX and PDF paths parse this markup with
        htmldocx and have their own page fitting.
        """
        out = []
        depth = 0
        start = 0
        for tag in re.finditer(r'<!--.*?-->|<table\b[^>]*>|</table\s*>', html,
                               flags=re.DOTALL):
            if tag.group().startswith('<!--'):
                continue
            if tag.group().startswith('</'):
                if depth == 0:
                    continue
                depth -= 1
                if depth == 0:
                    out.append(html[start:tag.end()])
                    out.append('</div>')
                    start = tag.end()
            else:
                if depth == 0:
                    out.append(html[start:tag.start()])
                    out.append('<div class="table-scroll">')
                    start = tag.start()
                depth += 1
        out.append(html[start:])
        return ''.join(out)

    def style_html_alert_boxes(self, html: str, show_labels: bool = False) -> str:
        """Replace alert markers in HTML with styled div elements.

        Converts zero-width space markers into colored alert boxes with
        left border and background shading matching the DOCX alert style.

        HTML patterns from markdown conversion:
        - show_labels=True:  <p><strong>\u200bNOTE\u200b:</strong> content</p>
        - show_labels=False: <p>\u200bNOTE\u200bcontent</p>
        """
        for alert_type, colors in self.ALERT_COLORS.items():
            marker = f'\u200b{alert_type}\u200b'
            if marker not in html:
                continue

            border = f'#{colors["border"]}'
            shading = f'#{colors["shading"]}'

            if show_labels:
                # Pattern: <p><strong>\u200bTYPE\u200b:</strong> content</p>
                pattern = re.compile(
                    r'<p><strong>' + re.escape(marker) + r':?</strong>\s*(.*?)</p>',
                    re.DOTALL
                )

                def make_alert(m, _border=border, _shading=shading, _type=alert_type):
                    content = m.group(1).strip()
                    return (
                        f'<div style="border-left:4px solid {_border};'
                        f'background:{_shading};padding:12px 16px;'
                        f'border-radius:4px;margin:1em 0">'
                        f'<strong>{_type}:</strong> {content}</div>'
                    )
            else:
                # Pattern: <p>\u200bTYPE\u200bcontent</p>
                pattern = re.compile(
                    r'<p>' + re.escape(marker) + r'(.*?)</p>',
                    re.DOTALL
                )

                def make_alert(m, _border=border, _shading=shading):
                    content = m.group(1).strip()
                    return (
                        f'<div style="border-left:4px solid {_border};'
                        f'background:{_shading};padding:12px 16px;'
                        f'border-radius:4px;margin:1em 0">'
                        f'{content}</div>'
                    )

            html = pattern.sub(make_alert, html)

        # Clean any remaining stray markers
        html = html.replace('\u200b', '')
        return html

    def render_math_to_png(self, latex: str, width: int = 800, display: bool = False) -> str:
        """Render a LaTeX math expression to a PNG base64 data URI.

        Uses matplotlib.mathtext to render LaTeX to PNG with a transparent
        background, scaled so the result is approximately `width` pixels
        wide. The PNG's DPI metadata is then set so reportlab places it at
        the right physical size in the PDF (12pt tall inline, 16pt display).

        Args:
            latex: LaTeX math expression (without delimiters)
            width: Target pixel width for the rendered image
            display: If True, use larger font size for display math
        """
        from matplotlib.mathtext import math_to_image
        from matplotlib.font_manager import FontProperties
        from PIL import Image as PILImage

        fontsize = 16 if display else 12
        prop = FontProperties(size=fontsize)

        # matplotlib mathtext requires $ delimiters
        tex = f'${latex.strip()}$'

        # Probe render at a reference DPI to learn the equation's natural width,
        # then re-render at a DPI that lands close to the requested pixel width
        # (clamped so tiny inline equations don't get absurd DPIs).
        ref_dpi = 200
        probe = io.BytesIO()
        math_to_image(tex, probe, dpi=ref_dpi, format='png', prop=prop)
        probe.seek(0)
        probe_w = max(1, PILImage.open(probe).width)
        render_dpi = ref_dpi * width / probe_w
        render_dpi = max(72.0, min(1200.0, render_dpi))

        buf = io.BytesIO()
        math_to_image(tex, buf, dpi=render_dpi, format='png', prop=prop)
        buf.seek(0)

        # Fix DPI metadata so reportlab sizes the image at fontsize points tall.
        img = PILImage.open(buf)
        target_height_inches = fontsize / 72  # points to inches
        effective_dpi = img.height / target_height_inches
        corrected_buf = io.BytesIO()
        img.save(corrected_buf, format='png', dpi=(effective_dpi, effective_dpi))
        corrected_buf.seek(0)

        b64 = base64.b64encode(corrected_buf.read()).decode('ascii')
        return f'data:image/png;base64,{b64}'

    def replace_math_with_images(self, content: str, width: int = 800) -> str:
        """Replace LaTeX math delimiters with rendered PNG images (PDF export).

        Protects code blocks (fenced and inline) before processing.
        Matches $$...$$ (display) and $...$ (inline) while avoiding
        false positives on currency amounts like $100.

        Args:
            content: Markdown content with LaTeX math expressions
            width: Target pixel width for each rendered math image
        """
        # Protect code blocks by replacing with placeholders
        code_placeholders = []

        # Protect fenced code blocks (```...```)
        def protect_fenced(match):
            code_placeholders.append(match.group(0))
            return f'[[MATH_CODE_BLOCK_{len(code_placeholders) - 1}]]'

        content = re.sub(r'```[\s\S]*?```', protect_fenced, content)

        # Protect inline code (`...`)
        def protect_inline(match):
            code_placeholders.append(match.group(0))
            return f'[[MATH_CODE_BLOCK_{len(code_placeholders) - 1}]]'

        content = re.sub(r'`[^`]+`', protect_inline, content)

        # Protect escaped dollars from being matched as math delimiters
        # Handle \\$ (double backslash + dollar, common in markdown/LaTeX) first,
        # then \$ (single backslash + dollar). Both render as literal $ in output.
        content = content.replace('\\\\$', '[[ESCAPED_DOLLAR]]')
        content = content.replace('\\$', '[[ESCAPED_DOLLAR]]')

        # Replace display math $$...$$ first (greedy within single expression)
        def replace_display(match):
            latex = match.group(1)
            try:
                data_uri = self.render_math_to_png(latex, width=width, display=True)
                return f'\n\n<div style="text-align:center"><img src="{data_uri}" alt="{latex}" style="max-width:100%"></div>\n\n'
            except Exception:
                return match.group(0)

        content = re.sub(r'\$\$(.+?)\$\$', replace_display, content, flags=re.DOTALL)

        # Replace inline math $...$ (require non-space after opening and before closing $)
        def replace_inline(match):
            latex = match.group(1)
            try:
                data_uri = self.render_math_to_png(latex, width=width, display=False)
                return f'<img src="{data_uri}" alt="{latex}" style="vertical-align:middle">'
            except Exception:
                return match.group(0)

        content = re.sub(r'(?<!\$)\$(\S(?:[^$]*?\S)?)\$(?!\$)', replace_inline, content)

        # Restore escaped dollars as literal $ (both \$ and \\$ mean literal dollar)
        content = content.replace('[[ESCAPED_DOLLAR]]', '$')

        # Restore code blocks
        for i, block in enumerate(code_placeholders):
            content = content.replace(f'[[MATH_CODE_BLOCK_{i}]]', block)

        return content

    def latex_to_omml(self, latex: str):
        """Convert LaTeX to OMML element for Word documents.

        Pipeline: LaTeX -> MathML (latex2mathml) -> OMML (XSLT mml2omml.xsl)
        Returns an lxml Element with the oMath OMML structure.

        Post-processes nary operators (sum, prod, int) whose <e/> element
        is empty - moves all following siblings into <e> so Word renders
        the summand/integrand inside the operator rather than showing
        a dashed placeholder square.
        """
        import latex2mathml.converter
        from lxml import etree

        mathml = latex2mathml.converter.convert(latex)
        tree = etree.fromstring(mathml.encode('utf-8'))

        xsl_path = Path(__file__).parent / 'mml2omml.xsl'
        xslt = etree.parse(str(xsl_path))
        transform = etree.XSLT(xslt)
        omml = transform(tree)
        root = omml.getroot()

        # Fix nary operators with empty <e/> by moving following siblings in
        OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
        for nary in root.iter(f'{{{OMML_NS}}}nary'):
            e_elem = nary.find(f'{{{OMML_NS}}}e')
            if e_elem is not None and len(e_elem) == 0 and (e_elem.text is None or not e_elem.text.strip()):
                parent = nary.getparent()
                if parent is not None:
                    siblings = list(parent)
                    nary_idx = siblings.index(nary)
                    # Move all elements after nary into <e>
                    for sibling in siblings[nary_idx + 1:]:
                        parent.remove(sibling)
                        e_elem.append(sibling)

        return root

    def replace_math_with_markers(self, content: str):
        """Replace LaTeX math delimiters with text markers for DOCX post-processing.

        Instead of rendering math as images, inserts zero-width joiner markers
        that are later replaced with native OMML equations after htmldocx processing.

        Returns (content, inline_math_list, display_math_list).
        """
        inline_math = []
        display_math = []

        # Protect code blocks by replacing with placeholders
        code_placeholders = []

        def protect_fenced(match):
            code_placeholders.append(match.group(0))
            return f'[[MATH_CODE_BLOCK_{len(code_placeholders) - 1}]]'

        content = re.sub(r'```[\s\S]*?```', protect_fenced, content)

        def protect_inline(match):
            code_placeholders.append(match.group(0))
            return f'[[MATH_CODE_BLOCK_{len(code_placeholders) - 1}]]'

        content = re.sub(r'`[^`]+`', protect_inline, content)

        # Protect escaped dollars from being matched as math delimiters
        # Handle \\$ (double backslash + dollar, common in markdown/LaTeX) first,
        # then \$ (single backslash + dollar). Both render as literal $ in output.
        content = content.replace('\\\\$', '[[ESCAPED_DOLLAR]]')
        content = content.replace('\\$', '[[ESCAPED_DOLLAR]]')

        # Replace display math $$...$$ first
        def replace_display(match):
            latex = match.group(1)
            idx = len(display_math)
            display_math.append(latex)
            return f'\n\n\u200dMATH_DISPLAY_{idx}\u200d\n\n'

        content = re.sub(r'\$\$(.+?)\$\$', replace_display, content, flags=re.DOTALL)

        # Replace inline math $...$
        def replace_inline(match):
            latex = match.group(1)
            idx = len(inline_math)
            inline_math.append(latex)
            return f'\u200dMATH_INLINE_{idx}\u200d'

        content = re.sub(r'(?<!\$)\$(\S(?:[^$]*?\S)?)\$(?!\$)', replace_inline, content)

        # Restore escaped dollars as literal $ (both \$ and \\$ mean literal dollar)
        content = content.replace('[[ESCAPED_DOLLAR]]', '$')

        # Restore code blocks
        for i, block in enumerate(code_placeholders):
            content = content.replace(f'[[MATH_CODE_BLOCK_{i}]]', block)

        return content, inline_math, display_math

    def merge_inline_math_omml(self, document, inline_math, display_math):
        """Post-process DOCX to replace math markers with native OMML equations.

        Scans all paragraph runs for MATH_INLINE_N and MATH_DISPLAY_N markers,
        splits runs at marker boundaries, and inserts OMML elements.
        """
        from docx.oxml.ns import qn
        from lxml import etree
        import copy

        OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

        for paragraph in document.paragraphs:
            full_text = paragraph.text

            # Handle display math (marker in its own paragraph)
            for idx, latex in enumerate(display_math):
                marker = f'\u200dMATH_DISPLAY_{idx}\u200d'
                if marker in full_text:
                    try:
                        omml_elem = self.latex_to_omml(latex)
                        # Create oMathPara wrapper for display (centered) math
                        omath_para = etree.SubElement(
                            paragraph._p, f'{{{OMML_NS}}}oMathPara'
                        )
                        omath_para.append(omml_elem)
                        # Remove all existing runs (marker text). The link
                        # walk, not `paragraph.runs`: a marker inside a link
                        # label sits in a run under `w:hyperlink`, which that
                        # property never returns, and it would print raw.
                        for run in self.docx_paragraph_runs(paragraph):
                            run._r.getparent().remove(run._r)
                    except Exception:
                        # Fallback: leave marker text (will show raw LaTeX)
                        for run in self.docx_paragraph_runs(paragraph):
                            if marker in run.text:
                                run.text = run.text.replace(marker, latex)

            # Handle inline math (marker within text runs)
            for idx, latex in enumerate(inline_math):
                marker = f'\u200dMATH_INLINE_{idx}\u200d'
                if marker not in full_text:
                    continue

                # Find the run containing the marker
                for run in self.docx_paragraph_runs(paragraph):
                    if marker not in run.text:
                        continue

                    try:
                        omml_elem = self.latex_to_omml(latex)
                    except Exception:
                        run.text = run.text.replace(marker, latex)
                        continue

                    parts = run.text.split(marker, 1)
                    parent = run._r.getparent()
                    run_index = list(parent).index(run._r)

                    # Set before-text in current run
                    run.text = parts[0]

                    # Insert OMML element after current run
                    parent.insert(run_index + 1, omml_elem)

                    # Create after-text run if needed
                    if parts[1]:
                        after_run = copy.deepcopy(run._r)
                        # CT_R's own setter, not a reach for a w:t that may
                        # not be there: when the marker opens the run, the
                        # `run.text = parts[0]` above cleared its content, so
                        # the copy has no w:t at all and .find returned None
                        after_run.text = parts[1]
                        parent.insert(run_index + 2, after_run)

                    break  # Only one marker per run expected

    def style_docx_alert_boxes(self, document, show_labels: bool = False) -> list:
        """Replace alert paragraphs with styled single-cell tables.

        Scans for zero-width space markers inserted by preprocess_github_alerts()
        and wraps each alert in a one-cell table with colored left border,
        background shading, and cell margins for padding control.
        Moves the original paragraph XML (preserving hyperlinks, bold, etc.)
        into the table cell rather than rebuilding it.

        A callout the author drew by hand - a ``<div>`` carrying a border or a
        background, marked by restructure_html_for_docx() - takes the same box
        in its own colours, so the two kinds of callout in a document are one
        construct downstream and the PDF pass reads them both the same way.

        Returns the ``w:tbl`` elements created, which carry their own styling
        and must be left out of the content-table pass.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy

        # Collect paragraphs to replace (can't modify while iterating)
        alert_tables = []
        replacements = []
        for paragraph in document.paragraphs:
            text = paragraph.text
            box = self._BOX_MARKER_RE.search(text)
            if box is not None:
                # No alert type: the colours are the author's own, read off the
                # div's CSS, and there is no label to strip
                replacements.append(([paragraph], '', {
                    'border': box.group(1).upper(),
                    'shading': box.group(2).upper(),
                }))
                continue
            if self._BOX_MORE_MARKER in text and replacements and not replacements[-1][1]:
                # The rest of a box already open: one table around the whole
                # run of paragraphs, not one table each with a spacer between
                replacements[-1][0].append(paragraph)
                continue
            for alert_type, colors in self.ALERT_COLORS.items():
                marker = f'\u200b{alert_type}\u200b'
                if marker not in text:
                    continue
                replacements.append(([paragraph], alert_type, colors))
                break

        for paragraphs, alert_type, colors in replacements:
            first = paragraphs[0]
            parent = first._p.getparent()

            # Clean zero-width markers from runs; strip type label only when hidden
            for run in (r for paragraph in paragraphs for r in paragraph.runs):
                text = run.text
                if not alert_type:
                    if '⁣' in text:
                        run.text = self._BOX_MARKER_RE.sub('', text).replace(
                            self._BOX_MORE_MARKER, '')
                    continue
                if '\u200b' in text:
                    cleaned = text.replace('\u200b', '')
                    if not show_labels:
                        for at in self.ALERT_COLORS:
                            if cleaned.startswith(at):
                                cleaned = cleaned[len(at):]
                                if cleaned.startswith(':'):
                                    cleaned = cleaned[1:]
                                cleaned = cleaned.lstrip()
                                break
                    run.text = cleaned

            # Build the one-cell table
            tbl = OxmlElement('w:tbl')

            # Table properties: 100% width via percentage
            tblPr = OxmlElement('w:tblPr')
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')
            tblPr.append(tblW)

            # Table borders: only left border colored, others invisible
            tblBorders = OxmlElement('w:tblBorders')
            for side in ['top', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tblBorders.append(border)
            left_border = OxmlElement('w:left')
            left_border.set(qn('w:val'), 'single')
            left_border.set(qn('w:sz'), '24')  # 3pt
            left_border.set(qn('w:space'), '0')
            left_border.set(qn('w:color'), colors['border'])
            tblBorders.append(left_border)
            tblPr.append(tblBorders)

            # Cell margins for internal padding
            tblCellMar = OxmlElement('w:tblCellMar')
            for side, val in [('top', '80'), ('bottom', '80'),
                              ('left', '180'), ('right', '120')]:
                margin = OxmlElement(f'w:{side}')
                margin.set(qn('w:w'), val)
                margin.set(qn('w:type'), 'dxa')
                tblCellMar.append(margin)
            tblPr.append(tblCellMar)

            tbl.append(tblPr)

            # Table grid (single column)
            tblGrid = OxmlElement('w:tblGrid')
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), '9360')
            tblGrid.append(gridCol)
            tbl.append(tblGrid)

            # Single row, single cell
            tr = OxmlElement('w:tr')
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), '5000')
            tcW.set(qn('w:type'), 'pct')
            tcPr.append(tcW)

            # Cell shading
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), colors['shading'])
            tcPr.append(shd)
            tc.append(tcPr)

            # Move the original paragraphs into the cell (preserves all
            # formatting, hyperlinks, bold, italic, line breaks, etc.)
            for paragraph in paragraphs:
                tc.append(copy.deepcopy(paragraph._p))
            tr.append(tc)
            tbl.append(tr)

            # Insert table then a spacer paragraph where the box began
            first._p.addnext(tbl)
            spacer = OxmlElement('w:p')
            tbl.addnext(spacer)

            # Remove the originals
            for paragraph in paragraphs:
                parent.remove(paragraph._p)
            alert_tables.append(tbl)

        return alert_tables

    @staticmethod
    def docx_run_shading(run) -> str:
        """Six-hex fill of a run's ``w:shd``, or '' when it carries none."""
        from docx.oxml.ns import qn

        rPr = run._r.find(qn('w:rPr'))
        shd = rPr.find(qn('w:shd')) if rPr is not None else None
        fill = shd.get(qn('w:fill')) if shd is not None else None
        if not fill or fill == 'auto' or len(fill) != 6:
            return ''
        return fill

    @staticmethod
    def docx_paragraph_runs(paragraph):
        """Every run under ``paragraph``, the ones inside a link included.

        ``Paragraph.runs`` is ``./w:r`` - direct children only - so a run Word
        nested in ``<w:hyperlink>`` is invisible to it. Anything that walks a
        paragraph run by run therefore skips the link: a sentinel marker left
        on it leaks into the visible text, and a rebuild from those runs drops
        the link label outright. The element walk reaches both, and the Run
        wrapper keeps the callers unchanged.
        """
        from docx.oxml.ns import qn
        from docx.text.run import Run

        return [Run(r, paragraph) for r in paragraph._p.iter(qn('w:r'))]

    @classmethod
    def docx_symbol_font(cls, char: str) -> str:
        """Font a DOCX run must name to draw ``char``, or '' for the body font.

        Only characters the body face has no glyph for get an answer. Word does
        substitute on its own, but it picks from whatever is installed, so the
        same document can come out with a star in one machine's fallback face
        and a hollow box on the next. Naming the font settles it.
        """
        if char in (cls.TASK_CHECKBOX_DONE, cls.TASK_CHECKBOX_OPEN):
            return cls.DOCX_CHECKBOX_FONT
        if char in cls.DOCX_SYMBOL_CHARS:
            return cls.DOCX_SYMBOL_FONT
        code = ord(char)
        if any(low <= code <= high for low, high in cls.DOCX_SYMBOL_RANGES):
            return cls.DOCX_SYMBOL_FONT
        return ''

    def style_docx_symbol_runs(self, document):
        """Name a font on every run character the DOCX body face cannot draw.

        Each stretch of symbol text is split into its own run and only that run
        gets the font, so the surrounding words keep the body typeface. Task
        checkboxes take MS Gothic and every other symbol DOCX_SYMBOL_FONT - see
        docx_symbol_font() for why the font is named rather than left to Word.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy

        def set_font(r_elem, name):
            rPr = r_elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r_elem.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                # rFonts must follow w:rStyle in the CT_RPr sequence
                rStyle = rPr.find(qn('w:rStyle'))
                if rStyle is not None:
                    rStyle.addnext(rFonts)
                else:
                    rPr.insert(0, rFonts)
            for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
                rFonts.set(qn(attr), name)

        def font_of(text, index):
            """Font for the character at ``index``, given what follows it."""
            char = text[index]
            if char in self.DOCX_VARIATION_SELECTORS:
                # No glyph of its own - it modifies the character before it
                return None
            if text[index + 1:index + 2] == self.DOCX_EMOJI_SELECTOR:
                # The author asked for the emoji face. Word's own fallback
                # draws it in colour; anything named here draws it flat.
                return ''
            return self.docx_symbol_font(char)

        def segment(text):
            """Split text into stretches of one font each, body font as ''."""
            out, buf, buf_font = [], '', ''
            for index in range(len(text)):
                font = font_of(text, index)
                if font is not None and buf and font != buf_font:
                    out.append((buf, buf_font))
                    buf = ''
                buf += text[index]
                if font is not None:
                    buf_font = font
            if buf:
                out.append((buf, buf_font))
            return out

        def text_node(text):
            node = OxmlElement('w:t')
            node.text = text
            node.set(qn('xml:space'), 'preserve')
            return node

        for r_elem in list(document.element.body.iter(qn('w:r'))):
            children = [c for c in r_elem if c.tag != qn('w:rPr')]
            texts = [c.text or '' for c in children if c.tag == qn('w:t')]
            if not any(self.docx_symbol_font(ch) for t in texts for ch in t):
                continue

            # Rebuild the whole run as a sequence of runs, one font each.
            # Content that carries no glyph - a line break, a tab, a picture -
            # keeps its place in that sequence: left behind on the original run
            # it would jump ahead of every split-off segment, moving a break to
            # the middle of the line and welding the next line onto it.
            groups, group, group_font = [], [], ''
            for child in children:
                if child.tag != qn('w:t') or not child.text:
                    group.append(child)
                    continue
                for seg_text, seg_font in segment(child.text):
                    if seg_font != group_font and group:
                        groups.append((group_font, group))
                        group = []
                    group_font = seg_font
                    group.append(text_node(seg_text))
            groups.append((group_font, group))

            # Formatting-only template: carries the run's rPr but none of its
            # content, so each cloned run starts empty and takes exactly the
            # nodes its group holds.
            template = copy.deepcopy(r_elem)
            for child in list(template):
                if child.tag != qn('w:rPr'):
                    template.remove(child)

            for child in children:
                r_elem.remove(child)
            first_font, first_nodes = groups[0]
            for node in first_nodes:
                r_elem.append(node)
            if first_font:
                set_font(r_elem, first_font)

            anchor = r_elem
            for group_font, group_nodes in groups[1:]:
                new_r = copy.deepcopy(template)
                for node in group_nodes:
                    new_r.append(node)
                if group_font:
                    set_font(new_r, group_font)
                anchor.addnext(new_r)
                anchor = new_r

    def remove_empty_paragraphs_before_images(self, document):
        """Remove empty paragraphs that immediately precede image paragraphs.

        htmldocx inserts blank paragraphs before images. This removes them
        to eliminate unnecessary vertical space above images.
        """
        from docx.oxml.ns import qn
        body = document.element.body
        paragraphs = body.findall(qn('w:p'))
        to_remove = []
        for i in range(len(paragraphs) - 1):
            current = paragraphs[i]
            nxt = paragraphs[i + 1]
            # Check if current paragraph is empty (no text content)
            if current.text and current.text.strip():
                continue
            runs = current.findall('.//' + qn('w:t'))
            if any(r.text and r.text.strip() for r in runs):
                continue
            # Check current has no image itself
            if (current.findall('.//' + qn('w:drawing')) or
                    current.findall('.//{urn:schemas-microsoft-com:vml}imagedata')):
                continue
            # Check next paragraph contains an image
            if (nxt.findall('.//' + qn('w:drawing')) or
                    nxt.findall('.//{urn:schemas-microsoft-com:vml}imagedata')):
                to_remove.append(current)
        for p in to_remove:
            p.getparent().remove(p)

    # Unicode sentinels for bookmark markers injected into HTML.
    # U+2063 (INVISIBLE SEPARATOR) is unlikely to appear in normal text and
    # survives htmldocx as literal text that we can regex-match later.
    _BOOKMARK_MARKER_RE = re.compile(r'⁣BM:([A-Za-z0-9_\-:.]+)⁣')
    _SAFE_BOOKMARK_ID_RE = re.compile(r'^[A-Za-z_][A-Za-z0-9_\-:.]*$')

    def inject_anchor_markers(self, html: str) -> str:
        """Inject sentinel markers for every HTML element with an id attribute.

        For ``<span id="D119">D119 ...</span>`` injects
        ``<span id="D119">⁣BM:D119⁣D119 ...</span>``. The marker
        survives htmldocx as plain text and is later converted into a Word
        bookmark by apply_anchor_bookmarks().
        """
        def replace(match):
            open_tag = match.group(0)
            anchor_id = match.group(1)
            if not self._SAFE_BOOKMARK_ID_RE.match(anchor_id):
                return open_tag
            return f'{open_tag}⁣BM:{anchor_id}⁣'

        return re.sub(r'<[^>]*?\bid="([^"]+)"[^>]*?>', replace, html)

    def apply_anchor_bookmarks(self, document):
        """Convert bookmark markers and external hash hyperlinks into internal anchors.

        Works in two passes on the finished DOCX tree:

        1. Replaces each ``⁣BM:NAME⁣`` marker in a run's text with a
           ``w:bookmarkStart`` / ``w:bookmarkEnd`` pair, preserving surrounding text.
        2. Rewrites every ``<w:hyperlink r:id="...">`` that points to a hash-prefix
           external relationship into ``<w:hyperlink w:anchor="NAME">`` and drops
           the now-unused relationship. htmldocx always writes anchor links as
           external, so this rescues them.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy

        body = document.element.body
        bookmark_id = 0

        for run_elem in list(body.iter(qn('w:r'))):
            t_elem = run_elem.find(qn('w:t'))
            if t_elem is None or not t_elem.text:
                continue
            text = t_elem.text
            match = self._BOOKMARK_MARKER_RE.search(text)
            if not match:
                continue

            anchor_name = match.group(1)
            before = text[:match.start()]
            after = text[match.end():]

            parent = run_elem.getparent()
            run_index = list(parent).index(run_elem)

            t_elem.text = before

            bm_start = OxmlElement('w:bookmarkStart')
            bm_start.set(qn('w:id'), str(bookmark_id))
            bm_start.set(qn('w:name'), anchor_name)
            bm_end = OxmlElement('w:bookmarkEnd')
            bm_end.set(qn('w:id'), str(bookmark_id))
            bookmark_id += 1

            parent.insert(run_index + 1, bm_start)
            parent.insert(run_index + 2, bm_end)

            if after:
                after_run = copy.deepcopy(run_elem)
                after_t = after_run.find(qn('w:t'))
                if after_t is not None:
                    after_t.text = after
                parent.insert(run_index + 3, after_run)

        HYPERLINK_RELTYPE = (
            'http://schemas.openxmlformats.org/officeDocument/2006/'
            'relationships/hyperlink'
        )
        rels = document.part.rels
        anchor_rels = {
            rel_id: rel.target_ref[1:]
            for rel_id, rel in rels.items()
            if rel.reltype == HYPERLINK_RELTYPE and rel.target_ref.startswith('#')
        }
        if not anchor_rels:
            return

        r_id_qn = qn('r:id')
        anchor_qn = qn('w:anchor')
        for hyperlink in body.iter(qn('w:hyperlink')):
            ref = hyperlink.get(r_id_qn)
            if ref in anchor_rels:
                del hyperlink.attrib[r_id_qn]
                hyperlink.set(anchor_qn, anchor_rels[ref])

        for rel_id in anchor_rels:
            del rels[rel_id]

    _BLOCKQUOTE_MARKER_RE = re.compile(r'⁣BQ:(\d+)⁣')
    _PILL_MARKER_RE = re.compile(r'⁣PILL:([0-9A-Fa-f]{6})⁣')
    #: Bar and fill of a callout box the author drew by hand - a ``<div>``
    #: given a border or a background - carried from restructure_html_for_docx
    #: to style_docx_alert_boxes, which boxes it the way it boxes an alert.
    _BOX_MARKER_RE = re.compile(r'⁣BOX:([0-9A-Fa-f]{6}):([0-9A-Fa-f]{6})⁣')

    #: A paragraph belonging to the box opened directly above it. A box is one
    #: table and a ``<div>`` holding blocks is one paragraph per block, so
    #: every block after the first says which box it is in: two asides written
    #: one after the other in the same colours are otherwise indistinguishable
    #: from one aside holding two blocks.
    _BOX_MORE_MARKER = '⁣BOX+⁣'

    # CSS named colours htmldocx fails to parse (it only understands hex),
    # so a `color: green` span renders black. Mapped to hex here. The whole
    # CSS list rather than a shortlist of likely names: a name that is missing
    # renders black with nothing to say it was ever understood.
    _CSS_NAMED_COLORS = {
    'aliceblue': 'F0F8FF', 'antiquewhite': 'FAEBD7', 'aqua': '00FFFF',
    'aquamarine': '7FFFD4', 'azure': 'F0FFFF', 'beige': 'F5F5DC',
    'bisque': 'FFE4C4', 'black': '000000', 'blanchedalmond': 'FFEBCD',
    'blue': '0000FF', 'blueviolet': '8A2BE2', 'brown': 'A52A2A',
    'burlywood': 'DEB887', 'cadetblue': '5F9EA0', 'chartreuse': '7FFF00',
    'chocolate': 'D2691E', 'coral': 'FF7F50', 'cornflowerblue': '6495ED',
    'cornsilk': 'FFF8DC', 'crimson': 'DC143C', 'cyan': '00FFFF',
    'darkblue': '00008B', 'darkcyan': '008B8B', 'darkgoldenrod': 'B8860B',
    'darkgray': 'A9A9A9', 'darkgreen': '006400', 'darkgrey': 'A9A9A9',
    'darkkhaki': 'BDB76B', 'darkmagenta': '8B008B',
    'darkolivegreen': '556B2F', 'darkorange': 'FF8C00',
    'darkorchid': '9932CC', 'darkred': '8B0000', 'darksalmon': 'E9967A',
    'darkseagreen': '8FBC8F', 'darkslateblue': '483D8B',
    'darkslategray': '2F4F4F', 'darkslategrey': '2F4F4F',
    'darkturquoise': '00CED1', 'darkviolet': '9400D3', 'deeppink': 'FF1493',
    'deepskyblue': '00BFFF', 'dimgray': '696969', 'dimgrey': '696969',
    'dodgerblue': '1E90FF', 'firebrick': 'B22222', 'floralwhite': 'FFFAF0',
    'forestgreen': '228B22', 'fuchsia': 'FF00FF', 'gainsboro': 'DCDCDC',
    'ghostwhite': 'F8F8FF', 'gold': 'FFD700', 'goldenrod': 'DAA520',
    'gray': '808080', 'green': '008000', 'greenyellow': 'ADFF2F',
    'grey': '808080', 'honeydew': 'F0FFF0', 'hotpink': 'FF69B4',
    'indianred': 'CD5C5C', 'indigo': '4B0082', 'ivory': 'FFFFF0',
    'khaki': 'F0E68C', 'lavender': 'E6E6FA', 'lavenderblush': 'FFF0F5',
    'lawngreen': '7CFC00', 'lemonchiffon': 'FFFACD', 'lightblue': 'ADD8E6',
    'lightcoral': 'F08080', 'lightcyan': 'E0FFFF',
    'lightgoldenrodyellow': 'FAFAD2', 'lightgray': 'D3D3D3',
    'lightgreen': '90EE90', 'lightgrey': 'D3D3D3', 'lightpink': 'FFB6C1',
    'lightsalmon': 'FFA07A', 'lightseagreen': '20B2AA',
    'lightskyblue': '87CEFA', 'lightslategray': '778899',
    'lightslategrey': '778899', 'lightsteelblue': 'B0C4DE',
    'lightyellow': 'FFFFE0', 'lime': '00FF00', 'limegreen': '32CD32',
    'linen': 'FAF0E6', 'magenta': 'FF00FF', 'maroon': '800000',
    'mediumaquamarine': '66CDAA', 'mediumblue': '0000CD',
    'mediumorchid': 'BA55D3', 'mediumpurple': '9370DB',
    'mediumseagreen': '3CB371', 'mediumslateblue': '7B68EE',
    'mediumspringgreen': '00FA9A', 'mediumturquoise': '48D1CC',
    'mediumvioletred': 'C71585', 'midnightblue': '191970',
    'mintcream': 'F5FFFA', 'mistyrose': 'FFE4E1', 'moccasin': 'FFE4B5',
    'navajowhite': 'FFDEAD', 'navy': '000080', 'oldlace': 'FDF5E6',
    'olive': '808000', 'olivedrab': '6B8E23', 'orange': 'FFA500',
    'orangered': 'FF4500', 'orchid': 'DA70D6', 'palegoldenrod': 'EEE8AA',
    'palegreen': '98FB98', 'paleturquoise': 'AFEEEE',
    'palevioletred': 'DB7093', 'papayawhip': 'FFEFD5', 'peachpuff': 'FFDAB9',
    'peru': 'CD853F', 'pink': 'FFC0CB', 'plum': 'DDA0DD',
    'powderblue': 'B0E0E6', 'purple': '800080', 'rebeccapurple': '663399',
    'red': 'FF0000', 'rosybrown': 'BC8F8F', 'royalblue': '4169E1',
    'saddlebrown': '8B4513', 'salmon': 'FA8072', 'sandybrown': 'F4A460',
    'seagreen': '2E8B57', 'seashell': 'FFF5EE', 'sienna': 'A0522D',
    'silver': 'C0C0C0', 'skyblue': '87CEEB', 'slateblue': '6A5ACD',
    'slategray': '708090', 'slategrey': '708090', 'snow': 'FFFAFA',
    'springgreen': '00FF7F', 'steelblue': '4682B4', 'tan': 'D2B48C',
    'teal': '008080', 'thistle': 'D8BFD8', 'tomato': 'FF6347',
    'turquoise': '40E0D0', 'violet': 'EE82EE', 'wheat': 'F5DEB3',
    'white': 'FFFFFF', 'whitesmoke': 'F5F5F5', 'yellow': 'FFFF00',
    'yellowgreen': '9ACD32',
    }

    #: ``rgb()`` / ``rgba()``, in the comma notation every devtools colour
    #: copy produces and the space notation CSS Color 4 added.
    _CSS_COLOR_FN_RE = re.compile(r'^(rgba?|hsla?)\(([^)]*)\)$')

    #: A colour token anywhere in a shorthand value. A function call is matched
    #: whole and first: its own commas and spaces would otherwise tear it into
    #: pieces, and the pieces of `url(assets/red-banner.png)` and
    #: `var(--blue-500)` include a bare colour NAME that the browser paints
    #: nowhere - one names a file and the other a custom property. `rgb()` and
    #: `rgba()` are function calls too, and resolve as themselves.
    _CSS_COLOR_TOKEN_RE = re.compile(
        r'[a-z-]*\([^()]*\)|#[0-9A-Fa-f]+|[a-z]+', re.I)

    #: The border properties that put a line on the page. `border-collapse`,
    #: `border-spacing`, `border-image` and `border-radius` open with the same
    #: word and draw none of them, so the set is written out rather than
    #: matched on the prefix.
    _CSS_BORDER_PROPS = frozenset(
        f'border{side}{part}'
        for side in ('', '-top', '-right', '-bottom', '-left')
        for part in ('', '-style', '-width', '-color'))

    #: A length that measures zero, in any unit or none. A border given one
    #: draws exactly what `none` draws.
    #: The four sides a border declaration can name, and the side a shorthand
    #: that names none of them sets - all of them.
    _CSS_BORDER_SIDES = ('top', 'right', 'bottom', 'left')

    #: The `border-style` keywords. A shorthand's three parts are order-free,
    #: so the style is whichever of its tokens names one of these; `none` and
    #: `hidden` are the two that draw no line, and a side with no style at all
    #: draws none either - `none` is what CSS starts from.
    _CSS_BORDER_STYLES = frozenset((
        'none', 'hidden', 'dotted', 'dashed', 'solid', 'double',
        'groove', 'ridge', 'inset', 'outset'))

    #: The three widths CSS names instead of measuring, and the shape of one it
    #: measures - a number in any unit, or in none.
    _CSS_BORDER_WIDTHS = frozenset(('thin', 'medium', 'thick'))
    _CSS_LENGTH_RE = re.compile(r'^-?\d*\.?\d+[a-z%]*$')

    #: Fill a ``<mark>`` gets when the author declared none. The HTML export
    #: writes no rule for ``mark``, so it takes the browser default - which is
    #: this same yellow, and the two formats agree without being told to.
    _HTML_MARK_FILL = _CSS_NAMED_COLORS['yellow']

    #: Inline tags htmldocx has no handler for, and the tag it does handle that
    #: means the same thing. ``<ins>`` is underline and ``<del>`` strike-through
    #: because that is how a browser draws them.
    _HTML_TAG_ALIASES = {
        'del': 's', 'strike': 's', 'ins': 'u', 'kbd': 'code', 'samp': 'code',
    }

    #: CSS declarations htmldocx drops, and the tag that carries the same
    #: meaning through it. Keyed by property; the value maps a declared CSS
    #: value to a tag name, or to '' when the declaration asks for nothing
    #: (``font-weight: normal``, ``text-decoration: none``). 600 is the weight
    #: CSS itself calls bold; a lambda in a class body cannot read a class
    #: constant, so the number stays inline.
    _CSS_STYLE_TAGS = {
        'font-weight': lambda v: 'b' if v in ('bold', 'bolder') or (
            v.isdigit() and int(v) >= 600) else '',
        'font-style': lambda v: 'i' if v in ('italic', 'oblique') else '',
        'text-decoration': lambda v: ''.join(
            t for t, k in (('u', 'underline'), ('s', 'line-through')) if k in v),
    }

    #: Elements that open a block of their own. A ``<div>`` wrapping any of
    #: these is a container, not a paragraph - see restructure_html_for_docx.
    _HTML_BLOCK_TAGS = ('p', 'div', 'table', 'ul', 'ol', 'pre', 'blockquote',
                        'hr', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6')

    #: Elements an inline tag may legally wrap. Everything else takes the
    #: emphasis around its CONTENTS instead: htmldocx reaches a cell with a
    #: direct-child walk, so a ``<b>`` between ``<tr>`` and ``<td>`` costs the
    #: cell - the row comes out short, or the export dies outright.
    _HTML_INLINE_TAGS = ('span', 'a', 'code', 'b', 'i', 'u', 's', 'em',
                         'strong', 'sub', 'sup', 'small', 'abbr', 'q',
                         'cite', 'time', 'var')

    #: Blocks that become exactly one body paragraph, so a marker put in one
    #: lands in the paragraph the box is built around. A ``<div>`` holding
    #: anything else takes no box: a table and a code block bring a structure
    #: of their own that stays outside the box the text is moved into, the way
    #: a picture does, and a list arrives as one paragraph per item, which the
    #: PDF rebuild redraws as callout body text with the bullet gone.
    _HTML_BOXABLE_BLOCKS = ('p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6')

    #: Elements that already own a run of text, so a <div> inside one is in
    #: inline position however it was written. Renaming it to <p> there nests
    #: a block in a block: htmldocx ends the list item, the heading or the
    #: cell early and the rest of the line becomes a paragraph of its own.
    _HTML_TEXT_HOLDERS = ('p', 'li', 'td', 'th',
                          'h1', 'h2', 'h3', 'h4', 'h5', 'h6')

    #: Table scaffolding, which holds no text of its own. An emphasis
    #: declaration here has nowhere to go - Word has no run inside a ``<tr>`` -
    #: and either placement breaks that same direct-child walk.
    _HTML_TABLE_TAGS = ('table', 'thead', 'tbody', 'tfoot', 'tr')

    @staticmethod
    def _css_declarations(style: str) -> list:
        """The declarations of a style attribute.

        Split on the ``;`` between declarations and not on one inside a
        quoted string or a ``url(...)``: torn at that one, ``url('tan;x.png')``
        leaves ``url('tan`` whose last word is a named colour.
        """
        out, cur, depth, quote = [], [], 0, ''
        for ch in style:
            if quote:
                if ch == quote:
                    quote = ''
            elif ch in '"\'':
                quote = ch
            elif ch == '(':
                depth += 1
            elif ch == ')':
                depth = max(0, depth - 1)
            elif ch == ';' and not depth:
                out.append(''.join(cur))
                cur = []
                continue
            cur.append(ch)
        out.append(''.join(cur))
        return out

    @classmethod
    def _css_text_align(cls, style: str) -> str:
        """Return the ``text-align`` value declared in ``style``, or ''."""
        for decl in cls._css_declarations(style):
            prop, _, value = decl.partition(':')
            if prop.strip().lower() == 'text-align':
                value = value.strip().lower()
                if value in ('left', 'right', 'center', 'justify'):
                    return value
        return ''

    @classmethod
    def _css_declaration(cls, style: str, prop: str) -> str:
        """The value ``style`` gives ``prop``, or '' - last declaration wins.

        Matched on the whole property name, so ``color`` is not read out of
        the ``background-color`` beside it.
        """
        found = ''
        for decl in cls._css_declarations(style):
            name, _, value = decl.partition(':')
            if name.strip().lower() == prop:
                found = value.replace('!important', ' ').strip()
        return found

    @staticmethod
    def _add_css(el, declaration: str) -> None:
        """Append one CSS declaration to an element's style attribute."""
        el['style'] = ';'.join(
            part for part in (el.get('style', '').strip().rstrip(';'),
                              declaration) if part)

    @classmethod
    def _normalize_css_color(cls, value: str):
        """Six-hex (no #) for a CSS colour value.

        '' for a colour that is there and invisible - `transparent`, or any
        notation whose alpha is zero - and None for a value that is not a
        colour this parser reads. The two are told apart because an invisible
        border is written off while an unreadable one keeps its box.
        """
        v = value.strip().lower()
        if v == 'transparent':
            return ''
        if v.startswith('#'):
            h = v[1:]
            if len(h) in (3, 4):
                h = ''.join(c * 2 for c in h)
            if len(h) not in (6, 8) or any(c not in '0123456789abcdef' for c in h):
                return None
            # `#RRGGBBAA` carries its alpha as a fourth byte, one spelling of
            # the value `rgba()` spells in numbers - composited the same way
            nums = [int(h[i:i + 2], 16) for i in range(0, len(h), 2)]
            if len(nums) == 4:
                nums[3] /= 255
        else:
            match = cls._CSS_COLOR_FN_RE.match(v)
            if not match:
                return cls._CSS_NAMED_COLORS.get(v)
            fn, args = match.groups()
            parts = args.replace(',', ' ').replace('/', ' ').split()
            if len(parts) < 3:
                return None
            try:
                if fn.startswith('hsl'):
                    # Hue in degrees, saturation and lightness as percentages;
                    # `colorsys` wants them as fractions, lightness first
                    import colorsys
                    hue = float(parts[0].removesuffix('deg')) / 360 % 1
                    sat = float(parts[1].rstrip('%')) / 100
                    lig = float(parts[2].rstrip('%')) / 100
                    nums = [c * 255 for c in colorsys.hls_to_rgb(hue, lig, sat)]
                    nums += [float(parts[3][:-1]) / 100 if parts[3].endswith('%')
                             else float(parts[3])] if len(parts) > 3 else []
                else:
                    # A channel is a number or a percentage of full strength,
                    # and the alpha that follows them the same of 1
                    nums = [float(p[:-1]) / 100 * (255 if i < 3 else 1)
                            if p.endswith('%') else float(p)
                            for i, p in enumerate(parts[:4])]
            except ValueError:
                return None
        if True:
            try:
                alpha = min(1.0, max(0.0, nums[3])) if len(nums) > 3 else 1.0
                if not alpha:
                    return ''  # what the keyword `transparent` spells in words
                channels = ''
                for number in nums[:3]:
                    # Word has no translucent run or cell and reportlab paints
                    # a callout opaque, so the colour is composited onto the
                    # white page rather than arriving at full strength: alpha
                    # is what makes a colour light, and dropping it turns a 5%
                    # black wash into a solid black bar over black text.
                    number = number * alpha + 255 * (1 - alpha)
                    # Clamped as a browser clamps it, and because a component
                    # past 255 would format to three hex digits and corrupt the
                    # colour. NaN and the infinities are values float() accepts
                    # and round() does not, so the guard has to cover both
                    channels += f'{min(255, max(0, round(number))):02X}'
            except (ValueError, OverflowError):
                return None
            return channels

    @classmethod
    def _css_color_invisible(cls, value: str) -> bool:
        """True when a colour token in ``value`` is there and invisible."""
        return any(cls._normalize_css_color(token) == ''
                   for token in cls._CSS_COLOR_TOKEN_RE.findall(value))

    @classmethod
    def _css_color_in(cls, value: str) -> str:
        """Six-hex of the first colour a CSS value names, or '' if it names none.

        A shorthand carries its colour among other words (``2px dashed #ccc``,
        ``#fff url(bg.png) no-repeat``), so every token is tried in turn rather
        than the colour being assumed to come first.
        """
        for token in cls._CSS_COLOR_TOKEN_RE.findall(value):
            resolved = cls._normalize_css_color(token)
            if resolved:
                return resolved
        return ''

    @classmethod
    def _css_border_parts(cls, value: str) -> tuple:
        """``(style, width, colour)`` a ``border`` shorthand declares, '' for absent.

        The three are order-free in CSS, so each token is placed by what it is
        rather than by where it stands: a style keyword, something measuring a
        length, and whatever is left over, which is the colour - rejoined,
        because ``rgb(1, 2, 3)`` is several whitespace tokens of one value.
        """
        line_style, width, color = '', '', []
        # A function call is one token whatever whitespace it holds:
        # split on spaces, `hsl(0 100% 50% / 0)` hands `100%` to the width
        for token in re.findall(r'[a-z-]*\([^()]*\)|\S+', value, re.I):
            if token in cls._CSS_BORDER_STYLES:
                line_style = token
            elif token in cls._CSS_BORDER_WIDTHS or cls._CSS_LENGTH_RE.match(token):
                width = token
            else:
                color.append(token)
        return line_style, width, ' '.join(color)

    @classmethod
    def _css_callout_box(cls, style: str):
        """``(bar_hex, fill_hex)`` when a style attribute draws a box, else None.

        A border or a background is how a document draws a callout by hand,
        markdown having no syntax for one. Both colours are always answered,
        never one of the two: the box is recognised again downstream BY those
        colours - the PDF reads them back off the finished DOCX in
        alert_info() - so a box that declares only a background takes its own
        fill for the bar, and one that declares only a border takes the
        blockquote grey and white.

        The border is read the way the cascade reads it, as DEF-DIAG-39 already
        reads emphasis: each declaration is folded onto the side and the part
        it names, last value winning, so a longhand can switch off a border the
        shorthand above it drew. A side draws only once it has a style, which
        is what keeps a colour on its own from being taken for a border.
        """
        fill = ''
        # Per side, the style, width and colour it was last given; '' where no
        # declaration named one and the CSS initial value stands
        sides = {side: ['', '', ''] for side in cls._CSS_BORDER_SIDES}
        for decl in cls._css_declarations(style):
            prop, _, value = decl.partition(':')
            prop = prop.strip().lower()
            # !important rides along on any value and is no part of it
            value = value.replace('!important', ' ').strip().lower()
            if prop in ('background', 'background-color'):
                # Last value wins, and one naming no colour (`none`,
                # `transparent`, a url) clears what stood above it; the other
                # `background-*` longhands never carry the colour
                fill = cls._css_color_in(value)
            elif prop in cls._CSS_BORDER_PROPS:
                part = prop[len('border'):]
                named = [s for s in cls._CSS_BORDER_SIDES
                         if part.startswith(f'-{s}')]
                if named:
                    part = part[len(named[0]) + 1:]
                for side in named or cls._CSS_BORDER_SIDES:
                    if part:
                        sides[side][
                            ('-style', '-width', '-color').index(part)] = value
                    else:
                        # A shorthand sets all three parts: the ones it leaves
                        # out go back to their initial values rather than
                        # keeping what a declaration above gave them
                        sides[side] = list(cls._css_border_parts(value))
        bar, bordered = '', False
        for side in cls._CSS_BORDER_SIDES:
            line_style, width, color = sides[side]
            # The three spellings that write a side off, and an author reaches
            # for each: no style or one that draws nothing, a width that is
            # not a positive length (`0`, `00px`, and `-1px`, which a browser
            # rejects with the whole declaration), and an invisible colour
            length = re.match(r'-?\d*\.?\d+', width)
            if (line_style in ('', 'none', 'hidden')
                    or (length and float(length.group()) <= 0)
                    # `transparent`, `#0000`, `rgba(...,0)`: an invisible
                    # colour in any spelling. A colour the parser cannot read
                    # is not invisible, and keeps the grey bar
                    or cls._css_color_invisible(color)):
                continue
            bordered = True
            bar = bar or cls._css_color_in(color)
        if not bordered and not fill:
            return None
        return (bar or fill or 'BBBBBB', fill or 'FFFFFF')

    @classmethod
    def _carry_box_color(cls, soup, el, blocks) -> None:
        """Move the box's own ``color`` onto the text standing inside it.

        htmldocx reads a colour off an inline tag only, so one declared on the
        div - or on the paragraph the div became - arrives nowhere, while the
        background beside it paints a whole cell: an author who set both would
        get one of the pair, where a browser paints both. The colour goes
        around each paragraph's contents rather than on the paragraph, which is
        the same placement the emphasis pass uses for a block.

        A heading is left out. The stylesheet gives one a colour of its own and
        that beats what it inherits, so a browser draws the heading in a
        red-texted box blue - and the Word heading style and the PDF heading
        ladder draw it blue too.
        """
        color = cls._css_declaration(el.get('style', ''), 'color')
        if not color:
            return
        for block in blocks:
            # `block is el` is the box drawn around one paragraph: the colour
            # being carried is that element's own, not a child overriding it
            if block.name != 'p' or (block is not el and cls._css_declaration(
                    block.get('style', ''), 'color')):
                continue
            span = soup.new_tag('span')
            cls._add_css(span, f'color:{color}')
            for node in list(block.contents):
                span.append(node.extract())
            block.append(span)

    @classmethod
    def _boxed_blocks(cls, el) -> list:
        """``el``'s block children, a run of loose inline content between them
        wrapped in a ``<p>`` of its own, in document order.

        Only an element can carry the box marker - it travels in the text of a
        block - so a bare text node between two blocks stayed at body level
        when the div was unwrapped, and htmldocx welded it onto the paragraph
        above: two authored blocks arrived as one, in the style of the first.
        A browser gives that text a block of its own, and consecutive inline
        nodes share one the way an anonymous block box gathers them. The
        newlines between tags are a run with no text and are left where they
        are rather than boxed as a blank line.
        """
        from bs4 import Comment, NavigableString, Tag

        def has_text(node):
            if isinstance(node, Comment):
                return False
            if isinstance(node, NavigableString):
                return bool(node.strip())
            return bool(node.get_text().strip())

        run = []
        for child in list(el.children) + [None]:
            if (child is not None
                    and getattr(child, 'name', None) not in cls._HTML_BLOCK_TAGS):
                run.append(child)
                continue
            # The newline that closed the block above and the one that opens
            # the block below belong to neither paragraph
            while run and not has_text(run[0]):
                run.pop(0)
            while run and not has_text(run[-1]):
                run.pop()
            if run:
                para = Tag(name='p')
                run[0].insert_before(para)
                for index, node in enumerate(run):
                    node.extract()
                    if isinstance(node, NavigableString):
                        # The line breaks the author wrote around the text
                        # collapse to a space, which a browser then drops at
                        # the edge of a block and Word would show
                        text = str(node)
                        text = text.lstrip() if not index else text
                        text = text.rstrip() if index == len(run) - 1 else text
                        node = NavigableString(text)
                    para.append(node)
            run = []
        return el.find_all(cls._HTML_BLOCK_TAGS, recursive=False)

    @classmethod
    def _mark_callout_box(cls, el, blocks=None):
        """Tag a bordered or shaded element as a callout for the DOCX passes.

        The marker travels in the text because text is the only thing htmldocx
        carries through: a declaration it has no handler for is read for
        emphasis and dropped, so border, background, padding and margin all
        arrive nowhere and the box exports as an ordinary paragraph.
        style_docx_alert_boxes() reads the marker back and moves the finished
        paragraph into the same box a GitHub alert gets.

        ``blocks`` are the block children the box is drawn around when the
        element itself will not survive to carry a marker - a ``<div>``
        wrapping blocks is unwrapped, and the marker has to be inside the
        blocks by then or it lands in whatever paragraph stood above the div.
        The first opens the box and the rest join it.

        The background declaration leaves with it. Left in, the pill pass would
        ALSO shade every run in the box - a highlighter band across text that
        is already sitting on the fill.

        Answers the blocks the box was drawn around, which is what the caller
        hands the rest of the div's style to before unwrapping it, or None
        when no box was drawn.
        """
        box = cls._css_callout_box(el.get('style', ''))
        if box is None:
            return None
        if el.find(['img', 'svg', 'picture', 'video', 'object', 'iframe']):
            # A picture leaves the box on both paths - htmldocx gives it a
            # paragraph of its own, which stays outside the table the text is
            # moved into and so lands BELOW the box it was written above, and
            # the PDF rebuilds a box from its runs alone. An unboxed paragraph
            # in the right order beats a box the picture fell out of.
            return None
        targets = [el] if blocks is None else blocks
        if any(block.name not in cls._HTML_BOXABLE_BLOCKS for block in targets):
            return None
        if blocks is not None:
            # Read again with the loose text gathered into paragraphs of its
            # own. After the guard above, which asks what the author WROTE in
            # the div - a paragraph this pass makes is boxable by construction
            # and would answer for a table that is not
            targets = cls._boxed_blocks(el)
        for block in targets[1:]:
            block.insert(0, cls._BOX_MORE_MARKER)
        targets[0].insert(0, f'⁣BOX:{box[0]}:{box[1]}⁣')
        kept = [d for d in cls._css_declarations(el.get('style', '')) if d.strip()
                and not d.partition(':')[0].strip().lower().startswith('background')]
        if kept:
            el['style'] = ';'.join(kept)
        else:
            del el['style']
        return targets

    def restructure_html_for_docx(self, html: str) -> str:
        """Fix htmldocx structural and styling blind spots before conversion.

        Handles six htmldocx limitations:

        1. Loose list items (``<li><p>text</p>...</li>``) - ``handle_li``
           opens a bullet paragraph, then the inner ``<p>`` opens a fresh
           ``Normal`` paragraph, so the bullet glyph ends up empty and the
           text loses its bullet. Unwrapping the leading ``<p>`` of each
           ``<li>`` puts the text back into the bullet paragraph.
        2. Blockquotes - there is no ``<blockquote>`` handler at all, so the
           inner ``<p>`` become plain paragraphs with no indent or bar. A
           sentinel marker ``⁣BQ:<indent>⁣`` is injected.
        3. Named CSS text colours (``color: green``) - htmldocx parses only
           hex, rendering named colours black. Normalised to hex in-place.
        4. Coloured pills (``background-color`` spans) - htmldocx maps a CSS
           background to a Word *highlight* from a tiny fixed palette, so an
           arbitrary hex becomes ``lightGray``. The background is stripped
           from the style and re-encoded as a ``⁣PILL:<hex>⁣`` marker
           that style_docx_color_runs() turns into true run shading.
        5. Inline HTML htmldocx has no handler for - ``<mark>``, ``<del>``,
           ``<kbd>``, ``<font color>``, an ``align=`` attribute, a bare
           ``<div>``, and the ``font-weight`` / ``font-style`` /
           ``text-decoration`` half of a style attribute. Each is rewritten
           into the markup htmldocx does read, so the formatting arrives in
           Word instead of the text landing there plain. A ``<div>`` wrapping
           blocks is the exception: it hands its children only its
           ``text-align`` and is then unwrapped, so the rest of its style
           attribute is lost.
        6. A ``<div>`` the author gave a border or a background - markdown
           having no syntax for a callout, that is how one is drawn. Nothing
           downstream reads border, background, padding or margin, so the box
           is re-encoded as a ``⁣BOX:<bar>:<fill>⁣`` marker that
           style_docx_alert_boxes() turns into the box a GitHub alert gets,
           in the div's own colours.
        """
        from bs4 import BeautifulSoup, NavigableString
        soup = BeautifulSoup(html, 'html.parser')

        # 2. Blockquote paragraphs: prefix an indent-encoding marker
        for bq in soup.find_all('blockquote'):
            depth_bq = len(bq.find_parents('blockquote')) + 1
            depth_list = len(bq.find_parents(['ul', 'ol']))
            indent_in = depth_list * 0.5 + depth_bq * 0.3
            marker = f'⁣BQ:{int(round(indent_in * 100))}⁣'
            paras = bq.find_all('p', recursive=False)
            if paras:
                for p in paras:
                    p.insert(0, marker)
            else:
                bq.insert(0, marker)

        # 5. Inline tags htmldocx cannot read, swapped for ones it can
        for el in soup.find_all(list(self._HTML_TAG_ALIASES)):
            el.name = self._HTML_TAG_ALIASES[el.name]
        for el in soup.find_all('mark'):
            el.name = 'span'
            # The style loop below takes the LAST background it sees, so adding
            # the default unconditionally would discard the author's own colour
            if 'background' not in el.get('style', '').lower():
                self._add_css(el, f'background-color:#{self._HTML_MARK_FILL}')
        # <font color> is the colouring a notebook markdown cell is most often
        # written with, and htmldocx reads colour off the style attribute only
        # An <a> carrying no href is a link TARGET, not a link - the
        # markdown-portable way to name a spot mid-document. htmldocx reads a
        # link's text through self.tags['a']['href'] without checking it is
        # there, so any such anchor holding text raises KeyError and the whole
        # export dies with HTTP 500. As a <span> it keeps its id for
        # inject_anchor_markers and htmldocx never sees an anchor at all.
        for el in soup.find_all('a', href=False):
            el.name = 'span'
            if el.has_attr('name') and not el.has_attr('id'):
                el['id'] = el['name']
            el.attrs.pop('name', None)
        for el in soup.find_all('font'):
            color = el.get('color', '')
            el.name = 'span'
            for attr in ('color', 'face', 'size'):
                el.attrs.pop(attr, None)
            if color:
                self._add_css(el, f'color:{color}')
        # `align="center"` is how a markdown file centres a block; htmldocx
        # reads the CSS property, so move the attribute onto the style
        for el in soup.find_all(align=True):
            align = str(el['align']).strip().lower()
            del el['align']
            if align in ('left', 'right', 'center', 'justify'):
                self._add_css(el, f'text-align:{align}')
        # A <div> has no handler either, so its text joins whatever paragraph
        # is already open and two blocks run together. One holding only inline
        # content becomes the paragraph it stands for; one wrapping blocks
        # hands them its alignment and steps out of the way.
        for div in soup.find_all('div'):
            children = div.find_all(self._HTML_BLOCK_TAGS, recursive=False)
            if not children:
                # Only when it stands on its own. Markdown wraps a div written
                # in inline position in a paragraph of its own making, and a
                # <p> nested in a <p> makes htmldocx open a second one - which
                # drops the content out of the blockquote or alert box whose
                # body it was, both of those being one paragraph per marker.
                if not div.find_parents(self._HTML_TEXT_HOLDERS):
                    div.name = 'p'
                    # A border or a background makes it a box the author drew,
                    # not a paragraph. Only here: a box is a table in Word, and
                    # Word has no table inside a list item, cell or heading -
                    # which is the branch below.
                    boxed = self._mark_callout_box(div)
                    if boxed:
                        self._carry_box_color(soup, div, boxed)
                else:
                    # A span, not unwrapped: the style loop below reads the
                    # style attribute off the element, and unwrapping takes it
                    # with the tag - so a div written inline arrived plain in
                    # the one pass whose whole purpose is that it does not
                    div.name = 'span'
                continue
            # A border or a background makes this one a box the author drew
            # too, and it is the same box - drawn around the paragraphs the
            # children become rather than around one. Same placement rule as
            # above: never inside a list item, a cell or a heading, Word
            # having no table there. It answers the blocks it boxed, one per
            # loose run of text as well as the children written as blocks, so
            # what the div hands down below reaches the whole box.
            boxed = None
            if not div.find_parents(self._HTML_TEXT_HOLDERS):
                boxed = self._mark_callout_box(div, children)
            children = boxed or children
            align = self._css_text_align(div.get('style', ''))
            if align:
                for child in children:
                    if not self._css_text_align(child.get('style', '')):
                        self._add_css(child, f'text-align:{align}')
            if boxed:
                # The div is unwrapped and its style attribute goes with it, so
                # what it holds for the text inside the box has to be handed
                # down first - the alignment above, and the colour here
                self._carry_box_color(soup, div, boxed)
            div.unwrap()

        # 1. Loose list items: merge a leading <p> into the <li> itself.
        # After the div pass above, which can turn a <li>'s <div> into exactly
        # that <p>.
        for li in soup.find_all('li'):
            first_el = next(
                (c for c in li.children if getattr(c, 'name', None)), None
            )
            if first_el is not None and first_el.name == 'p':
                first_el.unwrap()

        # 2. Where each ordered list begins. htmldocx writes every numbered
        # item as a 'List Number' paragraph on the template's one numbering
        # instance, so the Word file cannot tell a fresh list from an item
        # continuing after a table, a sample or a paragraph of its own -
        # both read as list paragraph, something, list paragraph. The mark
        # is read back and removed by restart_docx_list_numbering.
        for ol in soup.find_all('ol'):
            first_li = ol.find('li', recursive=False)
            if first_li is not None:
                first_li.insert(0, NavigableString(self._LIST_START))

        # 3 & 4. Inline colour / pill styling on any element with a style attr
        for el in soup.find_all(style=True):
            style = el.get('style', '')
            decls = [d.strip() for d in self._css_declarations(style) if d.strip()]
            kept, fg_hex, bg_hex, style_props = [], '', '', {}
            for d in decls:
                if ':' not in d:
                    kept.append(d)
                    continue
                prop, val = d.split(':', 1)
                prop = prop.strip().lower()
                if prop == 'text-decoration-line':
                    prop = 'text-decoration'  # the same declaration, longhand
                if prop in self._CSS_STYLE_TAGS:
                    # Dropped by htmldocx, so re-expressed as the tag that says
                    # the same thing. The declaration goes with it - left in, a
                    # `font-weight:bold` would still read as unstyled text.
                    # Last value per property, as the cascade says: a pasted
                    # `underline;none` pair must end with nothing, and a
                    # duplicated `bold` must not wrap twice
                    style_props[prop] = val.strip().lower()
                elif prop == 'color':
                    # A value this code could not resolve is dropped, whatever
                    # its notation. htmldocx reads a channel with int() and
                    # slices a hex blind, so handing one on either takes the
                    # whole export down - HTTP 500 and no document, for one
                    # malformed declaration anywhere in the file - or paints a
                    # colour that is on no screen: `#12345` arrives as its
                    # first two pairs and a nibble. A browser drops a
                    # declaration it cannot parse and the text keeps the body
                    # colour, which is what dropping it here leaves too
                    h = self._normalize_css_color(val)
                    if h:
                        fg_hex = h
                        kept.append(f'color:#{h}')
                elif prop == 'background-color' or prop == 'background':
                    # The shorthand puts the colour among an image and its
                    # repeat, and `rgb(244, 244, 245)` is several whitespace
                    # tokens of its own, so the value is scanned, not sliced.
                    # Unresolved, it goes the way of `color` above. Last value
                    # wins, and one naming no colour clears the fill, as the
                    # box reader reads the same declarations
                    bg_hex = self._css_color_in(val)  # re-encoded as a marker below
                else:
                    kept.append(d)
            wrappers = ''.join(self._CSS_STYLE_TAGS[prop](val)
                               for prop, val in style_props.items())
            if bg_hex or fg_hex or style_props or len(kept) != len(decls):
                if kept:
                    el['style'] = ';'.join(kept)
                else:
                    del el['style']
            if bg_hex:
                marker = f'⁣PILL:{bg_hex}⁣'
                # One marker per text node, not one per element: emphasis
                # inside the background splits it into several runs, and
                # style_docx_color_runs shades only the run it finds a marker
                # in - so a `<mark>` around bold text would shade nothing.
                # Whitespace counts: the space between two inline children
                # is a text node of its own, and skipping it leaves an
                # unshaded notch mid-highlight where a browser paints a
                # continuous band. Empty nodes are still skipped - a marker in
                # one would shade a run holding nothing.
                nested = [node for node in el.descendants
                          if isinstance(node, NavigableString) and node]
                for node in nested:
                    node.insert_before(marker)
                if not nested:
                    el.insert(0, marker)
            if el.name in self._HTML_TABLE_TAGS:
                # Nowhere to put the emphasis: Word has no run inside a row,
                # and either placement loses the cells
                wrappers = ''
            for tag in wrappers:
                if el.find_parent(tag) is not None:
                    # htmldocx keys its open tags by name and pops on the first
                    # close, so a nested twin ends the outer one early and the
                    # text after it arrives unformatted
                    continue
                # The same bookkeeping breaks from below: a literal twin
                # INSIDE the element closes the wrapper this pass is about
                # to open and the tail arrives plain. The wrapper subsumes
                # the twin - bold inside bold is bold - so only the NAME
                # goes: a rename to <span> keeps the twin's own style, id
                # and place in the style snapshot, where an unwrap took
                # its colour and its anchor with the tag. A twin across a
                # nested-table boundary stays whole - htmldocx scopes each
                # cell, so no early close crosses a table, and the outer
                # wrapper cannot reach that cell's runs to replace it
                for twin in el.find_all(tag):
                    if twin.find_parent('table') is el.find_parent('table'):
                        twin.name = 'span'
                wrapper = soup.new_tag(tag)
                if el.name in self._HTML_INLINE_TAGS:
                    el = el.wrap(wrapper)
                else:
                    # A block, a cell or a list item cannot sit inside <b> and
                    # still be found by the direct-child walk that reaches it,
                    # so the emphasis goes around its contents instead
                    for child in list(el.contents):
                        wrapper.append(child.extract())
                    el.append(wrapper)

        return str(soup)

    # Invisible times marks the first item of each ordered list from
    # restructure_html_for_docx to restart_docx_list_numbering (U+2063, the
    # invisible separator, is the bookmark and box marker)
    _LIST_START = '\u2062'

    def restart_docx_list_numbering(self, document):
        """Give every ordered list its own numbering instance, starting at 1.

        htmldocx numbers all 'List Number' paragraphs from the one instance
        the template's style names, so Word counted a second procedure on
        from the first and the PDF rebuild had to guess where a list ended -
        wrongly, whenever a step held a table, a sample or a paragraph of its
        own. Each list marked by restructure_html_for_docx gets a fresh
        instance of the style's abstract numbering with a start override of
        1; the paragraphs that follow at the same depth join it until the
        next mark at that depth. The mark is removed here. The PDF rebuild
        restarts its count when the instance changes and guesses nothing.
        """
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph as DocxParagraph
        from docx.shared import Inches

        numbering = document.part.numbering_part.element
        style_num = document.styles['List Number'].element.pPr.numPr
        abstract_id = numbering.num_having_numId(
            style_num.numId.val).abstractNumId.val

        def depth_of(para):
            # htmldocx indents 0.5in per nesting level, so the depth is in
            # the paragraph's own indent, not in the (single) style name
            indent = para.paragraph_format.left_indent or 0
            return max(0, round(indent / Inches(0.5)) - 1)

        current = {}
        for p in document.element.body.iter(qn('w:p')):
            para = DocxParagraph(p, document)
            starts = False
            for run in para.runs:
                if self._LIST_START in run.text:
                    run.text = run.text.replace(self._LIST_START, '')
                    starts = True
            if not para.style or para.style.name != 'List Number':
                continue
            depth = depth_of(para)
            if starts:
                num = numbering.add_num(abstract_id)
                num.add_lvlOverride(ilvl=0).add_startOverride(1)
                current[depth] = num.numId
            if depth not in current:
                continue
            num_pr = p.get_or_add_pPr().get_or_add_numPr()
            num_pr.get_or_add_ilvl().val = 0
            num_pr.get_or_add_numId().val = current[depth]

    def style_docx_color_runs(self, document):
        """Apply true run shading (``w:shd``) to runs carrying a pill marker
        from restructure_html_for_docx(). htmldocx can only express a CSS
        background as a Word highlight from a fixed palette; this gives the
        exact hex fill instead. Strips the marker afterwards.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph

        # Successors of w:shd in the CT_RPr sequence - w:shd is index 29 and
        # w:vertAlign 31, so appending it after the vertAlign a <sub>/<sup>
        # writes gives Word a file it refuses to open (mirrors set_font's
        # handling of w:rFonts in style_docx_symbol_runs).
        shd_succ = ('w:fitText', 'w:vertAlign', 'w:rtl', 'w:cs', 'w:em',
                    'w:lang', 'w:eastAsianLayout', 'w:specVanish', 'w:oMath')

        body = document.element.body
        for p_elem in body.iter(qn('w:p')):
            paragraph = Paragraph(p_elem, document)
            for run in self.docx_paragraph_runs(paragraph):
                text = run.text or ''
                m = self._PILL_MARKER_RE.search(text)
                if not m:
                    continue
                fill = m.group(1)
                run.text = self._PILL_MARKER_RE.sub('', text)

                rPr = run._r.get_or_add_rPr()
                # Remove any highlight htmldocx may have added for the bg
                for hl in rPr.findall(qn('w:highlight')):
                    rPr.remove(hl)
                existing = rPr.find(qn('w:shd'))
                if existing is not None:
                    rPr.remove(existing)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fill)
                succ = next((el for tag in shd_succ
                             for el in rPr.findall(qn(tag))), None)
                if succ is not None:
                    succ.addprevious(shd)
                else:
                    rPr.append(shd)

    def style_docx_blockquotes(self, document):
        """Apply indent, a gray left bar, light shading and muted italic text
        to paragraphs carrying the blockquote marker from
        restructure_html_for_docx(). Strips the marker afterwards.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Inches, RGBColor

        # Successors of w:pBdr / w:shd in the CT_PPr schema - inserting before
        # these keeps the element order valid (mirrors htmldocx's hr handler).
        pbdr_succ = (
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku',
            'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE',
            'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd', 'w:snapToGrid',
            'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:mirrorIndents',
            'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment',
            'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle',
            'w:rPr', 'w:sectPr', 'w:pPrChange',
        )
        shd_succ = pbdr_succ[1:]

        # Iterate every paragraph in the body, including those nested in table
        # cells (document.paragraphs only yields body-level paragraphs, so a
        # blockquote inside a table would otherwise keep its raw marker).
        from docx.text.paragraph import Paragraph
        body = document.element.body
        for p_elem in body.iter(qn('w:p')):
            paragraph = Paragraph(p_elem, document)
            match = self._BLOCKQUOTE_MARKER_RE.search(paragraph.text or '')
            if not match:
                continue
            indent_in = int(match.group(1)) / 100.0

            # Strip the marker from whichever run(s) carry it
            for run in paragraph.runs:
                if run.text and '⁣' in run.text:
                    run.text = self._BLOCKQUOTE_MARKER_RE.sub('', run.text)

            paragraph.paragraph_format.left_indent = Inches(indent_in)

            pPr = paragraph._p.get_or_add_pPr()

            existing = pPr.find(qn('w:pBdr'))
            if existing is not None:
                pPr.remove(existing)
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '18')      # ~2.25pt bar
            left.set(qn('w:space'), '12')   # gap bar-to-text
            left.set(qn('w:color'), 'BBBBBB')
            pBdr.append(left)
            pPr.insert_element_before(pBdr, *pbdr_succ)

            existing_shd = pPr.find(qn('w:shd'))
            if existing_shd is not None:
                pPr.remove(existing_shd)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F4F4F4')
            pPr.insert_element_before(shd, *shd_succ)

            for run in paragraph.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    def strip_monospace_from_unicode_runs(self, document):
        """Drop monospace font overrides from runs containing non-ASCII characters.

        htmldocx applies Courier font to every ``<code>``/``<pre>`` run, but
        Courier does not contain unicode arrows (U+2190 etc.) and many symbol
        glyphs. Stripping the rFonts override lets Word render these characters
        in the body font (Cambria).
        """
        from docx.oxml.ns import qn
        MONOSPACE_FONTS = {'Courier', 'Courier New', 'Consolas', 'Monaco'}
        body = document.element.body
        for run_elem in body.iter(qn('w:r')):
            text = ''.join((t.text or '') for t in run_elem.findall(qn('w:t')))
            if not text or text.isascii():
                continue
            rpr = run_elem.find(qn('w:rPr'))
            if rpr is None:
                continue
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                continue
            if (rfonts.get(qn('w:ascii')) or '') in MONOSPACE_FONTS:
                rpr.remove(rfonts)

    @staticmethod
    def fit_column_widths(natural: list, minimums: list, available: float) -> list:
        """Scale natural column widths down so the row fits ``available``.

        ``natural`` holds the width each column would take on a single line and
        ``minimums`` the width of its longest unbreakable word. When the natural
        widths already fit they are returned untouched. Otherwise every column
        keeps its minimum - so wrapping falls on word boundaries rather than
        mid-word - and the remaining space is shared in proportion to the
        excess, keeping wide columns wide. Falls back to equal columns when
        every column is already capped at its fair share.
        """
        ncols = len(natural)
        if ncols == 0:
            return []
        if sum(natural) <= available:
            return natural
        # No single column may claim more than its fair share as a minimum,
        # otherwise one very long word starves every other column.
        fair_share = available / ncols
        minimums = [min(m, fair_share) for m in minimums]
        slack = available - sum(minimums)
        excess = [max(0.0, n - m) for n, m in zip(natural, minimums)]
        if slack == 0 or sum(excess) == 0:
            return [fair_share] * ncols
        scale = slack / sum(excess)
        return [m + e * scale for m, e in zip(minimums, excess)]

    @classmethod
    def measured_column_widths(cls, table_data: list, available: float,
                               measure, floors: list | None = None) -> list:
        """Column widths for a table of cell strings that fit ``available``.

        ``measure(text, row_index)`` returns the rendered width of one string,
        letting the caller supply the font metrics. A cell's natural width is
        its longest line, its minimum the longest word it cannot break, and
        ``floors`` adds a per-column lower bound for content this cannot see
        (an image - only the DOCX path passes any; the PDF path renders a cell
        image but scales it to whatever column its text earns rather than
        flooring the column to it). Every column keeps at least the width of an
        empty cell, so none comes back zero-width, unless an equal share of
        ``available`` is itself smaller than that.
        """
        ncols = max((len(row) for row in table_data), default=0)
        natural, minimums = [], []
        for c in range(ncols):
            cells = [(r, row[c]) for r, row in enumerate(table_data)
                     if c < len(row)]
            # No floor may exceed an equal share, or one column starves the rest
            floor = min(max(floors[c] if floors else 0.0, measure('', 0)),
                        available / ncols)
            natural.append(max(
                [measure(line, r) for r, t in cells for line in t.split('\n')]
                + [floor]))
            minimums.append(max(
                [measure(word, r) for r, t in cells for word in t.split()]
                + [floor]))
        return cls.fit_column_widths(natural, minimums, available)

    @classmethod
    def pdf_table_column_layout(cls, table_data: list, available: float,
                                string_width, image_widths: list | None = None) -> tuple:
        """Cell padding and column widths for a PDF table of ``available`` width.

        ``string_width(text, row_index)`` measures rendered text. reportlab
        raises rather than rendering when a column is no wider than its own
        padding, so a table with very many columns gets a tighter padding
        instead of a width floor that would push it back off the page. The
        widths always sum to at most ``available``.

        ``image_widths`` is a per-column content width (points) for the widest
        cell image, so an image-only cell - which earns no width from its empty
        text - floors its column to the image instead of collapsing to a few
        points (the DOCX path floors the same way). measured_column_widths caps
        each floor at a fair share so an image cannot starve its text neighbours.

        Past roughly 150 columns a column is narrower than a single wide glyph
        at any font size used here, and reportlab paints such a glyph past its
        column - the borders stay inside the margin but the ink can pass it by
        a point or two. Such a table is unreadable regardless.
        """
        ncols = max((len(row) for row in table_data), default=1) or 1
        fair_share = available / ncols
        # Full padding while columns are roomy, then step down so reportlab
        # still has room to render as the fair share shrinks past the padding
        # (>= 16pt) and then past a hairline column (>= 3pt)
        full = cls.PDF_TABLE_CELL_PADDING
        side_padding = full if fair_share >= 16 else (1 if fair_share >= 3 else 0)
        # Both paddings, plus rounding slack so a word whose width matches its
        # column is not split mid-word
        cell_padding = 2 * side_padding + 2

        def measure(text, row_index):
            return string_width(text, row_index) + cell_padding

        # An image floor is the image content width plus the same padding a
        # text measurement carries, so the column holds the image at full size
        floors = None
        if image_widths and any(image_widths):
            floors = [w + cell_padding if w else 0.0 for w in image_widths]

        # The empty-cell floor already keeps every width above the padding
        return side_padding, cls.measured_column_widths(
            table_data, available, measure, floors)

    # Cell content a `w:t` run does not carry, so `cell.text` cannot see it.
    # The blank-row predicate is the only thing between the header delete below
    # and silent data loss, so it tests every kind this pipeline can emit.
    DOCX_CONTENT_ELEMENTS = (
        'w:drawing',                                  # inline/floating picture
        'w:pict', 'w:object',                         # VML / embedded object
        'm:oMath', 'm:oMathPara',                     # OMML equation
        '{urn:schemas-microsoft-com:vml}imagedata',   # VML image reference
    )

    @classmethod
    def docx_row_is_blank(cls, row) -> bool:
        """True when a table row holds neither text nor embedded content.

        Shared by the DOCX and PDF paths so the two cannot drift: 'no text'
        alone is not 'no content' - an image-on-top/caption-below grid puts its
        pictures in the header row, and deleting that row would drop them.
        """
        from docx.oxml.ns import qn

        if any(c.text.strip() for c in row.cells):
            return False
        tr = row._tr
        return not any(
            tr.findall('.//' + (tag if tag.startswith('{') else qn(tag)))
            for tag in cls.DOCX_CONTENT_ELEMENTS
        )

    def style_docx_table(self, table) -> None:
        """Style a content table, and drop its blank header row if it has one.

        Two jobs: the banded style plus 100% page width, and - because Markdown
        forces a header row that a borderless image/layout grid does not want -
        deleting that row when it is genuinely blank. The delete is structural,
        so it runs first; the PDF path applies the same rule separately in
        ``process_table`` because its intermediate DOCX skips this function.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        table.style = 'Light List Accent 1'

        # A Markdown table needs a header row, so a borderless image/layout
        # grid is written with an empty one (`|  |  |  |`). The rendered
        # markdown shows nothing for it, so drop the row outright - left in, it
        # is a blank banded bar that Word repeats and strands above its image
        # rows.
        rows = table.rows
        first_row_empty = len(rows) >= 1 and self.docx_row_is_blank(rows[0])

        # Header CHROME is a separate question from whether the row is blank: a
        # picture-only first row is content and is kept, but it is a figure, not
        # a header, and banding or repeating it would be wrong. Read on the
        # ORIGINAL first row, so a body row left in position 0 by the delete is
        # not promoted into a header in its place. The PDF path decides the same
        # way (`has_header` reads text only), so both formats look alike.
        header_has_text = bool(rows) and any(
            c.text.strip() for c in rows[0].cells)

        if first_row_empty and len(rows) > 1:
            tr = rows[0]._tr
            tr.getparent().remove(tr)

        tblPr = table._tbl.tblPr
        # Disable first column emphasis, and first row emphasis for any table
        # whose first row is not a real (text-bearing) header
        tblLook = tblPr.find(qn('w:tblLook'))
        if tblLook is not None:
            tblLook.set(qn('w:firstColumn'), '0')
            if not header_has_text:
                tblLook.set(qn('w:firstRow'), '0')
        # Set table to 100% page width. Update in place, or insert at the slot
        # w:tblPr's schema sequence gives it - appending would put it after
        # w:tblLook, and a validating reader may then drop the width entirely
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.insert_element_before(
                tblW, 'w:jc', 'w:tblCellSpacing', 'w:tblInd', 'w:tblBorders',
                'w:shd', 'w:tblLayout', 'w:tblCellMar', 'w:tblLook',
                'w:tblCaption', 'w:tblDescription', 'w:tblPrChange')
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')
        # Mark the first row as a repeating header: Word then repeats it at the
        # top of each page the table spans and does not strand it alone at a
        # page bottom. Only a table with a real (text-bearing) header row to
        # repeat qualifies - a blank or picture-only first row is not a header.
        if len(table.rows) > 1 and header_has_text:
            trPr = table.rows[0]._tr.get_or_add_trPr()
            if trPr.find(qn('w:tblHeader')) is None:
                trPr.append(OxmlElement('w:tblHeader'))

        # Keep each row whole across a page break, the DOCX counterpart of the
        # PDF's conditional `splitInRow`: without it Word breaks a row wherever
        # the page ends, stranding a caption on one page and its image on the
        # next when the whole row would have fitted the next page. A row taller
        # than a page still breaks - `w:cantSplit` is a request Word drops when
        # it cannot be met, and nothing here sets an exact `w:trHeight`, which
        # is what would make Word clip the overflow instead.
        for row in table.rows:
            trPr = row._tr.get_or_add_trPr()
            if trPr.find(qn('w:cantSplit')) is None:
                trPr.insert(0, OxmlElement('w:cantSplit'))

    def fit_docx_table_to_page(self, table, available_twips: int,
                               base_pt: float = None) -> None:
        """Pin an over-wide table to the page with proportional columns.

        Word's default autofit layout widens a table past the right margin when
        a cell holds an unbreakable token (a long URL, a code identifier). A
        fixed layout with an explicit grid keeps the table inside the margins
        and wraps the cell text instead. A table that already fits is left on
        autofit, where Word's own font metrics beat the estimate below.
        """
        from docx.oxml.ns import qn
        from docx.shared import Emu, Twips

        ncols = len(table.columns)
        if ncols == 0:
            return

        # Nominal width of one character of Cambria body text at the document's
        # base size - a stand-in for the measurement only Word itself can make.
        # 120 twips is the figure for the template's 11pt.
        base_pt = self.DOCX_TEMPLATE_BASE_PT if base_pt is None else base_pt
        char_twips = 120 * base_pt / self.DOCX_TEMPLATE_BASE_PT
        cell_margin_twips = self.DOCX_CELL_MARGIN_TWIPS

        # Row 0 is only bold when the table actually has a header. A borderless
        # grid whose blank header row was dropped upstream has plain body text
        # in row 0, and style_docx_table records that by clearing tblLook's
        # firstRow flag - widening it by the bold factor would over-allocate
        # that column at its neighbours' expense.
        tblPr = table._tbl.tblPr
        tblLook = tblPr.find(qn('w:tblLook')) if tblPr is not None else None
        header_is_bold = not (
            tblLook is not None and tblLook.get(qn('w:firstRow')) == '0'
        )

        def measure(text, row_index):
            # The header row carries the table style's bold face
            width = len(text) * char_twips
            bold = row_index == 0 and header_is_bold
            return (width * 1.08 if bold else width) + cell_margin_twips

        # Walk the grid once, by row: _Column.cells rebuilds the whole cell
        # list on every access, which turns this into O(rows x cols x cols)
        grid = [list(row.cells) for row in table.rows]
        table_data = [[cell.text for cell in row] for row in grid]

        def image_twips(cell):
            """Width the widest inline image in a cell asks for, 0 if none.

            measured_column_widths caps this at an equal share, so an oversized
            image - it is scaled to whatever column it lands in - cannot starve
            the text columns beside it.
            """
            widest = max((Emu(int(ext.get('cx') or 0)).twips
                          for ext in cell._tc.iter(qn('wp:extent'))), default=0)
            return widest + cell_margin_twips if widest else 0.0

        image_floors = [0.0] * ncols
        for row in grid:
            for index, cell in enumerate(row[:ncols]):
                image_floors[index] = max(image_floors[index], image_twips(cell))

        widths = self.measured_column_widths(
            table_data, float(available_twips), measure, image_floors)

        # Under the page width the estimate is not worth imposing - autofit
        # lays the table out with real metrics and it cannot overflow anyway.
        # The tolerance matters: a fitted result sums to `available` in float
        # arithmetic, which lands a hair under it often enough to matter
        if sum(widths) < available_twips - 1:
            return

        for col, width in zip(table.columns, widths):
            col.width = Twips(int(width))
        for row in grid:
            for index, cell in enumerate(row[:ncols]):
                cell.width = Twips(int(widths[index]))

        # autofit=False writes w:tblLayout type="fixed" in its schema-mandated
        # position; appending the element by hand puts it out of sequence.
        table.autofit = False

    def markdown_to_html(self, content: str, title: str = 'Exported Document',
                         compact: bool = False, math_support: bool = False,
                         theme: str = 'light',
                         dark_background: str = '#111111',
                         light_background: str = '#ffffff',
                         base_pt: float = None) -> str:
        """Convert markdown to standalone HTML.

        Args:
            content: Markdown content to convert
            title: Document title
            compact: If True, use tighter spacing (for PDF)
            math_support: If True, inject KaTeX CSS/JS for client-side math rendering
            theme: 'system' (auto-detect via prefers-color-scheme), 'light', or 'dark'
            dark_background: Background color for dark theme (hex)
            light_background: Background color for light theme (hex)
            base_pt: Base body text size in points; every other size in the
                stylesheet is expressed in `em`, so this scales the document
        """
        base_pt = self.DEFAULT_FONT_SIZE_PT if base_pt is None else base_pt
        import markdown

        # tab_length=2 supports 2-space nested lists (GitHub/CommonMark)
        base = ['tables', 'fenced_code', 'codehilite', 'toc']
        try:
            md = markdown.Markdown(
                extensions=base + [manual_break_aware_nl2br()], tab_length=2)
        except Exception:
            # The manual-break rule reads Markdown's internals to tell an
            # authored break from a generated one, and the wiring that touches
            # them runs HERE, inside the constructor - not in the factory
            # above. It is a cosmetic refinement either way, so any failure
            # falls back: one blank line too many beats a 500 on every export.
            md = markdown.Markdown(extensions=base + ['nl2br'], tab_length=2)
        body = md.convert(content)

        if compact:
            # PDF-optimized stylesheet with tighter spacing
            style = '''
        @page {
            size: A4;
            margin: 0.5in;
        }
        body {
            font-family: Calibri, "Noto Color Emoji", sans-serif;
            font-size: ''' + f'{base_pt:g}' + '''pt;
            line-height: 1.4;
            margin: 0;
            padding: 0;
            color: #333;
        }
        p {
            margin: 0.1em 0 0.4em 0;
        }
        pre {
            background: #f4f4f4;
            padding: 8px;
            margin: 0.5em 0;
            font-size: 9pt;
        }
        code {
            background: #f4f4f4;
            padding: 1px 3px;
            font-family: Courier, monospace;
            font-size: 9pt;
        }
        pre code {
            background: none;
            padding: 0;
        }
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 0.5em 0;
            font-size: 9pt;
            table-layout: fixed;
            word-wrap: break-word;
        }
        th, td {
            border: 1px solid #ddd;
            padding: 3px 5px;
            text-align: left;
            overflow: hidden;
        }
        th {
            background: #dbe5f1;
            color: #365F91;
            font-weight: bold;
        }
        img {
            max-width: 100%;
            height: auto;
        }
        blockquote {
            border-left: 3px solid #4F81BD;
            margin: 0.5em 0;
            padding-left: 10px;
            color: #666;
        }
        h1, h2, h3, h4, h5, h6 {
            font-family: Calibri, "Noto Color Emoji", sans-serif;
        }
        h1 {
            font-size: 18pt;
            margin: 0.6em 0 0.15em 0;
            color: #365F91;
        }
        h2 {
            font-size: 14pt;
            margin: 0.5em 0 0.1em 0;
            color: #4F81BD;
        }
        h3 {
            font-size: 12pt;
            margin: 0.4em 0 0.1em 0;
            color: #4F81BD;
        }
        h4 {
            font-size: 11pt;
            margin: 0.3em 0 0.1em 0;
            color: #4F81BD;
        }
        h5, h6 {
            font-size: 11pt;
            margin: 0.3em 0 0.1em 0;
            color: #243F60;
        }
        ul, ol {
            margin: 0.3em 0;
            padding-left: 1.5em;
        }
        li {
            margin: 0.1em 0;
        }
        a {
            color: #0563C1;
            text-decoration: underline;
            text-underline-offset: 2px;
        }'''
        else:
            # Standard HTML stylesheet with Pygments syntax highlighting
            pygments_css = self.get_pygments_css()
            pygments_dark_css = self.get_pygments_css(dark=True)

            use_dark = theme == 'dark'
            use_light = theme == 'light'
            use_system = theme == 'system'

            def adjust_color(hex_color: str, amount: int) -> str:
                """Lighten (positive) or darken (negative) a hex color."""
                hex_color = hex_color.lstrip('#')
                r = max(0, min(255, int(hex_color[0:2], 16) + amount))
                g = max(0, min(255, int(hex_color[2:4], 16) + amount))
                b = max(0, min(255, int(hex_color[4:6], 16) + amount))
                return f'#{r:02x}{g:02x}{b:02x}'

            # Determine base styles based on theme
            if use_dark:
                # Dark theme as base (no media query)
                bg_color = dark_background
                text_color = '#e6edf3'
                link_color = '#58a6ff'
                code_bg = '#161b22'
                border_color = '#30363d'
                heading_color = '#e6edf3'
                blockquote_color = '#8b949e'
                th_bg = adjust_color(dark_background, 20)
                tr_stripe = adjust_color(dark_background, 14)
                tr_hover = adjust_color(dark_background, 22)
                code_bg = adjust_color(dark_background, 20)
                border_color = adjust_color(dark_background, 40)
                active_pygments = pygments_dark_css
            else:
                # Light theme as base (for both 'light' and 'system')
                bg_color = light_background
                text_color = '#1a1a1a'
                link_color = '#0969da'
                code_bg = adjust_color(light_background, -6)
                border_color = adjust_color(light_background, -46)
                heading_color = '#1a1a1a'
                blockquote_color = '#656d76'
                th_bg = adjust_color(light_background, -6)
                tr_stripe = adjust_color(light_background, -6)
                tr_hover = adjust_color(light_background, -14)
                active_pygments = pygments_css

            style = f'''
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
            font-size: {base_pt:g}pt;
            line-height: 1.6;
            max-width: 50em;
            margin: 0 auto;
            padding: 20px;
            color: {text_color};
            background-color: {bg_color};
        }}
        a {{
            color: {link_color};
        }}
        pre {{
            background: {code_bg};
            padding: 10px;
            border-radius: 6px;
            overflow-x: auto;
            font-size: 0.85em;
            border: 1px solid {border_color};
        }}
        code {{
            background: {code_bg};
            padding: 2px 6px;
            border-radius: 4px;
            font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace;
            font-size: 0.85em;
        }}
        pre code {{
            background: none;
            padding: 0;
            border: none;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }}
        .table-scroll {{
            /* Wrapping alone cannot shrink a table past one character per
               column, so a table with very many columns scrolls inside this
               box rather than pushing the document sideways. It has to be a
               wrapper: overflow-x on the table itself needs display:block,
               which makes the grid shrink-to-fit and voids its width */
            overflow-x: auto;
        }}
        @media print {{
            /* Paper cannot scroll, so a scroll box crops what it hides -
               fewer columns than the unwrapped table would have printed */
            .table-scroll {{
                overflow-x: visible;
            }}
            /* Keep a table header with a body row so it is never stranded
               alone at a page bottom, and repeat it on each page */
            thead {{
                display: table-header-group;
                break-inside: avoid;
                break-after: avoid;
            }}
            tr {{
                break-inside: avoid;
            }}
            /* Cap a tall image (a mermaid diagram) to one printed page so it
               is not split across pages. A Letter page less ~1in of margins
               is about 9in; vh units are unreliable in print and width:auto
               defeats the cap, so a concrete bound is used. The base
               max-width:100% and height:auto keep the width and aspect ratio. */
            img {{
                max-height: 9in;
            }}
        }}
        th, td {{
            border: 1px solid {border_color};
            padding: 8px;
            text-align: left;
            overflow-wrap: anywhere;
            /* anywhere, not break-word: only anywhere shrinks a column's
               min-content width, which is what keeps a table with a long
               unbreakable token inside the page */
        }}
        th {{
            background: {th_bg};
            font-weight: 600;
        }}
        tbody tr:nth-child(even) {{
            background: {tr_stripe};
        }}
        tbody tr:hover {{
            background: {tr_hover};
        }}
        img {{
            max-width: 100%;
            height: auto;
        }}
        blockquote {{
            border-left: 4px solid {border_color};
            margin: 0;
            padding-left: 16px;
            color: {blockquote_color};
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 1.5em;
            margin-bottom: 0.5em;
            color: {heading_color};
        }}
        h1 {{
            padding-bottom: 0.3em;
            border-bottom: 1px solid {border_color};
        }}
        hr {{
            border: none;
            border-top: 1px solid {border_color};
            margin: 1.5em 0;
        }}
        /* Pygments syntax highlighting */
        {active_pygments}'''

            # System theme: add dark overrides via media query
            # Derive dark UI colors from dark_background
            if use_system:
                dk_code = adjust_color(dark_background, 20)
                dk_border = adjust_color(dark_background, 40)
                dk_stripe = adjust_color(dark_background, 14)
                dk_hover = adjust_color(dark_background, 22)
                style += f'''
        @media (prefers-color-scheme: dark) {{
            body {{
                background-color: {dark_background};
                color: #e6edf3;
            }}
            a {{
                color: #58a6ff;
            }}
            pre {{
                background: {dk_code};
                border-color: {dk_border};
            }}
            code {{
                background: {dk_code};
            }}
            pre code {{
                background: none;
            }}
            th, td {{
                border-color: {dk_border};
            }}
            th {{
                background: {dk_code};
                color: #e6edf3;
            }}
            tbody tr:nth-child(even) {{
                background: {dk_stripe};
            }}
            tbody tr:hover {{
                background: {dk_hover};
            }}
            blockquote {{
                border-left-color: {dk_border};
                color: #8b949e;
            }}
            h1, h2, h3, h4, h5, h6 {{
                color: #e6edf3;
            }}
            h1 {{
                border-bottom-color: {dk_border};
            }}
            hr {{
                border-top-color: {dk_border};
            }}
            img {{
                opacity: 0.9;
            }}
            /* Pygments dark syntax highlighting */
            {pygments_dark_css}
        }}'''

        katex_head = ''
        katex_script = ''
        if math_support:
            katex_head = '''
    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.21/dist/katex.min.css">
    <script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.21/dist/katex.min.js"></script>
    <script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.21/dist/contrib/auto-render.min.js"></script>'''
            katex_script = '''
    <script>
      document.addEventListener("DOMContentLoaded", function() {
        renderMathInElement(document.body, {
          delimiters: [
            {left: "$$", right: "$$", display: true},
            {left: "$", right: "$", display: false}
          ],
          ignoredTags: ["script", "noscript", "style", "textarea", "pre", "code"],
          throwOnError: false
        });
      });
    </script>'''

        # Force color-scheme on <html> element so embedded SVGs with
        # @media (prefers-color-scheme) follow the chosen theme
        if use_dark:
            html_style = ' style="color-scheme: dark"'
        elif use_light:
            html_style = ' style="color-scheme: light"'
        else:
            html_style = ' style="color-scheme: light dark"'

        html = f'''<!DOCTYPE html>
<html{html_style}>
<head>
    <meta charset="utf-8">
    <title>{title}</title>
    <style>{style}
    </style>{katex_head}
</head>
<body>
{body}{katex_script}
</body>
</html>'''
        return html

    @staticmethod
    def pdf_face_covers(font_name: str, text: str) -> bool:
        """Whether the registered PDF face ``font_name`` can draw every
        character of ``text``.

        reportlab paints a character the face has no glyph for as a blank
        advance - silently, with nothing in the output to say anything was
        dropped. It matters for the italic slots: DejaVu ships no oblique on
        most Linux boxes, so those are filled from Liberation, which carries
        no star, check mark or ballot box. Slant is decoration and the glyph
        is content, so a run the italic face cannot draw is rendered upright
        rather than blank.

        The answer is per RUN, not per character - the PDF path does no run
        splitting - so one star costs its whole paragraph the slant.
        """
        from reportlab.pdfbase import pdfmetrics

        try:
            face = pdfmetrics.getFont(font_name).face
        except Exception:
            # A name nothing registered - keep the styling rather than
            # suppress it on a guess
            return True
        cmap = getattr(face, 'charToGlyph', None)
        if cmap is None:
            # A core Type1 face, which reportlab reaches for when no oblique
            # file exists at all. It is WinAnsi-only and above that swaps in
            # the symbol encoding, so a star comes out as a dingbat rather
            # than as a gap - worse than the blank this guard exists to stop
            return all(ord(ch) < 256 for ch in text if ch not in '\n\r\t')
        # A break or a tab is layout, not a glyph: no cmap carries U+000A, and
        # reportlab turns both into markup before it ever looks a glyph up. Ask
        # about them and every multi-line run answers False for the wrong reason
        return all(ord(ch) in cmap for ch in text if ch not in '\n\r\t')

    def _register_unicode_fonts(self):
        """Register Unicode-supporting fonts from system paths with font family support."""
        from reportlab.pdfbase.pdfmetrics import registerFontFamily

        # Define font sets with normal, bold, italic, bolditalic variants
        font_sets = [
            # DejaVu fonts (most common, excellent Unicode support)
            {
                'normal': '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                'bold': '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
                'italic': '/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf',
                'boldItalic': '/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf',
            },
            # Liberation fonts (alternative)
            {
                'normal': '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
                'bold': '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf',
                'italic': '/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf',
                'boldItalic': '/usr/share/fonts/truetype/liberation/LiberationSans-BoldItalic.ttf',
            },
            # FreeSans (GNU FreeFont)
            {
                'normal': '/usr/share/fonts/truetype/freefont/FreeSans.ttf',
                'bold': '/usr/share/fonts/truetype/freefont/FreeSansBold.ttf',
                'italic': '/usr/share/fonts/truetype/freefont/FreeSansOblique.ttf',
                'boldItalic': '/usr/share/fonts/truetype/freefont/FreeSansBoldOblique.ttf',
            },
        ]

        font_names = {
            'normal': 'UnicodeSans',
            'bold': 'UnicodeSansBold',
            'italic': 'UnicodeSansItalic',
            'boldItalic': 'UnicodeSansBoldItalic',
        }

        registered_fonts = set()

        # Fill each variant from the first family that ships it, rather than
        # committing to one family and stopping. DejaVu (first, best Unicode
        # coverage) has no oblique files on many boxes, so the old loop left
        # the italic faces unregistered and every italic - body emphasis and
        # now the H4/H6 heading faces - fell back to a Helvetica core font, a
        # visible typeface switch mid-document. Normal and bold still come from
        # the highest-priority family present, so the dominant text keeps that
        # family's glyph coverage; only the italic slots reach past it.
        for variant in ('normal', 'bold', 'italic', 'boldItalic'):
            font_name = font_names[variant]
            for font_set in font_sets:
                path = font_set.get(variant)
                if path and os.path.exists(path):
                    try:
                        pdfmetrics.registerFont(TTFont(font_name, path))
                        registered_fonts.add(font_name)
                    except Exception:
                        continue
                    break

        # Register font family to enable <b> and <i> tags in Paragraph
        # What the italic slots actually resolved to. pdf_face_covers has to
        # ask about the face reportlab will draw with: when no oblique file
        # exists at all the slot falls back to a core Helvetica, and asking
        # about the unregistered `UnicodeSansItalic` returns True and lets the
        # blank glyph through - the guard inert on exactly the box it is for.
        self._pdf_italic_face = (
            'UnicodeSansItalic' if 'UnicodeSansItalic' in registered_fonts
            else 'Helvetica-Oblique')
        self._pdf_bold_italic_face = (
            'UnicodeSansBoldItalic' if 'UnicodeSansBoldItalic' in registered_fonts
            else 'Helvetica-BoldOblique')

        if 'UnicodeSans' in registered_fonts:
            try:
                italic_font = self._pdf_italic_face
                bold_italic_font = self._pdf_bold_italic_face

                registerFontFamily(
                    'UnicodeSans',
                    normal='UnicodeSans',
                    bold='UnicodeSansBold' if 'UnicodeSansBold' in registered_fonts else 'UnicodeSans',
                    italic=italic_font,
                    boldItalic=bold_italic_font
                )
            except Exception:
                pass

    def convert_docx_to_pdf(self, docx_bytes: bytes, code_blocks: list = None,
                            base_pt: float = None) -> bytes:
        """
        Convert DOCX document to PDF using python-docx + reportlab.

        Args:
            docx_bytes: DOCX file content as bytes
            code_blocks: Optional list of code blocks extracted from markdown

        Returns:
            PDF file content as bytes
        """
        if code_blocks is None:
            code_blocks = []
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab is required for PDF export")

        from docx import Document

        # Register Unicode fonts
        self._register_unicode_fonts()

        # Read the DOCX from bytes
        doc = Document(io.BytesIO(docx_bytes))

        # Determine which font to use
        registered = pdfmetrics.getRegisteredFontNames()
        font_name = 'UnicodeSans' if 'UnicodeSans' in registered else 'Helvetica'
        font_name_bold = ('UnicodeSansBold' if 'UnicodeSansBold' in registered
                          else 'Helvetica-Bold')
        # The minor headings need faces the body and H1-H3 never ask for: H5 is
        # regular (already `font_name`), H6 italic, H4 bold italic. Only those
        # two italic faces are looked up here; a missing oblique file falls to
        # the Helvetica core face, which _register_unicode_fonts now avoids by
        # taking the italics from a family that ships them.
        heading_faces = {
            (False, False): font_name,
            (False, True): self._pdf_italic_face,
            (True, True): self._pdf_bold_italic_face,
        }

        # Create PDF in memory
        pdf_buffer = io.BytesIO()
        pdf_doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=letter,
            rightMargin=self.PDF_PAGE_MARGIN,
            leftMargin=self.PDF_PAGE_MARGIN,
            topMargin=self.PDF_PAGE_MARGIN,
            bottomMargin=self.PDF_PAGE_MARGIN
        )

        # Get default styles
        styles = getSampleStyleSheet()

        # Every size below is a proportion of the base body size, so the
        # document scales as one and matches what the DOCX renders at
        base_pt = self.DEFAULT_FONT_SIZE_PT if base_pt is None else base_pt
        body_type = self.pdf_type('body', base_pt)
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=font_name,
            **body_type,
            spaceAfter=0.6 * base_pt
        )

        # Level-specific list styles for bullets
        list_bullet_style = ParagraphStyle(
            'CustomListBullet',
            parent=styles['Normal'],
            fontName=font_name,
            **body_type,
            spaceAfter=0.3 * base_pt,
            leftIndent=18,
            bulletIndent=6
        )

        list_bullet_2_style = ParagraphStyle(
            'CustomListBullet2',
            parent=styles['Normal'],
            fontName=font_name,
            **body_type,
            spaceAfter=0.3 * base_pt,
            leftIndent=36,
            bulletIndent=24
        )

        # Level-specific list styles for numbered lists
        list_number_style = ParagraphStyle(
            'CustomListNumber',
            parent=styles['Normal'],
            fontName=font_name,
            **body_type,
            spaceAfter=0.3 * base_pt,
            leftIndent=18,
            bulletIndent=6
        )

        list_number_2_style = ParagraphStyle(
            'CustomListNumber2',
            parent=styles['Normal'],
            fontName=font_name,
            **body_type,
            spaceAfter=0.3 * base_pt,
            leftIndent=36,
            bulletIndent=24
        )

        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontName=font_name_bold,
            **self.pdf_type('heading1', base_pt),
            spaceAfter=0.6 * base_pt,
            spaceBefore=1.0 * base_pt,
            textColor=colors.HexColor('#365F91')
        )

        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontName=font_name_bold,
            **self.pdf_type('heading2', base_pt),
            spaceAfter=0.4 * base_pt,
            spaceBefore=0.8 * base_pt,
            textColor=colors.HexColor('#4F81BD')
        )

        heading3_style = ParagraphStyle(
            'CustomHeading3',
            parent=styles['Heading3'],
            fontName=font_name_bold,
            **self.pdf_type('heading3', base_pt),
            spaceAfter=0.3 * base_pt,
            spaceBefore=0.6 * base_pt,
            textColor=colors.HexColor('#4F81BD')
        )

        heading_styles = {1: heading1_style, 2: heading2_style,
                          3: heading3_style}
        for level, (bold, italic, hex_color) in \
                self.PDF_MINOR_HEADING_FACES.items():
            heading_styles[level] = ParagraphStyle(
                f'CustomHeading{level}',
                parent=styles[f'Heading{level}'],
                fontName=heading_faces[(bold, italic)],
                **self.pdf_type(f'heading{level}', base_pt),
                spaceAfter=0.25 * base_pt,
                spaceBefore=0.5 * base_pt,
                textColor=colors.HexColor(hex_color)
            )

        # Callout body (blockquotes, alert boxes). Same size as body text but
        # no trailing space - the surrounding callout table supplies padding.
        callout_body_style = ParagraphStyle(
            'CustomCalloutBody',
            parent=normal_style,
            spaceAfter=0,
            spaceBefore=0,
        )

        # Table cells are Paragraphs so their text wraps inside the column
        table_cell_style = ParagraphStyle(
            'CustomTableCell',
            parent=styles['Normal'],
            fontName=font_name,
            **self.pdf_type('table', base_pt)
        )

        table_header_style = ParagraphStyle(
            'CustomTableHeader',
            parent=table_cell_style,
            fontName=font_name_bold,
            textColor=colors.HexColor('#365F91')
        )

        # Build the story (content) - iterate body elements in document order
        story = []

        from docx.text.paragraph import Paragraph as DocxParagraph
        from docx.table import Table as DocxTable
        from docx.oxml.ns import qn

        def get_list_info(para):
            """Get list type and level from paragraph style and indentation.

            Returns: (list_type, level) where list_type is 'number', 'bullet', or None
            """
            style_name = para.style.name if para.style else ''

            # Determine list type from style name
            list_type = None
            if 'List Number' in style_name:
                list_type = 'number'
            elif 'List Bullet' in style_name:
                list_type = 'bullet'
            elif 'List' in style_name:
                list_type = 'bullet'

            if list_type is None:
                # Check numPr for lists without explicit style
                try:
                    if para._element.pPr is not None and para._element.pPr.numPr is not None:
                        list_type = 'bullet'
                except AttributeError:
                    pass

            if list_type is None:
                return (None, 0)

            # Determine level from style name first (List Number 2, List Bullet 2)
            if '2' in style_name or '3' in style_name:
                level = 1
            else:
                # Check leftIndent for nesting level (720 = level 0, 1440+ = level 1+)
                level = 0
                try:
                    pPr = para._element.pPr
                    if pPr is not None:
                        ind = pPr.find(qn('w:ind'))
                        if ind is not None:
                            left_val = ind.get(qn('w:left'))
                            if left_val:
                                left_indent = int(left_val)
                                # 720 twips = level 0, 1440+ = level 1+
                                if left_indent > 720:
                                    level = 1
                except (AttributeError, ValueError):
                    pass

            return (list_type, level)

        # Track numbering for ordered lists
        number_counters = {0: 0, 1: 0, 2: 0}
        last_list_level = -1
        list_ids = {}  # depth -> numbering instance of the list open there

        def is_horizontal_rule(para):
            """Check if paragraph represents a horizontal divider line."""
            try:
                pPr = para._element.pPr
                if pPr is not None:
                    pBdr = pPr.find(qn('w:pBdr'))
                    if pBdr is not None:
                        bottom = pBdr.find(qn('w:bottom'))
                        top = pBdr.find(qn('w:top'))
                        if bottom is not None or top is not None:
                            if not para.text.strip():
                                return True
            except (AttributeError, TypeError):
                pass
            return False

        def format_run(run):
            """Format a single run with all its styling."""
            run_text = run.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            if not run_text:
                return run_text

            # Convert newlines to HTML line breaks
            run_text = run_text.replace('\n', '<br/>')
            result = run_text

            # Apply formatting tags (can be combined)
            if run.bold:
                result = f'<b>{result}</b>'
            # Slant only when the italic face can actually draw the text:
            # see pdf_face_covers - an italic table cell or blockquote would
            # otherwise delete a star rather than un-slant it
            if run.italic and self.pdf_face_covers(
                    self._pdf_bold_italic_face if run.bold
                    else self._pdf_italic_face, run.text):
                result = f'<i>{result}</i>'
            if run.underline:
                result = f'<u>{result}</u>'
            if run.font.strike:
                result = f'<strike>{result}</strike>'
            if run.font.subscript:
                result = f'<sub>{result}</sub>'
            if run.font.superscript:
                result = f'<super>{result}</super>'

            # Check for text color
            try:
                if run.font.color and run.font.color.rgb:
                    color_hex = str(run.font.color.rgb)
                    if color_hex and color_hex != 'None' and len(color_hex) == 6:
                        result = f'<font color="#{color_hex}">{result}</font>'
            except (AttributeError, TypeError):
                pass

            # Run shading - what a `<mark>` and a coloured pill become in the
            # intermediate DOCX. python-docx has no accessor for it, and
            # without this the highlight the reader sees in Word and in the
            # browser is simply absent from the PDF.
            fill = self.docx_run_shading(run)
            if fill:
                result = f'<font backColor="#{fill}">{result}</font>'

            return result

        # Word's alignment enum -> reportlab's. LEFT is the default in both,
        # so it is left out and an unaligned paragraph keeps its own style.
        pdf_align = {1: 1, 2: 2, 3: 4}

        def aligned(style, para):
            """``style``, or a copy of it carrying the paragraph's alignment."""
            ta = pdf_align.get(para.alignment)
            if ta is None:
                return style
            return ParagraphStyle(f'{style.name}-align{ta}', parent=style,
                                  alignment=ta)

        def heading_style_for(para):
            """The heading style a DOCX paragraph maps onto, or None when it is
            not a heading.

            One ladder, read by the body and by the inside of a callout alike:
            a heading routed into a box used to be painted as callout body
            text, so the same line of source read as a heading in Word and as
            body text in the PDF.
            """
            style_name = para.style.name if para.style else ''
            if not style_name.startswith('Heading'):
                return None
            level = re.match(r'Heading (\d+)', style_name)
            # An unnumbered or out-of-range 'Heading ...' keeps the old
            # catch-all target rather than losing its heading face
            style = heading_styles.get(
                int(level.group(1)) if level else 0, heading3_style)
            # A minor heading takes its slant from the STYLE, not from a
            # run, so format_run's gate never sees it - and the italic
            # faces are the ones with no star, check mark or ballot box.
            # Same trade as a run: upright and visible beats slanted and
            # blank.
            upright = (font_name_bold if 'Bold' in style.fontName
                       else font_name)
            if (not self.pdf_face_covers(style.fontName, para.text)
                    and self.pdf_face_covers(upright, para.text)):
                style = ParagraphStyle(f'{style.name}-upright',
                                       parent=style, fontName=upright)
            return style

        def process_paragraph(para):
            """Process a single paragraph and return reportlab element(s)."""
            nonlocal last_list_level

            # Check for horizontal rule/divider first
            if is_horizontal_rule(para):
                from reportlab.platypus import HRFlowable
                return [HRFlowable(width="100%", thickness=0.5, color=colors.grey, spaceBefore=3, spaceAfter=6)]

            text = para.text.strip()
            if not text and get_list_info(para)[0] != 'number':
                # Empty paragraph - render as actual blank line (not invisible
                # spacer). An empty numbered item still takes its number,
                # as Word gives it one
                return Paragraph("&nbsp;", normal_style)

            # Check for code block placeholder [[CODE_BLOCK_N]]
            code_match = re.match(r'\[\[CODE_BLOCK_(\d+)\]\]', text)
            if code_match and code_blocks:
                idx = int(code_match.group(1))
                if idx < len(code_blocks):
                    block = code_blocks[idx]
                    return self.highlight_code_for_pdf(
                        block['code'], block['lang'], base_pt,
                        avail_width=frame_width)
                return Paragraph("&nbsp;", normal_style)

            # Escape XML special characters for base text
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            text = text.replace('\n', '<br/>')

            # Check if any run has formatting that needs processing.
            # docx_paragraph_runs, not para.runs: the rebuild below replaces
            # para.text, which does include a link's label, so a walk that
            # stops at direct children would silently delete every hyperlink
            # from a paragraph that happens to carry any formatting.
            has_formatting = False
            for run in self.docx_paragraph_runs(para):
                if run.text.strip():
                    if (run.bold or run.italic or run.underline or
                        run.font.strike or run.font.subscript or run.font.superscript):
                        has_formatting = True
                        break
                    # Shading is the one run property with no python-docx
                    # accessor, so it has to be read off the XML - and a run
                    # whose only styling is a highlight is exactly the case
                    # this probe used to miss, sending the paragraph down the
                    # plain-text path where format_run never runs
                    if self.docx_run_shading(run):
                        has_formatting = True
                        break
                    try:
                        if run.font.color and run.font.color.rgb:
                            has_formatting = True
                            break
                    except (AttributeError, TypeError):
                        pass

            if has_formatting:
                formatted_parts = [format_run(run)
                                   for run in self.docx_paragraph_runs(para)]
                text = ''.join(formatted_parts)

            # Detect heading styles
            heading = heading_style_for(para)
            if heading is not None:
                return Paragraph(text, aligned(heading, para))

            # Blockquote: style_docx_blockquotes gave it a left border, shading
            # and indent that process_paragraph would otherwise drop. Render it
            # as a left-barred, shaded, indented callout so the PDF matches the
            # DOCX. The runs already carry the muted italic colour.
            bq = blockquote_info(para)
            if bq is not None:
                indent_pts, bar_hex, shd_hex = bq
                return make_callout(
                    [Paragraph(text, callout_body_style)],
                    bar_hex, shd_hex, left_pad=12 + max(0.0, indent_pts),
                    trailing=0.05,
                )

            # Check for list items
            list_type, level = get_list_info(para)

            if list_type == 'number':
                # A list ends where restart_docx_list_numbering began the
                # next one - the instance on the paragraph changes. Nothing
                # else is a boundary: a table, a sample or a paragraph
                # inside a step used to restart the count
                num_pr = para._element.pPr.numPr if para._element.pPr is not None else None
                num_id = num_pr.numId.val if num_pr is not None and num_pr.numId is not None else None
                if list_ids.get(level) != num_id:
                    list_ids[level] = num_id
                    for l in range(level, 3):
                        number_counters[l] = 0
                # Reset lower levels when moving up, increment current level
                if level <= last_list_level:
                    for l in range(level + 1, 3):
                        number_counters[l] = 0
                number_counters[level] += 1
                last_list_level = level

                prefix = f"{number_counters[level]}. "
                style = list_number_2_style if level > 0 else list_number_style
                return Paragraph(f'{prefix}{text}', style)

            elif list_type == 'bullet':
                last_list_level = level
                style = list_bullet_2_style if level > 0 else list_bullet_style
                return Paragraph(f'• {text}', style)

            else:
                return Paragraph(text, aligned(normal_style, para))

        # Printable frame: the page area less the frame's own 6pt padding each
        # side. Sizing a flowable to pdf_doc.width/height instead would put its
        # edge past the page margin. One source for both tables and images.
        frame_width = pdf_doc.width - 2 * self.PDF_FRAME_PADDING
        frame_height = pdf_doc.height - 2 * self.PDF_FRAME_PADDING

        def make_callout(flowables, bar_hex, shd_hex, left_pad=12, trailing=0.12):
            """A full-width box with a coloured left bar and background shading.

            Shared by blockquotes and GitHub alert boxes. ``splitInRow`` lets a
            callout taller than a page carry over instead of raising
            LayoutError - but it is enabled only for such a callout. Left on,
            it also tears a box that merely runs past the space left on the
            current page, splitting a quote or alert mid-frame when the whole
            box would have fit the next page.
            """
            t = Table([[flowables]], colWidths=[frame_width],
                      hAlign='LEFT', splitInRow=1)
            t.setStyle(TableStyle([
                ('LINEBEFORE', (0, 0), (0, -1), 3, colors.HexColor('#' + bar_hex)),
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#' + shd_hex)),
                ('LEFTPADDING', (0, 0), (-1, -1), left_pad),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            t.wrap(frame_width, frame_height)
            heights = getattr(t, '_rowHeights', None)
            if heights:
                t.splitInRow = 1 if max(heights) > frame_height else 0
            return [t, Spacer(1, trailing * inch)]

        def blockquote_info(para):
            """(indent_pts, bar_hex, shd_hex) if the paragraph is a styled
            blockquote (a left ``w:pBdr`` from style_docx_blockquotes), else None.
            Horizontal rules use a bottom/top border and empty text, already
            handled before this is reached."""
            pPr = para._element.pPr
            if pPr is None:
                return None
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is None:
                return None
            left = pBdr.find(qn('w:left'))
            if left is None:
                return None
            bar = left.get(qn('w:color')) or 'BBBBBB'
            if bar in ('auto', ''):
                bar = 'BBBBBB'
            shd = pPr.find(qn('w:shd'))
            shd_fill = shd.get(qn('w:fill')) if shd is not None else None
            if not shd_fill or shd_fill == 'auto':
                shd_fill = 'F4F4F4'
            indent_pts = 0.0
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                lv = ind.get(qn('w:left'))
                if lv:
                    try:
                        indent_pts = int(lv) / 20.0
                    except ValueError:
                        pass
            return (indent_pts, bar, shd_fill)

        def alert_info(tbl):
            """(bar_hex, shd_hex) if the table is a GitHub alert box (a single
            1x1 cell with a coloured left border, the other sides none, from
            style_docx_alert_boxes), else None. The 1x1 shape is required so a
            normal multi-row/column table that happens to share the border
            signature is never routed to the single-cell callout (which would
            drop every other cell)."""
            if len(tbl.rows) != 1 or len(tbl.rows[0].cells) != 1:
                return None
            tblPr = tbl._tbl.find(qn('w:tblPr'))
            if tblPr is None:
                return None
            borders = tblPr.find(qn('w:tblBorders'))
            if borders is None:
                return None
            left = borders.find(qn('w:left'))
            if left is None or left.get(qn('w:val')) != 'single':
                return None
            bar = left.get(qn('w:color'))
            # Black is a colour a hand-drawn box may be bordered with, so it is
            # not disqualifying; what keeps an ordinary bordered table out is
            # the 1x1 shape above and the blank other sides below, which no
            # content table has
            if not bar or bar == 'auto':
                return None
            top = borders.find(qn('w:top'))
            if top is not None and top.get(qn('w:val')) not in ('none', None):
                return None
            shd_fill = 'FFFFFF'
            try:
                tcPr = tbl.rows[0].cells[0]._tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None and shd.get(qn('w:fill')):
                        shd_fill = shd.get(qn('w:fill'))
            except (IndexError, AttributeError):
                pass
            return (bar, shd_fill)

        def process_alert(tbl, bar_hex, shd_hex):
            """Render an alert box as a coloured-left-bar shaded callout with
            normal (not table-header) text, preserving run formatting."""
            flow = []
            for p in tbl.rows[0].cells[0].paragraphs:
                markup = ''.join(format_run(run) for run in self.docx_paragraph_runs(p)).strip()
                if markup:
                    style = heading_style_for(p) or callout_body_style
                    if not flow and style.spaceBefore:
                        # The cell's own top padding is the space above the
                        # first line; the heading's would sit on top of it
                        style = ParagraphStyle(f'{style.name}-boxtop',
                                               parent=style, spaceBefore=0)
                    flow.append(Paragraph(markup, style))
            if not flow:
                flow = [Paragraph('&nbsp;', callout_body_style)]
            return make_callout(flow, bar_hex, shd_hex, left_pad=12)

        def process_table(tbl):
            """Process a single table and return reportlab elements."""
            # A GitHub alert box is a single-cell table with a coloured left
            # bar; render it as a callout rather than a header-styled table.
            alert = alert_info(tbl)
            if alert is not None:
                return process_alert(tbl, alert[0], alert[1])

            # Raw text, so the widths below are measured on what is rendered
            # rather than on the XML entities that stand in for it
            table_data = [[cell.text for cell in row.cells] for row in tbl.rows]

            if not table_data:
                return []

            # A borderless image/layout grid is written with an empty header row
            # (`|  |  |  |`) because Markdown requires a header. The rendered
            # markdown shows nothing for it, so drop it entirely rather than
            # emitting a blank bordered first row above the grid.
            rows = list(tbl.rows)
            dropped_empty_header = False

            if len(table_data) > 1 and self.docx_row_is_blank(rows[0]):
                table_data = table_data[1:]
                rows = rows[1:]
                dropped_empty_header = True

            # A Markdown layout grid carries an empty header row (`|  |  |  |`);
            # it must not be styled or repeated as a header, or a blank blue bar
            # detaches from its rows across a page break. Row 0 is a real header
            # only when it holds text - a picture-only row is a figure, and the
            # DOCX path reads it the same way. A grid whose empty header was
            # dropped has no header at all: the row now in position 0 is body
            # content and must not be promoted into one.
            has_header = (not dropped_empty_header
                          and any(text.strip() for text in table_data[0]))

            def cell_markup(text):
                markup = (text.replace('&', '&amp;')
                              .replace('<', '&lt;').replace('>', '&gt;'))
                # A Paragraph collapses newlines that a string cell honoured
                return markup.replace('\n', '<br/>')

            def string_width(text, row_index):
                # Measured in the bold face for every row, header or not. A
                # body cell now renders whatever formatting its runs carry, so
                # measuring it in the regular face under-allocates the column
                # and reportlab hard-splits a bold word mid-word. Widening one
                # column costs its neighbours nothing here: every cell is
                # measured the same way and fit_column_widths renormalises the
                # row back to the available width.
                return pdfmetrics.stringWidth(
                    text, font_name_bold, table_cell_style.fontSize)

            # Per-column widest image (points): an image-only cell earns no
            # column width from its empty text, so without this floor its column
            # collapses and the cell image renders a few points wide. Mirrors the
            # DOCX path's image_floors; the extent cx is in EMU.
            from docx.shared import Emu
            ncols = max((len(r) for r in table_data), default=0)
            image_widths = [0.0] * ncols
            for row in rows:
                for c, cell in enumerate(row.cells):
                    if c >= ncols:
                        continue
                    widest = max((Emu(int(ext.get('cx') or 0)).pt
                                  for ext in cell._tc.iter(qn('wp:extent'))),
                                 default=0)
                    if widest:
                        image_widths[c] = max(image_widths[c], widest)

            side_padding, col_widths = self.pdf_table_column_layout(
                table_data, frame_width, string_width, image_widths)

            # A cell adds 4pt top + 4pt bottom padding around its content, and
            # the header row overrides its bottom padding to 8pt (12pt total).
            # An RLImage is atomic (never splits), so a cell image capped at the
            # full frame height plus that padding would exceed one page and
            # raise LayoutError. Cap below the frame by the largest per-row
            # padding (the 12pt header row) so no row - header included - can
            # overflow.
            cell_image_max_height = frame_height - 12

            def cell_content(cell, col_width, style):
                """Flowables for one cell: its text plus any inline images,
                scaled to the column so a cell image renders instead of being
                dropped. Walks the cell's paragraphs so text and images keep
                document order."""
                content_width = max(1.0, col_width - 2 * side_padding)
                flow = []
                for p in cell.paragraphs:
                    # format_run per run, not cell_markup(p.text): a cell
                    # rebuilt from the plain-string projection drops every
                    # run property the body path keeps. docx_paragraph_runs
                    # so a link's label survives - p.runs would delete it.
                    ptext = ''.join(format_run(r)
                                    for r in self.docx_paragraph_runs(p))
                    if ptext.strip():
                        flow.append(Paragraph(ptext, style))
                    for drawing in p._p.findall('.//' + qn('w:drawing')):
                        img = process_image(drawing, doc, max_width=content_width,
                                            max_height=cell_image_max_height)
                        if img is not None:
                            flow.append(img)
                if not flow:
                    flow = [Paragraph(cell_markup(cell.text), style)]
                return flow

            # Paragraph/Image cells wrap and fit inside the column.
            wrapped = []
            for r, row in enumerate(rows):
                style = (table_header_style if (r == 0 and has_header)
                         else table_cell_style)
                wrapped.append([
                    cell_content(cell,
                                 col_widths[c] if c < len(col_widths) else frame_width,
                                 style)
                    for c, cell in enumerate(row.cells)
                ])

            # splitInRow lets a single row taller than the page carry over to
            # the next one - without it reportlab raises LayoutError and the
            # whole export fails, which wrapping made reachable. repeatRows=1
            # repeats a real header on each continuation page and, as a side
            # effect, defers the whole table to the next page rather than
            # stranding the header alone at a page bottom.
            t = Table(wrapped, colWidths=col_widths, hAlign='LEFT',
                      splitInRow=1, repeatRows=1 if has_header else 0)
            style = [
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), side_padding),
                ('RIGHTPADDING', (0, 0), (-1, -1), side_padding),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cccccc')),
            ]
            if has_header:
                style.append(
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#dbe5f1')))
                style.append(('BOTTOMPADDING', (0, 0), (-1, 0), 8))
            t.setStyle(TableStyle(style))
            # A repeated header re-emits on every continuation page, so each
            # body row must fit in the space it leaves (`frame_height -
            # header_h`). A text row splits fine under the header (splitInRow
            # carries the overflow) but cannot shrink below one line, and an
            # image is atomic and cannot shrink at all - either would raise
            # LayoutError if it can't fit under the repeated header. Drop the
            # repeat when the header leaves less than one body line (this also
            # covers a header taller than the whole frame), or when a body
            # row's tallest ATOMIC image can't fit. Keying off total row height
            # would needlessly drop the header for a tall *text* table.
            t.wrap(sum(col_widths), frame_height)
            if has_header and t._rowHeights:
                avail = frame_height - t._rowHeights[0]
                one_body_line = table_cell_style.leading + 8  # + cell padding

                def _atomic_floor(cells):
                    """Smallest height a row can shrink to when split: its
                    tallest atomic image plus cell padding, or 0 when the row
                    is all splittable text (bounded separately by one line)."""
                    tall = 0.0
                    for cell_flow in cells:
                        for f in cell_flow:
                            if isinstance(f, RLImage):
                                tall = max(tall, f.drawHeight)
                    return tall + 8 if tall else 0.0

                if avail < one_body_line or any(
                    _atomic_floor(r) > avail for r in wrapped[1:]
                ):
                    t.repeatRows = 0

            # Tear a row across pages ONLY when it cannot fit whole anywhere.
            # splitInRow is per-table, so leaving it on breaks any row that
            # merely runs past the space left on the current page - a cell's
            # caption stranded on one page and its image on the next - when
            # moving the whole row to the next page would have fit. The limit is
            # the space a row actually gets: a repeated header re-emits above
            # every continuation, so it comes off the top. Decided after the
            # repeat guard above so it sees the final repeatRows.
            row_heights = getattr(t, '_rowHeights', None)
            if row_heights:
                if t.repeatRows:
                    limit = frame_height - row_heights[0]
                    oversized = any(h > limit for h in row_heights[1:])
                else:
                    limit = frame_height
                    oversized = any(h > limit for h in row_heights)
                t.splitInRow = 1 if oversized else 0
            return [t, Spacer(1, 0.15 * inch)]

        def process_image(drawing_element, doc, max_width=None, max_height=None):
            """Extract image from drawing element and return reportlab Image.

            Scales to ``max_width`` then, if still too tall, ``max_height``,
            aspect preserved. Both default to the printable frame; a table cell
            passes its own column width so a cell image fits its column.
            """
            if max_width is None:
                max_width = frame_width
            if max_height is None:
                max_height = frame_height
            try:
                # Navigate to blip element containing image reference
                blip = drawing_element.find('.//' + qn('a:blip'))
                if blip is None:
                    return None

                # Get the relationship ID
                rId = blip.get(qn('r:embed'))
                if not rId:
                    return None

                # Get image data from document parts
                image_part = doc.part.related_parts.get(rId)
                if not image_part:
                    return None

                # Create reportlab Image from bytes
                img_buffer = io.BytesIO(image_part.blob)
                img = RLImage(img_buffer)

                # Scale to the width, then - if that leaves it taller than the
                # height - down to the height, aspect preserved. A tall mermaid
                # diagram at full width would otherwise run off the bottom.
                if max_width > 0 and img.drawWidth > max_width:
                    scale = max_width / img.drawWidth
                    img.drawWidth = max_width
                    img.drawHeight = img.drawHeight * scale
                if max_height > 0 and img.drawHeight > max_height:
                    scale = max_height / img.drawHeight
                    img.drawHeight = max_height
                    img.drawWidth = img.drawWidth * scale

                # Left-align image
                img.hAlign = 'LEFT'

                return img
            except Exception:
                return None

        def add_to_story(result):
            """Add paragraph result to story, handling lists or single elements."""
            if isinstance(result, list):
                story.extend(result)
            else:
                story.append(result)

        def build_blockquote_callout(paras, bq):
            """One callout for a run of consecutive blockquote paragraphs, so a
            multi-paragraph quote renders as a single shaded box with an
            unbroken left bar rather than a stack of separate boxes (Word
            merges them the same way in the DOCX)."""
            indent_pts, bar_hex, shd_hex = bq
            flow = []
            for p in paras:
                markup = ''.join(format_run(run) for run in self.docx_paragraph_runs(p)).strip()
                flow.append(Paragraph(markup or '&nbsp;', callout_body_style))
            return make_callout(flow, bar_hex, shd_hex,
                                left_pad=12 + max(0.0, indent_pts), trailing=0.12)

        # Iterate through body elements in document order
        body_elements = list(doc.element.body)
        i = 0
        while i < len(body_elements):
            element = body_elements[i]
            if element.tag == qn('w:p'):  # Paragraph
                para = DocxParagraph(element, doc)
                drawings = element.findall('.//' + qn('w:drawing'))

                # Group consecutive blockquote paragraphs (without images) into
                # a single callout.
                bq = blockquote_info(para) if not drawings else None
                if bq is not None:
                    group = [para]
                    j = i + 1
                    while (j < len(body_elements)
                           and body_elements[j].tag == qn('w:p')
                           and not body_elements[j].findall('.//' + qn('w:drawing'))):
                        nxt = DocxParagraph(body_elements[j], doc)
                        if blockquote_info(nxt) is None:
                            break
                        group.append(nxt)
                        j += 1
                    add_to_story(build_blockquote_callout(group, bq))
                    i = j
                    continue

                if drawings:
                    # Process paragraph text first (if any)
                    if para.text.strip():
                        add_to_story(process_paragraph(para))
                    # Then add images
                    for drawing in drawings:
                        img = process_image(drawing, doc)
                        if img:
                            story.append(img)
                            story.append(Spacer(1, 0.1 * inch))
                else:
                    add_to_story(process_paragraph(para))
            elif element.tag == qn('w:tbl'):  # Table
                tbl = DocxTable(element, doc)
                story.extend(process_table(tbl))
            i += 1

        # Build the PDF
        if not story:
            story.append(Paragraph("Document appears to be empty.", normal_style))

        pdf_doc.build(story)

        # Get the PDF bytes
        pdf_bytes = pdf_buffer.getvalue()
        pdf_buffer.close()

        return pdf_bytes


class ExportPdfHandler(ExportHandlerBase):
    """Handler for exporting markdown to PDF via DOCX intermediate."""

    @tornado.web.authenticated
    async def post(self):
        try:
            data = json.loads(self.request.body)
            relative_path = data.get('path')
            mermaid_diagrams = data.get('mermaidDiagrams', [])
            svg_pixel_width = self.svg_pixel_width(data.get("svgPixelWidth"))
            show_alert_labels = data.get('showAlertLabels', False)
            base_pt = self.font_size_pt(data.get('exportFontSize'))
            math_pixel_width = self.math_pixel_width(data.get('mathPixelWidth'))
            # DOCX/PDF SVG + Mermaid rasterization theme. Defaults to light
            # (Word docs are usually printed). 'auto' is resolved by the
            # frontend to a concrete light/dark before the request.
            docx_theme = data.get('docxTheme', data.get('htmlTheme', 'light'))
            svg_color_scheme = self.color_scheme_for(docx_theme)

            if not relative_path:
                self.set_status(400)
                self.finish(json.dumps({'error': 'No path provided'}))
                return

            file_path = self.get_absolute_path(relative_path)

            if not file_path.exists():
                self.set_status(404)
                self.finish(json.dumps({'error': 'File not found'}))
                return

            content = self.read_markdown_file(file_path)
            # Diagrams first, on the document as written - see the ordering
            # rule on render_mermaid_server_side.
            content, unrendered = self.replace_mermaid_with_images(
                content, mermaid_diagrams, use_png=True)
            content, export_warnings = await self.render_mermaid_server_side(
                content, color_scheme=svg_color_scheme,
                png_width=svg_pixel_width, diagram_indices=unrendered)
            content = self.preprocess_task_lists(content)
            content = self.preprocess_github_alerts(content, show_labels=show_alert_labels)
            content = self.replace_math_with_images(content, width=math_pixel_width)
            content = self.embed_images_as_base64(content, file_path.parent)
            # After every pass that rewrites the source and before the
            # format's own code-block handling: the pass measures what the
            # converter will render, so it must see the text the converter
            # gets. Measured on the raw source, a display-math item that the
            # math pass rewrites to a bare `- ` line was certified as a list
            # and rendered as the setext underline of the item above it.
            content = self.normalize_list_indentation(content)

            # Extract code blocks for PDF rendering (before DOCX conversion)
            content, code_blocks = self.extract_code_blocks(content)

            html = self.markdown_to_html(content, file_path.stem,
                                         base_pt=base_pt)

            # Step 1: Create DOCX in memory (same logic as DOCX export)
            from htmldocx import HtmlToDocx
            from docx import Document
            from docx.shared import Inches

            body_match = re.search(r'<body>(.*?)</body>', html, re.DOTALL)
            body_html = body_match.group(1) if body_match else html

            # Fix loose-list bullets and tag blockquotes before htmldocx
            body_html = self.restructure_html_for_docx(body_html)
            body_html = self.inject_anchor_markers(body_html)

            with tempfile.TemporaryDirectory() as temp_dir:
                body_html = await self.extract_data_uri_images(
                    body_html, temp_dir,
                    convert_svg=True,
                    svg_pixel_width=svg_pixel_width,
                    color_scheme=svg_color_scheme,
                )

                document = Document()

                for section in document.sections:
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.5)
                    section.left_margin = Inches(0.5)
                    section.right_margin = Inches(0.5)

                parser = HtmlToDocx()
                parser.add_html_to_document(body_html, document)

                self.apply_anchor_bookmarks(document)
                # One numbering instance per ordered list, mark removed
                self.restart_docx_list_numbering(document)
                # Style and de-marker blockquotes (also strips the marker so it
                # never leaks into the PDF text rebuild below)
                self.style_docx_blockquotes(document)
                # Apply true run shading to coloured pills (background spans)
                self.style_docx_color_runs(document)

                # Style GitHub alert boxes with colored borders and shading
                self.style_docx_alert_boxes(document, show_labels=show_alert_labels)

                # No table styling or fitting here: this DOCX is only an
                # intermediate for convert_docx_to_pdf, which reads cell text
                # and applies its own style and column widths

                while document.paragraphs and not document.paragraphs[0].text.strip():
                    p_elem = document.paragraphs[0]._element
                    # Keep paragraphs that contain images (drawings or VML)
                    from docx.oxml.ns import qn as _qn
                    has_image = bool(
                        p_elem.findall('.//' + _qn('w:drawing')) or
                        p_elem.findall('.//{urn:schemas-microsoft-com:vml}imagedata')
                    )
                    if has_image:
                        break
                    p_elem.getparent().remove(p_elem)

                # Remove empty paragraphs before images
                self.remove_empty_paragraphs_before_images(document)

                docx_buffer = io.BytesIO()
                document.save(docx_buffer)
                docx_bytes = docx_buffer.getvalue()

            # Step 2: Convert DOCX to PDF using reportlab (with code blocks)
            pdf_content = self.convert_docx_to_pdf(docx_bytes, code_blocks,
                                                   base_pt=base_pt)

            self.set_export_warnings(export_warnings)
            self.set_header('Content-Type', 'application/pdf')
            self.set_attachment_filename(f'{file_path.stem}.pdf')
            self.finish(pdf_content)

        except ChromiumUnavailableError as e:
            self.set_status(503)
            self.finish(json.dumps({
                'error': str(e),
                'errorCode': 'CHROMIUM_UNAVAILABLE',
                'message': (
                    'Chromium is required to render embedded SVG images. '
                    f'Run: {CHROMIUM_INSTALL_COMMAND}'
                ),
            }))
        except ImportError as e:
            self.set_status(500)
            self.finish(json.dumps({
                'error': f'Missing dependency: {e}. Install with: pip install reportlab python-docx htmldocx markdown'
            }))
        except Exception as e:
            self.set_status(500)
            self.finish(json.dumps({'error': str(e)}))


class ExportDocxHandler(ExportHandlerBase):
    """Handler for exporting markdown to DOCX."""

    @tornado.web.authenticated
    async def post(self):
        try:
            data = json.loads(self.request.body)
            relative_path = data.get('path')
            mermaid_diagrams = data.get('mermaidDiagrams', [])
            svg_pixel_width = self.svg_pixel_width(data.get("svgPixelWidth"))
            show_alert_labels = data.get('showAlertLabels', False)
            base_pt = self.font_size_pt(data.get('exportFontSize'))
            # DOCX/PDF SVG + Mermaid rasterization theme. Defaults to light
            # (Word docs are usually printed). 'auto' is resolved by the
            # frontend to a concrete light/dark before the request.
            docx_theme = data.get('docxTheme', data.get('htmlTheme', 'light'))
            svg_color_scheme = self.color_scheme_for(docx_theme)

            if not relative_path:
                self.set_status(400)
                self.finish(json.dumps({'error': 'No path provided'}))
                return

            file_path = self.get_absolute_path(relative_path)

            if not file_path.exists():
                self.set_status(404)
                self.finish(json.dumps({'error': 'File not found'}))
                return

            content = self.read_markdown_file(file_path)
            # Diagrams first, on the document as written - see the ordering
            # rule on render_mermaid_server_side. Use PNG for DOCX (better
            # Word compatibility).
            content, unrendered = self.replace_mermaid_with_images(
                content, mermaid_diagrams, use_png=True)
            content, export_warnings = await self.render_mermaid_server_side(
                content, color_scheme=svg_color_scheme,
                png_width=svg_pixel_width, diagram_indices=unrendered)
            content = self.preprocess_task_lists(content)
            content = self.preprocess_github_alerts(content, show_labels=show_alert_labels)
            # Use OMML markers for DOCX (native Word equations)
            content, inline_math, display_math = self.replace_math_with_markers(content)
            content = self.embed_images_as_base64(content, file_path.parent)
            # After every pass that rewrites the source and before the
            # format's own code-block handling: the pass measures what the
            # converter will render, so it must see the text the converter
            # gets. Measured on the raw source, a display-math item that the
            # math pass rewrites to a bare `- ` line was certified as a list
            # and rendered as the setext underline of the item above it.
            content = self.normalize_list_indentation(content)
            # Highlight code blocks with inline styles for DOCX
            content = self.highlight_code_blocks(content, use_inline_styles=True)
            html = self.markdown_to_html(content, file_path.stem,
                                         base_pt=base_pt)

            from htmldocx import HtmlToDocx
            from docx import Document
            from docx.shared import Emu, Inches, Twips

            # Extract just the body content for DOCX conversion
            body_match = re.search(r'<body>(.*?)</body>', html, re.DOTALL)
            body_html = body_match.group(1) if body_match else html

            # Fix loose-list bullets and tag blockquotes before htmldocx
            body_html = self.restructure_html_for_docx(body_html)

            # Inject bookmark sentinels for every id="..." anchor so Word
            # internal links (#anchor) resolve after htmldocx strips them.
            body_html = self.inject_anchor_markers(body_html)

            # Use temp directory for images (htmldocx can't handle data URIs)
            with tempfile.TemporaryDirectory() as temp_dir:
                body_html = await self.extract_data_uri_images(
                    body_html, temp_dir,
                    convert_svg=True,
                    svg_pixel_width=svg_pixel_width,
                    color_scheme=svg_color_scheme,
                )

                document = Document()

                # Set document margins (0.5 inch)
                for section in document.sections:
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.5)
                    section.left_margin = Inches(0.5)
                    section.right_margin = Inches(0.5)

                parser = HtmlToDocx()
                parser.add_html_to_document(body_html, document)

                # Insert native OMML equations replacing markers
                self.merge_inline_math_omml(document, inline_math, display_math)

                # Convert anchor markers to bookmarks and rewrite hash-prefix
                # hyperlinks from external to internal links.
                self.apply_anchor_bookmarks(document)
                # One numbering instance per ordered list, mark removed
                self.restart_docx_list_numbering(document)

                # Style blockquotes (left bar, indent, shading) and strip marker
                self.style_docx_blockquotes(document)
                # Apply true run shading to coloured pills (background spans)
                self.style_docx_color_runs(document)

                # Style GitHub alert boxes with colored borders and shading
                alert_tables = self.style_docx_alert_boxes(
                    document, show_labels=show_alert_labels)

                # Name a font on every glyph the body face cannot draw. After
                # the marker passes above, which read a sentinel out of a run's
                # text and would miss it once the run is split; before the
                # strip below, which is the other half of the font decision.
                self.style_docx_symbol_runs(document)

                # Drop Courier from runs holding a character it has no glyph
                # for. It runs after the split, so a symbol now sits in its own
                # run and its pure-ASCII neighbours keep their monospace -
                # inline code holding one star used to lose Courier entirely.
                self.strip_monospace_from_unicode_runs(document)

                # Style tables: banded rows (pale blue), no first column emphasis
                # Usable page area - one source for both the table grid below
                # and the image scaling further down
                section = document.sections[0]
                page_width = Emu(section.page_width - section.left_margin
                                 - section.right_margin)
                page_height = Emu(section.page_height - section.top_margin
                                  - section.bottom_margin)
                self.apply_docx_font_size(document, base_pt)
                for table in document.tables:
                    if table._tbl in alert_tables:
                        continue
                    self.style_docx_table(table)
                    self.fit_docx_table_to_page(table, page_width.twips, base_pt)

                # Remove empty paragraphs at the beginning
                while document.paragraphs and not document.paragraphs[0].text.strip():
                    p_elem = document.paragraphs[0]._element
                    # Keep paragraphs that contain images (drawings or VML)
                    from docx.oxml.ns import qn as _qn
                    has_image = bool(
                        p_elem.findall('.//' + _qn('w:drawing')) or
                        p_elem.findall('.//{urn:schemas-microsoft-com:vml}imagedata')
                    )
                    if has_image:
                        break
                    p_elem.getparent().remove(p_elem)

                # Remove empty paragraphs before images
                self.remove_empty_paragraphs_before_images(document)

                from docx.oxml.ns import qn as _qn

                def container_width(inline_el):
                    """Usable width for an inline shape - its table cell width
                    if nested in a table, otherwise the full page width.

                    Images in a multi-column table cell must be scaled to the
                    cell, not the page; a page-width image overflows its cell
                    and Word clips it at the cell boundary.
                    """
                    node = inline_el.getparent()
                    tc = None
                    while node is not None:
                        if node.tag == _qn('w:tc'):
                            tc = node
                            break
                        node = node.getparent()
                    if tc is None:
                        return page_width
                    tbl = tc.getparent()
                    while tbl is not None and tbl.tag != _qn('w:tbl'):
                        tbl = tbl.getparent()
                    cell_width = None
                    if tbl is not None:
                        grid = tbl.find(_qn('w:tblGrid'))
                        cols = grid.findall(_qn('w:gridCol')) if grid is not None else []
                        # Only a fitted table's grid states the real column
                        # widths. Elsewhere (the alert boxes, whose single
                        # gridCol is decorative next to their 100% tblW) the
                        # grid is fiction, so fall back to an equal share.
                        tbl_pr = tbl.find(_qn('w:tblPr'))
                        fitted = (tbl_pr is not None
                                  and tbl_pr.find(_qn('w:tblLayout')) is not None)
                        row = tc.getparent()
                        index = list(row.findall(_qn('w:tc'))).index(tc)
                        if fitted and index < len(cols):
                            # A fitted grid width is an outer width, so the
                            # allowance due is exactly the cell margin baked
                            # into it. No lower bound here: the column really
                            # is that narrow, and an image wider than its cell
                            # would push the fixed grid past the margin
                            outer = int(cols[index].get(_qn('w:w')) or 0)
                            return Twips(max(
                                1, outer - self.DOCX_CELL_MARGIN_TWIPS))
                        elif cols:
                            cell_width = int(page_width / len(cols))
                    if cell_width is None:
                        cell_width = page_width
                    # Per-cell content width, less the cell margin allowance
                    return max(Inches(0.5),
                               cell_width - Twips(self.DOCX_CELL_MARGIN_TWIPS))

                # Process images: scale every image down to fit its container
                # (page width, or table cell width when nested in a table).
                for shape in document.inline_shapes:
                    orig_width = shape.width
                    orig_height = shape.height
                    max_width = container_width(shape._inline)

                    ratio = 1.0
                    if orig_width > max_width:
                        ratio = min(ratio, max_width / orig_width)
                    if orig_height > page_height:
                        ratio = min(ratio, page_height / orig_height)

                    if ratio < 1.0:
                        # A column no wider than its own cell margin scales the
                        # image to a fraction of an EMU; truncating either side
                        # to zero writes a degenerate wp:extent
                        shape.width = max(1, int(orig_width * ratio))
                        shape.height = max(1, int(orig_height * ratio))

                docx_buffer = io.BytesIO()
                document.save(docx_buffer)
                docx_content = docx_buffer.getvalue()

            self.set_export_warnings(export_warnings)
            self.set_header('Content-Type',
                          'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.set_attachment_filename(f'{file_path.stem}.docx')
            self.finish(docx_content)

        except ChromiumUnavailableError as e:
            self.set_status(503)
            self.finish(json.dumps({
                'error': str(e),
                'errorCode': 'CHROMIUM_UNAVAILABLE',
                'message': (
                    'Chromium is required to render embedded SVG images. '
                    f'Run: {CHROMIUM_INSTALL_COMMAND}'
                ),
            }))
        except ImportError as e:
            self.set_status(500)
            self.finish(json.dumps({
                'error': f'Missing dependency: {e}. Install with: pip install python-docx htmldocx markdown'
            }))
        except Exception as e:
            self.set_status(500)
            self.finish(json.dumps({'error': str(e)}))


class ExportHtmlHandler(ExportHandlerBase):
    """Handler for exporting markdown to HTML."""

    @tornado.web.authenticated
    async def post(self):
        try:
            data = json.loads(self.request.body)
            relative_path = data.get('path')
            mermaid_diagrams = data.get('mermaidDiagrams', [])
            show_alert_labels = data.get('showAlertLabels', False)
            base_pt = self.font_size_pt(data.get('exportFontSize'))
            html_theme = data.get('htmlTheme', 'light')
            dark_background = data.get('htmlDarkBackground', '#111111')
            light_background = data.get('htmlLightBackground', '#ffffff')
            mermaid_color_scheme = self.color_scheme_for(html_theme)

            if not relative_path:
                self.set_status(400)
                self.finish(json.dumps({'error': 'No path provided'}))
                return

            file_path = self.get_absolute_path(relative_path)

            if not file_path.exists():
                self.set_status(404)
                self.finish(json.dumps({'error': 'File not found'}))
                return

            content = self.read_markdown_file(file_path)
            # Diagrams first, on the document as written - see the ordering
            # rule on render_mermaid_server_side.
            content, unrendered = self.replace_mermaid_with_images(
                content, mermaid_diagrams)
            content, export_warnings = await self.render_mermaid_server_side(
                content, color_scheme=mermaid_color_scheme,
                png_width=None, diagram_indices=unrendered)
            content = self.preprocess_task_lists(content)
            content = self.preprocess_github_alerts(content, show_labels=show_alert_labels)
            content = self.embed_images_as_base64(content, file_path.parent)
            # After every pass that rewrites the source and before the
            # format's own code-block handling: the pass measures what the
            # converter will render, so it must see the text the converter
            # gets. Measured on the raw source, a display-math item that the
            # math pass rewrites to a bare `- ` line was certified as a list
            # and rendered as the setext underline of the item above it.
            content = self.normalize_list_indentation(content)
            html = self.markdown_to_html(
                content, file_path.stem, math_support=True,
                theme=html_theme,
                dark_background=dark_background,
                light_background=light_background,
                base_pt=base_pt
            )

            # Style GitHub alert boxes with colored borders and shading
            html = self.style_html_alert_boxes(html, show_labels=show_alert_labels)
            html = self.drop_empty_table_headers(html)
            html = self.wrap_html_tables(html)

            self.set_export_warnings(export_warnings)
            self.set_header('Content-Type', 'text/html; charset=utf-8')
            self.set_attachment_filename(f'{file_path.stem}.html')
            self.finish(html.encode('utf-8'))

        except ImportError as e:
            self.set_status(500)
            self.finish(json.dumps({
                'error': f'Missing dependency: {e}. Install with: pip install markdown'
            }))
        except Exception as e:
            self.set_status(500)
            self.finish(json.dumps({'error': str(e)}))


def setup_route_handlers(web_app):
    """Register all route handlers for the extension."""
    host_pattern = ".*$"
    base_url = web_app.settings["base_url"]
    namespace = "jupyterlab-export-markdown-extension"

    handlers = [
        (url_path_join(base_url, namespace, "export/pdf"), ExportPdfHandler),
        (url_path_join(base_url, namespace, "export/docx"), ExportDocxHandler),
        (url_path_join(base_url, namespace, "export/html"), ExportHtmlHandler),
    ]

    web_app.add_handlers(host_pattern, handlers)
